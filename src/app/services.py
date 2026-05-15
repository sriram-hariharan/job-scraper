from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from types import SimpleNamespace
from typing import Any, Dict, List, Tuple
from urllib.parse import urlparse, urlunparse
import csv
import hashlib
import json
import logging
import threading
import os
import re
import subprocess
import sys
import tempfile
import shutil

from src.pipeline.post_run_notification import DEFAULT_NOTIFICATION_RECORDS_DIR

from src.config.consts import (
    _ALLOWED_REWRITE_REVIEW_STATES,
    _SCAN_DIMENSION_DISPLAY_LABELS,
    _SCAN_GROUP_LABELS,
    _SCAN_RECRUITER_TIPS_SIGNAL_KEYS,
    _SCAN_SEARCHABILITY_BULLET_WORD_LIMIT,
    _SCAN_SEARCHABILITY_DIMENSIONS,
    _SCAN_SEARCHABILITY_MIN_VISIBLE_TECH_TERMS,
    _SCAN_SIGNAL_DISPLAY_OVERRIDES,
    _SCAN_SKILLS_SIGNAL_KEYS,
    _SCAN_TITLE_SIGNAL_PATTERNS,
    DOMAIN_SIGNAL_PATTERNS,
)
from src.config.settings import (
    ACTIVE_APPLICATION_PLANNING_OUTPUT_DIR,
    SCHEDULER_RUN_HISTORY_PATH,
)
from src.matching.dimensions import get_dimension_weights
from src.matching.job_adapter import build_job_evidence
from src.matching.models import MatchPrefilterResult
from src.matching.prefilter import run_prefilter
from src.matching.scorer import score_resume_job_match
from src.tailoring.signal_family_matcher import (
    scan_issue_group_id_for_signals as _signal_family_group_id_for_signals,
    scan_issue_skill_type as _signal_family_skill_type,
    scan_issue_term_family as _signal_family_term_family,
)
from src.resume.evidence_builder import build_resume_evidence
from src.resume.models import ResumeDocument, ResumeExperienceEntry
from src.storage.application_actions.store import (
    application_action_db_row,
    insert_application_action_row_to_postgres,
)
from src.storage.application_actions.read_postgres import (
    get_latest_application_actions_rows,
)
from src.storage.patch_selections.store import (
    insert_patch_selection_row_to_postgres,
)
from src.storage.patch_selections.read_postgres import (
    get_patch_selections_postgres_status_payload,
)
from src.storage.operator_decisions.store import (
    insert_operator_decision_row_to_postgres,
    operator_decision_db_row,
)
from src.storage.operator_decisions.read_postgres import (
    get_operator_decisions_postgres_status_payload,
)
from src.storage.notification_state.store import (
    insert_notification_state_row_to_postgres,
)
from src.storage.notification_state.read_postgres import (
    get_notification_state_postgres_status_payload,
)
from src.storage.saved_scans.store import (
    insert_saved_scan_row_to_postgres,
    saved_scan_db_row,
)
from src.storage.saved_scans.read_postgres import (
    delete_saved_scan_postgres_payload,
    get_saved_scan_postgres_payload,
    get_saved_scans_postgres_payload,
    save_saved_scan_draft_postgres_payload,
)
from src.storage.profile_resumes.store import (
    delete_profile_resume_postgres_payload,
    get_profile_resume_blob_postgres_payload,
    get_profile_resumes_postgres_payload,
    upsert_profile_resume_postgres_payload,
)
from src.storage.auth.read_postgres import (
    delete_non_admin_auth_user_postgres_payload,
    get_non_admin_auth_users_postgres_payload,
    update_non_admin_auth_user_access_postgres_payload,
)
from src.storage.user_pipeline.read_postgres import (
    get_latest_user_pipeline_run_postgres_payload,
    get_user_pipeline_artifacts_postgres_payload,
    get_user_pipeline_runs_postgres_payload,
)
from src.storage.user_pipeline.store import (
    get_user_pipeline_active_run_postgres_payload,
    release_user_pipeline_active_run_postgres_payload,
    reserve_user_pipeline_active_run_postgres_payload,
    upsert_user_pipeline_artifact_postgres_payload,
    upsert_user_pipeline_run_postgres_payload,
    promote_user_seen_jobs_staging_postgres_payload,
    clear_user_seen_jobs_staging_postgres_payload,
)
from src.storage.redis_locks import (
    RedisLockHandle,
    acquire_redis_lock,
    redis_lock_key,
    release_redis_lock,
)
from src.pipeline.scheduler import (
    DEFAULT_LAUNCHD_AGENT_DIR,
    DEFAULT_LAUNCHD_INTERVAL_SECONDS,
    DEFAULT_LAUNCHD_LABEL_PREFIX,
    DEFAULT_LAUNCHD_LOG_DIR,
    DEFAULT_LAUNCHD_OUT_DIR,
    DEFAULT_LAUNCHD_TARGET,
    build_scheduled_job_command,
    build_scheduler_launchd_plist_payload,
    get_scheduled_job_definition,
    get_scheduled_job_definitions,
    get_scheduler_launchd_agent_status,
)
from src.storage.scheduler.contract import (
    scheduler_init_sql_generation_payload,
    scheduler_init_sql_payload,
    scheduler_job_definition_seed_rows,
    scheduler_postgres_table_specs,
    scheduler_schema_sql_payload,
    scheduler_seed_sql_generation_payload,
    scheduler_seed_sql_payload,
)

from src.storage.scheduler.read_postgres import (
    get_scheduler_postgres_status_payload,
)
from src.storage.scheduler.contract import (
    scheduler_contract_health_payload,
)
from src.tailoring.score_utils import (
    score_delta_to_points,
    score_to_points,
)

logger = logging.getLogger(__name__)

DEFAULT_OUTPUT_DIR = Path(
    os.environ.get("APPLICATION_PLANNING_OUTPUT_DIR", ACTIVE_APPLICATION_PLANNING_OUTPUT_DIR)
).expanduser()
DEFAULT_CORPUS_PATH = Path("data/rag/job_corpus.jsonl")
DEFAULT_PIPELINE_LOG_PATH = DEFAULT_OUTPUT_DIR / "live_pipeline_run.log"
DEFAULT_PIPELINE_STATUS_PATH = DEFAULT_OUTPUT_DIR / "live_pipeline_status.json"
DEFAULT_PROFILE_RESUME_DIR = Path(
    os.environ.get("RESUME_DIR", "data/profile_resumes")
).expanduser()
DEFAULT_SCAN_UPLOAD_DIR = Path(
    os.environ.get("SCAN_UPLOAD_DIR", "data/scan_uploads")
).expanduser()
DEFAULT_PIPELINE_SCRATCH_DIR = Path(
    os.environ.get("JOB_STACK_PIPELINE_SCRATCH_DIR", "tmp/pipeline_runs")
).expanduser()
DEFAULT_SCHEDULER_RUN_HISTORY_PATH = Path(SCHEDULER_RUN_HISTORY_PATH)
DEFAULT_NOTIFICATION_RECORDS_DIR = Path(DEFAULT_NOTIFICATION_RECORDS_DIR)

_PIPELINE_RUN_STATE: Dict[str, Any] = {
    "process": None,
    "log_handle": None,
    "status": "idle",
    "started_at": "",
    "finished_at": "",
    "return_code": None,
    "command": [],
    "output_dir": str(DEFAULT_OUTPUT_DIR),
    "log_path": str(DEFAULT_PIPELINE_LOG_PATH),
    "status_path": str(DEFAULT_PIPELINE_STATUS_PATH),
    "run_id": "",
    "child_pid": None,
    "error": "",
    "owner_user_id": "",
}

_PIPELINE_ACTIVE_RUNS: Dict[str, Dict[str, Any]] = {}
_PIPELINE_ACTIVE_RUNS_LOCK = threading.RLock()

_PIPELINE_ARTIFACT_INGESTED_RUN_KEYS: set[str] = set()
_PIPELINE_ARTIFACT_MAX_BYTES = int(
    os.environ.get("JOB_STACK_PIPELINE_ARTIFACT_MAX_BYTES", "5242880") or "5242880"
)

ALLOWED_APPLICATION_STATUSES = {
    "OPENED",
    "APPLIED",
    "SAVED",
    "NOT_APPLIED",
    "DISMISSED",
}

APPLICATION_ACTION_OVERLAY_FIELDS = [
    "application_status",
    "application_label",
    "is_applied",
]

ALLOWED_OPERATOR_DECISIONS = {
    "SELECT_RESUME",
}

_RESUME_PREVIEW_PATH_CACHE: Dict[str, str] = {}

_PATCH_SELECTION_OVERLAY_CACHE: Dict[str, Any] = {
    "version": 0,
    "data": None,
}
_JOB_METADATA_OVERLAY_CACHE: Dict[Tuple[str, int, int], Dict[str, Dict[str, Any]]] = {}
_TAILORING_WORKSPACE_BUTTON_STATE_CACHE: Dict[
    Tuple[str, int, int, int, int],
    Dict[str, Any],
] = {}

def _invalidate_patch_selection_overlay_cache() -> None:
    _PATCH_SELECTION_OVERLAY_CACHE["version"] = int(
        _PATCH_SELECTION_OVERLAY_CACHE.get("version", 0) or 0
    ) + 1
    _PATCH_SELECTION_OVERLAY_CACHE["data"] = None


def _patch_selection_overlay_cache_version() -> int:
    return int(_PATCH_SELECTION_OVERLAY_CACHE.get("version", 0) or 0)


def _job_metadata_overlay_cache_key(job_corpus: Path) -> Tuple[str, int, int]:
    resolved = Path(job_corpus).expanduser().resolve()
    if not resolved.exists():
        return (str(resolved), 0, 0)

    stat = resolved.stat()
    return (str(resolved), int(stat.st_mtime_ns), int(stat.st_size))


def _tailoring_workspace_button_state_cache_key(
    artifact_path: Path,
) -> Tuple[str, int, int, int, int]:
    resolved = artifact_path.expanduser().resolve()
    artifact_stat = resolved.stat()

    packet_path = _infer_packet_json_path_from_tailoring_artifact(resolved)
    packet_mtime_ns = 0
    packet_size = 0

    if packet_path and packet_path.exists() and packet_path.is_file():
        packet_stat = packet_path.stat()
        packet_mtime_ns = int(packet_stat.st_mtime_ns)
        packet_size = int(packet_stat.st_size)

    return (
        str(resolved),
        int(artifact_stat.st_mtime_ns),
        int(artifact_stat.st_size),
        packet_mtime_ns,
        packet_size,
    )

def _safe_owner_dir_name(owner_user_id: str = "") -> str:
    safe = re.sub(r"[^A-Za-z0-9_.-]+", "_", _clean_text(owner_user_id)).strip("._-")
    return safe[:80]


def _get_resume_dir(owner_user_id: str = "") -> Path:
    owner_dir = _safe_owner_dir_name(owner_user_id)
    resume_dir = (
        DEFAULT_PROFILE_RESUME_DIR / "users" / owner_dir
        if owner_dir
        else DEFAULT_PROFILE_RESUME_DIR
    )
    resume_dir.mkdir(parents=True, exist_ok=True)
    return resume_dir


def _get_scan_upload_dir() -> Path:
    upload_dir = DEFAULT_SCAN_UPLOAD_DIR
    upload_dir.mkdir(parents=True, exist_ok=True)
    return upload_dir


def _safe_run_dir_name(run_id: str = "") -> str:
    safe = re.sub(r"[^A-Za-z0-9_.-]+", "_", _clean_text(run_id)).strip("._-")
    return safe[:120] or "run"


def _pipeline_scratch_output_dir(
    *,
    owner_user_id: str,
    run_id: str,
) -> Path:
    owner_dir = _safe_owner_dir_name(owner_user_id) or "anonymous"
    run_dir = _safe_run_dir_name(run_id)
    output_dir = (
        DEFAULT_PIPELINE_SCRATCH_DIR
        / owner_dir
        / run_dir
        / "application_planning"
    )
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


def _sanitize_resume_filename(value: str) -> str:
    raw = str(value or "").strip()
    safe = Path(raw).name

    if not raw or not safe or safe != raw:
        raise ValueError("Invalid resume filename.")

    if Path(safe).suffix.lower() != ".pdf":
        raise ValueError("Only PDF resumes are supported.")

    if re.search(r"[/\\\\]", safe):
        raise ValueError("Invalid resume filename.")

    return safe


def _sanitize_scan_upload_filename(value: str) -> str:
    raw = str(value or "").strip()
    safe = Path(raw).name

    if not raw or not safe or safe != raw:
        raise ValueError("Invalid scan upload filename.")

    suffix = Path(safe).suffix.lower()
    if suffix not in {".pdf", ".docx", ".txt"}:
        raise ValueError("Only PDF, DOCX, and TXT scan uploads are supported.")

    if re.search(r"[/\\\\]", safe):
        raise ValueError("Invalid scan upload filename.")

    return safe


def profile_resumes_payload(
    *,
    owner_user_id: str = "",
) -> Dict[str, Any]:
    payload = get_profile_resumes_postgres_payload(
        owner_user_id=_clean_text(owner_user_id),
        database_url="",
        database_url_env="DATABASE_URL",
        psql_bin="psql",
        print_only=False,
    )
    resumes = list(payload.get("resumes", []) or [])
    return {
        "ok": True,
        "resume_dir": "",
        "storage": "postgres",
        "count": len(resumes),
        "resumes": resumes,
    }


def _public_admin_user_row(user: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "user_id": _clean_text(user.get("user_id")),
        "email": _clean_text(user.get("email")),
        "normalized_email": _clean_text(user.get("normalized_email")),
        "display_name": _clean_text(user.get("display_name")),
        "access_level": _clean_text(user.get("access_level")) or "user",
        "is_active": bool(user.get("is_active", False)),
        "is_admin": bool(user.get("is_admin", False)),
        "created_at": _clean_text(user.get("created_at")),
        "updated_at": _clean_text(user.get("updated_at")),
        "last_login_at": _clean_text(user.get("last_login_at")),
    }


def admin_profile_users_payload(limit: int = 100) -> Dict[str, Any]:
    payload = get_non_admin_auth_users_postgres_payload(
        limit=limit,
        database_url="",
        database_url_env="DATABASE_URL",
        psql_bin="psql",
        print_only=False,
        ensure_schema=True,
    )
    users = [_public_admin_user_row(user) for user in list(payload.get("users", []) or [])]
    return {
        "ok": True,
        "count": len(users),
        "total_count": int(payload.get("total_count", len(users)) or len(users)),
        "users": users,
    }


def admin_profile_update_user_access_payload(user_id: str, is_active: bool) -> Dict[str, Any]:
    payload = update_non_admin_auth_user_access_postgres_payload(
        user_id=_clean_text(user_id),
        is_active=bool(is_active),
        database_url="",
        database_url_env="DATABASE_URL",
        psql_bin="psql",
        print_only=False,
        ensure_schema=True,
    )
    if not bool(payload.get("updated", False)):
        raise ValueError("User not found or admin users cannot be modified here.")

    return {
        "ok": True,
        "updated": True,
        "user": _public_admin_user_row(dict(payload.get("user", {}) or {})),
        "revoked_session_count": int(payload.get("revoked_session_count", 0) or 0),
    }


def admin_profile_delete_user_payload(user_id: str) -> Dict[str, Any]:
    payload = delete_non_admin_auth_user_postgres_payload(
        user_id=_clean_text(user_id),
        database_url="",
        database_url_env="DATABASE_URL",
        psql_bin="psql",
        print_only=False,
        ensure_schema=True,
    )
    if not bool(payload.get("deleted", False)):
        raise ValueError("User not found or admin users cannot be deleted here.")

    return {
        "ok": True,
        "deleted": True,
        "user": _public_admin_user_row(dict(payload.get("user", {}) or {})),
    }


def user_pipeline_gate_payload(
    *,
    owner_user_id: str = "",
) -> Dict[str, Any]:
    owner = _clean_text(owner_user_id)
    if not owner:
        raise ValueError("Authenticated user is required.")

    resumes_payload = profile_resumes_payload(owner_user_id=owner)
    resume_count = int(resumes_payload.get("count", 0) or 0)
    has_resumes = resume_count > 0

    successful_runs_payload = get_user_pipeline_runs_postgres_payload(
        owner_user_id=owner,
        status="succeeded",
        limit=1,
        database_url="",
        database_url_env="DATABASE_URL",
        psql_bin="psql",
        print_only=False,
        ensure_schema=True,
    )
    has_successful_pipeline_run = (
        int(successful_runs_payload.get("total_row_count", 0) or 0) > 0
        or bool(successful_runs_payload.get("rows", []) or [])
    )

    can_run_live_pipeline = has_resumes
    can_delete_seen_data = has_resumes and has_successful_pipeline_run

    return {
        "ok": True,
        "owner_user_id": owner,
        "resume_count": resume_count,
        "has_resumes": has_resumes,
        "requires_resume_upload": not has_resumes,
        "can_run_live_pipeline": can_run_live_pipeline,
        "can_delete_seen_data": can_delete_seen_data,
        "has_successful_pipeline_run": has_successful_pipeline_run,
        "profile_resume_upload_url": "/profile?onboarding=resume_upload",
        "live_pipeline_block_reason": (
            "" if can_run_live_pipeline else "Upload at least one resume before running Live Pipeline."
        ),
        "delete_seen_data_block_reason": (
            ""
            if can_delete_seen_data
            else "Delete seen data is disabled until this user has completed at least one successful Live Pipeline run."
        ),
    }


def profile_upload_resume_payload(
    filename: str,
    file_bytes: bytes,
    *,
    owner_user_id: str = "",
) -> Dict[str, Any]:
    safe_name = _sanitize_resume_filename(filename)

    if not file_bytes:
        raise ValueError("Uploaded file is empty.")

    payload = upsert_profile_resume_postgres_payload(
        owner_user_id=_clean_text(owner_user_id),
        resume_name=safe_name,
        original_filename=safe_name,
        content_type="application/pdf",
        file_bytes=file_bytes,
        database_url="",
        database_url_env="DATABASE_URL",
        psql_bin="psql",
        print_only=False,
    )
    if not bool(payload.get("inserted", False)):
        raise ValueError(f"Resume already exists: {safe_name}")

    return {
        "ok": True,
        "message": "Resume uploaded.",
        "resume": payload.get("resume", {}),
    }


def profile_delete_resume_payload(
    resume_name: str,
    *,
    owner_user_id: str = "",
) -> Dict[str, Any]:
    safe_name = _sanitize_resume_filename(resume_name)
    payload = delete_profile_resume_postgres_payload(
        owner_user_id=_clean_text(owner_user_id),
        resume_name=safe_name,
        database_url="",
        database_url_env="DATABASE_URL",
        psql_bin="psql",
        print_only=False,
    )
    if not bool(payload.get("deleted", False)):
        raise ValueError(f"Resume not found: {safe_name}")

    return {
        "ok": True,
        "message": "Resume deleted.",
        "resume_name": safe_name,
    }


def profile_resume_file_payload(
    resume_name: str,
    *,
    owner_user_id: str = "",
) -> Dict[str, Any]:
    safe_name = _sanitize_resume_filename(resume_name)
    payload = get_profile_resume_blob_postgres_payload(
        owner_user_id=_clean_text(owner_user_id),
        resume_name=safe_name,
        database_url="",
        database_url_env="DATABASE_URL",
        psql_bin="psql",
        print_only=False,
    )
    file_bytes = bytes(payload.get("file_bytes", b"") or b"")
    if not file_bytes:
        raise ValueError(f"Resume not found: {safe_name}")

    resume = dict(payload.get("resume", {}) or {})
    return {
        "ok": True,
        "resume_name": safe_name,
        "content_type": _clean_text(resume.get("content_type")) or "application/pdf",
        "size_bytes": len(file_bytes),
        "file_bytes": file_bytes,
        "resume": resume,
    }


def _materialize_profile_resume_temp_path(
    resume_name: str,
    *,
    owner_user_id: str = "",
) -> Path:
    safe_name = _sanitize_resume_filename(resume_name)
    payload = profile_resume_file_payload(safe_name, owner_user_id=owner_user_id)
    file_bytes = bytes(payload.get("file_bytes", b"") or b"")
    digest = hashlib.sha256(
        (_clean_text(owner_user_id) + "\0" + safe_name).encode("utf-8") + file_bytes
    ).hexdigest()
    temp_dir = Path(tempfile.gettempdir()) / "applylens_profile_resume_previews" / _safe_owner_dir_name(owner_user_id)
    temp_dir.mkdir(parents=True, exist_ok=True)
    temp_path = temp_dir / f"{digest[:24]}_{safe_name}"
    if not temp_path.exists() or temp_path.stat().st_size != len(file_bytes):
        temp_path.write_bytes(file_bytes)
    return temp_path


def _scan_upload_mime_type(filename: str, fallback: str = "") -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return "application/pdf"
    if suffix == ".docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if suffix == ".txt":
        return "text/plain"
    return _clean_text(fallback) or "application/octet-stream"


def _scan_upload_target_path(filename: str, scan_timestamp: str) -> Path:
    upload_dir = _get_scan_upload_dir()
    safe_name = _sanitize_scan_upload_filename(filename)
    stamp = re.sub(r"[^0-9A-Za-z]+", "_", scan_timestamp).strip("_")[:40] or "scan"
    digest = hashlib.sha1(f"{scan_timestamp}:{safe_name}".encode("utf-8")).hexdigest()[:10]
    return upload_dir / f"{stamp}_{digest}_{safe_name}"


def _dual_write_saved_scan_postgres(row: Dict[str, Any]) -> Dict[str, Any]:
    database_url = str(os.environ.get("DATABASE_URL", "") or "").strip()
    if not database_url:
        return {
            "attempted": False,
            "ok": False,
            "skipped": "missing_database_url",
            "table_name": "saved_scans",
        }

    try:
        payload = insert_saved_scan_row_to_postgres(
            record=row,
            database_url="",
            database_url_env="DATABASE_URL",
            psql_bin="psql",
            print_only=False,
            allow_contract_drift=False,
            ensure_schema=True,
        )
        return {
            "attempted": True,
            "ok": True,
            "table_name": payload.get("table_name", "saved_scans"),
            "scan_id": str(payload.get("row", {}).get("scan_id", "") or ""),
            "contract_health_ok": bool(payload.get("contract_health_ok", False)),
            "command_text": str(payload.get("command_text", "") or ""),
        }
    except SystemExit as exc:
        return {
            "attempted": True,
            "ok": False,
            "table_name": "saved_scans",
            "error_type": "SystemExit",
            "error": str(exc),
        }
    except Exception as exc:
        return {
            "attempted": True,
            "ok": False,
            "table_name": "saved_scans",
            "error_type": exc.__class__.__name__,
            "error": str(exc),
        }


def _new_scan_resume_document(
    *,
    resume_name: str,
    resume_file_path: str = "",
    resume_text: str = "",
) -> ResumeDocument:
    safe_name = _clean_text(resume_name) or "New scan resume"
    raw_text = _clean_text(resume_text)
    if not raw_text and resume_file_path:
        try:
            resume_path = Path(resume_file_path)
            suffix = resume_path.suffix.lower()
            if suffix == ".pdf":
                raw_text = _extract_scan_upload_text_from_pdf(resume_path)
            elif suffix == ".docx":
                raw_text = _extract_scan_upload_text_from_docx(resume_path)
            elif suffix == ".txt":
                raw_text = resume_path.read_text(encoding="utf-8", errors="replace")
        except Exception:
            raw_text = ""

    return ResumeDocument(
        resume_id=safe_name,
        resume_name=safe_name,
        path=_clean_text(resume_file_path),
        raw_text=raw_text,
        normalized_text=re.sub(r"\s+", " ", raw_text).strip(),
    )


def _build_new_scan_document_preview(
    *,
    resume_text: str,
    personal_details: Dict[str, str] | None = None,
) -> Dict[str, Any]:
    safe_text = _clean_text(resume_text)
    if not safe_text:
        return {
            "ok": False,
            "pages": [],
            "error_message": "Resume text could not be reconstructed for preview.",
        }

    details = _normalize_workspace_personal_details(personal_details or {})
    raw_lines = [line.rstrip() for line in safe_text.splitlines()]
    lines = [line for line in raw_lines if _clean_text(line)]
    if not lines:
        lines = [safe_text]

    rows: List[Dict[str, Any]] = []
    seen_name = False
    seen_contact = False
    section_headings = {
        "summary",
        "professional summary",
        "experience",
        "professional experience",
        "work experience",
        "projects",
        "skills",
        "technical skills",
        "education",
        "certifications",
    }

    for index, line in enumerate(lines[:90]):
        clean_line = _clean_text(line)
        lower_line = clean_line.lower().strip(":")
        is_bullet = bool(re.match(r"^\s*(?:[-*•]|\d+[.)])\s+", line))
        is_heading = lower_line in section_headings or (
            clean_line.isupper() and len(clean_line.split()) <= 5 and not is_bullet
        )
        presentation_role = ""
        if not seen_name and index <= 2 and details.get("name") and details["name"].lower() in clean_line.lower():
            presentation_role = "header_name"
            seen_name = True
        elif (
            not seen_contact
            and index <= 5
            and (
                "@" in clean_line
                or re.search(r"\d{3}", clean_line)
                or "linkedin" in clean_line.lower()
            )
        ):
            presentation_role = "header_contact"
            seen_contact = True

        rows.append(
            {
                "text": clean_line,
                "display_text": clean_line,
                "is_bullet": is_bullet,
                "is_heading": is_heading,
                "is_section_heading": is_heading,
                "presentation_role": presentation_role,
                "left_indent_pt": 18 if is_bullet else 0,
                "gap_before": 8 if is_heading else 4,
            }
        )

    return {
        "ok": True,
        "source": "new_scan_resume_text",
        "pages": [
            {
                "page_number": 1,
                "rows": rows,
            }
        ],
    }


def _new_scan_job_record(
    *,
    company: str,
    role: str,
    job_description_text: str,
    job_url: str = "",
    job_doc_id: str = "",
) -> Dict[str, Any]:
    digest_source = "|".join([
        _clean_text(company),
        _clean_text(role),
        _clean_text(job_description_text),
        _clean_text(job_url),
    ])
    fallback_doc_id = "new_scan_" + hashlib.sha256(digest_source.encode("utf-8")).hexdigest()[:16]
    return {
        "doc_id": _clean_text(job_doc_id) or fallback_doc_id,
        "job_doc_id": _clean_text(job_doc_id) or fallback_doc_id,
        "company": _clean_text(company),
        "title": _clean_text(role),
        "job_title": _clean_text(role),
        "job_url": _clean_text(job_url),
        "source": "scan_workspace_new_scan",
        "preview": _clean_text(job_description_text),
        "retrieval_text": _clean_text(job_description_text),
        "description_text": _clean_text(job_description_text),
    }


def _new_scan_tailoring_summary(
    *,
    job_evidence: Any,
    match_result: Any,
) -> Dict[str, Any]:
    prefilter = getattr(match_result, "prefilter", None)
    matched_required = list(getattr(prefilter, "matched_required_terms", []) or [])
    matched_preferred = list(getattr(prefilter, "matched_preferred_terms", []) or [])
    matched_any = list(getattr(prefilter, "matched_any_terms", []) or [])
    missing_required = list(getattr(prefilter, "missing_requirements", []) or [])

    required_terms = list(getattr(job_evidence, "required_skills", []) or [])
    preferred_terms = list(getattr(job_evidence, "preferred_skills", []) or [])
    all_terms = _unique_scan_terms(
        list(getattr(job_evidence, "all_skills", []) or [])
        + required_terms
        + preferred_terms
    )
    if not missing_required:
        matched_keys = {
            _scan_issue_canonical_term(term)
            for term in matched_required + matched_preferred + matched_any
        }
        missing_required = [
            term for term in (required_terms or all_terms)
            if _scan_issue_canonical_term(term) not in matched_keys
        ]
    matched_terms = _unique_scan_terms(matched_required + matched_preferred + matched_any)

    return {
        "target_job_title": _clean_text(getattr(job_evidence, "title", "")),
        "job_title": _clean_text(getattr(job_evidence, "title", "")),
        "matched_required": matched_required,
        "matched_preferred": matched_preferred,
        "matched_terms": matched_terms,
        "missing_required": missing_required,
        "missing_preferred": [
            term for term in preferred_terms
            if _scan_issue_canonical_term(term) not in {
                _scan_issue_canonical_term(item)
                for item in matched_preferred + matched_any + missing_required
            }
        ],
        "required_terms": required_terms,
        "preferred_terms": preferred_terms,
        "all_terms": all_terms,
        "missing_terms": missing_required,
        "match_bucket": _clean_text(getattr(match_result, "match_bucket", "")),
    }


def _unique_scan_terms(values: List[Any]) -> List[str]:
    terms: List[str] = []
    seen: set[str] = set()
    for value in values:
        text = _clean_text(value)
        key = _scan_issue_canonical_term(text)
        if not text or not key or key in seen:
            continue
        seen.add(key)
        terms.append(text)
    return terms


def _enrich_new_scan_resume_evidence_for_scoring(resume_evidence: Any) -> Any:
    if resume_evidence is None:
        return resume_evidence

    aggregate_skills = _unique_scan_terms(
        list(getattr(resume_evidence, "skills", []) or [])
        + list(getattr(resume_evidence, "tools", []) or [])
        + list(getattr(resume_evidence, "tooling_signals", []) or [])
    )
    if not aggregate_skills:
        return resume_evidence

    entries = list(getattr(resume_evidence, "experience_entries", []) or [])
    if not entries:
        entries = [
            ResumeExperienceEntry(
                entry_id="new_scan_resume_text",
                entry_index=0,
                bullets=list(getattr(resume_evidence, "quantified_bullets", []) or []),
            )
        ]
        resume_evidence.experience_entries = entries

    first_entry = entries[0]
    first_entry.normalized_skills = _unique_scan_terms(
        list(getattr(first_entry, "normalized_skills", []) or []) + aggregate_skills
    )
    return resume_evidence


def _relaxed_new_scan_prefilter(resume_evidence: Any, job_evidence: Any) -> Any:
    original = run_prefilter(resume_evidence, job_evidence)
    if getattr(original, "passed", False):
        return original

    resume_terms = {
        _scan_issue_canonical_term(term)
        for term in (
            list(getattr(resume_evidence, "skills", []) or [])
            + list(getattr(resume_evidence, "tools", []) or [])
            + list(getattr(resume_evidence, "tooling_signals", []) or [])
            + [
                skill
                for entry in list(getattr(resume_evidence, "experience_entries", []) or [])
                for skill in list(getattr(entry, "normalized_skills", []) or [])
            ]
        )
        if _scan_issue_canonical_term(term)
    }
    required_terms = _unique_scan_terms(list(getattr(job_evidence, "required_skills", []) or []))
    preferred_terms = _unique_scan_terms(list(getattr(job_evidence, "preferred_skills", []) or []))
    all_terms = _unique_scan_terms(
        list(getattr(job_evidence, "all_skills", []) or [])
        + required_terms
        + preferred_terms
    )

    matched_required = [
        term for term in required_terms
        if _scan_issue_canonical_term(term) in resume_terms
    ]
    matched_preferred = [
        term for term in preferred_terms
        if _scan_issue_canonical_term(term) in resume_terms
    ]
    matched_any = [
        term for term in all_terms
        if _scan_issue_canonical_term(term) in resume_terms
    ]
    title_score = float(getattr(original, "best_title_score", 0.0) or 0.0)
    has_meaningful_overlap = (
        len(matched_required) >= 1
        or len(matched_any) >= 2
        or title_score >= 0.25
    )
    if not has_meaningful_overlap:
        return original

    return MatchPrefilterResult(
        passed=True,
        reasons=list(getattr(original, "reasons", []) or []) + [
            "New Scan relaxed deterministic prefilter because extracted resume/JD evidence has meaningful overlap."
        ],
        matched_terms=_unique_scan_terms(
            list(getattr(original, "matched_terms", []) or [])
            + matched_required
            + matched_preferred
            + matched_any
        ),
        missing_requirements=[
            term for term in required_terms
            if _scan_issue_canonical_term(term) not in {
                _scan_issue_canonical_term(item) for item in matched_required
            }
        ],
        best_title_score=title_score,
        best_title=_clean_text(getattr(original, "best_title", "")),
        matched_required_terms=matched_required,
        matched_preferred_terms=matched_preferred,
        matched_any_terms=matched_any,
        matched_required_count=len(matched_required),
        matched_preferred_count=len(matched_preferred),
        matched_any_count=len(matched_any),
    )


def _build_new_scan_review_payload(
    *,
    scan_id: str,
    scan_timestamp: str,
    resume_name: str,
    resume_file_path: str,
    resume_text: str,
    job_record: Dict[str, Any],
    owner_user_id: str = "",
) -> Dict[str, Any]:
    from src.tailoring.llm import tailoring_llm_model_config_payload

    resume_document = _new_scan_resume_document(
        resume_name=resume_name,
        resume_file_path=resume_file_path,
        resume_text=resume_text,
    )
    resume_evidence = _enrich_new_scan_resume_evidence_for_scoring(
        build_resume_evidence(resume_document)
    )
    job_evidence = build_job_evidence(job_record)
    prefilter_result = _relaxed_new_scan_prefilter(resume_evidence, job_evidence)
    match_result = score_resume_job_match(
        resume_evidence,
        job_evidence,
        prefilter=prefilter_result,
    )
    score = max(0, min(100, round(float(getattr(match_result, "final_score", 0.0) or 0.0) * 100)))
    tailoring_summary = _new_scan_tailoring_summary(
        job_evidence=job_evidence,
        match_result=match_result,
    )
    personal_details = _extract_resume_personal_details(
        resume_evidence,
        owner_user_id=owner_user_id,
    )
    if not personal_details.get("linkedin") and resume_file_path:
        try:
            for link_item in _workspace_export_extract_pdf_header_link_items(Path(resume_file_path)):
                label = _clean_text(link_item.get("label")).lower()
                uri = _clean_text(link_item.get("uri"))
                if uri and ("linkedin" in label or "linkedin.com" in uri.lower()):
                    personal_details["linkedin"] = uri
                    break
        except Exception:
            pass
    scan_issue_contract = _build_tailoring_scan_issue_contract(
        trusted_ready=[],
        trusted_optional=[],
        ai_optimize_optional=[],
        directional_guidance=[],
        resume_evidence=resume_evidence,
        tailoring_summary=tailoring_summary,
        jd_record=job_record,
    )
    counts = dict(scan_issue_contract.get("counts", {}) or {})
    selected_resume = _clean_text(resume_name) or "New scan resume"

    return {
        "ok": True,
        "preload_mode": "new_scan",
        "scan_entry_source": "scan_workspace_new_scan",
        "job_company": _clean_text(job_record.get("company") or job_record.get("job_company")),
        "job_title": _clean_text(job_record.get("title") or job_record.get("job_title")),
        "resume_name": selected_resume,
        "selected_resume": selected_resume,
        "selected_jd_record": job_record,
        "job": job_record,
        "job_snapshot": job_record,
        "selection": {
            "selected_resume": selected_resume,
            "selected_score": score / 100,
            "company": _clean_text(job_record.get("company") or job_record.get("job_company")),
            "title": _clean_text(job_record.get("title") or job_record.get("job_title")),
        },
        "score_snapshot": {
            "original_score": score / 100,
            "projected_score": score / 100,
            "projected_delta": 0,
            "draft_preview_status": "new_scan_ready",
            "draft_preview_note": "New Scan generated a deterministic match report from the submitted resume and job description.",
        },
        "scan_score": {
            "score": score,
            "source": "new_scan_match_score",
            "matched_count": int(counts.get("matched", 0) or 0),
            "missing_count": int(counts.get("missing", 0) or 0),
            "ai_count": int(counts.get("ai", 0) or 0),
            "total_count": int(counts.get("total", 0) or 0),
            "actionable_count": int(counts.get("actionable", 0) or 0),
            "label": "Match score",
        },
        "scan_session": {
            "session_id": scan_id,
            "session_version": 1,
            "selected_resume": selected_resume,
            "job_doc_id": _clean_text(job_record.get("job_doc_id") or job_record.get("doc_id")),
            "created_at": scan_timestamp,
            "issue_counts": counts,
        },
        "score_preview": _build_workspace_score_preview_contract(
            preview_status="new_scan_ready",
            preview_note="New Scan scoring is ready.",
            original_score=score / 100,
            projected_score=score / 100,
            projected_delta=0,
        ),
        "trusted_suggestions": {
            "direct_apply_ready": [],
            "direct_apply_optional": [],
        },
        "ai_optimize_suggestions": [],
        "directional_guidance": [],
        "lane_counts": {
            "direct_apply_ready": 0,
            "direct_apply_optional": 0,
            "ai_optimize_optional": 0,
            "direction_only": 0,
        },
        "scan_issue_contract": scan_issue_contract,
        "llm_model_config": tailoring_llm_model_config_payload(),
        "final_replacement_summary": {},
        "rewrite_review_summary": {},
        "rewrite_review_groups": [],
        "personal_details": {
            "extracted": personal_details,
            "saved": _normalize_workspace_personal_details({}),
            "current": personal_details,
        },
        "document_preview": _build_new_scan_document_preview(
            resume_text=resume_document.raw_text,
            personal_details=personal_details,
        ),
        "draft_status": "new_scan_ready",
        "has_saved_draft": False,
        "draft": {
            "selected_resume": selected_resume,
            "selected_patch_candidate_ids": [],
            "rewrite_review_decisions": {},
            "excluded_scan_issue_ids": [],
            "personal_details": personal_details,
            "draft_status": "new_scan_ready",
        },
        "new_scan": {
            "scan_id": scan_id,
            "scan_timestamp": scan_timestamp,
            "match_bucket": _clean_text(getattr(match_result, "match_bucket", "")),
            "dimension_scores": [
                {
                    "name": _clean_text(getattr(dimension, "name", "")),
                    "score": getattr(dimension, "score", None),
                    "weight": getattr(dimension, "weight", None),
                    "weighted_score": getattr(dimension, "weighted_score", None),
                    "reason": _clean_text(getattr(dimension, "reason", "")),
                    "evidence": list(getattr(dimension, "evidence", []) or []),
                }
                for dimension in list(getattr(match_result, "dimension_scores", []) or [])
            ],
        },
    }


def create_saved_scan_payload(
    *,
    scan_id: str = "",
    owner_user_id: str = "",
    owner_email: str = "",
    company: str = "",
    role: str = "",
    job_description_text: str = "",
    job_url: str = "",
    job_doc_id: str = "",
    saved_resume_name: str = "",
    resume_text: str = "",
    upload_filename: str = "",
    upload_content_type: str = "",
    upload_bytes: bytes | None = None,
    tailoring_json_path: str = "",
) -> Dict[str, Any]:
    safe_company = _clean_text(company)
    safe_role = _clean_text(role)
    safe_job_description = _clean_text(job_description_text)
    safe_job_url = _clean_text(job_url)
    safe_job_doc_id = _clean_text(job_doc_id)
    safe_scan_id = _clean_text(scan_id)
    safe_resume_text = _clean_text(resume_text)
    safe_tailoring_json_path = _clean_text(tailoring_json_path)
    should_write_postgres = bool(str(os.environ.get("DATABASE_URL", "") or "").strip())

    if not safe_job_description:
        raise ValueError("Job description is required for a new scan.")
    if not safe_company:
        raise ValueError("Company is required for a new scan.")
    if not safe_role:
        raise ValueError("Role is required for a new scan.")
    if not safe_job_url:
        raise ValueError("Job posting URL is required for a new scan.")

    scan_timestamp = datetime.now(timezone.utc).isoformat(timespec="microseconds")
    resume_source = "pasted_text"
    resume_name = ""
    resume_filename = ""
    resume_file_path = ""
    resume_processing_path = ""
    resume_file_mime_type = ""
    resume_size_bytes = 0

    if upload_bytes:
        safe_upload_name = _sanitize_scan_upload_filename(upload_filename)
        target_path = _scan_upload_target_path(safe_upload_name, scan_timestamp)
        target_path.write_bytes(upload_bytes)
        resume_source = "uploaded_file"
        resume_name = Path(safe_upload_name).stem
        resume_filename = safe_upload_name
        resume_file_path = str(target_path)
        resume_processing_path = str(target_path)
        resume_file_mime_type = _scan_upload_mime_type(safe_upload_name, upload_content_type)
        resume_size_bytes = len(upload_bytes)
    elif _clean_text(saved_resume_name):
        safe_saved_resume = _sanitize_resume_filename(saved_resume_name)
        saved_resume_payload = profile_resume_file_payload(
            safe_saved_resume,
            owner_user_id=owner_user_id,
        )
        saved_resume_temp_path = _materialize_profile_resume_temp_path(
            safe_saved_resume,
            owner_user_id=owner_user_id,
        )
        resume_source = "saved_resume"
        resume_name = safe_saved_resume
        resume_filename = safe_saved_resume
        resume_file_path = f"postgres://profile_resumes/{_clean_text(owner_user_id)}/{safe_saved_resume}"
        resume_processing_path = str(saved_resume_temp_path)
        resume_file_mime_type = _clean_text(saved_resume_payload.get("content_type")) or "application/pdf"
        resume_size_bytes = int(saved_resume_payload.get("size_bytes", 0) or 0)
    elif safe_resume_text:
        resume_source = "pasted_text"
        resume_name = "Pasted resume"
        resume_filename = ""
    else:
        raise ValueError("A saved resume, uploaded file, or pasted resume text is required.")

    job_record = _new_scan_job_record(
        company=safe_company,
        role=safe_role,
        job_description_text=safe_job_description,
        job_url=safe_job_url,
        job_doc_id=safe_job_doc_id,
    )

    row = saved_scan_db_row(
        {
            **({"scan_id": safe_scan_id} if safe_scan_id else {}),
            "owner_user_id": _clean_text(owner_user_id),
            "owner_email": _clean_text(owner_email),
            "scan_timestamp": scan_timestamp,
            "scan_source": "scan_workspace_new_scan",
            "scan_status": "processing",
            "resume_source": resume_source,
            "resume_name": resume_name,
            "resume_filename": resume_filename,
            "resume_file_path": resume_file_path,
            "resume_file_mime_type": resume_file_mime_type,
            "resume_size_bytes": resume_size_bytes,
            "resume_text": safe_resume_text,
            "job_doc_id": safe_job_doc_id,
            "job_url": safe_job_url,
            "job_company": safe_company,
            "job_title": safe_role,
            "job_description_text": safe_job_description,
            "match_rate": None,
            "tailoring_json_path": safe_tailoring_json_path,
            "note": "Scan processing started.",
            "payload_json": {
                "version": "saved_scan_intake_v1",
                "company": safe_company,
                "role": safe_role,
                "has_resume_text": bool(safe_resume_text),
                "has_uploaded_file": bool(upload_bytes),
                "job_description_length": len(safe_job_description),
                "resume_text_length": len(safe_resume_text),
            },
        }
    )

    review_payload = _build_new_scan_review_payload(
        scan_id=row["scan_id"],
        scan_timestamp=scan_timestamp,
        resume_name=resume_name or resume_filename or "New scan resume",
        resume_file_path=resume_processing_path or resume_file_path,
        resume_text=safe_resume_text,
        job_record=job_record,
        owner_user_id=owner_user_id,
    )
    scan_score = int(dict(review_payload.get("scan_score", {}) or {}).get("score", 0) or 0)
    row = saved_scan_db_row(
        {
            **row,
            "scan_status": "ready",
            "match_rate": scan_score,
            "note": "Scan report generated from New Scan.",
            "payload_json": {
                "version": "saved_scan_report_v1",
                "scan_review_payload": review_payload,
            },
            "owner_user_id": _clean_text(owner_user_id),
            "owner_email": _clean_text(owner_email),
        }
    )
    postgres_write = (
        _dual_write_saved_scan_postgres(row)
        if should_write_postgres
        else {
            "attempted": False,
            "ok": False,
            "skipped": "missing_database_url",
            "table_name": "saved_scans",
        }
    )

    return {
        "ok": bool(postgres_write.get("ok", False)),
        "scan_status": row["scan_status"],
        "scan": row,
        "scan_review_payload": review_payload,
        "postgres_write": postgres_write,
    }


def _scan_job_description_text_for_display(record: Dict[str, Any] | None = None) -> str:
    if not isinstance(record, dict):
        return ""

    values: List[Any] = []
    for key in (
        "description_text",
        "description",
        "job_description",
        "raw_description",
        "requirements",
        "responsibilities",
        "qualifications",
    ):
        raw = record.get(key)
        if isinstance(raw, list):
            values.extend(raw)
        elif isinstance(raw, dict):
            values.extend(raw.values())
        elif raw:
            values.append(raw)

    direct_text = "\n".join(_clean_text(value) for value in values if _clean_text(value)).strip()
    if direct_text:
        return direct_text

    return _clean_text(record.get("retrieval_text"))


def _scan_job_context_from_record(record: Dict[str, Any] | None = None) -> Dict[str, str]:
    if not isinstance(record, dict):
        return {}

    metadata = record.get("metadata", {}) if isinstance(record.get("metadata"), dict) else {}
    return {
        "company": _clean_text(
            record.get("company")
            or record.get("job_company")
            or metadata.get("company")
            or metadata.get("job_company")
        ),
        "title": _clean_text(
            record.get("title")
            or record.get("job_title")
            or metadata.get("title")
            or metadata.get("job_title")
        ),
        "job_url": _clean_text(
            record.get("job_url")
            or record.get("url")
            or record.get("link")
            or metadata.get("job_url")
            or metadata.get("url")
            or metadata.get("link")
        ),
        "job_doc_id": _clean_text(
            record.get("job_doc_id")
            or record.get("doc_id")
            or metadata.get("job_doc_id")
            or metadata.get("doc_id")
        ),
        "job_description_text": _scan_job_description_text_for_display(record),
    }


def scan_workspace_job_context_payload(
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    *,
    tailoring_json_path: str = "",
    job_doc_id: str = "",
    company: str = "",
    title: str = "",
) -> Dict[str, str]:
    context = {
        "company": _clean_text(company),
        "title": _clean_text(title),
        "job_url": "",
        "job_doc_id": _clean_text(job_doc_id),
        "job_description_text": "",
    }

    artifact_path_text = _clean_text(tailoring_json_path)
    if artifact_path_text:
        try:
            artifact_path = _resolve_planning_artifact_path(
                artifact_path_text,
                output_dir=output_dir,
            )
            payload_data = json.loads(artifact_path.read_text(encoding="utf-8"))
            if isinstance(payload_data, dict):
                job = payload_data.get("job", {}) if isinstance(payload_data.get("job"), dict) else {}
                job_snapshot = (
                    payload_data.get("job_snapshot", {})
                    if isinstance(payload_data.get("job_snapshot"), dict)
                    else {}
                )
                artifact_context = _scan_job_context_from_record(job_snapshot or job)
                for key, value in artifact_context.items():
                    if value:
                        context[key] = value
        except Exception:
            pass

    if not context["job_description_text"] and context["job_doc_id"]:
        try:
            corpus_context = _scan_job_context_from_record(
                _load_job_record_for_workspace_preview(context["job_doc_id"])
            )
            for key, value in corpus_context.items():
                if value:
                    context[key] = value
        except Exception:
            pass

    return context


def profile_saved_scans_payload(
    limit: int = 25,
    *,
    owner_user_id: str = "",
) -> Dict[str, Any]:
    try:
        payload = get_saved_scans_postgres_payload(
            limit=limit,
            owner_user_id=owner_user_id,
        )
        return {
            "ok": True,
            "source": "postgres",
            "count": int(payload.get("count", 0) or 0),
            "total_count": int(payload.get("postgres", {}).get("total_row_count", 0) or 0),
            "saved_scans": list(payload.get("rows", []) or []),
            "postgres": payload.get("postgres", {}),
        }
    except BaseException as exc:
        return {
            "ok": False,
            "source": "postgres",
            "count": 0,
            "total_count": 0,
            "saved_scans": [],
            "error_type": exc.__class__.__name__,
            "error": str(exc),
        }


def saved_scan_report_payload(
    scan_id: str,
    *,
    owner_user_id: str = "",
) -> Dict[str, Any]:
    safe_scan_id = _clean_text(scan_id)
    if not safe_scan_id:
        raise ValueError("scan_id is required.")

    payload = get_saved_scan_postgres_payload(
        scan_id=safe_scan_id,
        owner_user_id=owner_user_id,
    )
    row = dict(payload.get("scan", {}) or {})
    if not row:
        raise ValueError("Saved scan was not found.")

    report_payload = dict(row.get("payload_json", {}) or {})
    review_payload = report_payload.get("scan_review_payload")
    if not isinstance(review_payload, dict):
        raise ValueError("Saved scan does not contain a restorable report payload.")

    return {
        "ok": True,
        "scan": row,
        "scan_review_payload": review_payload,
    }


def delete_saved_scan_payload(
    scan_id: str,
    *,
    owner_user_id: str = "",
) -> Dict[str, Any]:
    safe_scan_id = _clean_text(scan_id)
    if not safe_scan_id:
        raise ValueError("scan_id is required.")

    payload = delete_saved_scan_postgres_payload(
        scan_id=safe_scan_id,
        owner_user_id=owner_user_id,
    )
    return {
        "ok": bool(payload.get("ok", False)),
        "scan_id": safe_scan_id,
    }


def save_saved_scan_state_payload(
    *,
    scan_id: str,
    owner_user_id: str = "",
    selected_patch_candidate_ids: Any = None,
    manual_bullet_edits: Any = None,
    rewrite_review_decisions: Any = None,
    excluded_scan_issue_ids: Any = None,
    personal_details: Any = None,
) -> Dict[str, Any]:
    safe_scan_id = _clean_text(scan_id)
    if not safe_scan_id:
        raise ValueError("scan_id is required.")

    draft = {
        "selected_resume": "",
        "selected_patch_candidate_ids": [
            _clean_text(value)
            for value in list(selected_patch_candidate_ids or [])
            if _clean_text(value)
        ],
        "manual_bullet_edits": dict(manual_bullet_edits or {})
        if isinstance(manual_bullet_edits, dict)
        else {},
        "rewrite_review_decisions": dict(rewrite_review_decisions or {})
        if isinstance(rewrite_review_decisions, dict)
        else {},
        "excluded_scan_issue_ids": _normalize_workspace_excluded_scan_issue_ids(
            excluded_scan_issue_ids or []
        ),
        "personal_details": _normalize_workspace_personal_details(personal_details or {}),
        "draft_status": "saved_scan_state",
        "saved_at": _utc_now(),
    }

    payload = save_saved_scan_draft_postgres_payload(
        scan_id=safe_scan_id,
        draft=draft,
        owner_user_id=owner_user_id,
    )
    return {
        "ok": bool(payload.get("ok", False)),
        "scan_id": safe_scan_id,
        "draft": draft,
        "has_saved_draft": True,
        "score_preview": {},
    }


def _extract_scan_upload_text_from_pdf(path: Path) -> str:
    from pdfminer.high_level import extract_text

    return _clean_text(extract_text(str(path)))


def _extract_scan_upload_text_from_docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    lines: List[str] = []
    for paragraph in document.paragraphs:
        text = _clean_text(paragraph.text)
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [_clean_text(cell.text) for cell in row.cells]
            line = " | ".join(cell for cell in cells if cell)
            if line:
                lines.append(line)
    return _clean_text("\n".join(lines))


def extract_scan_resume_upload_text_payload(
    *,
    filename: str = "",
    content_type: str = "",
    file_bytes: bytes | None = None,
) -> Dict[str, Any]:
    safe_name = _sanitize_scan_upload_filename(filename)
    if not file_bytes:
        raise ValueError("Uploaded resume file is empty.")

    suffix = Path(safe_name).suffix.lower()
    if suffix == ".txt":
        text = file_bytes.decode("utf-8", errors="replace")
    else:
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as handle:
            handle.write(file_bytes)
            temp_path = Path(handle.name)
        try:
            if suffix == ".pdf":
                text = _extract_scan_upload_text_from_pdf(temp_path)
            elif suffix == ".docx":
                text = _extract_scan_upload_text_from_docx(temp_path)
            else:
                raise ValueError("Only PDF, DOCX, and TXT scan uploads are supported.")
        finally:
            try:
                temp_path.unlink()
            except FileNotFoundError:
                pass

    clean_text = _clean_text(text)
    if not clean_text:
        raise ValueError("Could not extract readable resume text from this file.")

    return {
        "ok": True,
        "filename": safe_name,
        "content_type": _scan_upload_mime_type(safe_name, content_type),
        "char_count": len(clean_text),
        "text": clean_text,
    }

def _utc_now() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds")


def _derive_pipeline_log_path(output_dir: Path) -> Path:
    return output_dir / "live_pipeline_run.log"


def _derive_pipeline_status_path(output_dir: Path) -> Path:
    return output_dir / "live_pipeline_status.json"


def _new_run_id() -> str:
    return datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%S%fZ")


def _load_runtime_status_file(path: str) -> Dict[str, Any]:
    status_path = Path(str(path or "")).expanduser()
    if not status_path.exists():
        return {}

    try:
        return json.loads(status_path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def _job_app():
    import job_app
    return job_app


def _make_args(**kwargs):
    return SimpleNamespace(**kwargs)


def _normalize_pipeline_llm_actions(value: Any) -> str:
    if isinstance(value, list):
        actions = [str(item).strip() for item in value if str(item).strip()]
    else:
        raw = str(value or "").strip()
        actions = [part.strip() for part in raw.split(",") if part.strip()]

    if not actions:
        actions = ["APPLY", "APPLY_REVIEW_VARIANTS"]

    seen: List[str] = []
    for action in actions:
        if action not in seen:
            seen.append(action)

    return ",".join(seen)


def _normalize_delete_seen_data(value: Any) -> str:
    raw = str(value or "").strip().lower()
    if raw in {"yes", "y", "true", "1"}:
        return "yes"
    if raw in {"ask", "prompt"}:
        return "ask"
    return "no"



def _new_pipeline_run_state() -> Dict[str, Any]:
    return {
        "process": None,
        "log_handle": None,
        "status": "idle",
        "started_at": "",
        "finished_at": "",
        "return_code": None,
        "command": [],
        "output_dir": str(DEFAULT_OUTPUT_DIR),
        "log_path": str(DEFAULT_PIPELINE_LOG_PATH),
        "status_path": str(DEFAULT_PIPELINE_STATUS_PATH),
        "run_id": "",
        "child_pid": None,
        "error": "",
        "owner_user_id": "",
    }


def _max_concurrent_user_pipeline_runs() -> int:
    raw = _clean_text(os.environ.get("JOB_STACK_MAX_CONCURRENT_USER_PIPELINES")) or "2"
    try:
        return max(1, int(raw))
    except ValueError:
        return 2


def _active_pipeline_reservation_ttl_seconds() -> int:
    raw = _clean_text(os.environ.get("JOB_STACK_ACTIVE_PIPELINE_TTL_SECONDS")) or "21600"
    try:
        return max(60, int(raw))
    except ValueError:
        return 21600


def _active_pipeline_worker_id() -> str:
    return f"pid:{os.getpid()}"


def _user_pipeline_redis_admission_lock_enabled() -> bool:
    raw = _clean_text(os.environ.get("JOB_STACK_USER_PIPELINE_REDIS_ADMISSION_LOCK_ENABLED"))
    if raw:
        return raw.lower() in {"1", "true", "yes", "y", "on"}

    return bool(_clean_text(os.environ.get("REDIS_URL")))


def _user_pipeline_redis_admission_lock_payload(
    *,
    owner_user_id: str,
    run_id: str,
    ttl_seconds: int,
) -> Dict[str, Any]:
    owner = _clean_text(owner_user_id)
    safe_run_id = _clean_text(run_id)

    if not owner or not safe_run_id:
        return {
            "attempted": False,
            "acquired": False,
            "skipped": "missing_owner_or_run_id",
        }

    if not _user_pipeline_redis_admission_lock_enabled():
        return {
            "attempted": False,
            "acquired": False,
            "skipped": "disabled",
        }

    lock_key = redis_lock_key("user_pipeline_admission", owner)

    try:
        handle = acquire_redis_lock(
            lock_key,
            ttl_seconds=ttl_seconds,
            token=safe_run_id,
        )
    except Exception as exc:
        logger.warning(
            "Redis user pipeline admission lock skipped owner_user_id=%s run_id=%s: %s",
            owner,
            safe_run_id,
            exc,
        )
        return {
            "attempted": True,
            "acquired": False,
            "skipped": "redis_error",
            "error": repr(exc),
            "key": lock_key,
        }

    return {
        "attempted": True,
        "acquired": bool(handle.acquired),
        "skipped": "" if handle.acquired else "lock_held",
        "key": handle.key,
        "token": handle.token,
        "ttl_seconds": handle.ttl_seconds,
    }


def _release_user_pipeline_redis_admission_lock_payload(lock_payload: Dict[str, Any]) -> Dict[str, Any]:
    if not isinstance(lock_payload, dict):
        return {
            "attempted": False,
            "released": False,
            "skipped": "missing_payload",
        }

    if not bool(lock_payload.get("acquired", False)):
        return {
            "attempted": False,
            "released": False,
            "skipped": "not_acquired",
        }

    key = _clean_text(lock_payload.get("key"))
    token = _clean_text(lock_payload.get("token"))
    ttl_seconds = lock_payload.get("ttl_seconds") or _active_pipeline_reservation_ttl_seconds()

    if not key or not token:
        return {
            "attempted": False,
            "released": False,
            "skipped": "missing_key_or_token",
        }

    try:
        handle = RedisLockHandle(
            key=key,
            token=token,
            acquired=True,
            ttl_seconds=int(ttl_seconds),
        )
        released = release_redis_lock(handle)
        return {
            "attempted": True,
            "released": bool(released),
            "key": key,
        }
    except Exception as exc:
        logger.warning("Failed to release Redis user pipeline admission lock key=%s: %s", key, exc)
        return {
            "attempted": True,
            "released": False,
            "key": key,
            "error": repr(exc),
        }


def _active_run_redis_admission_lock_payload(active_run: Dict[str, Any]) -> Dict[str, Any]:
    if not isinstance(active_run, dict):
        return {}

    metadata = active_run.get("metadata_json") or {}
    if not isinstance(metadata, dict):
        return {}

    lock_payload = metadata.get("redis_admission_lock") or {}
    return dict(lock_payload) if isinstance(lock_payload, dict) else {}


def _release_user_pipeline_active_run(
    *,
    owner_user_id: str,
    run_id: str,
    terminal_status: str = "",
) -> None:
    owner = _clean_text(owner_user_id)
    safe_run_id = _clean_text(run_id)

    if not owner:
        return

    active_run: Dict[str, Any] = {}
    try:
        active_payload = get_user_pipeline_active_run_postgres_payload(
            owner_user_id=owner,
            database_url="",
            database_url_env="DATABASE_URL",
            psql_bin="psql",
            print_only=False,
            ensure_schema=True,
        )
        active_run = dict(active_payload.get("active_run", {}) or {})
    except Exception as exc:
        logger.warning(
            "Failed to read user pipeline active run before release owner_user_id=%s run_id=%s: %s",
            owner,
            safe_run_id,
            exc,
        )

    try:
        release_user_pipeline_active_run_postgres_payload(
            owner_user_id=owner,
            run_id=safe_run_id,
            terminal_status=terminal_status,
            database_url="",
            database_url_env="DATABASE_URL",
            psql_bin="psql",
            print_only=False,
            ensure_schema=True,
        )
    except Exception as exc:
        logger.warning(
            "Failed to release user pipeline active run owner_user_id=%s run_id=%s: %s",
            owner,
            safe_run_id,
            exc,
        )

    _release_user_pipeline_redis_admission_lock_payload(
        _active_run_redis_admission_lock_payload(active_run)
    )

def _clear_owner_active_pipeline_state(owner_user_id: str, run_id: str = "") -> None:
    owner = _clean_text(owner_user_id)
    safe_run_id = _clean_text(run_id)

    if not owner:
        return

    with _PIPELINE_ACTIVE_RUNS_LOCK:
        state = _PIPELINE_ACTIVE_RUNS.get(owner)
        if not isinstance(state, dict):
            return

        if safe_run_id and _clean_text(state.get("run_id")) != safe_run_id:
            return

        _PIPELINE_ACTIVE_RUNS.pop(owner, None)


def _active_user_pipeline_run_count() -> int:
    count = 0

    with _PIPELINE_ACTIVE_RUNS_LOCK:
        states = list(_PIPELINE_ACTIVE_RUNS.values())

    for state in states:
        if _is_pipeline_process_running(state.get("process")):
            count += 1

    return count


def _owner_active_pipeline_state(owner_user_id: str) -> Dict[str, Any]:
    owner = _clean_text(owner_user_id)
    if not owner:
        return {}

    with _PIPELINE_ACTIVE_RUNS_LOCK:
        state = _PIPELINE_ACTIVE_RUNS.get(owner)

    if not isinstance(state, dict):
        return {}

    return state


def _set_owner_active_pipeline_state(owner_user_id: str, state: Dict[str, Any]) -> None:
    owner = _clean_text(owner_user_id)
    if not owner:
        return

    with _PIPELINE_ACTIVE_RUNS_LOCK:
        _PIPELINE_ACTIVE_RUNS[owner] = state

def _pipeline_status_snapshot(state: Any = None) -> Dict[str, Any]:
    run_state = state if isinstance(state, dict) else _PIPELINE_RUN_STATE

    process = run_state.get("process")
    log_handle = run_state.get("log_handle")

    if process is not None:
        return_code = process.poll()
        if return_code is None:
            run_state["status"] = "running"
        else:
            run_state["status"] = "succeeded" if return_code == 0 else "failed"
            run_state["finished_at"] = run_state.get("finished_at") or _utc_now()
            run_state["return_code"] = return_code
            run_state["process"] = None
            run_state["child_pid"] = None

            if log_handle is not None:
                try:
                    log_handle.close()
                except Exception:
                    pass
                run_state["log_handle"] = None

    status = run_state.get("status", "idle")
    return {
        "status": status,
        "started_at": run_state.get("started_at", ""),
        "finished_at": run_state.get("finished_at", ""),
        "return_code": run_state.get("return_code"),
        "command": run_state.get("command", []),
        "output_dir": run_state.get("output_dir", str(DEFAULT_OUTPUT_DIR)),
        "log_path": run_state.get("log_path", str(DEFAULT_PIPELINE_LOG_PATH)),
        "status_path": run_state.get("status_path", str(DEFAULT_PIPELINE_STATUS_PATH)),
        "run_id": run_state.get("run_id", ""),
        "child_pid": run_state.get("child_pid"),
        "error": run_state.get("error", ""),
        "owner_user_id": run_state.get("owner_user_id", ""),
        "is_running": status == "running",
    }

def _write_runtime_status_file(path: Path, payload: Dict[str, Any]) -> None:
    resolved = Path(path).expanduser()
    resolved.parent.mkdir(parents=True, exist_ok=True)
    tmp_path = resolved.with_suffix(resolved.suffix + ".tmp")
    tmp_path.write_text(
        json.dumps(payload, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    tmp_path.replace(resolved)

def _pid_exists(pid: Any) -> bool:
    try:
        normalized = int(pid)
    except (TypeError, ValueError):
        return False

    if normalized <= 0:
        return False

    try:
        os.kill(normalized, 0)
    except ProcessLookupError:
        return False
    except PermissionError:
        return True
    except OSError:
        return False

    return True


def _heal_stale_running_runtime_status(
    snapshot: Dict[str, Any],
    runtime_status: Dict[str, Any],
) -> Dict[str, Any]:
    if snapshot.get("is_running"):
        return runtime_status

    if str(runtime_status.get("status", "") or "").strip().lower() != "running":
        return runtime_status

    child_pid = runtime_status.get("child_pid")
    if child_pid and _pid_exists(child_pid):
        return runtime_status

    healed = dict(runtime_status)
    healed["status"] = "failed"
    healed["finished_at"] = healed.get("finished_at") or _utc_now()
    healed["return_code"] = 1

    reason = (
        "Live pipeline is no longer running. "
        "The API server likely restarted or the child process exited unexpectedly."
    )
    healed["error"] = healed.get("error") or reason
    healed["summary_message"] = healed.get("summary_message") or reason
    healed["stage_message"] = reason

    status_path_raw = str(
        snapshot.get("status_path")
        or healed.get("status_path")
        or ""
    ).strip()
    if status_path_raw:
        _write_runtime_status_file(Path(status_path_raw), healed)

    return healed

def stop_live_pipeline_for_server_shutdown(
    reason: str = "Live pipeline stopped because the API server shut down.",
) -> Dict[str, Any]:
    process = _PIPELINE_RUN_STATE.get("process")
    log_handle = _PIPELINE_RUN_STATE.get("log_handle")
    status_path_raw = str(_PIPELINE_RUN_STATE.get("status_path", "") or "").strip()
    status_path = Path(status_path_raw).expanduser() if status_path_raw else None
    finished_at = _utc_now()

    if process is None:
        if log_handle is not None:
            try:
                log_handle.close()
            except Exception:
                pass
            _PIPELINE_RUN_STATE["log_handle"] = None

        return {
            "ok": True,
            "stopped": False,
            "reason": "no_active_process",
        }

    return_code = process.poll()
    was_running = return_code is None

    if was_running:
        process.terminate()
        try:
            return_code = process.wait(timeout=5)
        except subprocess.TimeoutExpired:
            process.kill()
            return_code = process.wait(timeout=5)

    if log_handle is not None:
        try:
            log_handle.close()
        except Exception:
            pass

    _PIPELINE_RUN_STATE["process"] = None
    _PIPELINE_RUN_STATE["child_pid"] = None
    _PIPELINE_RUN_STATE["log_handle"] = None
    _PIPELINE_RUN_STATE["finished_at"] = finished_at
    _PIPELINE_RUN_STATE["return_code"] = return_code if return_code is not None else 1

    if was_running:
        _PIPELINE_RUN_STATE["status"] = "failed"
        _PIPELINE_RUN_STATE["error"] = reason
    else:
        _PIPELINE_RUN_STATE["status"] = "succeeded" if return_code == 0 else "failed"
        _PIPELINE_RUN_STATE["error"] = "" if return_code == 0 else reason

    if status_path and status_path.exists():
        runtime_payload = _load_runtime_status_file(str(status_path))
        if runtime_payload:
            runtime_payload["finished_at"] = finished_at
            runtime_payload["return_code"] = return_code if return_code is not None else 1

            if was_running:
                runtime_payload["status"] = "failed"
                runtime_payload["error"] = reason
                runtime_payload["summary_message"] = reason
                runtime_payload["stage_message"] = reason
            else:
                runtime_payload["status"] = "succeeded" if return_code == 0 else "failed"
                if return_code != 0:
                    runtime_payload["error"] = reason
                    runtime_payload["summary_message"] = runtime_payload.get("summary_message") or reason
                    runtime_payload["stage_message"] = runtime_payload.get("stage_message") or reason

            _write_runtime_status_file(status_path, runtime_payload)

    return {
        "ok": True,
        "stopped": was_running,
        "return_code": return_code,
        "status_path": str(status_path) if status_path else "",
    }


def _runtime_status_is_stale_startup(
    snapshot: Dict[str, Any],
    runtime_status: Dict[str, Any],
) -> bool:
    if snapshot.get("is_running"):
        return False

    if str(runtime_status.get("status", "") or "").strip().lower() != "running":
        return False

    if str(runtime_status.get("current_stage", "") or "").strip().lower() != "startup":
        return False

    if str(runtime_status.get("stage_message", "") or "").strip() != "Launching pipeline subprocess":
        return False

    if list(runtime_status.get("completed_stages", []) or []):
        return False

    if dict(runtime_status.get("counts", {}) or {}):
        return False

    if runtime_status.get("final_job_count") not in (None, "", 0):
        return False

    return True


def _is_pipeline_process_running(process: Any) -> bool:
    if process is None:
        return False

    try:
        return process.poll() is None
    except Exception:
        return False


def _pipeline_terminal_statuses() -> set:
    return {"succeeded", "failed", "cancelled"}


def _latest_owner_pipeline_status_payload(*, owner_user_id: str) -> Dict[str, Any]:
    owner = _clean_text(owner_user_id)
    if not owner:
        return {
            "ok": True,
            "pipeline_gate": {},
            "status": "idle",
            "running": False,
            "can_start": True,
            "started_at": "",
            "finished_at": "",
            "return_code": None,
            "command": [],
            "output_dir": "",
            "log_path": "",
            "status_path": "",
            "run_id": "",
            "child_pid": None,
            "error": "",
        }

    latest_payload = get_latest_user_pipeline_run_postgres_payload(
        owner_user_id=owner,
        database_url="",
        database_url_env="DATABASE_URL",
        psql_bin="psql",
        print_only=False,
        ensure_schema=True,
    )
    run = dict(latest_payload.get("run", {}) or {})
    status = _clean_text(run.get("status")) or "idle"
    running = status in {"queued", "running", "starting"}

    return {
        "ok": True,
        "pipeline_gate": user_pipeline_gate_payload(owner_user_id=owner),
        "status": status,
        "running": running,
        "can_start": not running,
        "blocked_by_other_user": False,
        "message": _clean_text(run.get("summary_message")),
        "current_stage": _clean_text(run.get("current_stage")),
        "stage_message": _clean_text(run.get("stage_message")),
        "summary_message": _clean_text(run.get("summary_message")),
        "started_at": _clean_text(run.get("started_at")),
        "finished_at": _clean_text(run.get("completed_at")),
        "return_code": run.get("return_code"),
        "command": [],
        "output_dir": "",
        "log_path": "",
        "status_path": "",
        "run_id": _clean_text(run.get("run_id")),
        "child_pid": None,
        "error": _clean_text(run.get("error")),
        "status_json": run.get("status_json") or {},
    }


def _persist_user_pipeline_status_snapshot(
    *,
    owner_user_id: str,
    status_payload: Dict[str, Any],
) -> None:
    owner = _clean_text(owner_user_id)
    if not owner:
        return

    if isinstance(status_payload.get("pipeline"), dict):
        status_payload = dict(status_payload.get("pipeline") or {})

    run_id = _clean_text(status_payload.get("run_id"))
    if not run_id:
        return

    status = _clean_text(status_payload.get("status")) or "running"
    finished_at = _clean_text(status_payload.get("finished_at"))

    try:
        upsert_user_pipeline_run_postgres_payload(
            record={
                "run_id": run_id,
                "owner_user_id": owner,
                "status": status,
                "current_stage": _clean_text(status_payload.get("current_stage")),
                "stage_message": _clean_text(status_payload.get("stage_message")),
                "summary_message": _clean_text(status_payload.get("summary_message"))
                or _clean_text(status_payload.get("message")),
                "return_code": status_payload.get("return_code"),
                "started_at": _clean_text(status_payload.get("started_at")) or _utc_now(),
                "updated_at": _utc_now(),
                "completed_at": finished_at if status in _pipeline_terminal_statuses() else "",
                "config_json": {
                    "command": status_payload.get("command") or [],
                    "output_dir": _clean_text(status_payload.get("output_dir")),
                    "log_path": _clean_text(status_payload.get("log_path")),
                    "status_path": _clean_text(status_payload.get("status_path")),
                    "child_pid": status_payload.get("child_pid"),
                    "config": status_payload.get("config") or {},
                    "artifact_ingestion": status_payload.get("artifact_ingestion") or {},
                },
                "status_json": status_payload,
                "error": _clean_text(status_payload.get("error")),
            },
            database_url="",
            database_url_env="DATABASE_URL",
            psql_bin="psql",
            print_only=False,
            ensure_schema=True,
        )
    except Exception as exc:
        logger.warning("Failed to persist user pipeline status snapshot: %s", exc)

    if status in _pipeline_terminal_statuses():
        _release_user_pipeline_active_run(
            owner_user_id=owner,
            run_id=run_id,
            terminal_status=status,
        )
        _clear_owner_active_pipeline_state(owner, run_id)



def _pipeline_artifact_ingestion_key(*, owner_user_id: str, run_id: str) -> str:
    return f"{_clean_text(owner_user_id)}:{_clean_text(run_id)}"


def _pipeline_artifact_candidate_paths(output_dir: Path) -> List[Path]:
    root = Path(output_dir).expanduser()
    if not root.exists() or not root.is_dir():
        return []

    candidates: List[Path] = []
    root_names = {
        "live_pipeline_status.json",
        "live_pipeline_run.log",
        "current_run_job_corpus.jsonl",
        "best_resume_variant_by_job.csv",
        "application_shortlist_by_job.csv",
        "application_execution_queue.csv",
        "job_packet_manifest.csv",
    }

    for name in sorted(root_names):
        candidate = root / name
        if candidate.exists() and candidate.is_file():
            candidates.append(candidate)

    packet_dir = root / "job_packets"
    if packet_dir.exists() and packet_dir.is_dir():
        for candidate in sorted(packet_dir.rglob("*")):
            if candidate.is_file() and candidate.suffix.lower() in {".json", ".jsonl", ".md", ".txt"}:
                candidates.append(candidate)

    return candidates


def _pipeline_artifact_name(*, output_dir: Path, path: Path) -> str:
    try:
        return path.relative_to(output_dir).as_posix()
    except ValueError:
        return path.name


def _pipeline_artifact_kind(*, output_dir: Path, path: Path) -> str:
    name = _pipeline_artifact_name(output_dir=output_dir, path=path)
    filename = path.name
    suffix = path.suffix.lower()

    root_kind_by_name = {
        "live_pipeline_status.json": "live_pipeline_status",
        "live_pipeline_run.log": "live_pipeline_log",
        "current_run_job_corpus.jsonl": "current_run_job_corpus",
        "best_resume_variant_by_job.csv": "best_resume_variant_by_job",
        "application_shortlist_by_job.csv": "application_shortlist_by_job",
        "application_execution_queue.csv": "application_execution_queue",
        "job_packet_manifest.csv": "job_packet_manifest",
    }

    if name in root_kind_by_name:
        return root_kind_by_name[name]

    if name.startswith("job_packets/"):
        if filename.endswith("__tailoring_llm.json"):
            return "job_packet_tailoring_llm_json"
        if filename.endswith("__tailoring.json"):
            return "job_packet_tailoring_json"
        if filename.endswith("__tailoring.md"):
            return "job_packet_tailoring_markdown"
        if suffix == ".json":
            return "job_packet_json"
        if suffix == ".jsonl":
            return "job_packet_jsonl"
        if suffix == ".md":
            return "job_packet_markdown"
        return "job_packet_text"

    if suffix == ".json":
        return "json_artifact"
    if suffix == ".jsonl":
        return "jsonl_artifact"
    if suffix == ".csv":
        return "csv_artifact"
    if suffix in {".log", ".txt", ".md"}:
        return "text_artifact"
    return "pipeline_artifact"


def _pipeline_artifact_content_type(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".json":
        return "application/json"
    if suffix == ".jsonl":
        return "application/x-jsonlines"
    if suffix == ".csv":
        return "text/csv"
    if suffix == ".md":
        return "text/markdown"
    if suffix in {".log", ".txt"}:
        return "text/plain"
    return "application/octet-stream"


def _read_pipeline_artifact_record(
    *,
    owner_user_id: str,
    run_id: str,
    output_dir: Path,
    path: Path,
) -> Dict[str, Any]:
    size_bytes = path.stat().st_size
    if size_bytes > _PIPELINE_ARTIFACT_MAX_BYTES:
        raise ValueError(
            f"artifact too large: {path} ({size_bytes} bytes > {_PIPELINE_ARTIFACT_MAX_BYTES})"
        )

    raw_text = path.read_text(encoding="utf-8", errors="replace")
    content_json = None

    if path.suffix.lower() == ".json":
        try:
            content_json = json.loads(raw_text)
        except json.JSONDecodeError:
            content_json = None

    return {
        "owner_user_id": owner_user_id,
        "run_id": run_id,
        "artifact_kind": _pipeline_artifact_kind(output_dir=output_dir, path=path),
        "artifact_name": _pipeline_artifact_name(output_dir=output_dir, path=path),
        "content_type": _pipeline_artifact_content_type(path),
        "content_json": content_json,
        "content_text": "" if content_json is not None else raw_text,
        "created_at": _utc_now(),
    }


def _ingest_pipeline_run_artifacts_to_postgres(
    *,
    owner_user_id: str,
    run_id: str,
    output_dir: str,
    status_payload: Dict[str, Any],
) -> Dict[str, Any]:
    owner = _clean_text(owner_user_id)
    safe_run_id = _clean_text(run_id)

    if not owner or not safe_run_id:
        return {
            "ok": False,
            "attempted": False,
            "reason": "missing_owner_or_run_id",
            "ingested_count": 0,
            "skipped_count": 0,
            "error_count": 0,
        }

    ingestion_key = _pipeline_artifact_ingestion_key(
        owner_user_id=owner,
        run_id=safe_run_id,
    )

    if ingestion_key in _PIPELINE_ARTIFACT_INGESTED_RUN_KEYS:
        return {
            "ok": True,
            "attempted": False,
            "already_ingested": True,
            "ingested_count": 0,
            "skipped_count": 0,
            "error_count": 0,
        }

    root = Path(_clean_text(output_dir)).expanduser()
    if not root.exists() or not root.is_dir():
        return {
            "ok": False,
            "attempted": False,
            "reason": "output_dir_missing",
            "output_dir": str(root),
            "ingested_count": 0,
            "skipped_count": 0,
            "error_count": 0,
        }

    ingested: List[Dict[str, Any]] = []
    skipped: List[Dict[str, Any]] = []
    errors: List[Dict[str, Any]] = []

    for artifact_path in _pipeline_artifact_candidate_paths(root):
        artifact_name = _pipeline_artifact_name(output_dir=root, path=artifact_path)

        try:
            record = _read_pipeline_artifact_record(
                owner_user_id=owner,
                run_id=safe_run_id,
                output_dir=root,
                path=artifact_path,
            )
            payload = upsert_user_pipeline_artifact_postgres_payload(
                record=record,
                database_url="",
                database_url_env="DATABASE_URL",
                psql_bin="psql",
                print_only=False,
                ensure_schema=True,
            )
            artifact = dict(payload.get("artifact", {}) or {})
            ingested.append(
                {
                    "artifact_id": _clean_text(artifact.get("artifact_id")),
                    "artifact_kind": _clean_text(artifact.get("artifact_kind")) or record["artifact_kind"],
                    "artifact_name": _clean_text(artifact.get("artifact_name")) or record["artifact_name"],
                    "content_type": record["content_type"],
                    "size_bytes": artifact_path.stat().st_size,
                }
            )
        except ValueError as exc:
            skipped.append({"artifact_name": artifact_name, "reason": str(exc)})
        except Exception as exc:
            errors.append(
                {
                    "artifact_name": artifact_name,
                    "error_type": exc.__class__.__name__,
                    "error": str(exc),
                }
            )

    result = {
        "ok": len(errors) == 0,
        "attempted": True,
        "already_ingested": False,
        "owner_user_id": owner,
        "run_id": safe_run_id,
        "output_dir": str(root),
        "ingested_count": len(ingested),
        "skipped_count": len(skipped),
        "error_count": len(errors),
        "ingested": ingested,
        "skipped": skipped,
        "errors": errors,
        "ingested_at": _utc_now(),
    }

    try:
        upsert_user_pipeline_artifact_postgres_payload(
            record={
                "owner_user_id": owner,
                "run_id": safe_run_id,
                "artifact_kind": "artifact_ingestion_manifest",
                "artifact_name": "artifact_ingestion_manifest.json",
                "content_type": "application/json",
                "content_json": result,
                "content_text": "",
                "created_at": _utc_now(),
            },
            database_url="",
            database_url_env="DATABASE_URL",
            psql_bin="psql",
            print_only=False,
            ensure_schema=True,
        )
        result["manifest_written"] = True
    except Exception as exc:
        result["manifest_written"] = False
        result["manifest_error"] = str(exc)
        result["ok"] = False

    _PIPELINE_ARTIFACT_INGESTED_RUN_KEYS.add(ingestion_key)
    return result


def _truthy_env_value(value: Any) -> bool:
    return _clean_text(value).lower() in {"1", "true", "yes", "y", "on"}


def _finalize_seen_jobs_staging_payload(
    *,
    owner_user_id: str,
    run_id: str,
    terminal_status: str,
    artifact_ingestion: Dict[str, Any],
) -> Dict[str, Any]:
    owner = _clean_text(owner_user_id)
    safe_run_id = _clean_text(run_id)
    status = _clean_text(terminal_status).lower()
    ingestion_ok = bool(dict(artifact_ingestion or {}).get("ok", False))

    if not owner or not safe_run_id:
        return {
            "ok": False,
            "attempted": False,
            "action": "none",
            "reason": "missing_owner_or_run_id",
        }

    try:
        if status == "succeeded" and ingestion_ok:
            payload = promote_user_seen_jobs_staging_postgres_payload(
                owner_user_id=owner,
                run_id=safe_run_id,
                database_url="",
                database_url_env="DATABASE_URL",
                psql_bin="psql",
                print_only=False,
                ensure_schema=True,
            )
            return {
                "ok": True,
                "attempted": True,
                "action": "promoted",
                "status": status,
                "artifact_ingestion_ok": ingestion_ok,
                "staged_count": int(payload.get("staged_count", 0) or 0),
                "promoted_count": int(payload.get("promoted_count", 0) or 0),
                "deleted_count": int(payload.get("deleted_count", 0) or 0),
            }

        payload = clear_user_seen_jobs_staging_postgres_payload(
            owner_user_id=owner,
            run_id=safe_run_id,
            database_url="",
            database_url_env="DATABASE_URL",
            psql_bin="psql",
            print_only=False,
            ensure_schema=True,
        )
        return {
            "ok": True,
            "attempted": True,
            "action": "cleared",
            "status": status,
            "artifact_ingestion_ok": ingestion_ok,
            "deleted_count": int(payload.get("deleted_count", 0) or 0),
        }
    except Exception as exc:
        return {
            "ok": False,
            "attempted": True,
            "action": "error",
            "status": status,
            "artifact_ingestion_ok": ingestion_ok,
            "error_type": exc.__class__.__name__,
            "error": str(exc),
        }


def _pipeline_scratch_cleanup_payload(
    *,
    owner_user_id: str,
    run_id: str,
    output_dir: str,
    artifact_ingestion: Dict[str, Any],
) -> Dict[str, Any]:
    owner = _clean_text(owner_user_id)
    safe_run_id = _clean_text(run_id)
    raw_output_dir = _clean_text(output_dir)

    if _truthy_env_value(os.environ.get("JOB_STACK_KEEP_PIPELINE_SCRATCH")):
        return {
            "ok": True,
            "attempted": False,
            "skipped": "keep_scratch_enabled",
            "scratch_deleted": False,
        }

    if not bool(dict(artifact_ingestion or {}).get("ok", False)):
        return {
            "ok": False,
            "attempted": False,
            "skipped": "artifact_ingestion_not_ok",
            "scratch_deleted": False,
        }

    if not owner or not safe_run_id or not raw_output_dir:
        return {
            "ok": False,
            "attempted": False,
            "skipped": "missing_owner_run_or_output_dir",
            "scratch_deleted": False,
        }

    output_path = Path(raw_output_dir).expanduser()
    parts = output_path.parts

    # Only delete the canonical run-scoped scratch shape:
    # tmp/pipeline_runs/<owner_user_id>/<run_id>/application_planning
    if len(parts) < 5:
        return {
            "ok": False,
            "attempted": False,
            "skipped": "unsafe_output_dir_shape",
            "output_dir": str(output_path),
            "scratch_deleted": False,
        }

    if not (
        parts[-1] == "application_planning"
        and parts[-2] == safe_run_id
        and parts[-3] == owner
        and parts[-4] == "pipeline_runs"
        and parts[-5] == "tmp"
    ):
        return {
            "ok": False,
            "attempted": False,
            "skipped": "non_canonical_output_dir",
            "output_dir": str(output_path),
            "scratch_deleted": False,
        }

    run_root = output_path.parent

    if not run_root.exists():
        return {
            "ok": True,
            "attempted": False,
            "skipped": "scratch_already_missing",
            "scratch_dir": str(run_root),
            "scratch_deleted": False,
        }

    try:
        shutil.rmtree(run_root)
    except Exception as exc:
        return {
            "ok": False,
            "attempted": True,
            "scratch_dir": str(run_root),
            "scratch_deleted": False,
            "error_type": exc.__class__.__name__,
            "error": str(exc),
        }

    return {
        "ok": True,
        "attempted": True,
        "scratch_dir": str(run_root),
        "scratch_deleted": True,
    }


def _owner_db_active_pipeline_status_payload(*, owner_user_id: str) -> Dict[str, Any]:
    owner = _clean_text(owner_user_id)
    if not owner:
        return {}

    try:
        active_payload = get_user_pipeline_active_run_postgres_payload(
            owner_user_id=owner,
            database_url="",
            database_url_env="DATABASE_URL",
            psql_bin="psql",
            print_only=False,
            ensure_schema=True,
        )
    except Exception as exc:
        logger.warning(
            "Failed to read Postgres active pipeline run for owner_user_id=%s: %s",
            owner,
            exc,
        )
        return {}

    active_run = dict(active_payload.get("active_run", {}) or {})
    if not active_run:
        return {}

    state = _new_pipeline_run_state()
    state["status"] = _clean_text(active_run.get("status")) or "running"
    state["started_at"] = _clean_text(active_run.get("reserved_at"))
    state["finished_at"] = ""
    state["return_code"] = None
    state["command"] = []
    state["output_dir"] = _clean_text(active_run.get("output_dir"))
    state["log_path"] = ""
    state["status_path"] = _clean_text(active_run.get("status_path"))
    state["run_id"] = _clean_text(active_run.get("run_id"))
    state["child_pid"] = _clean_text(active_run.get("process_pid")) or None
    state["error"] = ""
    state["owner_user_id"] = owner

    payload = pipeline_status_payload(owner_user_id=owner, state=state)
    pipeline = dict(payload.get("pipeline", {}) or {})

    _persist_user_pipeline_status_snapshot(
        owner_user_id=owner,
        status_payload=pipeline,
    )

    if _clean_text(pipeline.get("status")) in _pipeline_terminal_statuses():
        artifact_ingestion = _ingest_pipeline_run_artifacts_to_postgres(
            owner_user_id=owner,
            run_id=_clean_text(pipeline.get("run_id")),
            output_dir=_clean_text(pipeline.get("output_dir")),
            status_payload=pipeline,
        )
        pipeline["artifact_ingestion"] = artifact_ingestion
        pipeline["scratch_cleanup"] = _pipeline_scratch_cleanup_payload(
            owner_user_id=owner,
            run_id=_clean_text(pipeline.get("run_id")),
            output_dir=_clean_text(pipeline.get("output_dir")),
            artifact_ingestion=artifact_ingestion,
        )
        pipeline["seen_job_finalization"] = _finalize_seen_jobs_staging_payload(
            owner_user_id=owner,
            run_id=_clean_text(pipeline.get("run_id")),
            terminal_status=_clean_text(pipeline.get("status")),
            artifact_ingestion=artifact_ingestion,
        )
        payload["pipeline"] = pipeline

        _persist_user_pipeline_status_snapshot(
            owner_user_id=owner,
            status_payload=pipeline,
        )

        _release_user_pipeline_active_run(
            owner_user_id=owner,
            run_id=_clean_text(pipeline.get("run_id")),
            terminal_status=_clean_text(pipeline.get("status")),
        )

    return payload

def owner_pipeline_status_payload(*, owner_user_id: str = "") -> Dict[str, Any]:
    owner = _clean_text(owner_user_id)
    if not owner:
        return pipeline_status_payload(owner_user_id=owner_user_id)

    state = _owner_active_pipeline_state(owner)
    if not state:
        db_active_payload = _owner_db_active_pipeline_status_payload(owner_user_id=owner)
        if db_active_payload:
            return db_active_payload
        return _latest_owner_pipeline_status_payload(owner_user_id=owner)

    payload = pipeline_status_payload(owner_user_id=owner, state=state)
    pipeline = dict(payload.get("pipeline", {}) or {})

    _persist_user_pipeline_status_snapshot(
        owner_user_id=owner,
        status_payload=pipeline,
    )

    if _clean_text(pipeline.get("status")) in _pipeline_terminal_statuses():
        artifact_ingestion = _ingest_pipeline_run_artifacts_to_postgres(
            owner_user_id=owner,
            run_id=_clean_text(pipeline.get("run_id")),
            output_dir=_clean_text(pipeline.get("output_dir")),
            status_payload=pipeline,
        )
        pipeline["artifact_ingestion"] = artifact_ingestion
        pipeline["scratch_cleanup"] = _pipeline_scratch_cleanup_payload(
            owner_user_id=owner,
            run_id=_clean_text(pipeline.get("run_id")),
            output_dir=_clean_text(pipeline.get("output_dir")),
            artifact_ingestion=artifact_ingestion,
        )
        pipeline["seen_job_finalization"] = _finalize_seen_jobs_staging_payload(
            owner_user_id=owner,
            run_id=_clean_text(pipeline.get("run_id")),
            terminal_status=_clean_text(pipeline.get("status")),
            artifact_ingestion=artifact_ingestion,
        )
        payload["pipeline"] = pipeline

        _persist_user_pipeline_status_snapshot(
            owner_user_id=owner,
            status_payload=pipeline,
        )

        _release_user_pipeline_active_run(
            owner_user_id=owner,
            run_id=_clean_text(pipeline.get("run_id")),
            terminal_status=_clean_text(pipeline.get("status")),
        )
        _clear_owner_active_pipeline_state(
            owner_user_id=owner,
            run_id=_clean_text(pipeline.get("run_id")),
        )

    return payload

def pipeline_status_payload(*, owner_user_id: str = "", state: Any = None) -> Dict[str, Any]:
    snapshot = _pipeline_status_snapshot(state=state)
    runtime_status = _load_runtime_status_file(snapshot.get("status_path", ""))

    if runtime_status and _runtime_status_is_stale_startup(snapshot, runtime_status):
        runtime_status = {}

    if runtime_status:
        runtime_status = _heal_stale_running_runtime_status(snapshot, runtime_status)

    merged = dict(snapshot)
    if runtime_status:
        if merged.get("status") == "idle":
            merged["status"] = runtime_status.get("status", merged["status"])

        if not merged.get("started_at"):
            merged["started_at"] = runtime_status.get("started_at", "")
        if not merged.get("finished_at"):
            merged["finished_at"] = runtime_status.get("finished_at", "")
        if merged.get("return_code") is None:
            merged["return_code"] = runtime_status.get("return_code")
        if not merged.get("error"):
            merged["error"] = runtime_status.get("error", "")

        merged.update({
            "run_id": runtime_status.get("run_id", merged.get("run_id", "")),
            "output_dir": runtime_status.get("output_dir", merged.get("output_dir", "")),
            "log_path": runtime_status.get("log_path", merged.get("log_path", "")),
            "status_path": runtime_status.get("status_path", merged.get("status_path", "")),
            "current_stage": runtime_status.get("current_stage", ""),
            "completed_stages": runtime_status.get("completed_stages", []),
            "stage_order": runtime_status.get("stage_order", []),
            "stage_started_at": runtime_status.get("stage_started_at", ""),
            "stage_message": runtime_status.get("stage_message", ""),
            "counts": runtime_status.get("counts", {}),
            "summary_message": runtime_status.get("summary_message", ""),
            "final_job_count": runtime_status.get("final_job_count"),
            "config": runtime_status.get("config", {}),
        })

        merged["is_running"] = merged.get("status") == "running"

    return {
        "ok": True,
        "pipeline_gate": user_pipeline_gate_payload(owner_user_id=owner_user_id) if _clean_text(owner_user_id) else {},
        "pipeline": merged,
    }

def scheduler_jobs_payload() -> Dict[str, Any]:
    return {
        "ok": True,
        "jobs": get_scheduled_job_definitions(),
    }


def scheduler_job_command_payload(
    *,
    job_name: str,
    planning_only: bool = False,
    run_application_planning: bool = True,
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    job_limit: int = 50,
    job_packet_limit: int = 0,
    llm_actions: Any = "APPLY,APPLY_REVIEW_VARIANTS",
    generate_tailoring: bool = False,
    generate_llm_tailoring: bool = False,
    refresh_llm_tailoring: bool = False,
    generate_llm_fallback: bool = False,
    delete_seen_data: str = "no",
) -> Dict[str, Any]:
    definition = get_scheduled_job_definition(job_name)

    command = build_scheduled_job_command(
        job_name,
        run_application_planning=bool(run_application_planning),
        planning_only=bool(planning_only),
        output_dir=output_dir,
        job_limit=int(job_limit),
        job_packet_limit=int(job_packet_limit),
        llm_actions=llm_actions,
        generate_tailoring=bool(generate_tailoring),
        generate_llm_tailoring=bool(generate_llm_tailoring),
        refresh_llm_tailoring=bool(refresh_llm_tailoring),
        generate_llm_fallback=bool(generate_llm_fallback),
        delete_seen_data=str(delete_seen_data or "no"),
    )

    return {
        "ok": True,
        "job": definition,
        "command": command,
        "command_text": " ".join(command),
        "options": {
            "planning_only": bool(planning_only),
            "run_application_planning": bool(run_application_planning),
            "output_dir": str(output_dir),
            "job_limit": int(job_limit),
            "job_packet_limit": int(job_packet_limit),
            "llm_actions": llm_actions,
            "generate_tailoring": bool(generate_tailoring),
            "generate_llm_tailoring": bool(generate_llm_tailoring),
            "refresh_llm_tailoring": bool(refresh_llm_tailoring),
            "generate_llm_fallback": bool(generate_llm_fallback),
            "delete_seen_data": str(delete_seen_data or "no"),
        },
    }

def _augment_launchd_config_exists(payload: Dict[str, Any]) -> Dict[str, Any]:
    item = dict(payload)
    item["plist_exists"] = Path(item["plist_path"]).expanduser().exists()
    item["stdout_log_exists"] = Path(item["stdout_log_path"]).expanduser().exists()
    item["stderr_log_exists"] = Path(item["stderr_log_path"]).expanduser().exists()
    return item

def scheduler_launchd_config_payload(
    *,
    job_name: str = "",
    planning_only: bool = False,
    run_application_planning: bool = True,
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    job_limit: int = 50,
    job_packet_limit: int = 0,
    llm_actions: Any = "APPLY,APPLY_REVIEW_VARIANTS",
    generate_tailoring: bool = False,
    generate_llm_tailoring: bool = False,
    refresh_llm_tailoring: bool = False,
    generate_llm_fallback: bool = False,
    delete_seen_data: str = "no",
    sync_postgres_run_history: bool = False,
    require_postgres_run_history_sync: bool = False,
    database_url_env: str = "DATABASE_URL",
    psql_bin: str = "psql",
    allow_contract_drift: bool = False,
    launchd_interval_seconds: int = DEFAULT_LAUNCHD_INTERVAL_SECONDS,
    launchd_out_dir: Path = DEFAULT_LAUNCHD_OUT_DIR,
    launchd_log_dir: Path = DEFAULT_LAUNCHD_LOG_DIR,
    launchd_label_prefix: str = DEFAULT_LAUNCHD_LABEL_PREFIX,
) -> Dict[str, Any]:
    common_kwargs = {
        "run_application_planning": bool(run_application_planning),
        "output_dir": output_dir,
        "job_limit": int(job_limit),
        "job_packet_limit": int(job_packet_limit),
        "llm_actions": llm_actions,
        "generate_tailoring": bool(generate_tailoring),
        "generate_llm_tailoring": bool(generate_llm_tailoring),
        "refresh_llm_tailoring": bool(refresh_llm_tailoring),
        "generate_llm_fallback": bool(generate_llm_fallback),
        "delete_seen_data": str(delete_seen_data or "no"),
        "history_path": DEFAULT_SCHEDULER_RUN_HISTORY_PATH,
        "sync_postgres_run_history": bool(sync_postgres_run_history),
        "require_postgres_run_history_sync": bool(require_postgres_run_history_sync),
        "database_url_env": database_url_env,
        "psql_bin": psql_bin,
        "allow_contract_drift": bool(allow_contract_drift),
        "launchd_interval_seconds": int(launchd_interval_seconds),
        "launchd_out_dir": launchd_out_dir,
        "launchd_log_dir": launchd_log_dir,
        "launchd_label_prefix": launchd_label_prefix,
    }

    if str(job_name or "").strip():
        item = _augment_launchd_config_exists(
            build_scheduler_launchd_plist_payload(
                job_name=job_name,
                planning_only=bool(planning_only),
                **common_kwargs,
            )
        )
        return {
            "ok": True,
            "mode": "single",
            "item": item,
        }

    items = [
        _augment_launchd_config_exists(
            build_scheduler_launchd_plist_payload(
                job_name="agent_discovery",
                planning_only=False,
                **common_kwargs,
            )
        ),
        _augment_launchd_config_exists(
            build_scheduler_launchd_plist_payload(
                job_name="live_pipeline",
                planning_only=True,
                **common_kwargs,
            )
        ),
    ]

    return {
        "ok": True,
        "mode": "default_pair",
        "items": items,
    }

def scheduler_launchd_agent_status_payload(
    *,
    job_name: str = "",
    planning_only: bool = False,
    run_application_planning: bool = True,
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    job_limit: int = 50,
    job_packet_limit: int = 0,
    llm_actions: Any = "APPLY,APPLY_REVIEW_VARIANTS",
    generate_tailoring: bool = False,
    generate_llm_tailoring: bool = False,
    refresh_llm_tailoring: bool = False,
    generate_llm_fallback: bool = False,
    delete_seen_data: str = "no",
    sync_postgres_run_history: bool = False,
    require_postgres_run_history_sync: bool = False,
    database_url_env: str = "DATABASE_URL",
    psql_bin: str = "psql",
    allow_contract_drift: bool = False,
    launchd_interval_seconds: int = DEFAULT_LAUNCHD_INTERVAL_SECONDS,
    launchd_out_dir: Path = DEFAULT_LAUNCHD_OUT_DIR,
    launchd_log_dir: Path = DEFAULT_LAUNCHD_LOG_DIR,
    launchd_label_prefix: str = DEFAULT_LAUNCHD_LABEL_PREFIX,
    launchd_agent_dir: Path = DEFAULT_LAUNCHD_AGENT_DIR,
    launchd_target: str = DEFAULT_LAUNCHD_TARGET,
) -> Dict[str, Any]:
    common_kwargs = {
        "run_application_planning": bool(run_application_planning),
        "output_dir": output_dir,
        "job_limit": int(job_limit),
        "job_packet_limit": int(job_packet_limit),
        "llm_actions": llm_actions,
        "generate_tailoring": bool(generate_tailoring),
        "generate_llm_tailoring": bool(generate_llm_tailoring),
        "refresh_llm_tailoring": bool(refresh_llm_tailoring),
        "generate_llm_fallback": bool(generate_llm_fallback),
        "delete_seen_data": str(delete_seen_data or "no"),
        "history_path": DEFAULT_SCHEDULER_RUN_HISTORY_PATH,
        "sync_postgres_run_history": bool(sync_postgres_run_history),
        "require_postgres_run_history_sync": bool(require_postgres_run_history_sync),
        "database_url_env": database_url_env,
        "psql_bin": psql_bin,
        "allow_contract_drift": bool(allow_contract_drift),
        "launchd_interval_seconds": int(launchd_interval_seconds),
        "launchd_out_dir": launchd_out_dir,
        "launchd_log_dir": launchd_log_dir,
        "launchd_label_prefix": launchd_label_prefix,
        "launchd_agent_dir": launchd_agent_dir,
        "launchd_target": launchd_target,
    }

    if str(job_name or "").strip():
        return {
            "ok": True,
            "mode": "single",
            "item": get_scheduler_launchd_agent_status(
                job_name=job_name,
                planning_only=bool(planning_only),
                **common_kwargs,
            ),
        }

    items = [
        get_scheduler_launchd_agent_status(
            job_name="agent_discovery",
            planning_only=False,
            **common_kwargs,
        ),
        get_scheduler_launchd_agent_status(
            job_name="live_pipeline",
            planning_only=True,
            **common_kwargs,
        ),
    ]

    return {
        "ok": True,
        "mode": "default_pair",
        "items": items,
    }

def scheduler_postgres_status_payload(
    *,
    limit: int = 10,
    database_url_env: str = "DATABASE_URL",
    psql_bin: str = "psql",
) -> Dict[str, Any]:
    postgres_payload = get_scheduler_postgres_status_payload(
        limit=limit,
        database_url="",
        database_url_env=database_url_env,
        psql_bin=psql_bin,
        print_only=False,
    )

    jsonl_rows = _load_scheduler_history_rows(DEFAULT_SCHEDULER_RUN_HISTORY_PATH)
    postgres_block = dict(postgres_payload.get("postgres", {}) or {})
    postgres_history_row_count = int(postgres_block.get("run_history_count", 0) or 0)

    return {
        "ok": True,
        "query_limit": int(limit),
        "history_jsonl_path": str(DEFAULT_SCHEDULER_RUN_HISTORY_PATH),
        "history_jsonl_row_count": len(jsonl_rows),
        "history_postgres_row_count": postgres_history_row_count,
        "history_count_matches_jsonl": postgres_history_row_count == len(jsonl_rows),
        "postgres_command_text": postgres_payload["command_text"],
        "postgres": postgres_block,
    }

def scheduler_operator_summary_payload(
    *,
    limit: int = 5,
    database_url_env: str = "DATABASE_URL",
    psql_bin: str = "psql",
) -> Dict[str, Any]:
    normalized_limit = max(int(limit), 1)

    contract = scheduler_contract_health_payload()
    postgres_payload = scheduler_postgres_status_payload(
        limit=normalized_limit,
        database_url_env=database_url_env,
        psql_bin=psql_bin,
    )

    jsonl_rows = _load_scheduler_history_rows(DEFAULT_SCHEDULER_RUN_HISTORY_PATH)
    latest_jsonl_rows = jsonl_rows[:normalized_limit]

    postgres_block = dict(postgres_payload.get("postgres", {}) or {})

    return {
        "ok": True,
        "limit": normalized_limit,
        "contract_health": contract,
        "history": {
            "jsonl_path": str(DEFAULT_SCHEDULER_RUN_HISTORY_PATH),
            "jsonl_row_count": postgres_payload["history_jsonl_row_count"],
            "postgres_row_count": postgres_payload["history_postgres_row_count"],
            "count_matches": postgres_payload["history_count_matches_jsonl"],
        },
        "latest_runs_by_job": postgres_block.get("latest_runs_by_job", []),
        "recent_postgres_runs": postgres_block.get("recent_runs", []),
        "recent_jsonl_runs": latest_jsonl_rows,
        "postgres_summary": {
            "job_definition_count": postgres_block.get("job_definition_count", 0),
            "active_job_count": postgres_block.get("active_job_count", 0),
            "run_history_count": postgres_block.get("run_history_count", 0),
            "success_count": postgres_block.get("success_count", 0),
            "failure_count": postgres_block.get("failure_count", 0),
        },
        "postgres_command_text": postgres_payload["postgres_command_text"],
    }


def _normalize_scheduler_filter_text(value: Any) -> str:
    return str(value or "").strip().lower().replace("-", "_").replace(" ", "_")


def _load_scheduler_history_rows(
    history_path: Path = DEFAULT_SCHEDULER_RUN_HISTORY_PATH,
) -> List[Dict[str, Any]]:
    if not history_path.exists():
        return []

    rows: List[Dict[str, Any]] = []

    with history_path.open("r", encoding="utf-8") as f:
        for line_number, raw_line in enumerate(f, start=1):
            line = raw_line.strip()
            if not line:
                continue

            try:
                payload = json.loads(line)
            except Exception:
                continue

            if not isinstance(payload, dict):
                continue

            row = dict(payload)
            row["_line_number"] = line_number
            rows.append(row)

    rows.sort(
        key=lambda row: (
            str(row.get("started_at", "") or ""),
            str(row.get("run_id", "") or ""),
            int(row.get("_line_number", 0) or 0),
        ),
        reverse=True,
    )
    return rows


def scheduler_history_payload(
    history_path: Path = DEFAULT_SCHEDULER_RUN_HISTORY_PATH,
    job_name: str = "",
    status: str = "",
    limit: int = 20,
) -> Dict[str, Any]:
    rows = _load_scheduler_history_rows(history_path)

    normalized_job_name = _normalize_scheduler_filter_text(job_name)
    normalized_status = _normalize_scheduler_filter_text(status)

    if normalized_job_name:
        rows = [
            row for row in rows
            if _normalize_scheduler_filter_text(row.get("job_name", "")) == normalized_job_name
        ]

    if normalized_status:
        rows = [
            row for row in rows
            if _normalize_scheduler_filter_text(row.get("status", "")) == normalized_status
        ]

    selected = rows[: max(int(limit), 0)]

    return {
        "ok": True,
        "history_path": str(history_path),
        "filters": {
            "job_name": job_name,
            "status": status,
            "limit": limit,
        },
        "total_matching_rows": len(rows),
        "rows": selected,
        "count": len(selected),
    }

def _load_notification_rows(
    notification_dir: Path = DEFAULT_NOTIFICATION_RECORDS_DIR,
) -> List[Dict[str, Any]]:
    if not notification_dir.exists() or not notification_dir.is_dir():
        return []

    rows: List[Dict[str, Any]] = []

    for path in notification_dir.glob("*.json"):
        if not path.is_file():
            continue

        try:
            payload = json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            continue

        if not isinstance(payload, dict):
            continue

        row = dict(payload)
        row["_path"] = str(path)
        rows.append(row)

    rows.sort(
        key=lambda row: (
            str(row.get("created_at", "") or ""),
            str(row.get("notification_id", "") or ""),
        ),
        reverse=True,
    )
    return rows


def _normalize_notification_read_flag(value: Any) -> bool:
    if isinstance(value, bool):
        return value

    if value is None:
        raw = ""
    else:
        raw = str(value).strip().lower()

    if raw in {"1", "true", "yes", "y", "read"}:
        return True

    if raw in {"0", "false", "no", "n", "off", "unread"}:
        return False

    raise ValueError("is_read must be a boolean-like value.")


def _normalize_optional_notification_read_filter(value: Any) -> Any:
    raw = str(value or "").strip()
    if not raw:
        return None
    return _normalize_notification_read_flag(raw)

def _load_latest_notification_state_overlay() -> Dict[str, Dict[str, Any]]:
    meta_payload = get_notification_state_postgres_status_payload(
        limit=1,
        database_url="",
        database_url_env="DATABASE_URL",
        psql_bin="psql",
        print_only=False,
    )
    meta_block = dict(meta_payload.get("postgres", {}) or {})
    query_limit = max(int(meta_block.get("latest_state_count", 0) or 0), 1)

    postgres_payload = get_notification_state_postgres_status_payload(
        limit=query_limit,
        database_url="",
        database_url_env="DATABASE_URL",
        psql_bin="psql",
        print_only=False,
    )
    postgres_block = dict(postgres_payload.get("postgres", {}) or {})
    postgres_rows = list(postgres_block.get("latest_rows", []) or [])

    latest_overlay: Dict[str, Dict[str, Any]] = {}
    for row in postgres_rows:
        notification_id = _clean_text(row.get("notification_id"))
        if not notification_id:
            continue

        latest_overlay[notification_id] = {
            "is_read": bool(row.get("is_read", False)),
            "state_timestamp": _clean_text(row.get("state_timestamp")),
        }

    return latest_overlay

def _apply_notification_state_overlay(
    rows: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    latest_by_notification_id = _load_latest_notification_state_overlay()
    overlaid_rows: List[Dict[str, Any]] = []

    for row in rows:
        merged = dict(row)
        merged["is_read"] = bool(merged.get("is_read", False))
        merged["read_state_timestamp"] = ""

        notification_id = str(merged.get("notification_id", "") or "").strip()
        overlay = latest_by_notification_id.get(notification_id)

        if overlay:
            merged["is_read"] = bool(overlay.get("is_read", False))
            merged["read_state_timestamp"] = str(overlay.get("state_timestamp", "") or "").strip()

        overlaid_rows.append(merged)

    return overlaid_rows


def notifications_payload(
    notification_dir: Path = DEFAULT_NOTIFICATION_RECORDS_DIR,
    job_name: str = "",
    level: str = "",
    delivery_status: str = "",
    is_read: str = "",
    limit: int = 20,
) -> Dict[str, Any]:
    rows = _apply_notification_state_overlay(
        _load_notification_rows(notification_dir),
    )

    normalized_job_name = _normalize_scheduler_filter_text(job_name)
    normalized_level = _normalize_scheduler_filter_text(level)
    normalized_delivery_status = _normalize_scheduler_filter_text(delivery_status)
    normalized_is_read = _normalize_optional_notification_read_filter(is_read)

    if normalized_job_name:
        rows = [
            row for row in rows
            if _normalize_scheduler_filter_text(row.get("job_name", "")) == normalized_job_name
        ]

    if normalized_level:
        rows = [
            row for row in rows
            if _normalize_scheduler_filter_text(row.get("level", "")) == normalized_level
        ]

    if normalized_delivery_status:
        rows = [
            row for row in rows
            if _normalize_scheduler_filter_text(row.get("delivery_status", "")) == normalized_delivery_status
        ]

    if normalized_is_read is not None:
        rows = [
            row for row in rows
            if bool(row.get("is_read", False)) == normalized_is_read
        ]

    selected = rows[: max(int(limit), 0)]

    return {
        "ok": True,
        "notification_dir": str(notification_dir),
        "filters": {
            "job_name": job_name,
            "level": level,
            "delivery_status": delivery_status,
            "is_read": is_read,
            "limit": limit,
        },
        "total_matching_rows": len(rows),
        "rows": selected,
        "count": len(selected),
    }


def notifications_summary_payload(
    notification_dir: Path = DEFAULT_NOTIFICATION_RECORDS_DIR,
    limit: int = 10,
) -> Dict[str, Any]:
    rows = _apply_notification_state_overlay(
        _load_notification_rows(notification_dir),
    )
    selected = rows[: max(int(limit), 0)]

    level_counts = Counter(
        str(row.get("level", "") or "<empty>")
        for row in rows
    )
    delivery_status_counts = Counter(
        str(row.get("delivery_status", "") or "<empty>")
        for row in rows
    )
    job_name_counts = Counter(
        str(row.get("job_name", "") or "<empty>")
        for row in rows
    )

    unread_count = sum(1 for row in rows if not bool(row.get("is_read", False)))
    read_count = len(rows) - unread_count

    return {
        "ok": True,
        "total_rows": len(rows),
        "read_count": read_count,
        "unread_count": unread_count,
        "level_counts": dict(sorted(level_counts.items())),
        "delivery_status_counts": dict(sorted(delivery_status_counts.items())),
        "job_name_counts": dict(sorted(job_name_counts.items())),
        "recent_notifications": selected,
        "count": len(selected),
    }


def notifications_unread_count_payload(
    notification_dir: Path = DEFAULT_NOTIFICATION_RECORDS_DIR,
) -> Dict[str, Any]:
    rows = _apply_notification_state_overlay(
        _load_notification_rows(notification_dir),
    )

    unread_count = sum(1 for row in rows if not bool(row.get("is_read", False)))
    read_count = len(rows) - unread_count

    return {
        "ok": True,
        "total_rows": len(rows),
        "read_count": read_count,
        "unread_count": unread_count,
    }

def _dual_write_notification_state_postgres(row: Dict[str, Any]) -> Dict[str, Any]:
    database_url = str(os.environ.get("DATABASE_URL", "") or "").strip()
    if not database_url:
        return {
            "attempted": False,
            "ok": False,
            "skipped": "missing_database_url",
            "table_name": "notification_state_events",
        }

    try:
        payload = insert_notification_state_row_to_postgres(
            record=row,
            database_url="",
            database_url_env="DATABASE_URL",
            psql_bin="psql",
            print_only=False,
            allow_contract_drift=False,
        )
        return {
            "attempted": True,
            "ok": True,
            "table_name": payload.get("table_name", "notification_state_events"),
            "state_id": str(payload.get("row", {}).get("state_id", "") or ""),
            "contract_health_ok": bool(payload.get("contract_health_ok", False)),
            "command_text": str(payload.get("command_text", "") or ""),
        }
    except SystemExit as exc:
        return {
            "attempted": True,
            "ok": False,
            "table_name": "notification_state_events",
            "error_type": "SystemExit",
            "error": str(exc),
        }
    except Exception as exc:
        return {
            "attempted": True,
            "ok": False,
            "table_name": "notification_state_events",
            "error_type": exc.__class__.__name__,
            "error": str(exc),
        }
    
def record_notification_read_state_payload(
    notification_dir: Path = DEFAULT_NOTIFICATION_RECORDS_DIR,
    *,
    notification_id: str = "",
    is_read: Any = True,
) -> Dict[str, Any]:
    clean_notification_id = str(notification_id or "").strip()
    if not clean_notification_id:
        raise ValueError("notification_id is required.")

    rows = _apply_notification_state_overlay(
        _load_notification_rows(notification_dir),
    )

    target_notification = None
    for row in rows:
        if str(row.get("notification_id", "") or "").strip() == clean_notification_id:
            target_notification = dict(row)
            break

    if target_notification is None:
        raise ValueError(f"Notification not found: {clean_notification_id}")

    normalized_is_read = _normalize_notification_read_flag(is_read)

    state_row = {
        "state_timestamp": datetime.now(timezone.utc).isoformat(timespec="microseconds"),
        "notification_id": clean_notification_id,
        "is_read": str(normalized_is_read),
    }

    postgres_write = _dual_write_notification_state_postgres(state_row)

    target_notification["is_read"] = normalized_is_read
    target_notification["read_state_timestamp"] = state_row["state_timestamp"]

    return {
        "ok": True,
        "state_row": state_row,
        "notification": target_notification,
    }

def scheduler_storage_contract_payload(
    *,
    include_sql: bool = False,
    include_generated_seed_sql: bool = False,
    include_generated_init_sql: bool = False,
) -> Dict[str, Any]:
    schema_payload = scheduler_schema_sql_payload()
    seed_payload = scheduler_seed_sql_payload()
    init_payload = scheduler_init_sql_payload()
    seed_generation_payload = scheduler_seed_sql_generation_payload()
    init_generation_payload = scheduler_init_sql_generation_payload()

    payload = {
        "ok": True,
        "tables": scheduler_postgres_table_specs(),
        "seed_rows": {
            "scheduler_job_definitions": scheduler_job_definition_seed_rows(),
        },
        "schema_sql_path": schema_payload["path"],
        "seed_sql_path": seed_payload["path"],
        "init_sql_path": init_payload["path"],
        "seed_sql_matches_artifact": seed_generation_payload["matches_artifact"],
        "init_sql_matches_artifact": init_generation_payload["matches_artifact"],
        "include_sql": bool(include_sql),
        "include_generated_seed_sql": bool(include_generated_seed_sql),
        "include_generated_init_sql": bool(include_generated_init_sql),
    }

    if include_sql:
        payload["schema_sql"] = schema_payload["sql"]
        payload["seed_sql"] = seed_payload["sql"]
        payload["init_sql"] = init_payload["sql"]

    if include_generated_seed_sql:
        payload["seed_sql_generated"] = seed_generation_payload["generated_sql"]

    if include_generated_init_sql:
        payload["init_sql_generated"] = init_generation_payload["generated_sql"]

    return payload

_PIPELINE_CHILD_ENV_EXACT_NAMES = {
    "APP_ENV",
    "DATABASE_URL",
    "REDIS_URL",
    "PATH",
    "HOME",
    "USER",
    "LOGNAME",
    "LANG",
    "LC_ALL",
    "TZ",
    "PYTHONPATH",
    "PYTHONUNBUFFERED",
    "SSL_CERT_FILE",
    "REQUESTS_CA_BUNDLE",
    "CURL_CA_BUNDLE",
    "TOKENIZERS_PARALLELISM",
    "HF_HOME",
    "TRANSFORMERS_CACHE",
    "SENTENCE_TRANSFORMERS_HOME",
}

_PIPELINE_CHILD_ENV_PREFIXES = (
    "JOB_STACK_",
    "JOB_APP_",
    "LLM_",
    "TAILOR_",
    "PATCH_REFINEMENT_",
    "GROQ_",
    "GEMINI_",
    "OPENAI_",
    "GOOGLE_",
    "HF_",
    "HUGGINGFACE_",
    "POSTGRES_",
    "DATABASE_",
    "REDIS_",
    "SMTP_",
    "EMAIL_",
    "AUTH_",
    "SESSION_",
)


def _pipeline_child_env_max_value_bytes() -> int:
    raw = _clean_text(os.environ.get("JOB_STACK_PIPELINE_CHILD_ENV_MAX_VALUE_BYTES"))
    try:
        max_bytes = int(raw)
    except Exception:
        max_bytes = 32768

    return max(1024, min(max_bytes, 262144))


def _pipeline_child_env(
    *,
    extra_env: Dict[str, Any] | None = None,
    base_env: Dict[str, Any] | None = None,
) -> Dict[str, str]:
    source = dict(base_env if base_env is not None else os.environ)
    child_env: Dict[str, str] = {}
    max_value_bytes = _pipeline_child_env_max_value_bytes()

    for raw_key, raw_value in source.items():
        key = _clean_text(raw_key)
        if not key:
            continue

        allowed = key in _PIPELINE_CHILD_ENV_EXACT_NAMES or any(
            key.startswith(prefix) for prefix in _PIPELINE_CHILD_ENV_PREFIXES
        )
        if not allowed:
            continue

        value = str(raw_value or "")
        if len(value.encode("utf-8")) > max_value_bytes:
            logger.warning(
                "Skipping oversized pipeline child env var key=%s size_bytes=%s max_value_bytes=%s",
                key,
                len(value.encode("utf-8")),
                max_value_bytes,
            )
            continue

        child_env[key] = value

    for raw_key, raw_value in dict(extra_env or {}).items():
        key = _clean_text(raw_key)
        if not key:
            continue

        value = str(raw_value or "")
        if len(value.encode("utf-8")) > max_value_bytes:
            logger.warning(
                "Skipping oversized extra pipeline child env var key=%s size_bytes=%s max_value_bytes=%s",
                key,
                len(value.encode("utf-8")),
                max_value_bytes,
            )
            continue

        child_env[key] = value

    child_env.setdefault("PYTHONUNBUFFERED", "1")

    if "PATH" not in child_env and _clean_text(os.environ.get("PATH")):
        child_env["PATH"] = _clean_text(os.environ.get("PATH"))

    return child_env


def _pipeline_child_env_summary(child_env: Dict[str, str]) -> Dict[str, Any]:
    values = [str(value or "") for value in dict(child_env or {}).values()]
    return {
        "mode": "allowlist",
        "env_var_count": len(child_env or {}),
        "env_size_bytes": sum(len(value.encode("utf-8")) for value in values),
        "max_value_bytes": _pipeline_child_env_max_value_bytes(),
    }

def run_live_pipeline_payload(
    owner_user_id: str = "",
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    log_path: Path = DEFAULT_PIPELINE_LOG_PATH,
    job_limit: int = 50,
    job_packet_limit: int = 0,
    llm_actions: Any = "APPLY,APPLY_REVIEW_VARIANTS",
    generate_tailoring: bool = False,
    generate_llm_tailoring: bool = False,
    refresh_llm_tailoring: bool = False,
    generate_llm_fallback: bool = False,
    generate_llm_adjudication: bool = False,
    planning_only: bool = False,
    delete_seen_data: str = "no",
) -> Dict[str, Any]:
    owner_for_pipeline_gate = _clean_text(owner_user_id)
    if owner_for_pipeline_gate:
        pipeline_gate = user_pipeline_gate_payload(owner_user_id=owner_for_pipeline_gate)

        if not bool(pipeline_gate.get("can_run_live_pipeline", False)):
            raise ValueError(
                _clean_text(pipeline_gate.get("live_pipeline_block_reason"))
                or "Upload at least one resume before running Live Pipeline."
            )

        requested_delete_seen_data = _clean_text(delete_seen_data).lower()
        if (
            requested_delete_seen_data in {"yes", "true", "1", "y"}
            and not bool(pipeline_gate.get("can_delete_seen_data", False))
        ):
            raise ValueError(
                _clean_text(pipeline_gate.get("delete_seen_data_block_reason"))
                or "Delete seen data is disabled until this user has completed at least one successful Live Pipeline run."
            )

        owner_state = _owner_active_pipeline_state(owner_for_pipeline_gate)
        if owner_state and _pipeline_status_snapshot(state=owner_state).get("is_running"):
            raise ValueError("A live pipeline run is already in progress for this user.")

    else:
        snapshot = _pipeline_status_snapshot()
        if snapshot.get("is_running"):
            raise ValueError("A live pipeline run is already in progress.")

    requested_output_dir = Path(output_dir).expanduser()
    run_id = _new_run_id()

    if owner_for_pipeline_gate:
        output_dir = _pipeline_scratch_output_dir(
            owner_user_id=owner_for_pipeline_gate,
            run_id=run_id,
        )
    else:
        output_dir = requested_output_dir
        output_dir.mkdir(parents=True, exist_ok=True)

    canonical_log_path = _derive_pipeline_log_path(output_dir)
    canonical_status_path = _derive_pipeline_status_path(output_dir)
    pipeline_job_corpus_path = output_dir / "current_run_job_corpus.jsonl"

    normalized_llm_actions = _normalize_pipeline_llm_actions(llm_actions)
    normalized_delete_seen_data = _normalize_delete_seen_data(delete_seen_data)

    ja = _job_app()
    effective_generate_llm_adjudication = bool(generate_llm_adjudication)

    args = _make_args(
        run_application_planning=True,
        job_limit=int(job_limit),
        job_packet_limit=int(job_packet_limit),
        output_dir=str(output_dir),
        llm_actions=normalized_llm_actions,
        generate_tailoring=bool(generate_tailoring),
        generate_llm_tailoring=bool(generate_llm_tailoring),
        refresh_llm_tailoring=bool(refresh_llm_tailoring),
        generate_llm_fallback=bool(generate_llm_fallback),
        generate_llm_adjudication=effective_generate_llm_adjudication,
        delete_seen_data=normalized_delete_seen_data,
    )
    cmd = ja._build_main_cmd(args, planning_only=bool(planning_only))

    runtime_payload = {
        "run_id": run_id,
        "child_pid": None,
        "status": "running",
        "started_at": _utc_now(),
        "finished_at": "",
        "current_stage": "startup",
        "completed_stages": [],
        "stage_order": [
            "startup",
            "scraping",
            "filtering",
            "dedupe",
            "ranking",
            "cache_filter",
            "details",
            "intelligence",
            "ai_evaluation_filter",
            "embedding_prefilter",
            "ai_evaluation",
            "resume_matching",
            "application_priority",
            "rag_export",
            "planning",
                        "finalization",
        ],
        "stage_started_at": _utc_now(),
        "stage_message": "Launching pipeline subprocess",
        "counts": {},
        "summary_message": "",
        "final_job_count": None,
        "return_code": None,
        "error": "",
        "output_dir": str(output_dir),
        "log_path": str(canonical_log_path),
        "status_path": str(canonical_status_path),
        "config": {
            "planning_only": bool(planning_only),
            "job_limit": int(job_limit),
            "job_packet_limit": int(job_packet_limit),
            "llm_actions": normalized_llm_actions.split(","),
            "generate_tailoring": bool(generate_tailoring),
            "generate_llm_tailoring": bool(generate_llm_tailoring),
            "refresh_llm_tailoring": bool(refresh_llm_tailoring),
            "generate_llm_fallback": bool(generate_llm_fallback),
            "generate_llm_adjudication": effective_generate_llm_adjudication,
            "delete_seen_data": normalized_delete_seen_data,
            "requested_output_dir": str(requested_output_dir),
            "storage_mode": "run_scoped_scratch" if owner_for_pipeline_gate else "legacy_output_dir",
            "owner_user_id": owner_for_pipeline_gate,
            "job_corpus_path": str(pipeline_job_corpus_path),
        },
    }
    canonical_status_path.write_text(
        json.dumps(runtime_payload, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )

    log_handle = canonical_log_path.open("w", encoding="utf-8", buffering=1)

    active_run_reserved = False
    redis_admission_lock: Dict[str, Any] = {}
    if owner_for_pipeline_gate:
        reservation_ttl_seconds = _active_pipeline_reservation_ttl_seconds()
        redis_admission_lock = _user_pipeline_redis_admission_lock_payload(
            owner_user_id=owner_for_pipeline_gate,
            run_id=run_id,
            ttl_seconds=reservation_ttl_seconds,
        )

        if (
            bool(redis_admission_lock.get("attempted", False))
            and not bool(redis_admission_lock.get("acquired", False))
            and _clean_text(redis_admission_lock.get("skipped")) == "lock_held"
        ):
            raise ValueError("A live pipeline run is already being admitted for this user. Try again in a few seconds.")

        reservation_payload = reserve_user_pipeline_active_run_postgres_payload(
            owner_user_id=owner_for_pipeline_gate,
            run_id=run_id,
            max_active_runs=_max_concurrent_user_pipeline_runs(),
            ttl_seconds=reservation_ttl_seconds,
            process_pid="",
            worker_id=_active_pipeline_worker_id(),
            output_dir=str(output_dir),
            status_path=str(canonical_status_path),
            metadata_json={
                "storage_mode": "postgres_active_run_admission",
                "requested_output_dir": str(requested_output_dir),
                "redis_admission_lock": redis_admission_lock,
            },
            database_url="",
            database_url_env="DATABASE_URL",
            psql_bin="psql",
            print_only=False,
            ensure_schema=True,
        )

        if not bool(reservation_payload.get("reserved", False)):
            _release_user_pipeline_redis_admission_lock_payload(redis_admission_lock)

            reason = _clean_text(reservation_payload.get("reason"))
            active_count = reservation_payload.get("active_count_before")
            max_count = reservation_payload.get("max_active_runs")

            if reason == "owner_already_running":
                raise ValueError("A live pipeline run is already in progress for this user.")

            if reason == "capacity_full":
                raise ValueError(
                    f"Live Pipeline capacity is currently full ({active_count}/{max_count}). Try again after a run finishes."
                )

            raise ValueError("Unable to reserve Live Pipeline capacity. Try again.")

        active_run_reserved = True
        runtime_payload["config"]["active_run_reservation"] = {
            "reserved": True,
            "active_count_after": reservation_payload.get("active_count_after"),
            "max_active_runs": reservation_payload.get("max_active_runs"),
            "ttl_seconds": reservation_payload.get("ttl_seconds"),
            "redis_admission_lock": {
                "attempted": bool(redis_admission_lock.get("attempted", False)),
                "acquired": bool(redis_admission_lock.get("acquired", False)),
                "skipped": _clean_text(redis_admission_lock.get("skipped")),
                "key": _clean_text(redis_admission_lock.get("key")),
                "ttl_seconds": redis_admission_lock.get("ttl_seconds"),
            },
        }


    child_env_extra = {
        "JOB_APP_PIPELINE_STATUS_PATH": str(canonical_status_path),
        "JOB_APP_PIPELINE_RUN_ID": run_id,
    }

    if owner_for_pipeline_gate:
        child_env_extra.update(
            {
                "JOB_STACK_OWNER_USER_ID": owner_for_pipeline_gate,
                "JOB_STACK_USER_PIPELINE_RUN_ID": run_id,
                "JOB_STACK_SEEN_JOBS_BACKEND": "postgres",
                "JOB_STACK_JOB_CORPUS_PATH": str(pipeline_job_corpus_path),
                "JOB_STACK_USER_PIPELINE_MODE": "true",
            }
        )

    child_env = _pipeline_child_env(extra_env=child_env_extra)
    runtime_payload["config"]["child_env"] = _pipeline_child_env_summary(child_env)

    try:
        process = subprocess.Popen(
            cmd,
            stdout=log_handle,
            stderr=subprocess.STDOUT,
            env=child_env,
        )
    except Exception as exc:
        log_handle.close()
        if owner_for_pipeline_gate and active_run_reserved:
            _release_user_pipeline_active_run(
                owner_user_id=owner_for_pipeline_gate,
                run_id=run_id,
                terminal_status="failed",
            )

        runtime_payload["status"] = "failed"
        runtime_payload["finished_at"] = _utc_now()
        runtime_payload["return_code"] = 1
        runtime_payload["error"] = repr(exc)
        runtime_payload["summary_message"] = "Failed to launch pipeline subprocess"
        canonical_status_path.write_text(
            json.dumps(runtime_payload, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )
        raise
    
    runtime_payload["child_pid"] = process.pid
    _write_runtime_status_file(canonical_status_path, runtime_payload)

    target_state = _new_pipeline_run_state()
    target_state["process"] = process
    target_state["log_handle"] = log_handle
    target_state["status"] = "running"
    target_state["started_at"] = _utc_now()
    target_state["finished_at"] = ""
    target_state["return_code"] = None
    target_state["command"] = cmd
    target_state["output_dir"] = str(output_dir)
    target_state["log_path"] = str(canonical_log_path)
    target_state["status_path"] = str(canonical_status_path)
    target_state["run_id"] = run_id
    target_state["child_pid"] = process.pid
    target_state["error"] = ""
    target_state["owner_user_id"] = owner_for_pipeline_gate

    if owner_for_pipeline_gate:
        _set_owner_active_pipeline_state(owner_for_pipeline_gate, target_state)
        started_payload = pipeline_status_payload(
            owner_user_id=owner_for_pipeline_gate,
            state=target_state,
        )
        _persist_user_pipeline_status_snapshot(
            owner_user_id=owner_for_pipeline_gate,
            status_payload=started_payload.get("pipeline", {}),
        )
    else:
        _PIPELINE_RUN_STATE.clear()
        _PIPELINE_RUN_STATE.update(target_state)
        started_payload = pipeline_status_payload()

    return {
        "ok": True,
        "message": "Live pipeline started.",
        "pipeline": started_payload["pipeline"],
    }

def _clean_text(value: Any) -> str:
    return str(value or "").strip()

def _scan_issue_display_label(value: Any) -> str:
    raw = _clean_text(value)
    if not raw:
        return ""

    key = raw.strip().lower()
    normalized_key = key.replace("-", "_").replace(" ", "_")

    if key in _SCAN_SIGNAL_DISPLAY_OVERRIDES:
        return _SCAN_SIGNAL_DISPLAY_OVERRIDES[key]

    if normalized_key in _SCAN_DIMENSION_DISPLAY_LABELS:
        return _SCAN_DIMENSION_DISPLAY_LABELS[normalized_key]

    cleaned = raw.replace("_", " ").replace("-", " ").strip()
    if cleaned.isupper() and len(cleaned) <= 5:
        return cleaned

    return cleaned[:1].upper() + cleaned[1:]

def _scan_issue_normalized_key(value: Any) -> str:
    return _clean_text(value).strip().lower().replace("-", "_").replace(" ", "_")


def _scan_issue_is_dimension_key(value: Any) -> bool:
    return _scan_issue_normalized_key(value) in _SCAN_DIMENSION_DISPLAY_LABELS


def _scan_issue_text_blob(row: Dict[str, Any]) -> str:
    values = [
        row.get("final_replacement_text"),
        row.get("rewrite_direction"),
        row.get("rewrite_instruction"),
        row.get("original_text"),
        row.get("current_evidence"),
        row.get("parent_bullet"),
        row.get("why_selected"),
        row.get("why_this_improves_match"),
        row.get("placement_guidance"),
    ]

    return re.sub(
        r"\s+",
        " ",
        " ".join(_clean_text(value).lower() for value in values if _clean_text(value)),
    ).strip()


def _scan_issue_text_contains_term(text: str, term: str) -> bool:
    clean_term = _clean_text(term).lower()
    if not text or not clean_term:
        return False

    pattern = r"(?<![a-z0-9])" + re.escape(clean_term) + r"(?![a-z0-9])"
    return bool(re.search(pattern, text))


def _scan_issue_text_signal_terms(row: Dict[str, Any]) -> List[str]:
    blob = _scan_issue_text_blob(row)
    if not blob:
        return []

    output: List[str] = []
    seen = set()

    for term in _SCAN_TITLE_SIGNAL_PATTERNS:
        if not _scan_issue_text_contains_term(blob, term):
            continue

        label = _scan_issue_display_label(term)
        key = label.lower()
        if not label or key in seen:
            continue

        seen.add(key)
        output.append(label)

        if len(output) >= 4:
            break

    return output

def _scan_issue_display_terms(row: Dict[str, Any]) -> List[str]:
    preferred_fields = (
        "supported_jd_signals",
        "jd_signal_terms",
        "supported_terms",
        "matched_signals",
    )

    output: List[str] = []
    seen = set()

    for field in preferred_fields:
        raw_values = list(row.get(field, []) or [])
        for value in raw_values:
            if _scan_issue_is_dimension_key(value):
                continue

            label = _scan_issue_display_label(value)
            key = label.lower()
            if not label or key in seen:
                continue

            seen.add(key)
            output.append(label)

    return output


def _scan_issue_dimension_labels(row: Dict[str, Any]) -> List[str]:
    output: List[str] = []
    seen = set()

    for value in list(row.get("likely_impacted_dimensions", []) or []):
        label = _scan_issue_display_label(value)
        key = label.lower()
        if not label or key in seen:
            continue
        seen.add(key)
        output.append(label)

    return output


def _scan_issue_group_label(group_id: str) -> str:
    return _SCAN_GROUP_LABELS.get(_clean_text(group_id), "Skills")

def _scan_issue_candidate_id(row: Dict[str, Any]) -> str:
    return (
        _clean_text(row.get("replacement_candidate_id"))
        or _clean_text(row.get("candidate_id"))
        or _clean_text(row.get("decision_id"))
        or _clean_text(row.get("source_bullet_id"))
    )


def _scan_issue_supported_signals(row: Dict[str, Any]) -> List[str]:
    raw_values = (
        list(row.get("supported_jd_signals", []) or [])
        + list(row.get("jd_signal_terms", []) or [])
        + list(row.get("supported_terms", []) or [])
        + list(row.get("matched_signals", []) or [])
        + list(row.get("likely_impacted_dimensions", []) or [])
    )

    output: List[str] = []
    seen = set()

    for value in raw_values:
        label = _scan_issue_display_label(value)
        key = label.lower()
        if not label or key in seen:
            continue
        seen.add(key)
        output.append(label)

    return output


def _scan_issue_is_deterministic_only_replacement(row: Dict[str, Any]) -> bool:
    source_blob = " ".join(
        _clean_text(row.get(field)).lower()
        for field in (
            "replacement_source",
            "patch_generation_method",
            "preferred_rewrite_source",
            "selected_source",
            "source_family",
            "resolved_to_source_family",
            "candidate_source",
            "rewrite_source",
            "final_replacement_source",
            "generation_source",
        )
        if _clean_text(row.get(field))
    )
    if not source_blob:
        return False
    if "llm" in source_blob or "live_ai" in source_blob:
        return False
    return "deterministic" in source_blob


def _scan_issue_title(row: Dict[str, Any], *, fallback_label: str) -> str:
    display_terms = _scan_issue_display_terms(row)
    if display_terms:
        return ", ".join(display_terms[:3])

    text_terms = _scan_issue_text_signal_terms(row)
    if text_terms:
        return ", ".join(text_terms[:3])

    dimension_labels = _scan_issue_dimension_labels(row)
    if dimension_labels:
        return ", ".join(dimension_labels[:2])

    for field in (
        "final_replacement_text",
        "rewrite_direction",
        "rewrite_instruction",
        "original_text",
        "current_evidence",
        "parent_bullet",
    ):
        text = _clean_text(row.get(field))
        if text:
            return text[:72] + ("..." if len(text) > 72 else "")

    return fallback_label

def _scan_issue_group_id_for_row(row: Dict[str, Any], *, lane: str) -> str:
    return _signal_family_group_id_for_signals(
        list(row.get("supported_jd_signals", []) or [])
        + list(row.get("jd_signal_terms", []) or [])
        + list(row.get("supported_terms", []) or [])
        + list(row.get("matched_signals", []) or []),
        likely_impacted_dimensions=list(row.get("likely_impacted_dimensions", []) or []),
        lane=lane,
    )

def _scan_issue_unique_display_labels(values: List[Any]) -> List[str]:
    output: List[str] = []
    seen = set()

    for value in values or []:
        label = _scan_issue_display_label(value)
        key = label.lower()
        if not label or key in seen:
            continue
        seen.add(key)
        output.append(label)

    return output


def _scan_issue_canonical_term(value: Any) -> str:
    label = _scan_issue_display_label(value)
    return re.sub(r"\s+", " ", label.strip().lower())


def _scan_issue_term_family(value: Any, dimensions: List[Any] | None = None) -> str:
    return _signal_family_term_family(value, dimensions or [])


_SCAN_SOFT_SKILL_TERMS = {
    "adaptability",
    "collaboration",
    "communication",
    "cross functional",
    "cross-functional",
    "decision making",
    "empathy",
    "leadership",
    "mentoring",
    "organization",
    "ownership",
    "problem solving",
    "stakeholder management",
    "stakeholders",
    "teamwork",
    "written communication",
}

_SCAN_OTHER_KEYWORD_TERMS = {
    "business impact",
    "customer",
    "customers",
    "domain",
    "industry",
    "metrics",
    "ownership scope",
    "scale",
    "seniority",
}


def _scan_issue_skill_type(value: Any, dimensions: List[Any] | None = None) -> Tuple[str, str]:
    return _signal_family_skill_type(value, dimensions or [])


def _scan_summary_has_advanced_degree_requirement(summary: Dict[str, Any] | None = None) -> bool:
    if not isinstance(summary, dict):
        return False

    values: List[Any] = []
    for key in (
        "required_education",
        "education_requirements",
        "matched_required",
        "missing_required",
        "required_terms",
        "missing_requirements",
    ):
        raw = summary.get(key)
        if isinstance(raw, dict):
            values.extend(raw.keys())
            values.extend(raw.values())
        elif isinstance(raw, list):
            values.extend(raw)
        elif raw:
            values.append(raw)

    blob = " ".join(_clean_text(value).lower() for value in values if _clean_text(value))
    return bool(
        re.search(
            r"\b(?:master|masters|ms|m\.s\.|mba|phd|ph\.d\.|doctorate|graduate degree|advanced degree)\b",
            blob,
        )
    )


def _scan_issue_score_priority(
    *,
    group_id: str,
    skill_type: str = "",
    dimensions: List[Any] | None = None,
    check_id: str = "",
    advanced_degree_required: bool = False,
) -> Tuple[int, str]:
    dimension_keys = {_scan_issue_normalized_key(item) for item in list(dimensions or [])}

    if group_id == "skills":
        if "title_alignment" in dimension_keys:
            return 3, "Job title"
        if skill_type == "soft_skill":
            return 4, "Soft skills"
        if skill_type == "other_keyword":
            return 5, "Other keywords"
        return 1, "Hard skills"

    if group_id == "searchability" and check_id == "education_info_searchable" and advanced_degree_required:
        return 2, "Education level"

    if group_id == "searchability" and check_id == "job_title_alignment":
        return 3, "Job title"

    return 0, "Non-scoring guidance"


def _scan_issue_score_priority_weight(
    *,
    group_id: str,
    skill_type: str = "",
    dimensions: List[Any] | None = None,
    check_id: str = "",
    advanced_degree_required: bool = False,
) -> float:
    weights = get_dimension_weights()
    dimension_keys = [
        _scan_issue_normalized_key(item)
        for item in list(dimensions or [])
        if _scan_issue_normalized_key(item) in weights
    ]
    if dimension_keys:
        return round(max(weights.get(key, 0.0) for key in dimension_keys), 6)

    if group_id == "skills":
        if skill_type == "soft_skill":
            return round(weights.get("stakeholder_translation_alignment", 0.0), 6)
        if skill_type == "other_keyword":
            return round(
                max(
                    weights.get("business_context_alignment", 0.0),
                    weights.get("domain_relevance", 0.0),
                ),
                6,
            )
        return round(
            max(
                weights.get("required_skills_alignment", 0.0),
                weights.get("tooling_alignment", 0.0),
                weights.get("analytics_ml_depth", 0.0),
            ),
            6,
        )

    if group_id == "searchability" and check_id == "job_title_alignment":
        return round(weights.get("title_alignment", 0.0), 6)

    if group_id == "searchability" and check_id == "education_info_searchable" and advanced_degree_required:
        return 0.0

    return 0.0


def _scan_issue_extract_summary_terms(summary: Dict[str, Any] | None, keys: List[str]) -> List[str]:
    if not isinstance(summary, dict):
        return []

    values: List[Any] = []
    for key in keys:
        raw = summary.get(key)
        if isinstance(raw, dict):
            values.extend(raw.keys())
        elif isinstance(raw, list):
            values.extend(raw)
        elif raw:
            values.append(raw)

    return _scan_issue_unique_display_labels(values)


def _scan_jd_text_from_record(record: Dict[str, Any] | None = None) -> str:
    if not isinstance(record, dict):
        return ""

    values: List[Any] = []
    for key in (
        "retrieval_text",
        "description",
        "job_description",
        "raw_description",
        "preview",
        "requirements",
        "responsibilities",
        "qualifications",
    ):
        raw = record.get(key)
        if isinstance(raw, list):
            values.extend(raw)
        elif isinstance(raw, dict):
            values.extend(raw.values())
        elif raw:
            values.append(raw)

    return "\n".join(_clean_text(value) for value in values if _clean_text(value)).strip()


def _scan_jd_context_chunk_is_metadata(chunk: str) -> bool:
    text = _clean_text(chunk)
    lower = text.lower()
    if not text:
        return True

    metadata_labels = re.findall(
        r"\b(?:company|title|location|source|role family|seniority|job id|posted|salary|employment type)\s*:",
        lower,
    )
    if len(metadata_labels) >= 2:
        return True

    if re.match(
        r"(?i)^(?:company|title|location|source|role family|seniority|job id|posted|salary|employment type)\s*:",
        text,
    ):
        return True

    if re.match(
        r"(?i)^[A-Z][A-Za-z .'-]+,\s*[A-Z][A-Za-z .'-]+,\s*(?:United States|USA)\b",
        text,
    ):
        return True

    return False


def _scan_jd_context_anchors_for_term(
    term: str,
    jd_record: Dict[str, Any] | None = None,
    *,
    limit: int = 2,
) -> List[Dict[str, str]]:
    jd_text = _scan_jd_text_from_record(jd_record)
    if not jd_text or not _clean_text(term):
        return []

    chunks = []
    for chunk in re.split(r"(?<=[.!?])\s+|(?:\s+[•\-]\s+)|(?:\n+)", jd_text):
        clean_chunk = _clean_text(chunk)
        if not clean_chunk:
            continue
        if _scan_jd_context_chunk_is_metadata(clean_chunk):
            continue
        chunks.append(clean_chunk)
    anchors: List[Dict[str, str]] = []
    seen = set()

    for chunk in chunks:
        if not _scan_issue_text_contains_term(chunk.lower(), term):
            continue
        key = _normalize_tailoring_workspace_text_key(chunk)
        if not key or key in seen:
            continue
        seen.add(key)
        anchors.append({
            "type": "job_description",
            "text": chunk[:260],
            "term": _scan_issue_display_label(term),
        })
        if len(anchors) >= limit:
            break

    return anchors


_SCAN_PREDICTED_SKILL_TERMS_BY_ROLE = {
    "machine_learning_engineer": [
        "Python",
        "PyTorch",
        "Tensorflow",
        "Spark",
        "Kubernetes",
        "Docker",
    ],
    "data_scientist": [
        "Python",
        "SQL",
        "Machine learning",
        "Statistical modeling",
        "A/B testing",
        "Causal inference",
    ],
    "analytics_engineer": [
        "SQL",
        "dbt",
        "Snowflake",
        "BigQuery",
        "Airflow",
    ],
}


def _scan_predicted_skill_role_key(
    *,
    tailoring_summary: Dict[str, Any] | None = None,
    jd_record: Dict[str, Any] | None = None,
) -> str:
    summary = dict(tailoring_summary or {})
    record = dict(jd_record or {})
    role_blob = " ".join(
        _clean_text(value).lower()
        for value in (
            _scan_searchability_target_title(summary),
            summary.get("role_archetype"),
            record.get("role_archetype"),
            record.get("role_family"),
            record.get("title"),
            record.get("job_title"),
        )
        if _clean_text(value)
    )

    if "machine learning engineer" in role_blob or "ml_engineer" in role_blob or "ml engineer" in role_blob:
        return "machine_learning_engineer"
    if "analytics engineer" in role_blob:
        return "analytics_engineer"
    if "data scientist" in role_blob or "data science" in role_blob:
        return "data_scientist"
    return ""


def _build_predicted_skill_scan_rows(
    *,
    existing_issues: List[Dict[str, Any]],
    resume_evidence: Any = None,
    tailoring_summary: Dict[str, Any] | None = None,
    jd_record: Dict[str, Any] | None = None,
) -> List[Dict[str, Any]]:
    role_key = _scan_predicted_skill_role_key(
        tailoring_summary=tailoring_summary,
        jd_record=jd_record,
    )
    if not role_key:
        return []

    existing_terms = {
        _scan_issue_canonical_term(issue.get("display_term") or issue.get("title"))
        for issue in existing_issues
        if _clean_text(issue.get("group_id")) == "skills"
    }
    jd_text = _scan_jd_text_from_record(jd_record).lower()

    rows: List[Dict[str, Any]] = []
    for term in _SCAN_PREDICTED_SKILL_TERMS_BY_ROLE.get(role_key, []):
        canonical = _scan_issue_canonical_term(term)
        if not canonical or canonical in existing_terms:
            continue
        if _scan_issue_text_contains_term(jd_text, term):
            continue

        matched_count = _scan_resume_term_match_count(term, resume_evidence)
        row_action_type = "predicted_skill"
        rows.append(
            {
                "issue_id": f"scan_issue:skills:predicted:{canonical.replace(' ', '_')}",
                "candidate_id": "",
                "replacement_candidate_id": "",
                "source_lane": "predicted_skill_check",
                "keyword_source": "predicted_skill",
                "group_id": "skills",
                "group_label": _scan_issue_group_label("skills"),
                "bucket": "predicted",
                "bucket_label": "Predicted skills",
                "title": _scan_issue_display_label(term),
                "display_term": _scan_issue_display_label(term),
                "canonical_term": canonical,
                "term_family": "skills",
                "skill_type": "hard_skill",
                "skill_type_label": "Hard skill",
                "score_priority_rank": 0,
                "score_priority_label": "Predicted role-adjacent skill",
                "score_priority_weight": 0.0,
                "score_priority_source": "role_based_prediction",
                "matched_count": matched_count,
                "required_count": 0,
                "coverage_label": "Predicted",
                "matched_count_label": "",
                "evidence_anchors": _scan_resume_evidence_anchors_for_term(term, resume_evidence),
                "jd_context_anchors": [],
                "jd_context_label": "",
                "has_ai_suggestion": False,
                "linked_candidate_ids": [],
                "best_candidate_id": "",
                "row_action_type": row_action_type,
                "row_action_label": "Predicted",
                "scan_issue_type": row_action_type,
                "severity": "low",
                "projected_score_delta_points": None,
                "is_visible_in_review": True,
                "can_accept": False,
                "can_accept_all": False,
                "can_focus_preview": False,
                "anchor_strategy": "none",
                "supported_jd_signals": [],
                "reason": "Common role-adjacent skill inferred from the target role, not explicitly detected in the JD.",
                "predicted_skill": True,
                "prediction_source": role_key,
                "raw": {
                    "source": "role_based_predicted_skill",
                    "role_key": role_key,
                },
            }
        )

    return rows


def _build_other_keyword_scan_rows(
    *,
    existing_issues: List[Dict[str, Any]],
    resume_evidence: Any = None,
    jd_record: Dict[str, Any] | None = None,
) -> List[Dict[str, Any]]:
    jd_text = _scan_jd_text_from_record(jd_record)
    if not jd_text:
        return []

    existing_terms = {
        _scan_issue_canonical_term(issue.get("display_term") or issue.get("title"))
        for issue in existing_issues
        if _clean_text(issue.get("group_id")) == "skills"
    }

    rows: List[Dict[str, Any]] = []
    for term in DOMAIN_SIGNAL_PATTERNS:
        canonical = _scan_issue_canonical_term(term)
        if not canonical or canonical in existing_terms:
            continue
        if not _scan_issue_text_contains_term(jd_text.lower(), term):
            continue

        rows.append(
            {
                "issue_id": f"scan_issue:skills:other_keyword:{canonical.replace(' ', '_')}",
                "candidate_id": "",
                "replacement_candidate_id": "",
                "source_lane": "other_keyword_check",
                "keyword_source": "other_keyword",
                "group_id": "skills",
                "group_label": _scan_issue_group_label("skills"),
                "bucket": "other_keyword",
                "bucket_label": "Other keywords",
                "title": _scan_issue_display_label(term),
                "display_term": _scan_issue_display_label(term),
                "canonical_term": canonical,
                "term_family": "domain",
                "skill_type": "other_keyword",
                "skill_type_label": "Other keyword",
                "score_priority_rank": 5,
                "score_priority_label": "Other keywords",
                "score_priority_weight": 0.0,
                "score_priority_source": "lower_impact_keyword_group",
                "matched_count": _scan_resume_term_match_count(term, resume_evidence),
                "required_count": 0,
                "coverage_label": "Keyword",
                "matched_count_label": "",
                "evidence_anchors": _scan_resume_evidence_anchors_for_term(term, resume_evidence),
                "jd_context_anchors": _scan_jd_context_anchors_for_term(term, jd_record),
                "jd_context_label": "",
                "has_ai_suggestion": False,
                "linked_candidate_ids": [],
                "best_candidate_id": "",
                "row_action_type": "other_keyword",
                "row_action_label": "Keyword",
                "scan_issue_type": "other_keyword",
                "severity": "low",
                "projected_score_delta_points": None,
                "is_visible_in_review": True,
                "can_accept": False,
                "can_accept_all": False,
                "can_focus_preview": False,
                "anchor_strategy": "none",
                "supported_jd_signals": [_scan_issue_display_label(term)],
                "reason": "Lower-impact industry or domain term detected in the JD. Use only if truthful and useful.",
                "other_keyword": True,
                "raw": {
                    "source": "domain_other_keyword_check",
                },
            }
        )

    return rows


def _scan_resume_visible_term_keys(resume_evidence: Any) -> set[str]:
    if resume_evidence is None:
        return set()

    keys = {
        _scan_issue_canonical_term(term)
        for term in _scan_resume_visible_search_terms(resume_evidence)
        if _scan_issue_canonical_term(term)
    }

    raw_text = _scan_resume_document_text(resume_evidence).lower()
    for term in list(keys):
        if term:
            keys.add(term.lower())

    for term in _SCAN_TITLE_SIGNAL_PATTERNS:
        label = _scan_issue_canonical_term(term)
        if label and _scan_issue_text_contains_term(raw_text, term):
            keys.add(label)

    return keys


def _scan_resume_term_match_count(term: str, resume_evidence: Any) -> int:
    if resume_evidence is None:
        return 0

    canonical = _scan_issue_canonical_term(term)
    if not canonical:
        return 0

    visible_term_hits = sum(
        1
        for value in _scan_resume_visible_search_terms(resume_evidence)
        if _scan_issue_canonical_term(value) == canonical
    )

    text = _scan_resume_document_text(resume_evidence).lower()
    text_hits = len(
        re.findall(
            r"(?<![a-z0-9])" + re.escape(canonical) + r"(?![a-z0-9])",
            text,
        )
    )

    return max(visible_term_hits, text_hits)


def _scan_resume_evidence_anchors_for_term(
    term: str,
    resume_evidence: Any,
    *,
    limit: int = 3,
) -> List[Dict[str, str]]:
    if resume_evidence is None:
        return []

    anchors: List[Dict[str, str]] = []
    seen = set()

    for bullet in _scan_resume_bullet_texts(resume_evidence):
        if not _scan_issue_text_contains_term(bullet.lower(), term):
            continue

        key = _normalize_tailoring_workspace_text_key(bullet)
        if not key or key in seen:
            continue

        seen.add(key)
        anchors.append({
            "type": "resume_bullet",
            "text": bullet,
        })

        if len(anchors) >= limit:
            break

    return anchors


def _scan_issue_best_candidate(issues: List[Dict[str, Any]]) -> Dict[str, Any] | None:
    if not issues:
        return None

    def sort_key(issue: Dict[str, Any]) -> Tuple[int, int, int, str]:
        delta = score_delta_to_points(issue.get("projected_overall_delta"))
        row_type = _clean_text(issue.get("scan_issue_type"))
        return (
            1 if row_type == "direct_replacement" else 0,
            int(delta or 0),
            1 if issue.get("can_accept") else 0,
            _clean_text(issue.get("candidate_id")),
        )

    return sorted(issues, key=sort_key, reverse=True)[0]


def _scan_keyword_issue_from_term(
    *,
    group_id: str,
    term: str,
    matched_count: int,
    required_count: int,
    source: str,
    linked_issues: List[Dict[str, Any]] | None = None,
    evidence_anchors: List[Dict[str, str]] | None = None,
    jd_context_anchors: List[Dict[str, str]] | None = None,
    fallback_bucket: str = "",
    reason: str = "",
) -> Dict[str, Any]:
    linked_issues = list(linked_issues or [])
    evidence_anchors = list(evidence_anchors or [])
    jd_context_anchors = list(jd_context_anchors or [])
    best_issue = _scan_issue_best_candidate(linked_issues)
    direct_issues = [
        issue for issue in linked_issues
        if _clean_text(issue.get("scan_issue_type")) == "direct_replacement"
        and issue.get("can_accept")
    ]
    direct_issue = _scan_issue_best_candidate(direct_issues)
    negative_only = bool(linked_issues) and all(
        _clean_text(issue.get("scan_issue_type")) == "rejected_by_score_gate"
        for issue in linked_issues
    )

    if negative_only:
        row_action_type = "hidden_rejected"
        bucket = "hidden"
    elif direct_issue:
        row_action_type = "direct_replacement"
        bucket = "ai"
        best_issue = direct_issue
    elif matched_count > 0:
        row_action_type = "matched"
        bucket = "matched"
    elif linked_issues and any(
        not _scan_issue_is_deterministic_only_replacement(dict(issue.get("raw", {}) or issue))
        for issue in linked_issues
    ):
        row_action_type = "phrase_generation"
        bucket = fallback_bucket or "missing"
    else:
        row_action_type = "manual_guidance"
        bucket = fallback_bucket or "missing"

    display_term = _scan_issue_display_label(term)
    canonical_term = _scan_issue_canonical_term(display_term)
    base = dict(best_issue or {})
    skill_type = ""
    skill_type_label = ""
    score_priority_rank = 0
    score_priority_label = "Non-scoring guidance"
    if group_id == "skills":
        skill_type, skill_type_label = _scan_issue_skill_type(
            display_term,
            base.get("likely_impacted_dimensions", []),
        )
        score_priority_rank, score_priority_label = _scan_issue_score_priority(
            group_id=group_id,
            skill_type=skill_type,
            dimensions=base.get("likely_impacted_dimensions", []),
        )
    required_count = max(1, int(required_count or 0))
    matched_count = max(0, int(matched_count or 0))
    if row_action_type == "matched":
        matched_prefix = "Backed" if evidence_anchors else "Seen"
        coverage_label = f"{matched_prefix} {matched_count}"
    else:
        coverage_label = f"{matched_count}/{required_count}"
    linked_candidate_ids = [
        _clean_text(issue.get("candidate_id"))
        for issue in linked_issues
        if _clean_text(issue.get("candidate_id"))
    ]
    best_candidate_id = _clean_text((best_issue or {}).get("candidate_id"))
    best_delta_points = score_delta_to_points((best_issue or {}).get("projected_overall_delta"))
    row_action_label = {
        "direct_replacement": "AI replacement",
        "phrase_generation": "Phrase",
        "manual_guidance": "Guidance",
        "matched": coverage_label,
        "hidden_rejected": "Hidden",
    }.get(row_action_type, "Guidance")
    severity = "low"
    if row_action_type == "direct_replacement":
        severity = "high"
    elif row_action_type in {"phrase_generation", "manual_guidance"}:
        severity = "medium"
    elif row_action_type == "hidden_rejected":
        severity = "diagnostic"

    base.update(
        {
            "issue_id": f"scan_issue:{group_id}:keyword:{canonical_term.replace(' ', '_')}",
            "candidate_id": best_candidate_id,
            "replacement_candidate_id": best_candidate_id,
            "source_lane": _clean_text((best_issue or {}).get("source_lane")) or "keyword_scan",
            "keyword_source": source,
            "group_id": group_id,
            "group_label": _scan_issue_group_label(group_id),
            "bucket": bucket,
            "bucket_label": (
                "AI Suggested"
                if row_action_type == "direct_replacement"
                else "Matched"
                if row_action_type == "matched"
                else "Missing / optimization opportunity"
            ),
            "title": display_term,
            "display_term": display_term,
            "canonical_term": canonical_term,
            "term_family": _scan_issue_term_family(display_term, base.get("likely_impacted_dimensions", [])),
            "skill_type": skill_type,
            "skill_type_label": skill_type_label,
            "score_priority_rank": score_priority_rank,
            "score_priority_label": score_priority_label,
            "score_priority_weight": _scan_issue_score_priority_weight(
                group_id=group_id,
                skill_type=skill_type,
                dimensions=base.get("likely_impacted_dimensions", []),
            ),
            "score_priority_source": "jobscan_match_report_order",
            "matched_count": matched_count,
            "required_count": required_count,
            "coverage_label": coverage_label,
            "matched_count_label": coverage_label if row_action_type == "matched" else "",
            "evidence_anchors": evidence_anchors,
            "jd_context_anchors": jd_context_anchors,
            "jd_context_label": _clean_text(jd_context_anchors[0].get("text")) if jd_context_anchors else "",
            "has_ai_suggestion": bool(linked_candidate_ids),
            "linked_candidate_ids": list(dict.fromkeys(linked_candidate_ids)),
            "best_candidate_id": best_candidate_id,
            "row_action_type": row_action_type,
            "row_action_label": row_action_label,
            "scan_issue_type": row_action_type,
            "severity": severity,
            "projected_score_delta_points": best_delta_points,
            "is_visible_in_review": row_action_type != "hidden_rejected",
            "can_accept": bool(row_action_type == "direct_replacement" and (best_issue or {}).get("can_accept")),
            "can_accept_all": bool(row_action_type == "direct_replacement" and (best_issue or {}).get("can_accept_all")),
            "can_focus_preview": bool(best_candidate_id and (best_issue or {}).get("can_focus_preview", True)),
            "anchor_strategy": "replacement_candidate" if best_candidate_id else "none",
            "supported_jd_signals": _scan_issue_unique_display_labels(
                [display_term]
                + list((best_issue or {}).get("supported_jd_signals", []) or [])
            ),
            "reason": reason or _clean_text((best_issue or {}).get("reason")),
            "raw": dict((best_issue or {}).get("raw", best_issue or {})),
        }
    )
    return base


def _scan_keyword_rows_from_replacement_issues(
    issues: List[Dict[str, Any]],
    *,
    resume_evidence: Any = None,
    tailoring_summary: Dict[str, Any] | None = None,
    jd_record: Dict[str, Any] | None = None,
) -> List[Dict[str, Any]]:
    resume_term_keys = _scan_resume_visible_term_keys(resume_evidence)
    matched_terms = _scan_issue_extract_summary_terms(
        tailoring_summary,
        ["matched_required", "matched_preferred", "matched_terms"],
    )
    missing_terms = _scan_issue_extract_summary_terms(
        tailoring_summary,
        ["missing_required", "missing_preferred", "missing_terms", "missing_requirements"],
    )

    matched_keys = {_scan_issue_canonical_term(term) for term in matched_terms if _scan_issue_canonical_term(term)}
    missing_keys = {_scan_issue_canonical_term(term) for term in missing_terms if _scan_issue_canonical_term(term)}

    grouped: Dict[str, List[Dict[str, Any]]] = {}
    display_by_key: Dict[str, str] = {}

    for issue in issues:
        if _clean_text(issue.get("group_id")) != "skills":
            continue

        terms = _scan_issue_display_terms(dict(issue.get("raw", {}) or issue)) or _scan_issue_display_terms(issue)
        if not terms:
            terms = [_clean_text(issue.get("title"))]

        for term in terms:
            canonical = _scan_issue_canonical_term(term)
            if not canonical:
                continue
            grouped.setdefault(canonical, []).append(issue)
            display_by_key.setdefault(canonical, _scan_issue_display_label(term))

    for term in matched_terms + missing_terms:
        canonical = _scan_issue_canonical_term(term)
        if canonical:
            grouped.setdefault(canonical, [])
            display_by_key.setdefault(canonical, _scan_issue_display_label(term))

    rows: List[Dict[str, Any]] = []
    for canonical in sorted(grouped.keys()):
        linked = grouped.get(canonical, [])
        display_term = display_by_key.get(canonical, canonical)
        is_matched = canonical in resume_term_keys or canonical in matched_keys
        is_missing = canonical in missing_keys
        required_count = max(1, len(linked), 1 if is_missing else 0)
        evidence_match_count = _scan_resume_term_match_count(display_term, resume_evidence)
        matched_count = max(1 if is_matched else 0, evidence_match_count)
        rows.append(
            _scan_keyword_issue_from_term(
                group_id="skills",
                term=display_term,
                matched_count=matched_count,
                required_count=required_count,
                source="tailoring_keyword_scan",
                linked_issues=linked,
                evidence_anchors=_scan_resume_evidence_anchors_for_term(
                    display_term,
                    resume_evidence,
                ),
                jd_context_anchors=_scan_jd_context_anchors_for_term(
                    display_term,
                    jd_record,
                ),
                fallback_bucket="missing",
                reason="Job-specific skill signal derived from scorer and replacement evidence.",
            )
        )

    return rows


def _scan_keyword_rows_from_non_skill_replacement_issues(
    issues: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    grouped: Dict[Tuple[str, str], List[Dict[str, Any]]] = {}
    display_by_key: Dict[Tuple[str, str], str] = {}

    for issue in issues:
        group_id = _clean_text(issue.get("group_id"))
        if not group_id or group_id == "skills":
            continue

        terms = _scan_issue_display_terms(dict(issue.get("raw", {}) or issue)) or _scan_issue_display_terms(issue)
        if not terms:
            terms = [_clean_text(issue.get("title"))]

        for term in terms:
            canonical = _scan_issue_canonical_term(term)
            if not canonical:
                continue
            key = (group_id, canonical)
            grouped.setdefault(key, []).append(issue)
            display_by_key.setdefault(key, _scan_issue_display_label(term))

    rows: List[Dict[str, Any]] = []
    for (group_id, canonical), linked in sorted(grouped.items()):
        rows.append(
            _scan_keyword_issue_from_term(
                group_id=group_id,
                term=display_by_key.get((group_id, canonical), canonical),
                matched_count=0,
                required_count=max(1, len(linked)),
                source="tailoring_keyword_scan",
                linked_issues=linked,
                fallback_bucket="missing",
                reason="Job-specific recruiter/search signal derived from tailoring guidance.",
            )
        )

    return rows


def _scan_keyword_rows_from_generic_issues(issues: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []

    for issue in issues:
        group_id = _clean_text(issue.get("group_id"))
        if group_id == "skills":
            continue

        row_action_type = "matched" if issue.get("bucket") == "matched" else "manual_guidance"
        display_term = _clean_text(issue.get("display_term")) or _clean_text(issue.get("title"))
        canonical = _scan_issue_canonical_term(display_term)
        is_check_group = group_id in {"searchability", "formatting", "recruiter_tips"}
        matched_count = 1 if row_action_type == "matched" else 0
        required_count = 1
        if row_action_type == "matched" and is_check_group:
            coverage_label = "Pass"
        elif row_action_type == "matched":
            coverage_label = f"Seen {matched_count}"
        else:
            coverage_label = f"{matched_count}/{required_count}"
        updated = dict(issue)
        updated.update(
            {
                "display_term": display_term,
                "canonical_term": canonical,
                "term_family": group_id,
                "score_priority_rank": int(issue.get("score_priority_rank") or 0),
                "score_priority_label": _clean_text(issue.get("score_priority_label")) or "Non-scoring guidance",
                "score_priority_weight": float(issue.get("score_priority_weight") or 0.0),
                "score_priority_source": _clean_text(issue.get("score_priority_source")) or "jobscan_match_report_order",
                "matched_count": matched_count,
                "required_count": required_count,
                "coverage_label": coverage_label,
                "matched_count_label": coverage_label if row_action_type == "matched" else "",
                "has_ai_suggestion": False,
                "linked_candidate_ids": [],
                "best_candidate_id": "",
                "row_action_type": row_action_type,
                "row_action_label": (
                    "Check"
                    if row_action_type == "matched" and is_check_group
                    else coverage_label
                    if row_action_type == "matched"
                    else "Guidance"
                ),
                "scan_issue_type": row_action_type,
                "severity": "low" if row_action_type == "matched" else "medium",
                "is_visible_in_review": True,
            }
        )
        rows.append(updated)

    return rows


def _scan_score_delta_points(value: Any) -> int | None:
    return score_delta_to_points(value)


def _scan_resume_document_text(resume_evidence: Any) -> str:
    document = getattr(resume_evidence, "document", None)
    return (
        _clean_text(getattr(document, "raw_text", ""))
        or _clean_text(getattr(document, "normalized_text", ""))
    )


_US_STATE_ABBREVIATIONS = {
    "AL", "AK", "AZ", "AR", "CA", "CO", "CT", "DE", "FL", "GA",
    "HI", "ID", "IL", "IN", "IA", "KS", "KY", "LA", "ME", "MD",
    "MA", "MI", "MN", "MS", "MO", "MT", "NE", "NV", "NH", "NJ",
    "NM", "NY", "NC", "ND", "OH", "OK", "OR", "PA", "RI", "SC",
    "SD", "TN", "TX", "UT", "VT", "VA", "WA", "WV", "WI", "WY",
    "DC",
}

_WORKSPACE_PERSONAL_DETAIL_FIELDS = (
    "name",
    "city",
    "state",
    "contact",
    "email",
    "linkedin",
    "github",
)


def _normalize_workspace_profile_url(value: Any, expected_host: str) -> str:
    raw = _clean_text(value)
    if not raw:
        return ""

    candidate = raw if re.match(r"^https?://", raw, flags=re.IGNORECASE) else f"https://{raw}"
    try:
        parsed = urlparse(candidate)
    except Exception:
        return ""

    host = (parsed.netloc or "").lower()
    if host.startswith("www."):
        host = host[4:]
    path = (parsed.path or "").rstrip("/")

    if expected_host == "linkedin":
        if host != "linkedin.com" or not re.match(r"^/in/[^/]+$", path, flags=re.IGNORECASE):
            return ""
        if "github" in path.lower():
            return ""
    elif expected_host == "github":
        if host != "github.com" or not re.match(r"^/[A-Za-z0-9-]+$", path):
            return ""

    return urlunparse(("https", host, path, "", "", ""))


def _normalize_workspace_personal_details(value: Any) -> Dict[str, str]:
    if isinstance(value, dict):
        raw_items = value
    else:
        raw_text = _clean_text(value)
        if not raw_text:
            raw_items = {}
        else:
            try:
                parsed = json.loads(raw_text)
            except Exception as exc:
                raise ValueError("personal_details must be a JSON object.") from exc
            if not isinstance(parsed, dict):
                raise ValueError("personal_details must be a JSON object.")
            raw_items = parsed

    normalized = {field: "" for field in _WORKSPACE_PERSONAL_DETAIL_FIELDS}
    for field in _WORKSPACE_PERSONAL_DETAIL_FIELDS:
        normalized[field] = _clean_text(raw_items.get(field))

    state = normalized["state"].upper()
    normalized["state"] = state if state in _US_STATE_ABBREVIATIONS else state[:2]
    normalized["linkedin"] = _normalize_workspace_profile_url(normalized["linkedin"], "linkedin")
    normalized["github"] = _normalize_workspace_profile_url(normalized["github"], "github")
    return normalized


def _extract_resume_personal_details(
    resume_evidence: Any,
    *,
    owner_user_id: str = "",
) -> Dict[str, str]:
    details = _normalize_workspace_personal_details({})
    if resume_evidence is None:
        return details

    document = getattr(resume_evidence, "document", None)
    resume_name = _clean_text(getattr(document, "resume_name", ""))
    if resume_name:
        try:
            resume_pdf_path = planning_resume_preview_path(
                resume_name,
                owner_user_id=owner_user_id,
            )
            for link_item in _workspace_export_extract_pdf_header_link_items(resume_pdf_path):
                label = _clean_text(link_item.get("label")).lower()
                uri = _clean_text(link_item.get("uri"))
                if uri and ("linkedin" in label or "linkedin.com" in uri.lower()):
                    details["linkedin"] = uri
                if uri and ("github" in label or "github.com" in uri.lower()):
                    details["github"] = uri
        except Exception:
            pass

    raw_text = _scan_resume_document_text(resume_evidence)
    if not raw_text:
        return details

    lines = [_clean_text(line) for line in raw_text.splitlines()]
    lines = [line for line in lines if line]
    header_lines = lines[:10]

    email_match = re.search(
        r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b",
        raw_text,
        flags=re.IGNORECASE,
    )
    if email_match:
        details["email"] = email_match.group(0)

    phone_match = re.search(
        r"(?:\+?1[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?)\d{3}[\s.-]?\d{4}\b",
        raw_text,
    )
    if phone_match:
        details["contact"] = phone_match.group(0).strip()

    linkedin_match = re.search(
        r"(?:https?://)?(?:www\.)?linkedin\.com/[^\s|,;]+",
        raw_text,
        flags=re.IGNORECASE,
    )
    if linkedin_match:
        details["linkedin"] = linkedin_match.group(0).strip().rstrip(".")
    else:
        compact_linkedin_match = re.search(
            r"\blinkedin\s*[:|-]\s*([A-Za-z0-9_.-]+)",
            raw_text,
            flags=re.IGNORECASE,
        )
        if compact_linkedin_match:
            details["linkedin"] = f"linkedin.com/in/{compact_linkedin_match.group(1).strip('/')}"

    github_match = re.search(
        r"(?:https?://)?(?:www\.)?github\.com/[A-Za-z0-9-]+",
        raw_text,
        flags=re.IGNORECASE,
    )
    if github_match:
        details["github"] = github_match.group(0).strip().rstrip(".")

    section_headings = {
        "summary", "experience", "work experience", "professional experience",
        "education", "skills", "technical skills", "projects", "certifications",
    }
    for line in header_lines:
        lower = line.lower().strip(":")
        if lower in section_headings:
            break
        if "@" in line or "linkedin.com" in lower or "github.com" in lower:
            continue
        if re.search(r"\d{3}", line):
            continue
        if len(line) > 80:
            continue
        if re.search(r"[A-Za-z]", line):
            details["name"] = line
            break

    state_pattern = "|".join(sorted(_US_STATE_ABBREVIATIONS))
    location_pattern = re.compile(
        rf"\b([A-Za-z][A-Za-z .'-]{{1,40}}?)\s*,\s*({state_pattern})(?:\s+\d{{5}}(?:-\d{{4}})?)?\b",
        flags=re.IGNORECASE,
    )
    for line in header_lines:
        match = location_pattern.search(line)
        if not match:
            continue
        city = _clean_text(match.group(1)).strip(" ,|-")
        state = _clean_text(match.group(2)).upper()
        if city and state in _US_STATE_ABBREVIATIONS:
            details["city"] = city
            details["state"] = state
            break

    return details


def _workspace_personal_details_contact_text(details: Dict[str, str]) -> str:
    safe_details = _normalize_workspace_personal_details(details)
    location = ", ".join(
        item for item in [safe_details.get("city", ""), safe_details.get("state", "")]
        if _clean_text(item)
    )
    return " | ".join(
        item for item in [
            location,
            safe_details.get("contact", ""),
            safe_details.get("email", ""),
            "LinkedIn" if _clean_text(safe_details.get("linkedin")) else "",
            "GitHub" if _clean_text(safe_details.get("github")) else "",
        ]
        if _clean_text(item)
    )


def _workspace_personal_details_link_items(details: Dict[str, str]) -> List[Dict[str, str]]:
    safe_details = _normalize_workspace_personal_details(details)
    link_items: List[Dict[str, str]] = []
    linkedin_url = _clean_text(safe_details.get("linkedin"))
    github_url = _clean_text(safe_details.get("github"))
    if linkedin_url:
        link_items.append({"label": "LinkedIn", "uri": linkedin_url})
    if github_url:
        link_items.append({"label": "GitHub", "uri": github_url})
    return link_items


def _scan_resume_bullet_texts(resume_evidence: Any) -> List[str]:
    bullets: List[str] = []

    for entry in list(getattr(resume_evidence, "experience_entries", []) or []):
        bullets.extend(_clean_text(bullet) for bullet in list(getattr(entry, "bullets", []) or []))

    for entry in list(getattr(resume_evidence, "project_entries", []) or []):
        bullets.extend(_clean_text(bullet) for bullet in list(getattr(entry, "bullets", []) or []))

    return [bullet for bullet in bullets if bullet]


def _scan_resume_visible_search_terms(resume_evidence: Any) -> List[str]:
    raw_terms: List[Any] = []

    for field_name in (
        "skills",
        "tools",
        "methods",
        "workflows",
        "tooling_signals",
        "analytics_ml_signals",
        "experimentation_signals",
        "domain_signals",
    ):
        raw_terms.extend(list(getattr(resume_evidence, field_name, []) or []))

    return _scan_issue_unique_display_labels(raw_terms)


def _scan_searchability_issue(
    *,
    check_id: str,
    bucket: str,
    title: str,
    reason: str,
    current_text: str = "",
    suggested_text: str = "",
    supported_jd_signals: List[str] | None = None,
    priority: str = "",
    score_priority_rank: int = 0,
    score_priority_label: str = "Non-scoring guidance",
) -> Dict[str, Any]:
    bucket_label = "Matched searchability signal" if bucket == "matched" else "Searchability opportunity"

    return {
        "issue_id": f"scan_issue:searchability:{check_id}",
        "candidate_id": "",
        "source_lane": "searchability_check",
        "group_id": "searchability",
        "group_label": _scan_issue_group_label("searchability"),
        "bucket": bucket,
        "bucket_label": bucket_label,
        "title": title,
        "status": "searchability_matched" if bucket == "matched" else "searchability_opportunity",
        "priority": priority,
        "score_priority_rank": int(score_priority_rank or 0),
        "score_priority_label": score_priority_label,
        "score_priority_weight": _scan_issue_score_priority_weight(
            group_id="searchability",
            check_id=check_id,
            advanced_degree_required=score_priority_label == "Education level",
        ),
        "score_priority_source": "jobscan_match_report_order",
        "confidence": "deterministic",
        "source": "resume_evidence",
        "section": "",
        "source_entry_id": "",
        "source_bullet_id": "",
        "source_bullet_ids": [],
        "original_text": current_text,
        "current_text": current_text,
        "suggested_text": suggested_text,
        "reason": reason,
        "supported_jd_signals": list(supported_jd_signals or []),
        "likely_impacted_dimensions": sorted(_SCAN_SEARCHABILITY_DIMENSIONS),
        "materiality_validation_status": "",
        "patch_generation_method": "deterministic_searchability_check",
        "projected_overall_delta": None,
        "original_final_score": None,
        "projected_final_score": None,
        "adjacent_risk_signals": [],
        "unsupported_risk_signals": [],
        "can_accept": False,
        "can_accept_all": False,
        "anchor_strategy": "none",
        "can_focus_preview": False,
        "raw": {
            "check_id": check_id,
            "source": "deterministic_searchability_check",
        },
    }


def _scan_recruiter_tip_issue(
    *,
    check_id: str,
    bucket: str,
    title: str,
    reason: str,
    current_text: str = "",
    suggested_text: str = "",
    priority: str = "",
) -> Dict[str, Any]:
    return {
        "issue_id": f"scan_issue:recruiter_tips:{check_id}",
        "candidate_id": "",
        "source_lane": "recruiter_tip_check",
        "group_id": "recruiter_tips",
        "group_label": _scan_issue_group_label("recruiter_tips"),
        "bucket": bucket,
        "bucket_label": "Strong recruiter-facing evidence" if bucket == "matched" else "Recruiter review opportunity",
        "title": title,
        "status": "recruiter_tip_matched" if bucket == "matched" else "recruiter_tip_opportunity",
        "priority": priority,
        "confidence": "deterministic",
        "source": "resume_evidence",
        "section": "",
        "source_entry_id": "",
        "source_bullet_id": "",
        "source_bullet_ids": [],
        "original_text": current_text,
        "current_text": current_text,
        "suggested_text": suggested_text,
        "reason": reason,
        "supported_jd_signals": [],
        "likely_impacted_dimensions": ["recruiter_clarity"],
        "materiality_validation_status": "",
        "patch_generation_method": "deterministic_recruiter_tip_check",
        "projected_overall_delta": None,
        "original_final_score": None,
        "projected_final_score": None,
        "adjacent_risk_signals": [],
        "unsupported_risk_signals": [],
        "can_accept": False,
        "can_accept_all": False,
        "anchor_strategy": "none",
        "can_focus_preview": False,
        "raw": {
            "check_id": check_id,
            "source": "deterministic_recruiter_tip_check",
        },
    }


def _scan_formatting_issue(
    *,
    check_id: str,
    bucket: str,
    title: str,
    reason: str,
    current_text: str = "",
    suggested_text: str = "",
    priority: str = "",
) -> Dict[str, Any]:
    return {
        "issue_id": f"scan_issue:formatting:{check_id}",
        "candidate_id": "",
        "source_lane": "formatting_check",
        "group_id": "formatting",
        "group_label": _scan_issue_group_label("formatting"),
        "bucket": bucket,
        "bucket_label": "ATS-friendly formatting signal" if bucket == "matched" else "Formatting opportunity",
        "title": title,
        "status": "formatting_matched" if bucket == "matched" else "formatting_opportunity",
        "priority": priority,
        "confidence": "deterministic",
        "source": "resume_text",
        "section": "",
        "source_entry_id": "",
        "source_bullet_id": "",
        "source_bullet_ids": [],
        "original_text": current_text,
        "current_text": current_text,
        "suggested_text": suggested_text,
        "reason": reason,
        "supported_jd_signals": [],
        "likely_impacted_dimensions": ["ats_parseability"],
        "materiality_validation_status": "",
        "patch_generation_method": "deterministic_formatting_check",
        "projected_overall_delta": None,
        "original_final_score": None,
        "projected_final_score": None,
        "adjacent_risk_signals": [],
        "unsupported_risk_signals": [],
        "can_accept": False,
        "can_accept_all": False,
        "anchor_strategy": "none",
        "can_focus_preview": False,
        "raw": {
            "check_id": check_id,
            "source": "deterministic_formatting_check",
        },
    }


def _scan_searchability_target_title(tailoring_summary: Dict[str, Any] | None = None) -> str:
    summary = dict(tailoring_summary or {})
    for key in (
        "target_job_title",
        "job_title",
        "role_title",
        "title",
        "position_title",
    ):
        value = _clean_text(summary.get(key))
        if value:
            return value

    job = summary.get("job")
    if isinstance(job, dict):
        return _clean_text(
            job.get("title")
            or job.get("job_title")
            or job.get("role_title")
        )

    return ""


def _scan_searchability_title_tokens(title: str) -> List[str]:
    stop_words = {
        "and",
        "the",
        "for",
        "with",
        "senior",
        "junior",
        "lead",
        "principal",
        "staff",
        "ii",
        "iii",
        "iv",
    }
    tokens = [
        token
        for token in re.findall(r"\b[a-z][a-z0-9+#.]{2,}\b", title.lower())
        if token not in stop_words
    ]
    return list(dict.fromkeys(tokens))


def _build_searchability_scan_issues(
    resume_evidence: Any = None,
    *,
    tailoring_summary: Dict[str, Any] | None = None,
) -> List[Dict[str, Any]]:
    if resume_evidence is None:
        return []

    raw_text = _scan_resume_document_text(resume_evidence)
    text_norm = raw_text.lower()
    bullets = _scan_resume_bullet_texts(resume_evidence)
    visible_terms = _scan_resume_visible_search_terms(resume_evidence)
    education_entries = list(getattr(resume_evidence, "education_entries", []) or [])
    target_title = _scan_searchability_target_title(tailoring_summary)
    advanced_degree_required = _scan_summary_has_advanced_degree_requirement(tailoring_summary)

    issues: List[Dict[str, Any]] = []

    contact_hits = []
    if re.search(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", raw_text, flags=re.IGNORECASE):
        contact_hits.append("email")
    if re.search(r"(?:\+?1[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?)\d{3}[\s.-]?\d{4}\b", raw_text):
        contact_hits.append("phone")
    if re.search(r"\b(?:linkedin\.com|github\.com|portfolio|www\.)\b", text_norm):
        contact_hits.append("profile link")

    if len(contact_hits) >= 2:
        issues.append(
            _scan_searchability_issue(
                check_id="contact_info_searchable",
                bucket="matched",
                title="Contact info searchable",
                reason=f"Detected parseable contact signals: {', '.join(contact_hits)}.",
                current_text=", ".join(contact_hits),
                priority="high",
            )
        )
    else:
        issues.append(
            _scan_searchability_issue(
                check_id="contact_info_searchable",
                bucket="missing",
                title="Make contact info searchable",
                reason="The scan could not detect at least two clear contact signals such as email, phone, or profile link.",
                current_text=", ".join(contact_hits),
                suggested_text="Use plain text for email, phone, LinkedIn, GitHub, or portfolio links in the resume header.",
                priority="high",
            )
        )

    education_terms = [
        "education",
        "university",
        "college",
        "bachelor",
        "master",
        "phd",
        "doctorate",
        "degree",
        "gpa",
    ]
    education_hits = [term for term in education_terms if re.search(rf"\b{re.escape(term)}\b", text_norm)]
    if education_entries or education_hits:
        issues.append(
            _scan_searchability_issue(
                check_id="education_info_searchable",
                bucket="matched",
                title="Education info searchable",
                reason="Detected parseable education text or structured education evidence.",
                current_text=", ".join(education_hits[:6]),
                priority="medium",
                score_priority_rank=2 if advanced_degree_required else 0,
                score_priority_label="Education level" if advanced_degree_required else "Non-scoring guidance",
            )
        )
    else:
        issues.append(
            _scan_searchability_issue(
                check_id="education_info_searchable",
                bucket="missing",
                title="Make education info searchable",
                reason="The scan could not detect a clear education section, degree, school, or related education signal.",
                suggested_text="Add a plain-text Education section with school, degree, and graduation details where applicable.",
                priority="medium",
                score_priority_rank=2 if advanced_degree_required else 0,
                score_priority_label="Education level" if advanced_degree_required else "Non-scoring guidance",
            )
        )

    date_pattern = re.compile(
        r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\.?\s+\d{4}\b"
        r"|\b(?:19|20)\d{2}\s*(?:-|–|to)\s*(?:present|current|(?:19|20)\d{2})\b"
        r"|\b(?:19|20)\d{2}\b",
        flags=re.IGNORECASE,
    )
    date_hits = date_pattern.findall(raw_text)
    if len(date_hits) >= 2:
        issues.append(
            _scan_searchability_issue(
                check_id="dates_searchable",
                bucket="matched",
                title="Dates searchable",
                reason=f"Detected {len(date_hits)} parseable date signal(s).",
                current_text=", ".join(_scan_issue_unique_display_labels(date_hits[:6])),
                priority="medium",
            )
        )
    else:
        issues.append(
            _scan_searchability_issue(
                check_id="dates_searchable",
                bucket="missing",
                title="Make dates searchable",
                reason="The scan found fewer than two parseable date signals for experience or education timelines.",
                current_text=", ".join(_scan_issue_unique_display_labels(date_hits[:6])),
                suggested_text="Use plain-text month/year or year ranges for roles and education dates.",
                priority="medium",
            )
        )

    section_patterns = {
        "experience": r"(?im)^\s*(professional\s+experience|experience|work\s+experience)\s*:?\s*$",
        "education": r"(?im)^\s*education\s*:?\s*$",
        "skills": r"(?im)^\s*(technical\s+skills|skills|core\s+skills)\s*:?\s*$",
        "projects": r"(?im)^\s*(projects|selected\s+projects)\s*:?\s*$",
    }
    parseable_sections = [
        label
        for label, pattern in section_patterns.items()
        if re.search(pattern, raw_text)
    ]
    if len(parseable_sections) >= 2:
        issues.append(
            _scan_searchability_issue(
                check_id="section_headings_parseable",
                bucket="matched",
                title="Section headings parseable",
                reason=f"Detected standard section headings: {', '.join(parseable_sections)}.",
                current_text=", ".join(parseable_sections),
                priority="high",
            )
        )
    else:
        issues.append(
            _scan_searchability_issue(
                check_id="section_headings_parseable",
                bucket="missing",
                title="Use standard section headings",
                reason="The scan detected fewer than two standard resume section headings.",
                current_text=", ".join(parseable_sections),
                suggested_text="Use plain section headings such as Experience, Skills, Projects, and Education.",
                priority="high",
            )
        )

    title_tokens = _scan_searchability_title_tokens(target_title)
    matched_title_tokens = [token for token in title_tokens if re.search(rf"\b{re.escape(token)}\b", text_norm)]
    if title_tokens:
        required_title_hits = max(1, min(len(title_tokens), 2))
        if len(matched_title_tokens) >= required_title_hits:
            issues.append(
                _scan_searchability_issue(
                    check_id="job_title_alignment",
                    bucket="matched",
                    title="Job title alignment visible",
                    reason=f"The resume includes target-role terms from {target_title}.",
                    current_text=", ".join(matched_title_tokens),
                    supported_jd_signals=[target_title],
                    priority="high",
                    score_priority_rank=3,
                    score_priority_label="Job title",
                )
            )
        else:
            issues.append(
                _scan_searchability_issue(
                    check_id="job_title_alignment",
                    bucket="missing",
                    title="Add target job title alignment",
                    reason=f"The scan found limited title alignment for the target role: {target_title}.",
                    current_text=", ".join(matched_title_tokens),
                    suggested_text="Mention the target role or truthful equivalent role language in the summary, headline, or relevant experience.",
                    supported_jd_signals=[target_title],
                    priority="high",
                    score_priority_rank=3,
                    score_priority_label="Job title",
                )
            )

    skills = list(getattr(resume_evidence, "skills", []) or [])
    has_skills_section = bool(re.search(r"(?im)^\s*(technical\s+skills|skills|core\s+skills)\s*:?\s*$", raw_text))
    has_parseable_skills = bool(skills)

    if has_skills_section or has_parseable_skills:
        issues.append(
            _scan_searchability_issue(
                check_id="skills_section_parseable",
                bucket="matched",
                title="Skills section parseable",
                reason="The resume exposes a skills signal that can be parsed deterministically.",
                current_text=", ".join(_scan_issue_unique_display_labels(skills[:12])),
                supported_jd_signals=_scan_issue_unique_display_labels(skills[:6]),
                priority="medium",
            )
        )
    else:
        issues.append(
            _scan_searchability_issue(
                check_id="skills_section_parseable",
                bucket="missing",
                title="Add a parseable skills section",
                reason="The scan could not detect a clear skills section or structured skills list.",
                suggested_text="Add a concise Technical Skills section with the strongest JD-relevant tools, methods, and platforms.",
                priority="high",
            )
        )

    if len(visible_terms) >= int(_SCAN_SEARCHABILITY_MIN_VISIBLE_TECH_TERMS):
        issues.append(
            _scan_searchability_issue(
                check_id="visible_technical_terms",
                bucket="matched",
                title="Technical terms visible",
                reason=f"{len(visible_terms)} technical/search terms are visible in the parsed resume evidence.",
                current_text=", ".join(visible_terms[:12]),
                supported_jd_signals=visible_terms[:6],
                priority="medium",
            )
        )
    else:
        issues.append(
            _scan_searchability_issue(
                check_id="visible_technical_terms",
                bucket="missing",
                title="Add more searchable technical terms",
                reason=(
                    f"Only {len(visible_terms)} technical/search term(s) were visible in the parsed resume evidence."
                ),
                current_text=", ".join(visible_terms[:12]),
                suggested_text="Surface the strongest supported JD keywords in Skills and experience bullets.",
                supported_jd_signals=visible_terms[:6],
                priority="high",
            )
        )

    quantified_bullets = list(getattr(resume_evidence, "quantified_bullets", []) or [])
    if quantified_bullets:
        issues.append(
            _scan_searchability_issue(
                check_id="quantified_impact_visible",
                bucket="matched",
                title="Quantified impact visible",
                reason=f"{len(quantified_bullets)} quantified bullet(s) were detected.",
                current_text=_clean_text(quantified_bullets[0])[:220],
                priority="medium",
            )
        )
    else:
        issues.append(
            _scan_searchability_issue(
                check_id="quantified_impact_visible",
                bucket="missing",
                title="Add quantified impact",
                reason="No quantified bullets were detected in the parsed resume evidence.",
                suggested_text="Add measurable impact where already supported by the original experience.",
                priority="medium",
            )
        )

    long_bullets = [
        bullet for bullet in bullets
        if len(re.findall(r"\w+", bullet)) > int(_SCAN_SEARCHABILITY_BULLET_WORD_LIMIT)
    ]

    if long_bullets:
        issues.append(
            _scan_searchability_issue(
                check_id="bullet_length_risk",
                bucket="missing",
                title="Shorten long bullets",
                reason=(
                    f"{len(long_bullets)} bullet(s) exceed "
                    f"{int(_SCAN_SEARCHABILITY_BULLET_WORD_LIMIT)} words, which can reduce scan readability."
                ),
                current_text=_clean_text(long_bullets[0])[:260],
                suggested_text="Split or tighten long bullets while preserving the supported evidence.",
                priority="medium",
            )
        )
    else:
        issues.append(
            _scan_searchability_issue(
                check_id="bullet_length_risk",
                bucket="matched",
                title="Bullet length looks scannable",
                reason=(
                    f"No parsed experience/project bullet exceeded "
                    f"{int(_SCAN_SEARCHABILITY_BULLET_WORD_LIMIT)} words."
                ),
                priority="low",
            )
        )

    return issues


def _build_recruiter_tip_scan_issues(resume_evidence: Any = None) -> List[Dict[str, Any]]:
    if resume_evidence is None:
        return []

    bullets = _scan_resume_bullet_texts(resume_evidence)
    raw_text = _scan_resume_document_text(resume_evidence)
    lower_text = raw_text.lower()
    action_verb_pattern = re.compile(
        r"\b(?:built|automated|developed|deployed|implemented|designed|created|improved|reduced|increased|partnered|launched|evaluated|optimized|analyzed)\b",
        flags=re.IGNORECASE,
    )

    action_bullets = [bullet for bullet in bullets if action_verb_pattern.search(bullet)]
    if action_bullets:
        yield_issue = _scan_recruiter_tip_issue(
            check_id="action_led_bullets",
            bucket="matched",
            title="Action-led bullets",
            reason=f"{len(action_bullets)} parsed bullet(s) start from recruiter-friendly accomplishment language.",
            current_text=action_bullets[0][:240],
            priority="medium",
        )
    else:
        yield_issue = _scan_recruiter_tip_issue(
            check_id="action_led_bullets",
            bucket="missing",
            title="Lead bullets with outcomes",
            reason="The scan did not find strong action-led accomplishment phrasing in parsed bullets.",
            suggested_text="Start bullets with owned actions and outcomes that are already supported by your experience.",
            priority="medium",
        )

    issues = [yield_issue]

    quantified_bullets = list(getattr(resume_evidence, "quantified_bullets", []) or [])
    if quantified_bullets:
        issues.append(
            _scan_recruiter_tip_issue(
                check_id="business_impact",
                bucket="matched",
                title="Business impact is measurable",
                reason=f"{len(quantified_bullets)} quantified impact bullet(s) help recruiters understand scope.",
                current_text=_clean_text(quantified_bullets[0])[:240],
                priority="medium",
            )
        )
    else:
        issues.append(
            _scan_recruiter_tip_issue(
                check_id="business_impact",
                bucket="missing",
                title="Add measurable business impact",
                reason="Recruiters will have an easier time evaluating fit when outcomes include scale, volume, time, cost, quality, or adoption metrics.",
                suggested_text="Add a truthful metric to one high-signal bullet where the evidence supports it.",
                priority="high",
            )
        )

    repeated_terms = [
        term
        for term, count in Counter(
            word
            for word in re.findall(r"\b[a-z][a-z0-9+#.]{3,}\b", lower_text)
            if word not in {"with", "from", "that", "this", "using", "into", "over"}
        ).items()
        if count >= 8
    ]
    if repeated_terms:
        issues.append(
            _scan_recruiter_tip_issue(
                check_id="repetition_risk",
                bucket="missing",
                title="Reduce repeated phrasing",
                reason=f"Repeated terms may make the resume feel less specific: {', '.join(sorted(repeated_terms)[:4])}.",
                suggested_text="Vary repeated phrasing only where it preserves the same truthful evidence.",
                priority="low",
            )
        )
    else:
        issues.append(
            _scan_recruiter_tip_issue(
                check_id="repetition_risk",
                bucket="matched",
                title="Repetition looks controlled",
                reason="No heavily repeated recruiter-facing phrasing was detected.",
                priority="low",
            )
        )

    return issues


def _build_formatting_scan_issues(resume_evidence: Any = None) -> List[Dict[str, Any]]:
    if resume_evidence is None:
        return []

    raw_text = _scan_resume_document_text(resume_evidence)
    lines = [_clean_text(line) for line in raw_text.splitlines()]
    non_empty_lines = [line for line in lines if line]
    issues: List[Dict[str, Any]] = []

    word_count = len(re.findall(r"\w+", raw_text))
    if word_count >= 80:
        issues.append(
            _scan_formatting_issue(
                check_id="text_extractable",
                bucket="matched",
                title="Resume text is extractable",
                reason=f"Detected {word_count} parseable word(s), which suggests the resume is not image-only.",
                current_text=f"{word_count} words parsed",
                priority="high",
            )
        )
    else:
        issues.append(
            _scan_formatting_issue(
                check_id="text_extractable",
                bucket="missing",
                title="Make resume text extractable",
                reason="Very little parseable text was detected, which can happen with image-only or heavily graphical resumes.",
                current_text=f"{word_count} words parsed",
                suggested_text="Use selectable text in an ATS-friendly PDF or DOCX instead of scanned images.",
                priority="high",
            )
        )

    def _looks_like_contact_separator_line(line: str) -> bool:
        lowered = line.lower()
        has_contact_signal = (
            "@" in line
            or "linkedin" in lowered
            or "github" in lowered
            or re.search(r"\+?\d[\d\s().-]{6,}\d", line)
        )
        return bool(has_contact_signal and line.count("|") >= 1)

    table_like_lines = [
        line for line in non_empty_lines
        if not _looks_like_contact_separator_line(line)
        and (line.count("|") >= 2 or line.count("\t") >= 2)
    ]
    wide_gap_lines = [
        line for line in non_empty_lines
        if re.search(r"\S\s{6,}\S", line)
    ]
    if table_like_lines or len(wide_gap_lines) >= 4:
        issues.append(
            _scan_formatting_issue(
                check_id="table_column_risk",
                bucket="missing",
                title="Reduce table or column formatting risk",
                reason="The parsed text shows table-like separators or wide spacing that can indicate columns/tables.",
                current_text=(table_like_lines or wide_gap_lines)[0][:220],
                suggested_text="Prefer a single-column layout with plain section headings and simple bullet lists.",
                priority="high",
            )
        )
    else:
        issues.append(
            _scan_formatting_issue(
                check_id="table_column_risk",
                bucket="matched",
                title="No obvious table or column risk",
                reason="The parsed text does not show repeated table separators or wide column spacing.",
                priority="high",
            )
        )

    bullet_lines = [
        line for line in non_empty_lines
        if re.match(r"^\s*(?:[-*•▪◦‣]|\d+[.)])\s+", line)
    ]
    unusual_bullets = [
        line for line in non_empty_lines
        if re.match(r"^\s*[➤✓✔★◆●]\s+", line)
    ]
    if unusual_bullets:
        issues.append(
            _scan_formatting_issue(
                check_id="bullet_formatting",
                bucket="missing",
                title="Use simple bullet symbols",
                reason="Decorative bullet symbols can be less reliable for ATS parsing.",
                current_text=unusual_bullets[0][:220],
                suggested_text="Use simple bullets such as -, *, or standard round bullets.",
                priority="medium",
            )
        )
    elif bullet_lines:
        issues.append(
            _scan_formatting_issue(
                check_id="bullet_formatting",
                bucket="matched",
                title="Bullet formatting is simple",
                reason=f"Detected {len(bullet_lines)} line(s) using common bullet/list markers.",
                current_text=bullet_lines[0][:220],
                priority="medium",
            )
        )
    else:
        issues.append(
            _scan_formatting_issue(
                check_id="bullet_formatting",
                bucket="missing",
                title="Use parseable bullet lists",
                reason="The scan did not detect common bullet/list markers in parsed resume text.",
                suggested_text="Use simple bullet lists for accomplishments and responsibilities.",
                priority="medium",
            )
        )

    replacement_char_count = raw_text.count("\ufffd")
    odd_symbol_count = len(re.findall(r"[^\x00-\x7F]", raw_text))
    if replacement_char_count or odd_symbol_count > 40:
        issues.append(
            _scan_formatting_issue(
                check_id="special_character_risk",
                bucket="missing",
                title="Reduce special character parsing risk",
                reason="The parsed text contains many non-standard characters or replacement glyphs.",
                current_text=f"{odd_symbol_count} non-ASCII character(s), {replacement_char_count} replacement glyph(s)",
                suggested_text="Use standard punctuation and avoid decorative symbols in core resume text.",
                priority="medium",
            )
        )
    else:
        issues.append(
            _scan_formatting_issue(
                check_id="special_character_risk",
                bucket="matched",
                title="Special character risk is low",
                reason="The parsed text does not show heavy use of decorative or corrupted characters.",
                priority="low",
            )
        )

    heading_order = []
    heading_patterns = [
        ("summary", r"(?im)^\s*(summary|profile|professional\s+summary)\s*:?\s*$"),
        ("skills", r"(?im)^\s*(technical\s+skills|skills|core\s+skills)\s*:?\s*$"),
        ("experience", r"(?im)^\s*(professional\s+experience|experience|work\s+experience)\s*:?\s*$"),
        ("education", r"(?im)^\s*education\s*:?\s*$"),
    ]
    for key, pattern in heading_patterns:
        match = re.search(pattern, raw_text)
        if match:
            heading_order.append((match.start(), key))
    heading_order.sort()
    detected_headings = [key for _, key in heading_order]
    if len(detected_headings) >= 3:
        issues.append(
            _scan_formatting_issue(
                check_id="standard_resume_structure",
                bucket="matched",
                title="Standard resume structure detected",
                reason=f"Detected standard headings: {', '.join(detected_headings)}.",
                current_text=", ".join(detected_headings),
                priority="medium",
            )
        )
    else:
        issues.append(
            _scan_formatting_issue(
                check_id="standard_resume_structure",
                bucket="missing",
                title="Use standard resume sections",
                reason="The scan detected fewer than three standard resume headings.",
                current_text=", ".join(detected_headings),
                suggested_text="Use common headings such as Summary, Skills, Experience, Projects, and Education.",
                priority="medium",
            )
        )

    return issues


def _build_scan_match_report_summary(
    *,
    issues: List[Dict[str, Any]],
    groups: List[Dict[str, Any]],
) -> Dict[str, Any]:
    visible_issues = [
        issue for issue in issues
        if issue.get("is_visible_in_review", True)
    ]

    def priority_row(
        *,
        key: str,
        label: str,
        rank: int,
        group_ids: List[str],
        skill_types: List[str] | None = None,
        score_priority_labels: List[str] | None = None,
    ) -> Dict[str, Any]:
        allowed_groups = set(group_ids)
        allowed_skill_types = set(skill_types or [])
        allowed_priority_labels = set(score_priority_labels or [])

        rows = []
        for issue in visible_issues:
            group_id = _clean_text(issue.get("group_id"))
            if group_id not in allowed_groups:
                continue
            if allowed_skill_types and _clean_text(issue.get("skill_type")) not in allowed_skill_types:
                continue
            if allowed_priority_labels and _clean_text(issue.get("score_priority_label")) not in allowed_priority_labels:
                continue
            rows.append(issue)

        matched = sum(1 for issue in rows if issue.get("bucket") == "matched")
        missing = sum(1 for issue in rows if issue.get("bucket") == "missing")
        ai = sum(1 for issue in rows if issue.get("bucket") == "ai")
        total = len(rows)

        return {
            "key": key,
            "label": label,
            "rank": rank,
            "total": total,
            "matched": matched,
            "missing": missing,
            "ai": ai,
            "complete": total > 0 and missing == 0,
        }

    priority_rows = [
        priority_row(
            key="hard_skills",
            label="Hard skills",
            rank=1,
            group_ids=["skills"],
            skill_types=["hard_skill"],
        ),
        priority_row(
            key="education_level",
            label="Education level",
            rank=2,
            group_ids=["searchability"],
            score_priority_labels=["Education level"],
        ),
        priority_row(
            key="job_title",
            label="Job title",
            rank=3,
            group_ids=["searchability"],
            score_priority_labels=["Job title"],
        ),
        priority_row(
            key="soft_skills",
            label="Soft skills",
            rank=4,
            group_ids=["skills"],
            skill_types=["soft_skill"],
        ),
        priority_row(
            key="other_keywords",
            label="Other keywords",
            rank=5,
            group_ids=["skills"],
            skill_types=["other_keyword"],
        ),
    ]

    group_counts = {
        _clean_text(group.get("group_id")): dict(group.get("counts", {}) or {})
        for group in groups
    }

    return {
        "version": "scan_match_report_v1",
        "priority_order": priority_rows,
        "group_counts": group_counts,
        "formatting": group_counts.get("formatting", {}),
        "searchability": group_counts.get("searchability", {}),
    }


def _scan_issue_from_replacement_row(
    row: Dict[str, Any],
    *,
    lane: str,
    group_id: str,
    group_label: str,
    bucket: str,
    bucket_label: str,
    can_accept: bool,
    can_accept_all: bool,
    index: int,
) -> Dict[str, Any]:
    candidate_id = _scan_issue_candidate_id(row)
    issue_id = candidate_id or f"{lane}:{index}"

    original_text = (
        _clean_text(row.get("original_text"))
        or _clean_text(row.get("current_evidence"))
        or _clean_text(row.get("parent_bullet"))
    )
    suggested_text = (
        _clean_text(row.get("final_replacement_text"))
        or _clean_text(row.get("rewrite_direction"))
        or _clean_text(row.get("rewrite_instruction"))
    )
    projected_overall_delta = row.get("projected_overall_delta", None)
    delta_points = score_delta_to_points(projected_overall_delta)
    score_gate = _clean_text(row.get("score_gate"))
    if not score_gate:
        if delta_points is not None and delta_points > 0 and can_accept:
            score_gate = "direct_replacement"
        elif delta_points is not None and delta_points < 0:
            score_gate = "rejected_by_score_gate"
        else:
            score_gate = "score_neutral_guidance"

    if score_gate == "direct_replacement" and _scan_issue_is_deterministic_only_replacement(row):
        score_gate = "score_neutral_guidance"

    can_direct_accept = score_gate == "direct_replacement"

    resolved_group_id = _scan_issue_group_id_for_row(row, lane=lane)
    resolved_group_label = _scan_issue_group_label(resolved_group_id)

    return {
        "issue_id": f"scan_issue:{lane}:{issue_id}",
        "candidate_id": candidate_id,
        "source_lane": lane,
        "group_id": resolved_group_id,
        "group_label": resolved_group_label,
        "bucket": bucket,
        "bucket_label": bucket_label,
        "title": _scan_issue_title(row, fallback_label=bucket_label),
        "status": _clean_text(row.get("replacement_status")) or lane,
        "priority": _clean_text(row.get("apply_priority")) or _clean_text(row.get("priority")),
        "confidence": _clean_text(row.get("confidence")),
        "source": _clean_text(row.get("source")),
        "section": _clean_text(row.get("section")),
        "source_entry_id": _clean_text(row.get("source_entry_id")),
        "source_bullet_id": _clean_text(row.get("source_bullet_id")),
        "source_bullet_ids": list(row.get("source_bullet_ids", []) or []),
        "original_text": original_text,
        "current_text": original_text,
        "suggested_text": suggested_text,
        "reason": (
            _clean_text(row.get("why_selected"))
            or _clean_text(row.get("why_this_improves_match"))
            or _clean_text(row.get("rewrite_direction"))
            or _clean_text(row.get("rewrite_instruction"))
            or _clean_text(row.get("placement_guidance"))
        ),
        "supported_jd_signals": _scan_issue_supported_signals(row),
        "likely_impacted_dimensions": list(row.get("likely_impacted_dimensions", []) or []),
        "materiality_validation_status": _clean_text(row.get("materiality_validation_status")),
        "patch_generation_method": _clean_text(row.get("patch_generation_method")),
        "projected_overall_delta": projected_overall_delta,
        "projected_score_delta_points": delta_points,
        "projected_dimension_deltas": dict(row.get("projected_dimension_deltas", {}) or {}),
        "precheck_projected_overall_delta": row.get("precheck_projected_overall_delta", None),
        "precheck_projected_dimension_deltas": dict(row.get("precheck_projected_dimension_deltas", {}) or {}),
        "original_final_score": row.get("original_final_score", None),
        "projected_final_score": row.get("projected_final_score", None),
        "scan_issue_type": score_gate,
        "is_visible_in_review": score_gate != "rejected_by_score_gate",
        "llm_judge_score_intent": _clean_text(row.get("llm_judge_score_intent")),
        "llm_judge_expected_dimensions": list(row.get("llm_judge_expected_dimensions", []) or []),
        "llm_judge_risk_flags": list(row.get("llm_judge_risk_flags", []) or []),
        "llm_judge_reason": _clean_text(row.get("llm_judge_reason")),
        "llm_substantive_judge_score_intent": _clean_text(row.get("llm_substantive_judge_score_intent")),
        "llm_substantive_judge_expected_dimensions": list(
            row.get("llm_substantive_judge_expected_dimensions", []) or []
        ),
        "llm_substantive_judge_risk_flags": list(row.get("llm_substantive_judge_risk_flags", []) or []),
        "llm_substantive_judge_reason": _clean_text(row.get("llm_substantive_judge_reason")),
        "adjacent_risk_signals": list(row.get("adjacent_risk_signals", []) or []),
        "unsupported_risk_signals": list(row.get("unsupported_risk_signals", []) or []),
        "can_accept": bool(can_accept and can_direct_accept),
        "can_accept_all": bool(can_accept_all and can_direct_accept),
        "anchor_strategy": "replacement_candidate",
        "can_focus_preview": bool(candidate_id),
        "raw": row,
    }


def _build_tailoring_scan_issue_contract(
    *,
    trusted_ready: List[Dict[str, Any]],
    trusted_optional: List[Dict[str, Any]],
    ai_optimize_optional: List[Dict[str, Any]],
    directional_guidance: List[Dict[str, Any]],
    resume_evidence: Any = None,
    tailoring_summary: Dict[str, Any] | None = None,
    jd_record: Dict[str, Any] | None = None,
) -> Dict[str, Any]:
    replacement_issues: List[Dict[str, Any]] = []

    lane_specs = [
        (
            "direct_apply_ready",
            trusted_ready,
            "matched",
            "Matched",
            True,
            True,
        ),
        (
            "direct_apply_optional",
            trusted_optional,
            "matched",
            "Matched",
            True,
            False,
        ),
        (
            "ai_optimize_optional",
            ai_optimize_optional,
            "ai",
            "AI Suggested",
            True,
            True,
        ),
        (
            "direction_only",
            directional_guidance,
            "missing",
            "Missing / optimization opportunity",
            False,
            False,
        ),
    ]

    for lane, rows, bucket, bucket_label, can_accept, can_accept_all in lane_specs:
        for index, row in enumerate(rows or [], start=1):
            if not isinstance(row, dict):
                continue

            replacement_issues.append(
                _scan_issue_from_replacement_row(
                    row,
                    lane=lane,
                    group_id="skills",
                    group_label="Skills",
                    bucket=bucket,
                    bucket_label=bucket_label,
                    can_accept=can_accept,
                    can_accept_all=can_accept_all,
                    index=index,
                )
            )

    deterministic_issues = (
        _build_searchability_scan_issues(
            resume_evidence,
            tailoring_summary=tailoring_summary,
        )
        + _build_formatting_scan_issues(resume_evidence)
        + _build_recruiter_tip_scan_issues(resume_evidence)
    )
    issues = (
        _scan_keyword_rows_from_replacement_issues(
            replacement_issues,
            resume_evidence=resume_evidence,
            tailoring_summary=tailoring_summary,
            jd_record=jd_record,
        )
        + _scan_keyword_rows_from_non_skill_replacement_issues(replacement_issues)
        + _scan_keyword_rows_from_generic_issues(deterministic_issues)
    )
    issues.extend(
        _build_predicted_skill_scan_rows(
            existing_issues=issues,
            resume_evidence=resume_evidence,
            tailoring_summary=tailoring_summary,
            jd_record=jd_record,
        )
    )
    issues.extend(
        _build_other_keyword_scan_rows(
            existing_issues=issues,
            resume_evidence=resume_evidence,
            jd_record=jd_record,
        )
    )

    hidden_replacement_issues = [
        issue for issue in replacement_issues
        if not issue.get("is_visible_in_review", True)
        and not any(
            issue.get("candidate_id") in list(keyword_issue.get("linked_candidate_ids", []) or [])
            for keyword_issue in issues
        )
    ]
    issues.extend(hidden_replacement_issues)

    visible_issues = [
        issue for issue in issues
        if issue.get("is_visible_in_review", True)
    ]

    counts = {
        "total": len(visible_issues),
        "hidden": len(issues) - len(visible_issues),
        "matched": sum(1 for issue in visible_issues if issue.get("bucket") == "matched"),
        "missing": sum(1 for issue in visible_issues if issue.get("bucket") == "missing"),
        "ai": sum(1 for issue in visible_issues if issue.get("bucket") == "ai"),
        "actionable": sum(1 for issue in visible_issues if issue.get("can_accept")),
        "accept_all_eligible": sum(1 for issue in visible_issues if issue.get("can_accept_all")),
    }

    def _group_counts(group_id: str) -> Dict[str, int]:
        group_issues = [
            issue for issue in issues
            if _clean_text(issue.get("group_id")) == group_id
            and issue.get("is_visible_in_review", True)
        ]

        return {
            "total": len(group_issues),
            "matched": sum(1 for issue in group_issues if issue.get("bucket") == "matched"),
            "missing": sum(1 for issue in group_issues if issue.get("bucket") == "missing"),
            "ai": sum(1 for issue in group_issues if issue.get("bucket") == "ai"),
            "predicted": sum(1 for issue in group_issues if issue.get("bucket") == "predicted"),
            "other_keyword": sum(1 for issue in group_issues if issue.get("bucket") == "other_keyword"),
        }

    def _skill_type_counts() -> Dict[str, int]:
        skill_issues = [
            issue for issue in issues
            if _clean_text(issue.get("group_id")) == "skills"
            and issue.get("is_visible_in_review", True)
        ]
        return {
            "hard_skill": sum(1 for issue in skill_issues if issue.get("skill_type") == "hard_skill"),
            "soft_skill": sum(1 for issue in skill_issues if issue.get("skill_type") == "soft_skill"),
            "other_keyword": sum(1 for issue in skill_issues if issue.get("skill_type") == "other_keyword"),
        }

    skills_counts = _group_counts("skills")
    searchability_counts = _group_counts("searchability")
    formatting_counts = _group_counts("formatting")
    recruiter_tips_counts = _group_counts("recruiter_tips")

    groups = [
        {
            "group_id": "skills",
            "label": _scan_issue_group_label("skills"),
            "description": "Skill and JD-signal alignment issues derived from tailoring replacement lanes.",
            "counts": skills_counts,
            "skill_type_counts": _skill_type_counts(),
            "buckets": [
                {
                    "bucket": "matched",
                    "label": "Matched skills",
                    "count": skills_counts["matched"],
                },
                {
                    "bucket": "missing",
                    "label": "Missing / optimization opportunities",
                    "count": skills_counts["missing"],
                },
                {
                    "bucket": "ai",
                    "label": "AI suggested",
                    "count": skills_counts["ai"],
                },
                {
                    "bucket": "predicted",
                    "label": "Predicted skills",
                    "count": skills_counts["predicted"],
                    "summary": "Role-adjacent skills inferred from the target title, not explicit JD requirements.",
                },
                {
                    "bucket": "other_keyword",
                    "label": "Other keywords",
                    "count": skills_counts["other_keyword"],
                    "summary": "Lower-impact domain or industry terms detected in the JD.",
                },
            ],
        },
        {
            "group_id": "searchability",
            "label": _scan_issue_group_label("searchability"),
            "description": "Deterministic ATS searchability and parse checks.",
            "counts": searchability_counts,
            "buckets": [
                {
                    "bucket": "matched",
                    "label": "Matched searchability signals",
                    "count": searchability_counts["matched"],
                },
                {
                    "bucket": "missing",
                    "label": "Searchability opportunities",
                    "count": searchability_counts["missing"],
                },
                {
                    "bucket": "ai",
                    "label": "AI suggested",
                    "count": searchability_counts["ai"],
                },
            ],
        },
        {
            "group_id": "formatting",
            "label": _scan_issue_group_label("formatting"),
            "description": "Deterministic ATS formatting checks for text extraction, tables, columns, bullets, and symbols.",
            "counts": formatting_counts,
            "buckets": [
                {
                    "bucket": "matched",
                    "label": "ATS-friendly formatting signals",
                    "count": formatting_counts["matched"],
                },
                {
                    "bucket": "missing",
                    "label": "Formatting opportunities",
                    "count": formatting_counts["missing"],
                },
                {
                    "bucket": "ai",
                    "label": "AI suggested",
                    "count": formatting_counts["ai"],
                },
            ],
        },
        {
            "group_id": "recruiter_tips",
            "label": _scan_issue_group_label("recruiter_tips"),
            "description": "Recruiter-facing clarity, positioning, and claim-safety notes derived from deterministic tailoring guidance.",
            "counts": recruiter_tips_counts,
            "buckets": [
                {
                    "bucket": "matched",
                    "label": "Strong recruiter-facing evidence",
                    "count": recruiter_tips_counts["matched"],
                },
                {
                    "bucket": "missing",
                    "label": "Recruiter review opportunities",
                    "count": recruiter_tips_counts["missing"],
                },
                {
                    "bucket": "ai",
                    "label": "AI suggested",
                    "count": recruiter_tips_counts["ai"],
                },
            ],
        },
    ]

    return {
        "version": "scan_issue_contract_v2",
        "source": "tailoring_keyword_scan",
        "groups": groups,
        "match_report": _build_scan_match_report_summary(issues=visible_issues, groups=groups),
        "issues": issues,
        "replacement_issues": replacement_issues,
        "counts": counts,
        "acceptance_policy": {
            "direct_apply_ready": {
                "can_accept": True,
                "can_accept_all": True,
                "trusted": True,
            },
            "direct_apply_optional": {
                "can_accept": True,
                "can_accept_all": False,
                "trusted": True,
            },
            "ai_optimize_optional": {
                "can_accept": True,
                "can_accept_all": True,
                "trusted": False,
            },
            "direction_only": {
                "can_accept": False,
                "can_accept_all": False,
                "trusted": False,
            },
        },
    }


def _coerce_scan_score_value(*values: Any) -> int | None:
    for value in values:
        score = score_to_points(value)
        if score is not None:
            return score

    return None


def _build_tailoring_scan_score_snapshot(
    *,
    selection: Dict[str, Any],
    preview_response: Dict[str, Any],
    scan_issue_contract: Dict[str, Any],
) -> Dict[str, Any]:
    counts = dict(scan_issue_contract.get("counts", {}) or {})
    total = int(counts.get("total", 0) or 0)
    matched = int(counts.get("matched", 0) or 0)
    actionable = int(counts.get("actionable", 0) or 0)

    score = _coerce_scan_score_value(
        preview_response.get("projected_score"),
        selection.get("selected_score"),
        preview_response.get("original_score"),
    )

    if score is None:
        score = round((matched / total) * 100) if total else 0

    return {
        "score": score,
        "source": "projected_score" if preview_response.get("projected_score") is not None else "scan_issue_contract",
        "matched_count": matched,
        "missing_count": int(counts.get("missing", 0) or 0),
        "ai_count": int(counts.get("ai", 0) or 0),
        "total_count": total,
        "actionable_count": actionable,
        "label": "Optimization score",
    }


def _build_tailoring_scan_session_snapshot(
    *,
    artifact_path: Path,
    draft: Dict[str, Any],
    selected_resume: str,
    scan_issue_contract: Dict[str, Any],
    preview_response: Dict[str, Any],
) -> Dict[str, Any]:
    job_doc_id = _clean_text(draft.get("job_doc_id"))
    artifact_signature = _clean_text(draft.get("artifact_signature"))
    session_source = "|".join([
        artifact_signature,
        str(artifact_path),
        selected_resume,
        job_doc_id,
    ])
    session_id = hashlib.sha256(session_source.encode("utf-8")).hexdigest()[:24]
    counts = dict(scan_issue_contract.get("counts", {}) or {})
    telemetry = dict(draft.get("rewrite_review_telemetry", {}) or {})

    return {
        "session_id": session_id,
        "session_version": 1,
        "artifact_signature": artifact_signature,
        "tailoring_json_path": str(artifact_path),
        "selected_resume": selected_resume,
        "job_doc_id": job_doc_id,
        "draft_status": _clean_text(draft.get("draft_status")),
        "saved_at": _clean_text(draft.get("saved_at")),
        "issue_counts": counts,
        "decision_counts": {
            "accepted": int(telemetry.get("accepted_count", 0) or 0),
            "accepted_as_is": int(telemetry.get("accepted_as_is_count", 0) or 0),
            "edited_after_accept": int(telemetry.get("edited_after_accept_count", 0) or 0),
            "rejected": int(telemetry.get("rejected_count", 0) or 0),
            "pending": int(telemetry.get("pending_count", 0) or 0),
            "manual_edits": int(telemetry.get("manual_edit_count", 0) or 0),
        },
        "score_preview_status": _clean_text(preview_response.get("preview_status")),
        "score_preview_note": _clean_text(preview_response.get("preview_note")),
    }


def _build_workspace_score_preview_contract(
    *,
    preview_status: str,
    preview_note: str,
    original_score: Any = None,
    projected_score: Any = None,
    projected_delta: Any = None,
    selected_candidate_ids: List[str] | None = None,
    manual_bullet_edits: Dict[str, str] | None = None,
    patch_specs: List[Dict[str, Any]] | None = None,
    unresolved_manual_edit_keys: List[str] | None = None,
    dimension_deltas: Dict[str, Any] | None = None,
) -> Dict[str, Any]:
    selected_ids = list(selected_candidate_ids or [])
    manual_map = dict(manual_bullet_edits or {})
    patches = list(patch_specs or [])
    unresolved_keys = list(unresolved_manual_edit_keys or [])
    delta_points = score_delta_to_points(projected_delta)

    dimension_delta_points: Dict[str, int] = {}
    for key, value in dict(dimension_deltas or {}).items():
        points = score_delta_to_points(value)
        if points is not None:
            dimension_delta_points[str(key)] = points

    return {
        "version": "workspace_score_preview_v1",
        "status": _clean_text(preview_status),
        "note": _clean_text(preview_note),
        "original_score": original_score,
        "projected_score": projected_score,
        "projected_delta": projected_delta,
        "original_score_points": score_to_points(original_score),
        "projected_score_points": score_to_points(projected_score),
        "delta_points": delta_points,
        "dimension_delta_points": dimension_delta_points,
        "selected_candidate_ids": selected_ids,
        "selected_patch_count": len(selected_ids),
        "manual_edit_count": len(manual_map),
        "patch_count": len(patches),
        "changed_bullet_keys": [
            _clean_text(patch.get("bullet_key"))
            for patch in patches
            if _clean_text(patch.get("bullet_key"))
        ],
        "unresolved_manual_edit_keys": unresolved_keys,
        "has_previewable_changes": bool(patches),
        "can_update_score": projected_score is not None or original_score is not None,
        "requires_full_preview_reload": False,
    }


def _build_workspace_draft_fragments_contract(
    *,
    preview_status: str,
    preview_note: str,
    patch_specs: List[Dict[str, Any]] | None = None,
    unresolved_manual_edit_keys: List[str] | None = None,
) -> Dict[str, Any]:
    changed_bullets: List[Dict[str, Any]] = []
    seen = set()

    for patch in list(patch_specs or []):
        bullet_key = _clean_text(patch.get("bullet_key"))
        patch_text = _clean_text(patch.get("patch_text"))
        if not bullet_key or not patch_text or bullet_key in seen:
            continue
        seen.add(bullet_key)

        changed_bullets.append({
            "bullet_key": bullet_key,
            "candidate_id": _clean_text(patch.get("candidate_id")),
            "source_bullet_id": _clean_text(patch.get("source_bullet_id")),
            "original_text": _clean_text(patch.get("source_raw_text")),
            "current_text": patch_text,
            "patch_text": patch_text,
            "patch_source": _clean_text(patch.get("patch_source")),
            "can_patch_in_place": True,
            "requires_full_preview_reload": False,
        })

    return {
        "version": "workspace_draft_fragments_v1",
        "status": _clean_text(preview_status),
        "note": _clean_text(preview_note),
        "changed_bullets": changed_bullets,
        "changed_bullet_keys": [row["bullet_key"] for row in changed_bullets],
        "fragment_count": len(changed_bullets),
        "can_patch_in_place": True,
        "requires_full_preview_reload": False,
        "unresolved_manual_edit_keys": list(unresolved_manual_edit_keys or []),
    }


def _normalize_selected_patch_candidate_ids(value: Any) -> List[str]:
    if isinstance(value, list):
        raw_items = value
    else:
        raw_text = _clean_text(value)
        if not raw_text:
            raw_items = []
        else:
            try:
                parsed = json.loads(raw_text)
                raw_items = parsed if isinstance(parsed, list) else [raw_text]
            except Exception:
                raw_items = [part.strip() for part in raw_text.split(",") if part.strip()]

    normalized: List[str] = []
    seen = set()

    for item in raw_items:
        candidate_id = _clean_text(item)
        if not candidate_id or candidate_id in seen:
            continue
        seen.add(candidate_id)
        normalized.append(candidate_id)

    return normalized


def _serialize_selected_patch_candidate_ids(value: Any) -> str:
    return json.dumps(
        _normalize_selected_patch_candidate_ids(value),
        ensure_ascii=False,
    )


def _load_tailoring_json_artifact(path: Path) -> Dict[str, Any]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError(f"Tailoring artifact must be a JSON object: {path}")
    return data

def _infer_packet_json_path_from_tailoring_artifact(artifact_path: Path) -> Path:
    name = artifact_path.name
    suffix = "__tailoring.json"
    if not name.endswith(suffix):
        return Path("")
    return artifact_path.with_name(f"{name[:-len(suffix)]}.json")


def _artifact_needs_operator_rehydration(payload_data: Dict[str, Any]) -> bool:
    replacement_candidates = list(payload_data.get("replacement_candidates", []) or [])
    if not replacement_candidates:
        return any(
            [
                list(payload_data.get("rewrite_candidates", []) or []),
                list(payload_data.get("bullet_reuse_candidates", []) or []),
                list(payload_data.get("edit_cards", []) or []),
                list(payload_data.get("top_edit_priorities", []) or []),
                list(payload_data.get("bullet_diagnoses", []) or []),
            ]
        )

    final_replacement_summary = dict(payload_data.get("final_replacement_summary", {}) or {})
    rewrite_review_groups = list(payload_data.get("rewrite_review_groups", []) or [])

    if "ai_optimize_optional_count" not in final_replacement_summary:
        return True

    if payload_data.get("ai_optimize_optional_replacements") is None:
        return True

    group_ids = {
        _clean_text(row.get("group_id"))
        for row in rewrite_review_groups
        if isinstance(row, dict)
    }

    has_ai_optimize_group = "ai_optimize_optional" in group_ids
    has_any_optional_ai_rows = bool(list(payload_data.get("ai_optimize_optional_replacements", []) or []))

    if has_any_optional_ai_rows and not has_ai_optimize_group:
        return True

    return False


def _rehydrate_legacy_tailoring_operator_payload(
    artifact_path: Path,
    payload_data: Dict[str, Any],
) -> Dict[str, Any]:
    data = dict(payload_data)

    if not _artifact_needs_operator_rehydration(data):
        return data

    packet_path = _infer_packet_json_path_from_tailoring_artifact(artifact_path)
    if not packet_path or not packet_path.exists() or not packet_path.is_file():
        return data

    try:
        from src.tailoring.packet_support import _load_packet
        from src.tailoring.rendering import _build_payload, _build_operator_markdown_payload

        packet = _load_packet(packet_path)
        rebuilt_base = _build_payload(packet, include_llm_prompts=False)
        rebuilt_operator = _build_operator_markdown_payload(rebuilt_base, None)
    except Exception:
        return data

    merged = dict(data)
    merged.update(rebuilt_operator)
    return merged

def _tailoring_artifact_candidate_ids(payload: Dict[str, Any]) -> List[str]:
    candidate_ids: List[str] = []
    seen = set()

    for row in list(payload.get("replacement_candidates", []) or []):
        candidate_id = _clean_text(row.get("candidate_id"))
        if not candidate_id or candidate_id in seen:
            continue
        seen.add(candidate_id)
        candidate_ids.append(candidate_id)

    return candidate_ids

def _default_selected_candidate_ids_from_replacement_plan(payload: Dict[str, Any]) -> List[str]:
    selected_ids: List[str] = []
    seen = set()

    for row in list(payload.get("app_ready_replacements", []) or []):
        candidate_id = _clean_text(row.get("replacement_candidate_id"))
        if not candidate_id or candidate_id in seen:
            continue
        seen.add(candidate_id)
        selected_ids.append(candidate_id)

    return selected_ids

def _tailoring_artifact_signature(payload: Dict[str, Any]) -> str:
    job = payload.get("job", {}) or {}
    selection = payload.get("selection", {}) or {}

    signature_payload = {
        "company": _clean_text(job.get("company")),
        "title": _clean_text(job.get("title")),
        "selected_resume": _clean_text(selection.get("selected_resume")),
        "replacement_candidate_ids": sorted(_tailoring_artifact_candidate_ids(payload)),
    }

    blob = json.dumps(signature_payload, sort_keys=True, ensure_ascii=False)
    return hashlib.sha1(blob.encode("utf-8")).hexdigest()[:16]

def _load_latest_patch_selection_overlay() -> Dict[str, Dict[str, Any]]:
    cached = _PATCH_SELECTION_OVERLAY_CACHE.get("data")
    if isinstance(cached, dict):
        return dict(cached)

    try:
        meta_payload = get_patch_selections_postgres_status_payload(
            limit=1,
            database_url="",
            database_url_env="DATABASE_URL",
            psql_bin="psql",
            print_only=False,
        )
    except (Exception, SystemExit):
        _PATCH_SELECTION_OVERLAY_CACHE["data"] = {}
        return {}

    meta_block = dict(meta_payload.get("postgres", {}) or {})
    query_limit = max(int(meta_block.get("latest_state_count", 0) or 0), 1)

    try:
        postgres_payload = get_patch_selections_postgres_status_payload(
            limit=query_limit,
            database_url="",
            database_url_env="DATABASE_URL",
            psql_bin="psql",
            print_only=False,
        )
    except (Exception, SystemExit):
        _PATCH_SELECTION_OVERLAY_CACHE["data"] = {}
        return {}

    postgres_block = dict(postgres_payload.get("postgres", {}) or {})
    postgres_rows = list(postgres_block.get("latest_rows", []) or [])

    latest_by_path: Dict[str, Dict[str, Any]] = {}
    for row in postgres_rows:
        artifact_path = _clean_text(row.get("tailoring_json_path"))
        if not artifact_path:
            continue

        latest_by_path[artifact_path] = {
            "selection_timestamp": _clean_text(row.get("selection_timestamp")),
            "job_doc_id": _clean_text(row.get("job_doc_id")),
            "queue_rank": _clean_text(row.get("queue_rank")),
            "job_company": _clean_text(row.get("job_company")),
            "job_title": _clean_text(row.get("job_title")),
            "selected_resume": _clean_text(row.get("selected_resume")),
            "tailoring_json_path": artifact_path,
            "artifact_signature": _clean_text(row.get("artifact_signature")),
            "selected_candidate_ids_json": row.get("selected_candidate_ids_json", []),
            "note": _clean_text(row.get("note")),
        }

    _PATCH_SELECTION_OVERLAY_CACHE["data"] = dict(latest_by_path)
    return dict(latest_by_path)

def _ensure_tailoring_preview_fields(payload_data: Dict[str, Any]) -> Dict[str, Any]:
    from src.tailoring.rendering import build_selected_patch_set_counterfactual_preview

    data = dict(payload_data)

    replacement_candidates = list(data.get("replacement_candidates", []) or [])
    if not replacement_candidates:
        return data

    if not isinstance(data.get("patch_set_counterfactual_preview"), dict):
        data["patch_set_counterfactual_preview"] = build_selected_patch_set_counterfactual_preview(
            data,
            selected_candidate_ids=None,
        )

    return data

def _derive_workspace_button_state_from_raw_payload(
    payload_data: Dict[str, Any],
) -> Dict[str, Any]:
    def _as_int(value: Any) -> int:
        try:
            return max(int(value or 0), 0)
        except (TypeError, ValueError):
            return 0

    app_ready = list(payload_data.get("app_ready_replacements", []) or [])
    direct_apply_optional = list(payload_data.get("direct_apply_optional_replacements", []) or [])
    ai_optimize_optional = list(payload_data.get("ai_optimize_optional_replacements", []) or [])
    direction_only = list(payload_data.get("direction_only_replacements", []) or [])
    final_replacement_decisions = list(payload_data.get("final_replacement_decisions", []) or [])
    anchor_cards = list(payload_data.get("anchor_cards", []) or [])
    top_anchor_priorities = list(payload_data.get("top_anchor_priorities", []) or [])

    # Fast modern path: use explicit operator payload lanes if they already exist.
    if any(
        [
            app_ready,
            direct_apply_optional,
            ai_optimize_optional,
            direction_only,
            final_replacement_decisions,
            anchor_cards,
            top_anchor_priorities,
        ]
    ):
        ready_count = len(app_ready)
        actionable_count = len(app_ready) + len(direct_apply_optional) + len(ai_optimize_optional)
        review_count = len(direction_only)
        anchor_count = len(anchor_cards) or len(top_anchor_priorities)

        has_replacement_plan = bool(
            final_replacement_decisions
            or app_ready
            or direct_apply_optional
            or ai_optimize_optional
            or direction_only
        )
        has_anchor_evidence = bool(anchor_cards or top_anchor_priorities)

        if actionable_count > 0:
            workspace_state = "ready"
        elif review_count > 0 or has_replacement_plan or has_anchor_evidence:
            workspace_state = "review"
        else:
            workspace_state = "empty"

        return {
            "tailoring_ready_replacement_count": ready_count,
            "tailoring_actionable_replacement_count": actionable_count,
            "tailoring_review_replacement_count": review_count + anchor_count,
            "tailoring_has_ready_replacements": actionable_count > 0,
            "tailoring_has_review_guidance": (review_count + anchor_count) > 0,
            "tailoring_workspace_state": workspace_state,
        }

    # Cheap strict fallback for legacy artifacts:
    # only trust explicit summary counts for ready/actionable; otherwise degrade to review.
    final_replacement_summary = dict(payload_data.get("final_replacement_summary", {}) or {})

    ready_count = max(
        _as_int(final_replacement_summary.get("app_ready_count")),
        _as_int(final_replacement_summary.get("direct_apply_ready_count")),
    )
    direct_apply_optional_count = _as_int(
        final_replacement_summary.get("direct_apply_optional_count")
    )
    ai_optimize_optional_count = _as_int(
        final_replacement_summary.get("ai_optimize_optional_count")
    )
    direction_only_count = _as_int(
        final_replacement_summary.get("direction_only_count")
    )

    actionable_count = ready_count + direct_apply_optional_count + ai_optimize_optional_count

    legacy_review_count = max(
        len(list(payload_data.get("top_edit_priorities", []) or [])),
        len(list(payload_data.get("edit_cards", []) or [])),
        len(list(payload_data.get("rewrite_candidates", []) or [])),
        len(list(payload_data.get("bullet_reuse_candidates", []) or [])),
        len(list(payload_data.get("bullet_diagnoses", []) or [])),
    )
    anchor_count = max(
        len(list(payload_data.get("anchor_cards", []) or [])),
        len(list(payload_data.get("top_anchor_priorities", []) or [])),
    )

    review_count = max(direction_only_count, legacy_review_count)

    if actionable_count > 0:
        workspace_state = "ready"
    elif review_count > 0 or anchor_count > 0:
        workspace_state = "review"
    else:
        workspace_state = "empty"

    return {
        "tailoring_ready_replacement_count": ready_count,
        "tailoring_actionable_replacement_count": actionable_count,
        "tailoring_review_replacement_count": review_count + anchor_count,
        "tailoring_has_ready_replacements": actionable_count > 0,
        "tailoring_has_review_guidance": (review_count + anchor_count) > 0,
        "tailoring_workspace_state": workspace_state,
    }


def _tailoring_workspace_button_state(
    row: Dict[str, Any],
    output_dir: Path = DEFAULT_OUTPUT_DIR,
) -> Dict[str, Any]:
    result = {
        "tailoring_ready_replacement_count": 0,
        "tailoring_actionable_replacement_count": 0,
        "tailoring_review_replacement_count": 0,
        "tailoring_has_ready_replacements": False,
        "tailoring_has_review_guidance": False,
        "tailoring_workspace_state": "empty",
    }

    raw_path = _clean_text(row.get("tailoring_json"))
    if not raw_path:
        return result

    try:
        artifact_path = _resolve_planning_artifact_path(raw_path, output_dir=output_dir)
        cache_key = _tailoring_workspace_button_state_cache_key(artifact_path)
        cached = _TAILORING_WORKSPACE_BUTTON_STATE_CACHE.get(cache_key)
        if cached is not None:
            return dict(cached)

        payload_data = _load_tailoring_json_artifact(artifact_path)

        result.update(_derive_workspace_button_state_from_raw_payload(payload_data))

        _TAILORING_WORKSPACE_BUTTON_STATE_CACHE[cache_key] = dict(result)
    except Exception:
        return result

    return result

def _normalize_tailoring_state_filter_values(value: Any) -> List[str]:
    if isinstance(value, list):
        raw_items = value
    else:
        raw_text = _clean_text(value)
        raw_items = [part.strip() for part in raw_text.split(",")] if raw_text else []

    normalized: List[str] = []
    seen = set()

    for item in raw_items:
        text = _clean_text(item).lower()
        if not text:
            continue

        if text == "empty":
            text = "unavailable"

        if text not in {"ready", "review", "unavailable"}:
            continue

        if text in seen:
            continue

        seen.add(text)
        normalized.append(text)

    return normalized


def _row_matches_tailoring_state_filter(
    row: Dict[str, Any],
    requested_states: List[str],
    *,
    output_dir: Path,
) -> Tuple[bool, Dict[str, Any]]:
    tailoring_state = _tailoring_workspace_button_state(row, output_dir=output_dir)
    workspace_state = _clean_text(tailoring_state.get("tailoring_workspace_state")).lower()

    normalized_state = "unavailable" if workspace_state == "empty" else workspace_state
    matches = not requested_states or normalized_state in set(requested_states)

    enriched_row = {
        **dict(row),
        **tailoring_state,
    }
    enriched_row["tailoring_workspace_state"] = normalized_state

    return matches, enriched_row

def _apply_saved_patch_selection_overlay(
    artifact_path: Path,
    payload_data: Dict[str, Any],
) -> Dict[str, Any]:
    from src.tailoring.rendering import build_selected_patch_set_counterfactual_preview

    data = _ensure_tailoring_preview_fields(payload_data)
    data.setdefault("selected_patch_candidate_ids", [])
    data.setdefault("selected_patch_selection_status", "none")
    data.setdefault("selected_patch_selection_note", "")
    data.setdefault("selected_patch_selection_timestamp", "")

    latest_by_path = _load_latest_patch_selection_overlay()
    selection_row = latest_by_path.get(str(artifact_path))

    if not selection_row:
        default_selected_ids = _default_selected_candidate_ids_from_replacement_plan(data)
        if not default_selected_ids:
            return data

        data["selected_patch_candidate_ids"] = default_selected_ids
        data["selected_patch_selection_status"] = "auto_from_replacement_plan"
        data["selected_patch_selection_note"] = (
            "No saved manual patch selection found. Defaulting to apply-now replacements from the final replacement plan."
        )
        data["selected_patch_set_counterfactual_preview"] = build_selected_patch_set_counterfactual_preview(
            data,
            selected_candidate_ids=default_selected_ids,
        )
        return data

    saved_signature = _clean_text(selection_row.get("artifact_signature"))
    current_signature = _tailoring_artifact_signature(data)

    data["selected_patch_selection_timestamp"] = _clean_text(selection_row.get("selection_timestamp"))

    if saved_signature and saved_signature != current_signature:
        data["selected_patch_selection_status"] = "stale_signature"
        data["selected_patch_selection_note"] = (
            "Saved patch selection was ignored because the tailoring artifact changed after the selection was recorded."
        )
        return data

    requested_ids = _normalize_selected_patch_candidate_ids(
        selection_row.get("selected_candidate_ids_json", "")
    )
    valid_candidate_ids = set(_tailoring_artifact_candidate_ids(data))
    selected_ids = [candidate_id for candidate_id in requested_ids if candidate_id in valid_candidate_ids]

    if not selected_ids:
        data["selected_patch_selection_status"] = "no_valid_candidates"
        data["selected_patch_selection_note"] = (
            "Saved patch selection exists, but none of the candidate IDs are valid for the current artifact."
        )
        return data

    data["selected_patch_candidate_ids"] = selected_ids
    data["selected_patch_selection_status"] = "applied"
    data["selected_patch_selection_note"] = "Saved patch selection applied to this tailoring artifact."
    data["selected_patch_set_counterfactual_preview"] = build_selected_patch_set_counterfactual_preview(
        data,
        selected_candidate_ids=selected_ids,
    )
    return data

def _sanitize_optional_resume_filename(value: Any) -> str:
    raw = _clean_text(value)
    return _sanitize_resume_filename(raw) if raw else ""


def _normalize_operator_decision(value: Any) -> str:
    normalized = _clean_text(value).upper().replace(" ", "_")
    if not normalized:
        raise ValueError("decision is required.")
    if normalized not in ALLOWED_OPERATOR_DECISIONS:
        allowed = ", ".join(sorted(ALLOWED_OPERATOR_DECISIONS))
        raise ValueError(f"Invalid decision={normalized!r}. Allowed values: {allowed}")
    return normalized


def _operator_decision_key(row: Dict[str, Any]) -> str:
    return _application_action_key(row)


def _validate_operator_decision_identity(row: Dict[str, Any]) -> None:
    if not _operator_decision_key(row):
        raise ValueError(
            "Operator decision requires job_doc_id, job_url, or job_company + job_title."
        )


def _resolve_resume_preview_path(
    resume_name: str,
    *,
    owner_user_id: str = "",
) -> Path:
    safe_name = _sanitize_resume_filename(resume_name)
    owner_dir_name = _safe_owner_dir_name(owner_user_id)
    cache_key = f"{owner_dir_name}:{safe_name}" if owner_dir_name else safe_name

    if owner_dir_name:
        cached = _RESUME_PREVIEW_PATH_CACHE.get(cache_key, "")
        if cached:
            cached_path = Path(cached)
            if cached_path.exists() and cached_path.is_file():
                return cached_path

        resolved = _materialize_profile_resume_temp_path(
            safe_name,
            owner_user_id=owner_user_id,
        )
        _RESUME_PREVIEW_PATH_CACHE[cache_key] = str(resolved)
        return resolved

    cached = _RESUME_PREVIEW_PATH_CACHE.get(cache_key, "")
    if cached:
        cached_path = Path(cached)
        if cached_path.exists() and cached_path.is_file():
            return cached_path

    candidate_paths: List[Path] = []
    candidate_paths.append(_get_resume_dir() / safe_name)

    for profile_path in candidate_paths:
        if profile_path.exists() and profile_path.is_file():
            resolved = profile_path.resolve()
            _RESUME_PREVIEW_PATH_CACHE[cache_key] = str(resolved)
            return resolved

    ignore_dirs = {
        ".git",
        "__pycache__",
        "node_modules",
        ".venv",
        "venv",
        "env",
        "job_scrap312",
    }

    matches: List[Path] = []
    search_root = Path.cwd().resolve()

    for path in search_root.rglob(safe_name):
        if not path.is_file():
            continue
        if path.suffix.lower() != ".pdf":
            continue
        if any(part in ignore_dirs for part in path.parts):
            continue
        matches.append(path.resolve())

    if not matches:
        raise ValueError(f"Resume preview file not found: {safe_name}")

    matches.sort(key=lambda path: (len(path.parts), len(str(path))))
    chosen = matches[0]
    _RESUME_PREVIEW_PATH_CACHE[cache_key] = str(chosen)
    return chosen

def planning_resume_preview_path(
    resume_name: str,
    *,
    owner_user_id: str = "",
) -> Path:
    return _resolve_resume_preview_path(
        resume_name,
        owner_user_id=owner_user_id,
    )

def _slugify_text(value: Any, max_len: int = 80) -> str:
    text = _clean_text(value).lower()
    text = re.sub(r"[^a-z0-9]+", "_", text).strip("_")
    if not text:
        text = "item"
    return text[:max_len]


def _load_job_doc_id_to_index(job_corpus_path: Path) -> Dict[str, int]:
    from src.matching.job_adapter import build_job_evidence

    if not job_corpus_path.exists():
        raise ValueError(f"Missing job corpus: {job_corpus_path}")

    mapping: Dict[str, int] = {}
    with job_corpus_path.open("r", encoding="utf-8") as f:
        for idx, line in enumerate(f):
            line = line.strip()
            if not line:
                continue

            record = json.loads(line)
            job_evidence = build_job_evidence(record)
            job_doc_id = _clean_text(getattr(job_evidence, "job_doc_id", ""))

            if job_doc_id:
                mapping[job_doc_id] = idx

    if not mapping:
        raise ValueError(f"No usable job_doc_id values found in {job_corpus_path}")

    return mapping

def _job_corpus_identity_values(record: Dict[str, Any]) -> List[str]:
    from src.matching.job_adapter import build_job_evidence

    metadata = record.get("metadata", {}) if isinstance(record.get("metadata"), dict) else {}

    adapter_job_doc_id = ""
    try:
        adapter_job_doc_id = _clean_text(getattr(build_job_evidence(record), "job_doc_id", ""))
    except Exception:
        adapter_job_doc_id = ""

    values = [
        adapter_job_doc_id,
        _clean_text(record.get("job_doc_id")),
        _clean_text(record.get("doc_id")),
        _clean_text(record.get("job_url")),
        _clean_text(metadata.get("job_doc_id")),
        _clean_text(metadata.get("doc_id")),
        _clean_text(metadata.get("job_url")),
    ]

    out: List[str] = []
    seen = set()
    for value in values:
        if not value or value in seen:
            continue
        seen.add(value)
        out.append(value)
    return out


def _job_corpus_company_title(record: Dict[str, Any]) -> Tuple[str, str]:
    metadata = record.get("metadata", {}) if isinstance(record.get("metadata"), dict) else {}

    company = _clean_text(
        record.get("job_company")
        or record.get("company")
        or metadata.get("job_company")
        or metadata.get("company")
    )
    title = _clean_text(
        record.get("job_title")
        or record.get("title")
        or metadata.get("job_title")
        or metadata.get("title")
    )
    return company, title


def _resolve_job_index_for_regeneration(
    job_corpus_path: Path,
    *,
    job_doc_id: str = "",
    job_url: str = "",
    job_company: str = "",
    job_title: str = "",
) -> int:
    ja = _job_app()

    clean_job_doc_id = _clean_text(job_doc_id)
    clean_job_url = _clean_text(job_url)
    target_company_norm = ja._normalize_text(job_company)
    target_title_norm = ja._normalize_text(job_title)

    if not job_corpus_path.exists():
        raise ValueError(f"Missing job corpus: {job_corpus_path}")

    exact_identity_hits: List[int] = []
    title_company_hits: List[int] = []

    with job_corpus_path.open("r", encoding="utf-8") as f:
        for idx, raw_line in enumerate(f):
            line = raw_line.strip()
            if not line:
                continue

            record = json.loads(line)

            identity_values = _job_corpus_identity_values(record)
            if clean_job_doc_id and clean_job_doc_id in identity_values:
                exact_identity_hits.append(idx)
                continue

            if clean_job_url and clean_job_url in identity_values:
                exact_identity_hits.append(idx)
                continue

            record_company, record_title = _job_corpus_company_title(record)
            record_company_norm = ja._normalize_text(record_company)
            record_title_norm = ja._normalize_text(record_title)

            if (
                target_company_norm
                and target_title_norm
                and record_company_norm == target_company_norm
                and record_title_norm == target_title_norm
            ):
                title_company_hits.append(idx)

    if len(exact_identity_hits) == 1:
        return exact_identity_hits[0]

    if len(exact_identity_hits) > 1:
        raise ValueError(
            f"Ambiguous corpus match for job identity: job_doc_id={clean_job_doc_id!r} job_url={clean_job_url!r}"
        )

    if len(title_company_hits) == 1:
        return title_company_hits[0]

    if len(title_company_hits) > 1:
        raise ValueError(
            f"Ambiguous corpus match for company/title fallback: company={job_company!r} title={job_title!r}"
        )

    raise ValueError(
        f"Could not resolve corpus index for regeneration: "
        f"job_doc_id={clean_job_doc_id!r} job_url={clean_job_url!r} "
        f"company={job_company!r} title={job_title!r}"
    )

def _load_csv_rows_with_fieldnames(path: Path) -> Tuple[List[Dict[str, str]], List[str]]:
    if not path.exists():
        raise ValueError(f"Missing CSV: {path}")

    with path.open("r", encoding="utf-8", newline="") as f:
        reader = csv.DictReader(f)
        rows = list(reader)
        fieldnames = list(reader.fieldnames or [])

    return rows, fieldnames


def _write_csv_rows(path: Path, fieldnames: List[str], rows: List[Dict[str, Any]]) -> None:
    with path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow(row)


def _run_checked_cmd(cmd: List[str]) -> None:
    subprocess.run(cmd, check=True)


def _read_regenerated_llm_status(llm_json_path: Path) -> Dict[str, str]:
    defaults = {
        "llm_tailoring_status": "disabled",
        "llm_cache_hit": "",
        "llm_parse_ok": "",
        "llm_provider": "",
        "llm_model": "",
        "llm_error_type": "",
        "llm_retryable": "",
        "llm_retry_used": "",
    }

    if not llm_json_path.exists():
        return defaults

    try:
        data = json.loads(llm_json_path.read_text(encoding="utf-8"))
    except Exception:
        out = dict(defaults)
        out["llm_tailoring_status"] = "unreadable"
        out["llm_error_type"] = "unreadable_json"
        out["llm_retryable"] = "False"
        return out

    parse_ok = bool(data.get("parse_ok"))
    cache_hit = bool(data.get("cache_hit"))
    retry_used = bool(data.get("retry_used"))

    if parse_ok and cache_hit:
        status = "cached"
    elif parse_ok:
        status = "generated"
    else:
        status = "failed"

    return {
        "llm_tailoring_status": status,
        "llm_cache_hit": str(cache_hit),
        "llm_parse_ok": str(parse_ok),
        "llm_provider": str(data.get("provider", "")),
        "llm_model": str(data.get("model", "")),
        "llm_error_type": "" if parse_ok else "llm_parse_failed",
        "llm_retryable": "" if parse_ok else "False",
        "llm_retry_used": str(retry_used),
    }


def _find_planning_row_for_regeneration(
    rows: List[Dict[str, Any]],
    *,
    job_doc_id: str = "",
    queue_rank: str = "",
) -> Dict[str, Any]:
    clean_job_doc_id = _clean_text(job_doc_id)
    clean_queue_rank = _clean_text(queue_rank)

    if clean_job_doc_id:
        for row in rows:
            if _clean_text(row.get("job_doc_id")) == clean_job_doc_id:
                return row

    if clean_queue_rank:
        for row in rows:
            if _clean_text(row.get("queue_rank")) == clean_queue_rank:
                return row

    raise ValueError("Could not find planning row for targeted regeneration.")


def regenerate_selected_resume_tailoring_payload(
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    job_corpus: Path = DEFAULT_CORPUS_PATH,
    *,
    job_doc_id: str = "",
    queue_rank: str = "",
    selected_resume: str = "",
    generate_llm_tailoring: bool = False,
    refresh_llm_tailoring: bool = False,
) -> Dict[str, Any]:
    ja = _job_app()
    merged_rows = ja._build_job_index(output_dir)
    target_row = _find_planning_row_for_regeneration(
        merged_rows,
        job_doc_id=job_doc_id,
        queue_rank=queue_rank,
    )

    chosen_resume = _sanitize_resume_filename(
        selected_resume or target_row.get("operator_selected_resume", "")
    )
    if not chosen_resume:
        raise ValueError("No operator-selected resume is available for regeneration.")

    allowed_resumes = {
        value
        for value in [
            _sanitize_optional_resume_filename(target_row.get("winner_resume")),
            _sanitize_optional_resume_filename(target_row.get("runner_up_resume")),
        ]
        if value
    }
    if chosen_resume not in allowed_resumes:
        allowed = ", ".join(sorted(allowed_resumes))
        raise ValueError(
            f"selected_resume must match the current winner or runner-up. Allowed: {allowed}"
        )

    job_doc_id_value = _clean_text(target_row.get("job_doc_id"))
    job_url_value = _clean_text(target_row.get("job_url"))
    company = _clean_text(target_row.get("job_company"))
    title = _clean_text(target_row.get("job_title"))
    action = _clean_text(target_row.get("action"))

    if not any([job_doc_id_value, job_url_value, company, title]):
        raise ValueError("Target planning row is missing usable job identity for regeneration.")

    job_index = _resolve_job_index_for_regeneration(
        job_corpus,
        job_doc_id=job_doc_id_value,
        job_url=job_url_value,
        job_company=company,
        job_title=title,
    )

    job_packets_dir = Path(output_dir) / "job_packets"
    job_packets_dir.mkdir(parents=True, exist_ok=True)

    training_log_jsonl_path = Path(output_dir) / "training_logs" / "tailoring_runs.jsonl"
    training_log_jsonl_path.parent.mkdir(parents=True, exist_ok=True)

    file_slug = (
        f"{_slugify_text(company, 30)}__"
        f"{_slugify_text(title, 60)}__"
        f"{_slugify_text(chosen_resume, 40)}"
    )

    packet_json_path = job_packets_dir / f"{file_slug}.json"
    tailoring_json_path = job_packets_dir / f"{file_slug}__tailoring.json"
    tailoring_md_path = job_packets_dir / f"{file_slug}__tailoring.md"
    tailoring_llm_json_path = job_packets_dir / f"{file_slug}__tailoring_llm.json"

    diff_cmd = [
        sys.executable,
        "jd_resume_diff_helper.py",
        "--job-corpus",
        str(job_corpus),
        "--job-index",
        str(job_index),
        "--resume-name-contains",
        chosen_resume,
        "--disable-semantic-evidence",
        "--output-json",
        str(packet_json_path),
    ]

    _run_checked_cmd(diff_cmd)

    tailoring_cmd = [
        sys.executable,
        "generate_tailoring_suggestions.py",
        "--packet-json",
        str(packet_json_path),
        "--output-json",
        str(tailoring_json_path),
        "--output-md",
        str(tailoring_md_path),
        "--training-log-jsonl",
        str(training_log_jsonl_path),
    ]

    llm_status = {
        "llm_tailoring_status": "disabled",
        "llm_cache_hit": "",
        "llm_parse_ok": "",
        "llm_provider": "",
        "llm_model": "",
        "llm_error_type": "",
        "llm_retryable": "",
        "llm_retry_used": "",
    }
    llm_json_value = ""

    if generate_llm_tailoring:
        tailoring_cmd.extend(
            [
                "--use-llm",
                "--output-llm-json",
                str(tailoring_llm_json_path),
            ]
        )
        if refresh_llm_tailoring:
            tailoring_cmd.append("--refresh-llm-cache")

    _run_checked_cmd(tailoring_cmd)

    if generate_llm_tailoring:
        llm_status = _read_regenerated_llm_status(tailoring_llm_json_path)
        llm_json_value = str(tailoring_llm_json_path)

    manifest_path = Path(output_dir) / "job_packet_manifest.csv"
    manifest_rows, fieldnames = _load_csv_rows_with_fieldnames(manifest_path)

    if "packet_status" not in fieldnames:
        insert_at = fieldnames.index("packet_json") if "packet_json" in fieldnames else len(fieldnames)
        fieldnames.insert(insert_at, "packet_status")
        for row in manifest_rows:
            row.setdefault("packet_status", "")

    updated = False
    for row in manifest_rows:
        if _clean_text(row.get("job_doc_id")) != job_doc_id_value:
            continue

        row["queue_rank"] = _clean_text(target_row.get("queue_rank"))
        row["needs_variant_review"] = _clean_text(target_row.get("needs_variant_review"))
        row["missing_requirement_count"] = _clean_text(target_row.get("missing_requirement_count"))
        row["queue_priority_reason"] = _clean_text(target_row.get("queue_priority_reason"))
        row["job_doc_id"] = job_doc_id_value
        row["job_company"] = company
        row["job_title"] = title
        row["action"] = action
        row["winner_resume"] = _clean_text(target_row.get("winner_resume"))
        row["winner_score"] = _clean_text(target_row.get("winner_score"))
        row["runner_up_resume"] = _clean_text(target_row.get("runner_up_resume"))
        row["runner_up_score"] = _clean_text(target_row.get("runner_up_score"))
        row["score_gap"] = _clean_text(target_row.get("score_gap"))
        row["is_tie"] = _clean_text(target_row.get("is_tie"))
        row["packet_status"] = "generated"
        row["packet_json"] = str(packet_json_path)
        row["tailoring_json"] = str(tailoring_json_path)
        row["tailoring_md"] = str(tailoring_md_path)
        row["tailoring_llm_json"] = llm_json_value
        row["llm_tailoring_status"] = llm_status["llm_tailoring_status"]
        row["llm_cache_hit"] = llm_status["llm_cache_hit"]
        row["llm_parse_ok"] = llm_status["llm_parse_ok"]
        row["llm_provider"] = llm_status["llm_provider"]
        row["llm_model"] = llm_status["llm_model"]
        row["llm_error_type"] = llm_status["llm_error_type"]
        row["llm_retryable"] = llm_status["llm_retryable"]
        row["llm_retry_used"] = llm_status["llm_retry_used"]
        updated = True
        break

    if not updated:
        manifest_row = {field: "" for field in fieldnames}
        manifest_row["queue_rank"] = _clean_text(target_row.get("queue_rank"))
        manifest_row["needs_variant_review"] = _clean_text(target_row.get("needs_variant_review"))
        manifest_row["missing_requirement_count"] = _clean_text(target_row.get("missing_requirement_count"))
        manifest_row["queue_priority_reason"] = _clean_text(target_row.get("queue_priority_reason"))
        manifest_row["job_doc_id"] = job_doc_id_value
        manifest_row["job_company"] = company
        manifest_row["job_title"] = title
        manifest_row["action"] = action
        manifest_row["winner_resume"] = _clean_text(target_row.get("winner_resume"))
        manifest_row["winner_score"] = _clean_text(target_row.get("winner_score"))
        manifest_row["runner_up_resume"] = _clean_text(target_row.get("runner_up_resume"))
        manifest_row["runner_up_score"] = _clean_text(target_row.get("runner_up_score"))
        manifest_row["score_gap"] = _clean_text(target_row.get("score_gap"))
        manifest_row["is_tie"] = _clean_text(target_row.get("is_tie"))
        manifest_row["packet_status"] = "generated"
        manifest_row["packet_json"] = str(packet_json_path)
        manifest_row["tailoring_json"] = str(tailoring_json_path)
        manifest_row["tailoring_md"] = str(tailoring_md_path)
        manifest_row["tailoring_llm_json"] = llm_json_value
        manifest_row["llm_tailoring_status"] = llm_status["llm_tailoring_status"]
        manifest_row["llm_cache_hit"] = llm_status["llm_cache_hit"]
        manifest_row["llm_parse_ok"] = llm_status["llm_parse_ok"]
        manifest_row["llm_provider"] = llm_status["llm_provider"]
        manifest_row["llm_model"] = llm_status["llm_model"]
        manifest_row["llm_error_type"] = llm_status["llm_error_type"]
        manifest_row["llm_retryable"] = llm_status["llm_retryable"]
        manifest_row["llm_retry_used"] = llm_status["llm_retry_used"]
        manifest_rows.append(manifest_row)

    _write_csv_rows(manifest_path, fieldnames, manifest_rows)

    return {
        "ok": True,
        "job_doc_id": job_doc_id_value,
        "job_company": company,
        "job_title": title,
        "action": action,
        "selected_resume": chosen_resume,
        "packet_json": str(packet_json_path),
        "tailoring_json": str(tailoring_json_path),
        "tailoring_md": str(tailoring_md_path),
        "tailoring_llm_json": llm_json_value,
        "training_log_jsonl": str(training_log_jsonl_path),
        "llm_tailoring_status": llm_status["llm_tailoring_status"],
        "manifest_path": str(manifest_path),
    }

def _normalize_application_status(value: Any) -> str:
    normalized = _clean_text(value).upper().replace(" ", "_")
    if not normalized:
        raise ValueError("application_status is required.")
    if normalized not in ALLOWED_APPLICATION_STATUSES:
        allowed = ", ".join(sorted(ALLOWED_APPLICATION_STATUSES))
        raise ValueError(f"Invalid application_status={normalized!r}. Allowed values: {allowed}")
    return normalized


def _application_action_key(row: Dict[str, Any]) -> str:
    ja = _job_app()

    job_doc_id = _clean_text(row.get("job_doc_id"))
    if job_doc_id:
        return f"job_doc_id::{job_doc_id}"

    job_url = _clean_text(row.get("job_url"))
    if job_url:
        return f"job_url::{job_url}"

    company = ja._normalize_text(row.get("job_company", ""))
    title = ja._normalize_text(row.get("job_title", ""))
    if company or title:
        return f"title::{company}||{title}"

    return ""


def _validate_application_identity(row: Dict[str, Any]) -> None:
    if not _application_action_key(row):
        raise ValueError(
            "Application action requires job_doc_id, job_url, or job_company + job_title."
        )

def _dual_write_application_action_postgres(row: Dict[str, Any]) -> Dict[str, Any]:
    database_url = str(os.environ.get("DATABASE_URL", "") or "").strip()
    if not database_url:
        return {
            "attempted": False,
            "ok": False,
            "skipped": "missing_database_url",
            "table_name": "application_actions",
        }

    try:
        payload = insert_application_action_row_to_postgres(
            record=row,
            database_url="",
            database_url_env="DATABASE_URL",
            psql_bin="psql",
            print_only=False,
            allow_contract_drift=False,
        )
        return {
            "attempted": True,
            "ok": True,
            "table_name": payload.get("table_name", "application_actions"),
            "action_id": str(payload.get("row", {}).get("action_id", "") or ""),
            "action_key": str(payload.get("row", {}).get("action_key", "") or ""),
            "contract_health_ok": bool(payload.get("contract_health_ok", False)),
            "command_text": str(payload.get("command_text", "") or ""),
        }
    except SystemExit as exc:
        return {
            "attempted": True,
            "ok": False,
            "table_name": "application_actions",
            "error_type": "SystemExit",
            "error": str(exc),
        }
    except Exception as exc:
        return {
            "attempted": True,
            "ok": False,
            "table_name": "application_actions",
            "error_type": exc.__class__.__name__,
            "error": str(exc),
        }

def _application_action_latest_sort_key(row: Dict[str, Any]) -> Tuple[str, str]:
    normalized = application_action_db_row(dict(row))
    return (
        str(normalized.get("action_timestamp", "") or ""),
        str(normalized.get("action_id", "") or ""),
    )


def _load_latest_application_actions(owner_user_id: str = "") -> List[Dict[str, str]]:
    postgres_payload = get_latest_application_actions_rows(
        database_url="",
        database_url_env="DATABASE_URL",
        psql_bin="psql",
        print_only=False,
        owner_user_id=_clean_text(owner_user_id),
    )
    postgres_rows = list(postgres_payload.get("rows", []) or [])

    normalized_rows: List[Dict[str, str]] = []
    for row in postgres_rows:
        normalized_rows.append({
            "owner_user_id": _clean_text(row.get("owner_user_id")),
            "action_timestamp": _clean_text(row.get("action_timestamp")),
            "job_doc_id": _clean_text(row.get("job_doc_id")),
            "job_url": _clean_text(row.get("job_url")),
            "job_company": _clean_text(row.get("job_company")),
            "job_title": _clean_text(row.get("job_title")),
            "application_status": _clean_text(row.get("application_status")),
            "source_view": _clean_text(row.get("source_view")),
            "note": _clean_text(row.get("note")),
        })

    normalized_rows.sort(
        key=lambda row: _application_action_latest_sort_key(row),
        reverse=True,
    )
    return normalized_rows

def _application_row_key_candidates(row: Dict[str, Any]) -> List[str]:
    ja = _job_app()

    key_candidates: List[str] = []

    direct_key = _application_action_key(row)
    if direct_key:
        key_candidates.append(direct_key)

    job_doc_id = _clean_text(row.get("job_doc_id"))
    if job_doc_id:
        key_candidates.append(f"job_doc_id::{job_doc_id}")

    job_url = _clean_text(row.get("job_url"))
    if job_url:
        key_candidates.append(f"job_url::{job_url}")

    company = ja._normalize_text(row.get("job_company", "") or row.get("company", ""))
    title = ja._normalize_text(row.get("job_title", "") or row.get("title", ""))
    if company or title:
        key_candidates.append(f"title::{company}||{title}")

    deduped: List[str] = []
    for key in key_candidates:
        if key and key not in deduped:
            deduped.append(key)
    return deduped


def _application_overlay_from_row(action_row: Dict[str, Any]) -> Dict[str, Any]:
    status = _clean_text(action_row.get("application_status")).upper()
    is_applied = status == "APPLIED"

    return {
        "application_status": status,
        "application_label": "Applied" if is_applied else "Apply",
        "is_applied": is_applied,
    }


def _load_latest_application_action_overlay(owner_user_id: str = "") -> Dict[str, Dict[str, Any]]:
    latest_rows = _load_latest_application_actions(owner_user_id=owner_user_id)
    latest_by_key: Dict[str, Dict[str, Any]] = {}

    for row in latest_rows:
        overlay = _application_overlay_from_row(row)
        for key in _application_row_key_candidates(row):
            latest_by_key[key] = overlay

    return latest_by_key


def _overlay_application_actions(
    rows: List[Dict[str, Any]],
    owner_user_id: str = "",
) -> List[Dict[str, Any]]:
    latest_by_key = _load_latest_application_action_overlay(owner_user_id=owner_user_id)

    overlaid_rows: List[Dict[str, Any]] = []
    for row in rows:
        merged = dict(row)

        for field in APPLICATION_ACTION_OVERLAY_FIELDS:
            if field == "application_label":
                merged.setdefault(field, "Apply")
            elif field == "is_applied":
                merged.setdefault(field, False)
            else:
                merged.setdefault(field, "")

        overlay = None
        for key in _application_row_key_candidates(row):
            if key in latest_by_key:
                overlay = latest_by_key[key]
                break

        if overlay:
            merged.update(overlay)

        overlaid_rows.append(merged)

    return overlaid_rows

def _exclude_applied_rows(
    rows: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    return [
        row for row in rows
        if not bool(row.get("is_applied", False))
    ]

def _load_job_metadata_overlay_from_corpus(
    job_corpus: Path = DEFAULT_CORPUS_PATH,
) -> Dict[str, Dict[str, Any]]:
    cache_key = _job_metadata_overlay_cache_key(job_corpus)
    cached = _JOB_METADATA_OVERLAY_CACHE.get(cache_key)
    if cached is not None:
        return cached

    latest_by_key: Dict[str, Dict[str, Any]] = {}
    resolved_corpus = Path(job_corpus).expanduser().resolve()

    if not resolved_corpus.exists():
        return latest_by_key

    with resolved_corpus.open("r", encoding="utf-8") as f:
        for line in f:
            raw = line.strip()
            if not raw:
                continue

            try:
                record = json.loads(raw)
            except Exception:
                continue

            metadata = record.get("metadata", {}) if isinstance(record.get("metadata"), dict) else {}

            job_doc_id = _clean_text(
                record.get("job_doc_id")
                or record.get("doc_id")
                or metadata.get("job_doc_id")
                or metadata.get("doc_id")
                or record.get("job_url")
                or metadata.get("job_url")
            )
            job_url = _clean_text(
                record.get("job_url")
                or metadata.get("job_url")
                or job_doc_id
            )
            job_company = _clean_text(
                record.get("job_company")
                or record.get("company")
                or metadata.get("job_company")
                or metadata.get("company")
            )
            job_title = _clean_text(
                record.get("job_title")
                or record.get("title")
                or metadata.get("job_title")
                or metadata.get("title")
            )
            posted_at = _clean_text(
                record.get("posted_at")
                or metadata.get("posted_at")
            )

            if not any([job_doc_id, job_url, job_company, job_title]):
                continue

            key_row = {
                "job_doc_id": job_doc_id,
                "job_url": job_url,
                "job_company": job_company,
                "job_title": job_title,
            }
            overlay = {
                "posted_at": posted_at,
                "job_url": job_url,
            }

            for key in _application_row_key_candidates(key_row):
                if key:
                    latest_by_key[key] = overlay

    _JOB_METADATA_OVERLAY_CACHE.clear()
    _JOB_METADATA_OVERLAY_CACHE[cache_key] = latest_by_key
    return latest_by_key

def _overlay_job_metadata(
    rows: List[Dict[str, Any]],
    job_corpus: Path = DEFAULT_CORPUS_PATH,
) -> List[Dict[str, Any]]:
    latest_by_key = _load_job_metadata_overlay_from_corpus(job_corpus)

    overlaid_rows: List[Dict[str, Any]] = []

    for row in rows:
        merged = dict(row)
        merged["posted_at"] = _clean_text(merged.get("posted_at"))
        merged["job_url"] = _clean_text(merged.get("job_url")) or _clean_text(merged.get("job_doc_id"))

        if latest_by_key:
            for key in _application_row_key_candidates(merged):
                overlay = latest_by_key.get(key)
                if not overlay:
                    continue

                if overlay.get("posted_at"):
                    merged["posted_at"] = overlay["posted_at"]
                if overlay.get("job_url") and not _clean_text(merged.get("job_url")):
                    merged["job_url"] = overlay["job_url"]
                break

        overlaid_rows.append(merged)

    return overlaid_rows

def _csv_rows_from_text(text: Any) -> List[Dict[str, str]]:
    raw = str(text or "")
    if not raw.strip():
        return []

    reader = csv.DictReader(raw.splitlines())
    return [dict(row or {}) for row in reader]


def _jsonl_row_count_from_text(text: Any) -> int:
    count = 0
    for line in str(text or "").splitlines():
        if line.strip():
            count += 1
    return count


def _artifact_row_by_name(rows: List[Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
    return {
        _clean_text(row.get("artifact_name")): dict(row or {})
        for row in rows
        if _clean_text(row.get("artifact_name"))
    }


def _artifact_text_by_name(rows: List[Dict[str, Any]], artifact_name: str) -> str:
    row = _artifact_row_by_name(rows).get(_clean_text(artifact_name), {})
    return str(row.get("content_text", "") or "")


def _build_job_index_from_planning_rows(
    ja: Any,
    *,
    best_rows: List[Dict[str, str]],
    queue_rows: List[Dict[str, str]],
    manifest_rows: List[Dict[str, str]],
) -> List[Dict[str, str]]:
    merged: Dict[str, Dict[str, str]] = {}

    for source_rows in [best_rows, queue_rows, manifest_rows]:
        for row in source_rows:
            job_doc_id = _clean_text(row.get("job_doc_id"))
            company = _clean_text(row.get("job_company"))
            title = _clean_text(row.get("job_title"))

            key = job_doc_id or f"{ja._normalize_text(company)}||{ja._normalize_text(title)}"
            if not key.strip("|"):
                continue

            if key not in merged:
                merged[key] = {}
            merged[key].update(row)

    merged_rows = list(merged.values())
    return ja._overlay_operator_decisions(merged_rows)


def _job_metadata_overlay_from_jsonl_text(text: Any) -> Dict[str, Dict[str, Any]]:
    latest_by_key: Dict[str, Dict[str, Any]] = {}

    for line in str(text or "").splitlines():
        raw = line.strip()
        if not raw:
            continue

        try:
            record = json.loads(raw)
        except Exception:
            continue

        metadata = record.get("metadata", {}) if isinstance(record.get("metadata"), dict) else {}

        job_doc_id = _clean_text(
            record.get("job_doc_id")
            or record.get("doc_id")
            or metadata.get("job_doc_id")
            or metadata.get("doc_id")
            or record.get("job_url")
            or metadata.get("job_url")
        )
        job_url = _clean_text(
            record.get("job_url")
            or metadata.get("job_url")
            or job_doc_id
        )
        job_company = _clean_text(
            record.get("job_company")
            or record.get("company")
            or metadata.get("job_company")
            or metadata.get("company")
        )
        job_title = _clean_text(
            record.get("job_title")
            or record.get("title")
            or metadata.get("job_title")
            or metadata.get("title")
        )
        posted_at = _clean_text(
            record.get("posted_at")
            or metadata.get("posted_at")
        )

        if not any([job_doc_id, job_url, job_company, job_title]):
            continue

        key_row = {
            "job_doc_id": job_doc_id,
            "job_url": job_url,
            "job_company": job_company,
            "job_title": job_title,
        }
        overlay = {
            "posted_at": posted_at,
            "job_url": job_url,
        }

        for key in _application_row_key_candidates(key_row):
            if key:
                latest_by_key[key] = overlay

    return latest_by_key


def _overlay_job_metadata_from_map(
    rows: List[Dict[str, Any]],
    latest_by_key: Dict[str, Dict[str, Any]],
) -> List[Dict[str, Any]]:
    overlaid_rows: List[Dict[str, Any]] = []

    for row in rows:
        merged = dict(row)
        merged["posted_at"] = _clean_text(merged.get("posted_at"))
        merged["job_url"] = _clean_text(merged.get("job_url")) or _clean_text(merged.get("job_doc_id"))

        if latest_by_key:
            for key in _application_row_key_candidates(merged):
                overlay = latest_by_key.get(key)
                if not overlay:
                    continue

                if overlay.get("posted_at"):
                    merged["posted_at"] = overlay["posted_at"]
                if overlay.get("job_url") and not _clean_text(merged.get("job_url")):
                    merged["job_url"] = overlay["job_url"]
                break

        overlaid_rows.append(merged)

    return overlaid_rows


def _latest_user_pipeline_artifact_context(
    *,
    owner_user_id: str = "",
) -> Dict[str, Any]:
    owner = _clean_text(owner_user_id)
    if not owner:
        return {}

    latest_payload = get_user_pipeline_runs_postgres_payload(
        owner_user_id=owner,
        limit=1,
        status="succeeded",
        database_url="",
        database_url_env="DATABASE_URL",
        psql_bin="psql",
        print_only=False,
        ensure_schema=True,
    )
    runs = list(latest_payload.get("rows", []) or [])
    if not runs:
        return {}

    run = dict(runs[0] or {})
    run_id = _clean_text(run.get("run_id"))
    if not run_id:
        return {}

    artifacts_payload = get_user_pipeline_artifacts_postgres_payload(
        owner_user_id=owner,
        run_id=run_id,
        limit=100000,
        database_url="",
        database_url_env="DATABASE_URL",
        psql_bin="psql",
        print_only=False,
        ensure_schema=True,
    )
    artifacts = list(artifacts_payload.get("rows", []) or [])
    if not artifacts:
        return {}

    best_text = _artifact_text_by_name(artifacts, "best_resume_variant_by_job.csv")
    shortlist_text = _artifact_text_by_name(artifacts, "application_shortlist_by_job.csv")
    queue_text = _artifact_text_by_name(artifacts, "application_execution_queue.csv")
    manifest_text = _artifact_text_by_name(artifacts, "job_packet_manifest.csv")
    corpus_text = _artifact_text_by_name(artifacts, "current_run_job_corpus.jsonl")

    if not any([best_text, queue_text, manifest_text]):
        return {}

    return {
        "ok": True,
        "artifact_source": "postgres:user_pipeline_artifacts",
        "owner_user_id": owner,
        "run_id": run_id,
        "run": run,
        "artifacts": artifacts,
        "best_rows": _csv_rows_from_text(best_text),
        "shortlist_rows": _csv_rows_from_text(shortlist_text),
        "queue_rows": _csv_rows_from_text(queue_text),
        "manifest_rows": _csv_rows_from_text(manifest_text),
        "current_run_job_corpus_text": corpus_text,
        "job_corpus_rows": _jsonl_row_count_from_text(corpus_text),
    }


def health_payload() -> Dict[str, Any]:
    from src.rag.retriever import get_semantic_status

    semantic_status = get_semantic_status()

    return {
        "ok": True,
        "service": "job-operator-api",
        "semantic_retrieval": semantic_status,
        "rag_answer_ready": bool(semantic_status.get("ready", False)),
    }


def user_workspace_state_payload(owner_user_id: str = "") -> Dict[str, Any]:
    safe_owner = _clean_text(owner_user_id)
    resume_count = 0
    saved_scan_count = 0
    application_action_count = 0
    operator_decision_count = 0

    if safe_owner:
        try:
            resume_payload = get_profile_resumes_postgres_payload(
                owner_user_id=safe_owner,
                database_url="",
                database_url_env="DATABASE_URL",
                psql_bin="psql",
                print_only=False,
            )
            resume_count = int(resume_payload.get("count", 0) or 0)
        except Exception:
            resume_count = 0

        try:
            scans_payload = get_saved_scans_postgres_payload(
                limit=1,
                database_url="",
                database_url_env="DATABASE_URL",
                psql_bin="psql",
                print_only=False,
                owner_user_id=safe_owner,
            )
            saved_scan_count = int(scans_payload.get("count", 0) or 0)
        except Exception:
            saved_scan_count = 0

        try:
            action_payload = get_latest_application_actions_rows(
                database_url="",
                database_url_env="DATABASE_URL",
                psql_bin="psql",
                print_only=False,
                owner_user_id=safe_owner,
            )
            application_action_count = int(action_payload.get("count", 0) or 0)
        except Exception:
            application_action_count = 0

        try:
            decision_payload = get_operator_decisions_postgres_status_payload(
                limit=1,
                database_url="",
                database_url_env="DATABASE_URL",
                psql_bin="psql",
                print_only=False,
                owner_user_id=safe_owner,
            )
            operator_decision_count = int(
                dict(decision_payload.get("postgres", {}) or {}).get("latest_state_count", 0) or 0
            )
        except Exception:
            operator_decision_count = 0

    total_owned_items = (
        resume_count
        + saved_scan_count
        + application_action_count
        + operator_decision_count
    )

    return {
        "ok": True,
        "pipeline_gate": user_pipeline_gate_payload(owner_user_id=owner_user_id),
        "owner_user_id": safe_owner,
        "has_owned_data": total_owned_items > 0,
        "counts": {
            "resumes": resume_count,
            "saved_scans": saved_scan_count,
            "application_actions": application_action_count,
            "operator_decisions": operator_decision_count,
        },
    }


def status_payload(
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    job_corpus: Path = DEFAULT_CORPUS_PATH,
    top_k: int = 10,
    owner_user_id: str = "",
) -> Dict[str, Any]:
    ja = _job_app()
    artifact_context = _latest_user_pipeline_artifact_context(owner_user_id=owner_user_id)

    if artifact_context:
        best_rows = list(artifact_context.get("best_rows", []) or [])
        shortlist_rows = list(artifact_context.get("shortlist_rows", []) or [])
        queue_rows = list(artifact_context.get("queue_rows", []) or [])
        manifest_rows = list(artifact_context.get("manifest_rows", []) or [])
        merged_rows = _build_job_index_from_planning_rows(
            ja,
            best_rows=best_rows,
            queue_rows=queue_rows,
            manifest_rows=manifest_rows,
        )
        job_corpus_rows = int(artifact_context.get("job_corpus_rows", 0) or 0)
        planning_output_dir_value = (
            f"postgres:user_pipeline_artifacts/"
            f"{artifact_context.get('owner_user_id', '')}/"
            f"{artifact_context.get('run_id', '')}"
        )
    else:
        best_rows = ja._load_csv_rows(output_dir / "best_resume_variant_by_job.csv")
        shortlist_rows = ja._load_csv_rows(output_dir / "application_shortlist_by_job.csv")
        queue_rows = ja._load_csv_rows(output_dir / "application_execution_queue.csv")
        manifest_rows = ja._load_csv_rows(output_dir / "job_packet_manifest.csv")
        merged_rows = ja._build_job_index(output_dir)
        job_corpus_rows = ja._count_jsonl_rows(job_corpus)
        planning_output_dir_value = str(output_dir)

    decision_rows = _load_latest_operator_decision_rows(owner_user_id=owner_user_id)
    undecided_review_counts = ja._count_undecided_review_rows(merged_rows)

    winner_bucket_counts = Counter(
        str(row.get("winner_bucket", "") or "<empty>")
        for row in best_rows
    )
    fallback_status_counts = Counter(
        str(row.get("llm_fallback_status", "") or "<empty>")
        for row in best_rows
    )
    action_counts = Counter(
        str(row.get("action", "") or "<empty>")
        for row in queue_rows
    )
    llm_tailoring_counts = Counter(
        str(row.get("llm_tailoring_status", "") or "<empty>")
        for row in manifest_rows
    )
    decision_counts = Counter(
        str(row.get("decision", "") or "<empty>")
        for row in decision_rows
    )

    latest_by_key = ja._load_latest_decision_overlay()
    application_overlay_by_key = _load_latest_application_action_overlay(owner_user_id=owner_user_id)
    job_metadata_by_key = (
        _job_metadata_overlay_from_jsonl_text(artifact_context.get("current_run_job_corpus_text", ""))
        if artifact_context
        else _load_job_metadata_overlay_from_corpus(job_corpus)
    )

    top_rows = sorted(
        queue_rows,
        key=lambda row: (
            int(str(row.get("queue_rank", "999999") or "999999")),
            -ja._parse_float(row.get("winner_score", "0")),
        ),
    )[:top_k]

    top_queue = []
    for row in top_rows:
        overlay_row = dict(row)
        overlay_row["posted_at"] = _clean_text(overlay_row.get("posted_at"))
        overlay_row["job_url"] = _clean_text(overlay_row.get("job_url")) or _clean_text(overlay_row.get("job_doc_id"))
        for field in ja.OPERATOR_DECISION_OVERLAY_FIELDS:
            overlay_row.setdefault(field, "")

        key_candidates = [
            ja._decision_row_key(row),
            f"queue_rank::{str(row.get('queue_rank', '') or '').strip()}",
            (
                f"title::{ja._normalize_text(row.get('job_company', ''))}"
                f"||{ja._normalize_text(row.get('job_title', ''))}"
            ),
        ]

        for key in key_candidates:
            if key and key in latest_by_key:
                overlay_row.update(latest_by_key[key])
                break
        
        for key in _application_row_key_candidates(overlay_row):
            metadata = job_metadata_by_key.get(key)
            if not metadata:
                continue

            if metadata.get("posted_at"):
                overlay_row["posted_at"] = metadata["posted_at"]
            if metadata.get("job_url") and not _clean_text(overlay_row.get("job_url")):
                overlay_row["job_url"] = metadata["job_url"]
            break
        
        for field in APPLICATION_ACTION_OVERLAY_FIELDS:
            if field == "application_label":
                overlay_row.setdefault(field, "Apply")
            elif field == "is_applied":
                overlay_row.setdefault(field, False)
            else:
                overlay_row.setdefault(field, "")

        for key in _application_row_key_candidates(overlay_row):
            if key in application_overlay_by_key:
                overlay_row.update(application_overlay_by_key[key])
                break
        
        top_queue.append(overlay_row)

    return {
        "summary": {
            "job_corpus_path": str(job_corpus),
            "job_corpus_rows": job_corpus_rows,
            "planning_output_dir": planning_output_dir_value,
            "artifact_source": artifact_context.get("artifact_source", "filesystem") if artifact_context else "filesystem",
            "pipeline_run_id": artifact_context.get("run_id", "") if artifact_context else "",
            "best_variant_rows": len(best_rows),
            "shortlist_rows": len(shortlist_rows),
            "execution_queue_rows": len(queue_rows),
            "packet_manifest_rows": len(manifest_rows),
            "operator_decisions_storage": "postgres",
            "operator_decisions_rows": len(decision_rows),
        },
        "winner_bucket_counts": dict(sorted(winner_bucket_counts.items())),
        "llm_fallback_status_counts": dict(sorted(fallback_status_counts.items())),
        "queue_action_counts": dict(sorted(action_counts.items())),
        "operator_decision_counts": dict(sorted(decision_counts.items())),
        "undecided_review_counts": dict(sorted(undecided_review_counts.items())),
        "llm_tailoring_status_counts": dict(sorted(llm_tailoring_counts.items())),
        "top_queue_rows": top_queue,
    }


def browse_payload(
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    owner_user_id: str = "",
    **filters: Any,
) -> Dict[str, Any]:
    ja = _job_app()

    resolved_filters = {
        "action": "",
        "needs_review": "",
        "is_tie": "",
        "fallback_status": "",
        "winner_bucket": "",
        "tailoring_state": "",
        "company_contains": "",
        "title_contains": "",
        "limit": 15,
        "undecided_only": "",
        "page": 1,
    }
    resolved_filters.update(filters)

    requested_limit = max(int(resolved_filters.get("limit", 15) or 15), 1)
    current_page = max(int(resolved_filters.get("page", 1) or 1), 1)
    page_size = 15

    requested_tailoring_states = _normalize_tailoring_state_filter_values(
        resolved_filters.get("tailoring_state", [])
    )

    try:
        artifact_context = _latest_user_pipeline_artifact_context(owner_user_id=owner_user_id)
        if artifact_context:
            rows = _build_job_index_from_planning_rows(
                ja,
                best_rows=list(artifact_context.get("best_rows", []) or []),
                queue_rows=list(artifact_context.get("queue_rows", []) or []),
                manifest_rows=list(artifact_context.get("manifest_rows", []) or []),
            )
            job_metadata_by_key = _job_metadata_overlay_from_jsonl_text(
                artifact_context.get("current_run_job_corpus_text", "")
            )
        else:
            rows = ja._build_job_index(output_dir)
            job_metadata_by_key = {}

        selection_filters = dict(resolved_filters)
        selection_filters["limit"] = max(len(rows), 1)
        selection_filters.pop("page", None)
        selection_filters.pop("tailoring_state", None)

        args = _make_args(**selection_filters)
        selected = ja._select_browse_rows(rows, args)

        selected = _overlay_application_actions(selected, owner_user_id=owner_user_id)

        selected = _exclude_applied_rows(selected)

        if requested_tailoring_states:
            enriched_selected: List[Dict[str, Any]] = []
            for row in selected:
                matches, enriched_row = _row_matches_tailoring_state_filter(
                    row,
                    requested_tailoring_states,
                    output_dir=output_dir,
                )
                if matches:
                    enriched_selected.append(enriched_row)
            selected = enriched_selected

        selected = selected[:requested_limit]

        total_count = len(selected)
        total_pages = max((total_count + page_size - 1) // page_size, 1)
        current_page = min(current_page, total_pages)

        start = (current_page - 1) * page_size
        end = start + page_size
        page_rows = selected[start:end]

        page_rows = (
            _overlay_job_metadata_from_map(page_rows, job_metadata_by_key)
            if artifact_context
            else _overlay_job_metadata(page_rows, job_corpus=DEFAULT_CORPUS_PATH)
        )

        finalized_page_rows: List[Dict[str, Any]] = []
        for row in page_rows:
            if requested_tailoring_states:
                finalized_page_rows.append(row)
                continue

            _, enriched_row = _row_matches_tailoring_state_filter(
                row,
                [],
                output_dir=output_dir,
            )
            finalized_page_rows.append(enriched_row)

        payload = {
            "filters": {
                "action": resolved_filters.get("action", ""),
                "needs_review": resolved_filters.get("needs_review", ""),
                "is_tie": resolved_filters.get("is_tie", ""),
                "fallback_status": resolved_filters.get("fallback_status", ""),
                "winner_bucket": resolved_filters.get("winner_bucket", ""),
                "tailoring_state": requested_tailoring_states,
                "company_contains": resolved_filters.get("company_contains", ""),
                "title_contains": resolved_filters.get("title_contains", ""),
                "limit": requested_limit,
                "undecided_only": resolved_filters.get("undecided_only", ""),
                "page": current_page,
            },
            "rows": finalized_page_rows,
            "count": len(finalized_page_rows),
            "total_count": total_count,
            "page": current_page,
            "page_size": page_size,
            "total_pages": total_pages,
            "has_prev_page": current_page > 1,
            "has_next_page": current_page < total_pages,
        }

        return payload

    except Exception:
        raise

def review_payload(
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    owner_user_id: str = "",
    **filters: Any,
) -> Dict[str, Any]:
    ja = _job_app()
    artifact_context = _latest_user_pipeline_artifact_context(owner_user_id=owner_user_id)
    if artifact_context:
        rows = _build_job_index_from_planning_rows(
            ja,
            best_rows=list(artifact_context.get("best_rows", []) or []),
            queue_rows=list(artifact_context.get("queue_rows", []) or []),
            manifest_rows=list(artifact_context.get("manifest_rows", []) or []),
        )
        job_metadata_by_key = _job_metadata_overlay_from_jsonl_text(
            artifact_context.get("current_run_job_corpus_text", "")
        )
    else:
        rows = ja._build_job_index(output_dir)
        job_metadata_by_key = {}

    resolved_filters = {
        "queue_rank": None,
        "job_doc_id": "",
        "company_contains": "",
        "title_contains": "",
        "include_non_review": False,
        "limit": 5,
        "action": "",
        "undecided_only": "",
    }
    resolved_filters.update(filters)

    args = _make_args(**resolved_filters)
    selected = ja._select_review_rows(rows, args)
    selected = (
        _overlay_job_metadata_from_map(selected, job_metadata_by_key)
        if artifact_context
        else _overlay_job_metadata(selected, job_corpus=DEFAULT_CORPUS_PATH)
    )
    selected = _overlay_application_actions(selected, owner_user_id=owner_user_id)
    selected = _exclude_applied_rows(selected)

    return {
        "filters": resolved_filters,
        "rows": selected,
        "count": len(selected),
    }


def workflow_payload(
    view: str,
    limit: int = 20,
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    owner_user_id: str = "",
) -> Dict[str, Any]:
    ja = _job_app()
    artifact_context = _latest_user_pipeline_artifact_context(owner_user_id=owner_user_id)
    if artifact_context:
        rows = _build_job_index_from_planning_rows(
            ja,
            best_rows=list(artifact_context.get("best_rows", []) or []),
            queue_rows=list(artifact_context.get("queue_rows", []) or []),
            manifest_rows=list(artifact_context.get("manifest_rows", []) or []),
        )
        job_metadata_by_key = _job_metadata_overlay_from_jsonl_text(
            artifact_context.get("current_run_job_corpus_text", "")
        )
    else:
        rows = ja._build_job_index(output_dir)
        job_metadata_by_key = {}

    selected = ja._workflow_view_rows(rows, view)[:limit]
    selected = (
        _overlay_job_metadata_from_map(selected, job_metadata_by_key)
        if artifact_context
        else _overlay_job_metadata(selected, job_corpus=DEFAULT_CORPUS_PATH)
    )
    selected = _overlay_application_actions(selected, owner_user_id=owner_user_id)
    selected = _exclude_applied_rows(selected)
    return {
        "view": view,
        "rows": selected,
        "count": len(selected),
    }


def planner_payload(
    request: str,
    limit: int = 20,
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    owner_user_id: str = "",
) -> Dict[str, Any]:
    ja = _job_app()
    view = ja._infer_planner_view(request)
    artifact_context = _latest_user_pipeline_artifact_context(owner_user_id=owner_user_id)
    if artifact_context:
        rows = _build_job_index_from_planning_rows(
            ja,
            best_rows=list(artifact_context.get("best_rows", []) or []),
            queue_rows=list(artifact_context.get("queue_rows", []) or []),
            manifest_rows=list(artifact_context.get("manifest_rows", []) or []),
        )
        job_metadata_by_key = _job_metadata_overlay_from_jsonl_text(
            artifact_context.get("current_run_job_corpus_text", "")
        )
    else:
        rows = ja._build_job_index(output_dir)
        job_metadata_by_key = {}

    selected = ja._workflow_view_rows(rows, view)[:limit]
    selected = (
        _overlay_job_metadata_from_map(selected, job_metadata_by_key)
        if artifact_context
        else _overlay_job_metadata(selected, job_corpus=DEFAULT_CORPUS_PATH)
    )
    selected = _overlay_application_actions(selected, owner_user_id=owner_user_id)
    return {
        "request": request,
        "resolved_view": view,
        "rows": selected,
        "count": len(selected),
    }

def _resolve_planning_artifact_path(
    path: str,
    output_dir: Path = DEFAULT_OUTPUT_DIR,
) -> Path:
    raw = str(path or "").strip()
    if not raw:
        raise ValueError("Artifact path is required.")

    resolved = Path(raw).expanduser().resolve()
    output_root = Path(output_dir).expanduser().resolve()

    try:
        resolved.relative_to(output_root)
    except ValueError as exc:
        raise ValueError("Artifact path must stay inside the planning output directory.") from exc

    if not resolved.exists() or not resolved.is_file():
        raise ValueError(f"Artifact not found: {raw}")

    if resolved.suffix.lower() not in {".md", ".json"}:
        raise ValueError(f"Unsupported artifact type: {resolved.suffix}")

    return resolved


def planning_artifact_payload(
    path: str,
    output_dir: Path = DEFAULT_OUTPUT_DIR,
) -> Dict[str, Any]:
    artifact_path = _resolve_planning_artifact_path(path, output_dir=output_dir)
    suffix = artifact_path.suffix.lower()
    text = artifact_path.read_text(encoding="utf-8")

    payload: Dict[str, Any] = {
        "ok": True,
        "path": str(artifact_path),
        "kind": "json" if suffix == ".json" else "text",
    }

    if suffix == ".json":
        data = json.loads(text)
        if isinstance(data, dict):
            if artifact_path.name.endswith("__tailoring.json"):
                data = _rehydrate_legacy_tailoring_operator_payload(
                    artifact_path,
                    data,
                )

            if data.get("replacement_candidates") is not None:
                data = _apply_saved_patch_selection_overlay(
                    artifact_path,
                    data,
                )
        payload["data"] = data
    else:
        payload["text"] = text

    return payload

def tailoring_scan_preload_payload(
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    *,
    tailoring_json_path: str = "",
    selected_resume: str = "",
    owner_user_id: str = "",
) -> Dict[str, Any]:
    from src.tailoring.llm import tailoring_llm_model_config_payload

    artifact_path = _resolve_planning_artifact_path(
        tailoring_json_path,
        output_dir=output_dir,
    )

    if artifact_path.suffix.lower() != ".json":
        raise ValueError("Scan preload requires a tailoring JSON artifact.")

    artifact_response = planning_artifact_payload(
        str(artifact_path),
        output_dir=output_dir,
    )
    payload_data = dict(artifact_response.get("data", {}) or {})

    if not payload_data:
        raise ValueError("Scan preload could not load tailoring artifact payload.")

    draft_response = load_tailoring_workspace_draft_payload(
        output_dir=output_dir,
        tailoring_json_path=str(artifact_path),
        selected_resume=selected_resume,
    )
    draft = dict(draft_response.get("draft", {}) or {})

    preview_response = preview_tailoring_workspace_draft_payload(
        output_dir=output_dir,
        tailoring_json_path=str(artifact_path),
        selected_resume=selected_resume,
        owner_user_id=owner_user_id,
    )

    selection = dict(payload_data.get("selection", {}) or {})
    job = dict(payload_data.get("job", {}) or {})
    job_snapshot = dict(payload_data.get("job_snapshot", {}) or {})

    effective_selected_resume = (
        _sanitize_optional_resume_filename(selected_resume)
        or _sanitize_optional_resume_filename(draft.get("selected_resume"))
        or _sanitize_optional_resume_filename(selection.get("selected_resume"))
    )
    if not effective_selected_resume:
        raise ValueError("Scan preload requires a selected resume.")

    packet_path = _infer_packet_json_path_from_tailoring_artifact(artifact_path)

    trusted_ready = list(payload_data.get("app_ready_replacements", []) or [])
    trusted_optional = list(payload_data.get("direct_apply_optional_replacements", []) or [])
    ai_optimize_optional = list(payload_data.get("ai_optimize_optional_replacements", []) or [])
    directional_guidance = list(payload_data.get("direction_only_replacements", []) or [])

    final_replacement_summary = dict(payload_data.get("final_replacement_summary", {}) or {})
    rewrite_review_summary = dict(payload_data.get("rewrite_review_summary", {}) or {})
    rewrite_review_groups = list(payload_data.get("rewrite_review_groups", []) or [])

    selected_jd_record = dict(job_snapshot or job)

    try:
        scan_resume_evidence = _load_resume_evidence_for_workspace_preview(
            effective_selected_resume,
            owner_user_id=owner_user_id,
        )
    except Exception:
        scan_resume_evidence = None
    extracted_personal_details = _extract_resume_personal_details(
        scan_resume_evidence,
        owner_user_id=owner_user_id,
    )
    saved_personal_details = _normalize_workspace_personal_details(
        draft.get("personal_details", {})
    )
    has_saved_personal_details = any(saved_personal_details.values())
    current_personal_details = (
        saved_personal_details if has_saved_personal_details else extracted_personal_details
    )

    scan_issue_contract = _build_tailoring_scan_issue_contract(
        trusted_ready=trusted_ready,
        trusted_optional=trusted_optional,
        ai_optimize_optional=ai_optimize_optional,
        directional_guidance=directional_guidance,
        resume_evidence=scan_resume_evidence,
        tailoring_summary=dict(payload_data.get("summary", {}) or {}),
        jd_record=selected_jd_record,
    )
    scan_score_snapshot = _build_tailoring_scan_score_snapshot(
        selection=selection,
        preview_response=preview_response,
        scan_issue_contract=scan_issue_contract,
    )
    scan_session_snapshot = _build_tailoring_scan_session_snapshot(
        artifact_path=artifact_path,
        draft=draft,
        selected_resume=effective_selected_resume,
        scan_issue_contract=scan_issue_contract,
        preview_response=preview_response,
    )

    return {
        "ok": True,
        "preload_mode": "tailoring_artifact",
        "scan_entry_source": "tailoring_suggestions",
        "artifact_references": {
            "tailoring_json_path": str(artifact_path),
            "packet_json_path": str(packet_path) if packet_path and packet_path.exists() else "",
        },
        "selected_resume": effective_selected_resume,
        "selected_jd_record": selected_jd_record,
        "job": job,
        "job_snapshot": job_snapshot,
        "selection": selection,
        "score_snapshot": {
            "artifact_selected_score": selection.get("selected_score", None),
            "draft_preview_status": _clean_text(preview_response.get("preview_status")),
            "draft_preview_note": _clean_text(preview_response.get("preview_note")),
            "original_score": preview_response.get("original_score"),
            "projected_score": preview_response.get("projected_score"),
            "projected_delta": preview_response.get("projected_delta"),
            "selected_patch_set_counterfactual_preview": preview_response.get(
                "selected_patch_set_counterfactual_preview"
            ),
        },
        "scan_score": scan_score_snapshot,
        "scan_session": scan_session_snapshot,
        "score_preview": preview_response.get("score_preview", {}),
        "trusted_suggestions": {
            "direct_apply_ready": trusted_ready,
            "direct_apply_optional": trusted_optional,
        },
        "ai_optimize_suggestions": ai_optimize_optional,
        "directional_guidance": directional_guidance,
        "lane_counts": {
            "direct_apply_ready": len(trusted_ready),
            "direct_apply_optional": len(trusted_optional),
            "ai_optimize_optional": len(ai_optimize_optional),
            "direction_only": len(directional_guidance),
        },
        "scan_issue_contract": scan_issue_contract,
        "llm_model_config": tailoring_llm_model_config_payload(),
        "final_replacement_summary": final_replacement_summary,
        "rewrite_review_summary": rewrite_review_summary,
        "rewrite_review_groups": rewrite_review_groups,
        "personal_details": {
            "extracted": extracted_personal_details,
            "saved": saved_personal_details,
            "current": current_personal_details,
        },
        "draft_status": _clean_text(draft_response.get("draft_status")),
        "has_saved_draft": bool(draft_response.get("has_saved_draft", False)),
        "draft": draft,
    }

def _operator_decision_latest_sort_key(row: Dict[str, Any]) -> Tuple[str, str]:
    normalized = operator_decision_db_row(dict(row))
    return (
        str(normalized.get("decision_timestamp", "") or ""),
        str(normalized.get("decision_id", "") or ""),
    )

def _load_latest_operator_decision_rows(owner_user_id: str = "") -> List[Dict[str, Any]]:
    meta_payload = get_operator_decisions_postgres_status_payload(
        limit=1,
        database_url="",
        database_url_env="DATABASE_URL",
        psql_bin="psql",
        print_only=False,
        owner_user_id=_clean_text(owner_user_id),
    )
    meta_block = dict(meta_payload.get("postgres", {}) or {})
    query_limit = max(int(meta_block.get("latest_state_count", 0) or 0), 1)

    postgres_payload = get_operator_decisions_postgres_status_payload(
        limit=query_limit,
        database_url="",
        database_url_env="DATABASE_URL",
        psql_bin="psql",
        print_only=False,
        owner_user_id=_clean_text(owner_user_id),
    )
    postgres_block = dict(postgres_payload.get("postgres", {}) or {})
    postgres_rows = list(postgres_block.get("latest_rows", []) or [])

    normalized_rows: List[Dict[str, Any]] = []
    for row in postgres_rows:
        normalized = operator_decision_db_row({
            "owner_user_id": row.get("owner_user_id", ""),
            "decision_timestamp": row.get("decision_timestamp", ""),
            "queue_rank": row.get("queue_rank", ""),
            "job_doc_id": row.get("job_doc_id", ""),
            "job_company": row.get("job_company", ""),
            "job_title": row.get("job_title", ""),
            "planning_action": row.get("planning_action", ""),
            "winner_resume": row.get("winner_resume", ""),
            "winner_score": row.get("winner_score", ""),
            "runner_up_resume": row.get("runner_up_resume", ""),
            "runner_up_score": row.get("runner_up_score", ""),
            "selected_resume": row.get("selected_resume", ""),
            "decision": row.get("decision", ""),
            "note": row.get("note", ""),
        })
        normalized_rows.append(normalized)

    normalized_rows.sort(
        key=lambda row: _operator_decision_latest_sort_key(row),
        reverse=True,
    )
    return normalized_rows


def decisions_payload(
    queue_rank: int | None = None,
    decision: Any = "",
    selected_resume: str = "",
    company_contains: str = "",
    title_contains: str = "",
    limit: int = 15,
    page: int = 1,
    owner_user_id: str = "",
) -> Dict[str, Any]:
    ja = _job_app()
    rows = _load_latest_operator_decision_rows(owner_user_id=owner_user_id)

    resolved_filters = {
        "queue_rank": queue_rank,
        "decision": decision,
        "selected_resume": selected_resume,
        "company_contains": company_contains,
        "title_contains": title_contains,
        "limit": limit,
        "page": page,
    }

    requested_limit = max(int(limit or 15), 1)
    current_page = max(int(page or 1), 1)
    page_size = 15

    selection_filters = dict(resolved_filters)
    selection_filters["limit"] = max(len(rows), 1)
    selection_filters.pop("page", None)

    args = _make_args(**selection_filters)
    selected = ja._select_decision_rows(rows, args)
    selected = _overlay_job_metadata(selected, job_corpus=DEFAULT_CORPUS_PATH)
    selected = _overlay_application_actions(selected, owner_user_id=owner_user_id)

    selected = selected[:requested_limit]

    total_count = len(selected)
    total_pages = max((total_count + page_size - 1) // page_size, 1)
    current_page = min(current_page, total_pages)

    start_idx = (current_page - 1) * page_size
    end_idx = start_idx + page_size
    page_rows = selected[start_idx:end_idx]

    return {
        "filters": {
            "queue_rank": queue_rank,
            "decision": decision,
            "selected_resume": selected_resume,
            "company_contains": company_contains,
            "title_contains": title_contains,
            "limit": requested_limit,
            "page": current_page,
        },
        "rows": page_rows,
        "count": len(page_rows),
        "total_count": total_count,
        "page": current_page,
        "page_size": page_size,
        "total_pages": total_pages,
        "has_prev_page": current_page > 1,
        "has_next_page": current_page < total_pages,
    }

def _dual_write_operator_decision_postgres(row: Dict[str, Any]) -> Dict[str, Any]:
    database_url = str(os.environ.get("DATABASE_URL", "") or "").strip()
    if not database_url:
        return {
            "attempted": False,
            "ok": False,
            "skipped": "missing_database_url",
            "table_name": "operator_decisions",
        }

    try:
        payload = insert_operator_decision_row_to_postgres(
            record=row,
            database_url="",
            database_url_env="DATABASE_URL",
            psql_bin="psql",
            print_only=False,
            allow_contract_drift=False,
        )
        return {
            "attempted": True,
            "ok": True,
            "table_name": payload.get("table_name", "operator_decisions"),
            "decision_id": str(payload.get("row", {}).get("decision_id", "") or ""),
            "decision_key": str(payload.get("row", {}).get("decision_key", "") or ""),
            "contract_health_ok": bool(payload.get("contract_health_ok", False)),
            "command_text": str(payload.get("command_text", "") or ""),
        }
    except SystemExit as exc:
        return {
            "attempted": True,
            "ok": False,
            "table_name": "operator_decisions",
            "error_type": "SystemExit",
            "error": str(exc),
        }
    except Exception as exc:
        return {
            "attempted": True,
            "ok": False,
            "table_name": "operator_decisions",
            "error_type": exc.__class__.__name__,
            "error": str(exc),
        }
    
def record_operator_resume_selection_payload(
    *,
    queue_rank: str = "",
    job_doc_id: str = "",
    job_company: str = "",
    job_title: str = "",
    planning_action: str = "",
    decision: str = "SELECT_RESUME",
    selected_resume: str = "",
    winner_resume: str = "",
    winner_score: str = "",
    runner_up_resume: str = "",
    runner_up_score: str = "",
    note: str = "",
    owner_user_id: str = "",
) -> Dict[str, Any]:
    ja = _job_app()

    safe_selected = _sanitize_resume_filename(selected_resume)
    safe_winner = _sanitize_optional_resume_filename(winner_resume)
    safe_runner = _sanitize_optional_resume_filename(runner_up_resume)

    allowed_resumes = {name for name in [safe_winner, safe_runner] if name}
    if not allowed_resumes:
        raise ValueError("No eligible resume choices were provided.")

    if safe_selected not in allowed_resumes:
        allowed = ", ".join(sorted(allowed_resumes))
        raise ValueError(
            f"selected_resume must be one of the eligible choices. Allowed: {allowed}"
        )

    row = {
        "owner_user_id": _clean_text(owner_user_id),
        "decision_timestamp": datetime.now(timezone.utc).isoformat(timespec="microseconds"),
        "queue_rank": _clean_text(queue_rank),
        "job_doc_id": _clean_text(job_doc_id),
        "job_company": _clean_text(job_company),
        "job_title": _clean_text(job_title),
        "planning_action": _clean_text(planning_action),
        "winner_resume": safe_winner,
        "winner_score": _clean_text(winner_score),
        "runner_up_resume": safe_runner,
        "runner_up_score": _clean_text(runner_up_score),
        "selected_resume": safe_selected,
        "decision": _normalize_operator_decision(decision),
        "note": _clean_text(note),
    }

    _validate_operator_decision_identity(row)
    postgres_write = _dual_write_operator_decision_postgres(row)

    return {
        "ok": True,
        "row": row,
        "csv_write_disabled": True,
        "postgres_write": postgres_write,
    }

def preview_planning_patch_selection_payload(
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    *,
    tailoring_json_path: str = "",
    selected_candidate_ids: Any = None,
) -> Dict[str, Any]:
    from src.tailoring.rendering import build_selected_patch_set_counterfactual_preview

    artifact_path = _resolve_planning_artifact_path(
        tailoring_json_path,
        output_dir=output_dir,
    )

    if artifact_path.suffix.lower() != ".json":
        raise ValueError("Patch selection preview requires a tailoring JSON artifact.")

    payload_data = _load_tailoring_json_artifact(artifact_path)

    replacement_candidates = list(payload_data.get("replacement_candidates", []) or [])
    if not replacement_candidates:
        raise ValueError("Tailoring artifact does not contain replacement candidates.")

    normalized_ids = _normalize_selected_patch_candidate_ids(selected_candidate_ids)
    if not normalized_ids:
        normalized_ids = _default_selected_candidate_ids_from_replacement_plan(payload_data)
    valid_candidate_ids = set(_tailoring_artifact_candidate_ids(payload_data))
    unknown_candidate_ids = [candidate_id for candidate_id in normalized_ids if candidate_id not in valid_candidate_ids]
    if unknown_candidate_ids:
        raise ValueError(
            f"Unknown candidate IDs for this artifact: {', '.join(sorted(unknown_candidate_ids))}"
        )

    preview = build_selected_patch_set_counterfactual_preview(
        payload_data,
        selected_candidate_ids=normalized_ids,
    )

    return {
        "ok": True,
        "path": str(artifact_path),
        "selected_patch_candidate_ids": normalized_ids,
        "selected_patch_set_counterfactual_preview": preview,
    }

def _dual_write_patch_selection_postgres(row: Dict[str, Any]) -> Dict[str, Any]:
    database_url = str(os.environ.get("DATABASE_URL", "") or "").strip()
    if not database_url:
        return {
            "attempted": False,
            "ok": False,
            "skipped": "missing_database_url",
            "table_name": "patch_selections",
        }

    try:
        payload = insert_patch_selection_row_to_postgres(
            record=row,
            database_url="",
            database_url_env="DATABASE_URL",
            psql_bin="psql",
            print_only=False,
            allow_contract_drift=False,
        )
        return {
            "attempted": True,
            "ok": True,
            "table_name": payload.get("table_name", "patch_selections"),
            "selection_id": str(payload.get("row", {}).get("selection_id", "") or ""),
            "contract_health_ok": bool(payload.get("contract_health_ok", False)),
            "command_text": str(payload.get("command_text", "") or ""),
        }
    except SystemExit as exc:
        return {
            "attempted": True,
            "ok": False,
            "table_name": "patch_selections",
            "error_type": "SystemExit",
            "error": str(exc),
        }
    except Exception as exc:
        return {
            "attempted": True,
            "ok": False,
            "table_name": "patch_selections",
            "error_type": exc.__class__.__name__,
            "error": str(exc),
        }
    
def record_planning_patch_selection_payload(
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    *,
    tailoring_json_path: str = "",
    job_doc_id: str = "",
    queue_rank: str = "",
    selected_resume: str = "",
    selected_candidate_ids: Any = None,
    note: str = "",
) -> Dict[str, Any]:
    from src.tailoring.rendering import build_selected_patch_set_counterfactual_preview

    artifact_path = _resolve_planning_artifact_path(
        tailoring_json_path,
        output_dir=output_dir,
    )

    if artifact_path.suffix.lower() != ".json":
        raise ValueError("Patch selection requires a tailoring JSON artifact.")

    payload_data = _load_tailoring_json_artifact(artifact_path)

    replacement_candidates = list(payload_data.get("replacement_candidates", []) or [])
    if not replacement_candidates:
        raise ValueError("Tailoring artifact does not contain replacement candidates.")

    normalized_ids = _normalize_selected_patch_candidate_ids(selected_candidate_ids)
    if not normalized_ids:
        normalized_ids = _default_selected_candidate_ids_from_replacement_plan(payload_data)

    valid_candidate_ids = set(_tailoring_artifact_candidate_ids(payload_data))
    unknown_candidate_ids = [candidate_id for candidate_id in normalized_ids if candidate_id not in valid_candidate_ids]
    if unknown_candidate_ids:
        raise ValueError(
            f"Unknown candidate IDs for this artifact: {', '.join(sorted(unknown_candidate_ids))}"
        )

    job = payload_data.get("job", {}) or {}
    selection = payload_data.get("selection", {}) or {}

    artifact_selected_resume = _clean_text(selection.get("selected_resume"))
    requested_selected_resume = _sanitize_optional_resume_filename(selected_resume)

    if requested_selected_resume and artifact_selected_resume and requested_selected_resume != artifact_selected_resume:
        raise ValueError(
            f"selected_resume does not match the tailoring artifact resume. "
            f"artifact={artifact_selected_resume!r} request={requested_selected_resume!r}"
        )

    row = {
        "selection_timestamp": datetime.now(timezone.utc).isoformat(timespec="microseconds"),
        "job_doc_id": _clean_text(job_doc_id) or _clean_text(job.get("job_doc_id")),
        "queue_rank": _clean_text(queue_rank),
        "job_company": _clean_text(job.get("company")),
        "job_title": _clean_text(job.get("title")),
        "selected_resume": requested_selected_resume or artifact_selected_resume,
        "tailoring_json_path": str(artifact_path),
        "artifact_signature": _tailoring_artifact_signature(payload_data),
        "selected_candidate_ids_json": _serialize_selected_patch_candidate_ids(normalized_ids),
        "note": _clean_text(note),
    }

    ja = _job_app()
    postgres_write = _dual_write_patch_selection_postgres(row)

    preview = build_selected_patch_set_counterfactual_preview(
        payload_data,
        selected_candidate_ids=normalized_ids,
    )

    if bool(postgres_write.get("ok", False)):
        _invalidate_patch_selection_overlay_cache()

    return {
        "ok": True,
        "csv_write_disabled": True,
        "selected_patch_candidate_ids": normalized_ids,
        "selection": row,
        "postgres_write": postgres_write,
        "selected_patch_set_counterfactual_preview": preview,
    }

def _tailoring_workspace_draft_artifact_path(artifact_path: Path) -> Path:
    name = artifact_path.name
    suffix = "__tailoring.json"

    if name.endswith(suffix):
        prefix = name[:-len(suffix)]
        return artifact_path.with_name(f"{prefix}__tailoring_workspace_draft.json")

    return artifact_path.with_name(f"{artifact_path.stem}__tailoring_workspace_draft.json")


def _normalize_workspace_rewrite_review_decisions(value: Any) -> Dict[str, Dict[str, str]]:
    if isinstance(value, dict):
        raw_items = value
    else:
        raw_text = _clean_text(value)
        if not raw_text:
            return {}

        try:
            parsed = json.loads(raw_text)
        except Exception as exc:
            raise ValueError("rewrite_review_decisions must be a JSON object.") from exc

        if not isinstance(parsed, dict):
            raise ValueError("rewrite_review_decisions must be a JSON object.")

        raw_items = parsed

    normalized: Dict[str, Dict[str, str]] = {}

    for raw_key, raw_value in raw_items.items():
        candidate_id = _clean_text(raw_key)
        if not candidate_id:
            continue

        if isinstance(raw_value, dict):
            state = _clean_text(raw_value.get("state")).lower() or "pending"
            note = _clean_text(raw_value.get("note"))
        else:
            state = _clean_text(raw_value).lower() or "pending"
            note = ""

        if state not in _ALLOWED_REWRITE_REVIEW_STATES:
            allowed = ", ".join(sorted(_ALLOWED_REWRITE_REVIEW_STATES))
            raise ValueError(
                f"Invalid rewrite review state={state!r} for {candidate_id!r}. Allowed: {allowed}"
            )

        normalized[candidate_id] = {
            "state": state,
            "note": note,
        }

    return normalized

def _normalize_workspace_manual_bullet_edits(value: Any) -> Dict[str, str]:
    if isinstance(value, dict):
        raw_items = value
    else:
        raw_text = _clean_text(value)
        if not raw_text:
            return {}

        try:
            parsed = json.loads(raw_text)
        except Exception as exc:
            raise ValueError("manual_bullet_edits must be a JSON object.") from exc

        if not isinstance(parsed, dict):
            raise ValueError("manual_bullet_edits must be a JSON object.")

        raw_items = parsed

    normalized: Dict[str, str] = {}
    for raw_key, raw_value in raw_items.items():
        key = _clean_text(raw_key)
        if not key:
            continue
        normalized[key] = str(raw_value or "")

    return normalized


def _normalize_workspace_excluded_scan_issue_ids(value: Any) -> List[str]:
    if isinstance(value, list):
        raw_items = value
    else:
        raw_text = _clean_text(value)
        if not raw_text:
            return []

        try:
            parsed = json.loads(raw_text)
        except Exception as exc:
            raise ValueError("excluded_scan_issue_ids must be a JSON list.") from exc

        if not isinstance(parsed, list):
            raise ValueError("excluded_scan_issue_ids must be a JSON list.")

        raw_items = parsed

    output: List[str] = []
    seen = set()
    for value in raw_items:
        issue_id = _clean_text(value)
        if not issue_id or issue_id in seen:
            continue
        seen.add(issue_id)
        output.append(issue_id)
    return output


def _build_tailoring_workspace_default_draft_payload(
    artifact_path: Path,
    payload_data: Dict[str, Any],
    *,
    selected_resume: str = "",
) -> Dict[str, Any]:
    job = payload_data.get("job", {}) or {}
    selection = payload_data.get("selection", {}) or {}

    selected_candidate_ids = _normalize_selected_patch_candidate_ids(
        payload_data.get("selected_patch_candidate_ids", [])
    )
    if not selected_candidate_ids:
        selected_candidate_ids = _default_selected_candidate_ids_from_replacement_plan(
            payload_data
        )

    safe_selected_resume = (
        _sanitize_optional_resume_filename(selected_resume)
        or _sanitize_optional_resume_filename(selection.get("selected_resume"))
    )

    draft_path = _tailoring_workspace_draft_artifact_path(artifact_path)

    return {
        "draft_version": 1,
        "draft_status": "default",
        "saved_at": "",
        "tailoring_json_path": str(artifact_path),
        "draft_json_path": str(draft_path),
        "artifact_signature": _tailoring_artifact_signature(payload_data),
        "job_doc_id": _clean_text(job.get("job_doc_id")),
        "job_company": _clean_text(job.get("company")),
        "job_title": _clean_text(job.get("title")),
        "selected_resume": safe_selected_resume,
        "selected_patch_candidate_ids": selected_candidate_ids,
        "manual_bullet_edits": {},
        "excluded_scan_issue_ids": [],
        "personal_details": {},
        "note": "",
        "rewrite_review_decisions": {},
        "source_selected_patch_selection_status": _clean_text(
            payload_data.get("selected_patch_selection_status")
        ),
        "source_selected_patch_selection_note": _clean_text(
            payload_data.get("selected_patch_selection_note")
        ),
        "rewrite_review_telemetry": _build_workspace_rewrite_review_telemetry(
            payload_data,
            selected_candidate_ids=selected_candidate_ids,
            manual_bullet_edits={},
            rewrite_review_decisions={},
        ),
    }


def load_tailoring_workspace_draft_payload(
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    *,
    tailoring_json_path: str = "",
    selected_resume: str = "",
) -> Dict[str, Any]:
    artifact_path = _resolve_planning_artifact_path(
        tailoring_json_path,
        output_dir=output_dir,
    )

    if artifact_path.suffix.lower() != ".json":
        raise ValueError("Workspace draft loading requires a tailoring JSON artifact.")

    payload_data = _load_tailoring_json_artifact(artifact_path)

    if artifact_path.name.endswith("__tailoring.json"):
        payload_data = _rehydrate_legacy_tailoring_operator_payload(
            artifact_path,
            payload_data,
        )

    if payload_data.get("replacement_candidates") is not None:
        payload_data = _apply_saved_patch_selection_overlay(
            artifact_path,
            payload_data,
        )

    default_draft = _build_tailoring_workspace_default_draft_payload(
        artifact_path,
        payload_data,
        selected_resume=selected_resume,
    )

    draft_path = Path(default_draft["draft_json_path"])
    if not draft_path.exists():
        return {
            "ok": True,
            "draft_status": "default",
            "has_saved_draft": False,
            "draft": default_draft,
        }

    saved_data = _load_tailoring_json_artifact(draft_path)
    saved_signature = _clean_text(saved_data.get("artifact_signature"))
    current_signature = _clean_text(default_draft.get("artifact_signature"))

    if saved_signature and saved_signature != current_signature:
        stale_draft = dict(default_draft)
        stale_draft["draft_status"] = "stale_signature"
        stale_draft["note"] = (
            "Saved workspace draft was ignored because the tailoring artifact changed."
        )
        return {
            "ok": True,
            "draft_status": "stale_signature",
            "has_saved_draft": False,
            "draft": stale_draft,
        }

    valid_candidate_ids = set(_tailoring_artifact_candidate_ids(payload_data))
    saved_selected_ids = _normalize_selected_patch_candidate_ids(
        saved_data.get("selected_patch_candidate_ids", [])
    )
    saved_selected_ids = [
        candidate_id
        for candidate_id in saved_selected_ids
        if candidate_id in valid_candidate_ids
    ]
    saved_manual_edits = _normalize_workspace_manual_bullet_edits(
        saved_data.get("manual_bullet_edits", {})
    )
    saved_review_decisions = _derive_workspace_rewrite_review_decisions(
        payload_data,
        selected_candidate_ids=saved_selected_ids,
        manual_bullet_edits=saved_data.get("manual_bullet_edits", {}),
        rewrite_review_decisions=saved_data.get("rewrite_review_decisions", {}),
    )
    saved_review_telemetry = _build_workspace_rewrite_review_telemetry(
        payload_data,
        selected_candidate_ids=saved_selected_ids,
        manual_bullet_edits=saved_manual_edits,
        rewrite_review_decisions=saved_review_decisions,
    )
    saved_excluded_scan_issue_ids = _normalize_workspace_excluded_scan_issue_ids(
        saved_data.get("excluded_scan_issue_ids", [])
    )
    saved_personal_details = _normalize_workspace_personal_details(
        saved_data.get("personal_details", {})
    )

    merged = dict(default_draft)
    merged.update({
        "draft_status": "saved",
        "saved_at": _clean_text(saved_data.get("saved_at")),
        "selected_resume": (
            _sanitize_optional_resume_filename(saved_data.get("selected_resume"))
            or merged["selected_resume"]
        ),
        "selected_patch_candidate_ids": saved_selected_ids,
        "manual_bullet_edits": saved_manual_edits,
        "excluded_scan_issue_ids": saved_excluded_scan_issue_ids,
        "personal_details": saved_personal_details,
        "note": _clean_text(saved_data.get("note")),
        "rewrite_review_decisions": saved_review_decisions,
        "rewrite_review_telemetry": saved_review_telemetry,
    })

    return {
        "ok": True,
        "draft_status": "saved",
        "has_saved_draft": True,
        "draft": merged,
    }


def save_tailoring_workspace_draft_payload(
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    *,
    tailoring_json_path: str = "",
    selected_resume: str = "",
    selected_patch_candidate_ids: Any = None,
    manual_bullet_edits: Any = None,
    rewrite_review_decisions: Any = None,
    excluded_scan_issue_ids: Any = None,
    personal_details: Any = None,
    note: str = "",
) -> Dict[str, Any]:
    artifact_path = _resolve_planning_artifact_path(
        tailoring_json_path,
        output_dir=output_dir,
    )

    if artifact_path.suffix.lower() != ".json":
        raise ValueError("Workspace draft saving requires a tailoring JSON artifact.")

    payload_data = _load_tailoring_json_artifact(artifact_path)

    if artifact_path.name.endswith("__tailoring.json"):
        payload_data = _rehydrate_legacy_tailoring_operator_payload(
            artifact_path,
            payload_data,
        )

    if payload_data.get("replacement_candidates") is not None:
        payload_data = _apply_saved_patch_selection_overlay(
            artifact_path,
            payload_data,
        )

    draft_payload = _build_tailoring_workspace_default_draft_payload(
        artifact_path,
        payload_data,
        selected_resume=selected_resume,
    )

    valid_candidate_ids = set(_tailoring_artifact_candidate_ids(payload_data))
    requested_candidate_ids = _normalize_selected_patch_candidate_ids(
        selected_patch_candidate_ids
    )
    if not requested_candidate_ids:
        requested_candidate_ids = list(draft_payload["selected_patch_candidate_ids"])

    unknown_candidate_ids = [
        candidate_id
        for candidate_id in requested_candidate_ids
        if candidate_id not in valid_candidate_ids
    ]
    if unknown_candidate_ids:
        raise ValueError(
            f"Unknown candidate IDs for this artifact: {', '.join(sorted(unknown_candidate_ids))}"
        )

    manual_edit_map = _normalize_workspace_manual_bullet_edits(manual_bullet_edits)
    excluded_issue_ids = _normalize_workspace_excluded_scan_issue_ids(excluded_scan_issue_ids)
    personal_detail_map = _normalize_workspace_personal_details(personal_details)
    review_decision_map = _normalize_workspace_rewrite_review_decisions(
        rewrite_review_decisions
    )
    derived_review_decisions = _derive_workspace_rewrite_review_decisions(
        payload_data,
        selected_candidate_ids=requested_candidate_ids,
        manual_bullet_edits=manual_edit_map,
        rewrite_review_decisions=review_decision_map,
    )

    derived_review_telemetry = _build_workspace_rewrite_review_telemetry(
        payload_data,
        selected_candidate_ids=requested_candidate_ids,
        manual_bullet_edits=manual_edit_map,
        rewrite_review_decisions=derived_review_decisions,
    )
    
    saved_at = datetime.now(timezone.utc).isoformat(timespec="microseconds")
    draft_payload.update({
        "draft_status": "saved",
        "saved_at": saved_at,
        "selected_patch_candidate_ids": requested_candidate_ids,
        "manual_bullet_edits": manual_edit_map,
        "excluded_scan_issue_ids": excluded_issue_ids,
        "personal_details": personal_detail_map,
        "note": _clean_text(note),
        "rewrite_review_decisions": derived_review_decisions,
        "rewrite_review_telemetry": derived_review_telemetry,
    })

    draft_path = Path(draft_payload["draft_json_path"])
    draft_path.parent.mkdir(parents=True, exist_ok=True)
    draft_path.write_text(
        json.dumps(draft_payload, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )

    return {
        "ok": True,
        "draft_status": "saved",
        "has_saved_draft": True,
        "draft_json_path": str(draft_path),
        "draft": draft_payload,
    }

def _normalize_tailoring_workspace_compare_text(value: Any) -> str:
    return _normalize_tailoring_workspace_text_key(value)

def _normalize_tailoring_workspace_text_key(value: Any) -> str:
    text = str(value or "").lower()
    text = re.sub(r"[•▪◦·]", " ", text)
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[^a-z0-9%$&.,+/\- ]+", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _tailoring_workspace_candidate_id(item: Dict[str, Any]) -> str:
    return _clean_text(item.get("replacement_candidate_id") or item.get("candidate_id"))


def _tailoring_workspace_bullet_key(item: Dict[str, Any]) -> str:
    candidate_id = _tailoring_workspace_candidate_id(item)
    if candidate_id:
        return f"candidate:{candidate_id}"

    original_text = _clean_text(item.get("original_text"))
    if not original_text:
        return ""

    return f"text:{_normalize_tailoring_workspace_text_key(original_text)}"


def _tailoring_workspace_surfaced_items(payload_data: Dict[str, Any]) -> List[Dict[str, Any]]:
    return [
        *list(payload_data.get("app_ready_replacements", []) or []),
        *list(payload_data.get("direct_apply_optional_replacements", []) or []),
        *list(payload_data.get("ai_optimize_optional_replacements", []) or []),
        *list(payload_data.get("direction_only_replacements", []) or []),
    ]


def _derive_workspace_rewrite_review_decisions(
    payload_data: Dict[str, Any],
    *,
    selected_candidate_ids: Any,
    manual_bullet_edits: Any,
    rewrite_review_decisions: Any,
) -> Dict[str, Dict[str, str]]:
    selected_set = set(_normalize_selected_patch_candidate_ids(selected_candidate_ids))
    manual_map = _normalize_workspace_manual_bullet_edits(manual_bullet_edits)
    decision_map = _normalize_workspace_rewrite_review_decisions(rewrite_review_decisions)

    derived: Dict[str, Dict[str, str]] = {}

    for candidate_id, row in decision_map.items():
        derived[candidate_id] = {
            "state": _clean_text(row.get("state")).lower() or "pending",
            "note": _clean_text(row.get("note")),
        }

    for item in _tailoring_workspace_surfaced_items(payload_data):
        candidate_id = _tailoring_workspace_candidate_id(item)
        if not candidate_id:
            continue

        current = dict(derived.get(candidate_id, {"state": "pending", "note": ""}))
        state = _clean_text(current.get("state")).lower() or "pending"
        note = _clean_text(current.get("note"))

        if state not in {"accepted", "edited_after_accept"}:
            derived[candidate_id] = {
                "state": state,
                "note": note,
            }
            continue

        bullet_key = _tailoring_workspace_bullet_key(item)
        if not bullet_key:
            derived[candidate_id] = {
                "state": state,
                "note": note,
            }
            continue

        current_text = _clean_text(item.get("current_evidence") or item.get("original_text"))
        selected_patch_text = ""
        if candidate_id in selected_set:
            selected_patch_text = _clean_text(item.get("final_replacement_text"))

        effective_base_text = selected_patch_text or current_text
        manual_text = _clean_text(manual_map.get(bullet_key))

        if not manual_text or not effective_base_text:
            derived[candidate_id] = {
                "state": "accepted",
                "note": note,
            }
            continue

        manual_norm = _normalize_tailoring_workspace_compare_text(manual_text)
        base_norm = _normalize_tailoring_workspace_compare_text(effective_base_text)

        derived[candidate_id] = {
            "state": "edited_after_accept" if manual_norm and base_norm and manual_norm != base_norm else "accepted",
            "note": note,
        }

    return derived

def _build_workspace_rewrite_review_telemetry(
    payload_data: Dict[str, Any],
    *,
    selected_candidate_ids: Any,
    manual_bullet_edits: Any,
    rewrite_review_decisions: Any,
) -> Dict[str, Any]:
    selected_ids = _normalize_selected_patch_candidate_ids(selected_candidate_ids)
    selected_set = set(selected_ids)
    manual_map = _normalize_workspace_manual_bullet_edits(manual_bullet_edits)
    effective_decisions = _derive_workspace_rewrite_review_decisions(
        payload_data,
        selected_candidate_ids=selected_ids,
        manual_bullet_edits=manual_map,
        rewrite_review_decisions=rewrite_review_decisions,
    )

    surfaced_candidate_ids: List[str] = []
    seen = set()
    for item in _tailoring_workspace_surfaced_items(payload_data):
        candidate_id = _tailoring_workspace_candidate_id(item)
        if not candidate_id or candidate_id in seen:
            continue
        seen.add(candidate_id)
        surfaced_candidate_ids.append(candidate_id)

    pending_candidate_ids: List[str] = []
    reviewed_candidate_ids: List[str] = []

    pending_count = 0
    accepted_count = 0
    accepted_as_is_count = 0
    edited_after_accept_count = 0
    rejected_count = 0

    for candidate_id in surfaced_candidate_ids:
        state_row = effective_decisions.get(candidate_id, {"state": "pending", "note": ""})
        state = _clean_text(state_row.get("state")).lower() or "pending"

        if state == "pending":
            pending_count += 1
            pending_candidate_ids.append(candidate_id)
        else:
            reviewed_candidate_ids.append(candidate_id)

        if state == "accepted":
            accepted_count += 1
            accepted_as_is_count += 1
        elif state == "edited_after_accept":
            accepted_count += 1
            edited_after_accept_count += 1
        elif state == "rejected":
            rejected_count += 1

    reviewed_count = accepted_count + rejected_count
    remaining_to_review_count = pending_count

    return {
        "pending_count": pending_count,
        "accepted_count": accepted_count,
        "accepted_as_is_count": accepted_as_is_count,
        "edited_after_accept_count": edited_after_accept_count,
        "rejected_count": rejected_count,
        "reviewed_count": reviewed_count,
        "remaining_to_review_count": remaining_to_review_count,
        "selected_candidate_count": len(selected_set),
        "manual_edit_count": len(manual_map),
        "reviewed_candidate_ids": reviewed_candidate_ids,
        "pending_candidate_ids": pending_candidate_ids,
    }


def _build_tailoring_workspace_effective_patch_specs(
    payload_data: Dict[str, Any],
    *,
    selected_candidate_ids: List[str],
    manual_bullet_edits: Dict[str, str],
) -> Tuple[List[Dict[str, Any]], List[str]]:
    selected_set = set(_normalize_selected_patch_candidate_ids(selected_candidate_ids))
    manual_map = {
        _clean_text(key): str(value or "")
        for key, value in dict(manual_bullet_edits or {}).items()
        if _clean_text(key)
    }

    unresolved_manual_keys = set(manual_map.keys())
    patch_specs: List[Dict[str, Any]] = []
    seen_bullet_keys = set()

    for item in _tailoring_workspace_surfaced_items(payload_data):
        original_text = _clean_text(item.get("original_text"))
        if not original_text:
            continue

        bullet_key = _tailoring_workspace_bullet_key(item)
        if not bullet_key or bullet_key in seen_bullet_keys:
            continue
        seen_bullet_keys.add(bullet_key)

        candidate_id = _tailoring_workspace_candidate_id(item)

        selected_patch_text = ""
        if candidate_id and candidate_id in selected_set:
            selected_patch_text = _clean_text(item.get("final_replacement_text"))

        has_manual_override = bullet_key in manual_map
        if has_manual_override:
            unresolved_manual_keys.discard(bullet_key)

        effective_patch_text = (
            manual_map.get(bullet_key)
            if has_manual_override
            else selected_patch_text
        )

        if not effective_patch_text or effective_patch_text == original_text:
            continue

        patch_specs.append({
            "bullet_key": bullet_key,
            "candidate_id": candidate_id,
            "source_bullet_id": _clean_text(item.get("source_bullet_id")),
            "source_raw_text": original_text,
            "patch_text": effective_patch_text,
            "patch_source": "manual_edit" if has_manual_override else "selected_patch",
        })

    return patch_specs, sorted(unresolved_manual_keys)


def _load_job_record_for_workspace_preview(
    job_doc_id: str,
    job_corpus_path: Path = DEFAULT_CORPUS_PATH,
) -> Dict[str, Any]:
    from src.matching.job_adapter import build_job_evidence

    clean_job_doc_id = _clean_text(job_doc_id)
    if not clean_job_doc_id:
        raise ValueError("Workspace draft preview requires job_doc_id.")

    if not job_corpus_path.exists():
        raise ValueError(f"Missing job corpus: {job_corpus_path}")

    with job_corpus_path.open("r", encoding="utf-8") as f:
        for raw_line in f:
            line = raw_line.strip()
            if not line:
                continue

            record = json.loads(line)
            record_job_doc_id = _clean_text(getattr(build_job_evidence(record), "job_doc_id", ""))
            if record_job_doc_id == clean_job_doc_id:
                return record

    raise ValueError(f"Could not find job_doc_id in corpus: {clean_job_doc_id}")

_WORKSPACE_JOB_IDENTITY_KEYS = {
    "job_doc_id",
    "company",
    "title",
    "link",
    "url",
    "job_url",
}


def _workspace_job_record_has_substantive_context(value: Any, *, top_level: bool = True) -> bool:
    if isinstance(value, dict):
        for raw_key, child in value.items():
            key = _clean_text(raw_key).lower()
            if top_level and key in _WORKSPACE_JOB_IDENTITY_KEYS:
                continue
            if _workspace_job_record_has_substantive_context(child, top_level=False):
                return True
        return False

    if isinstance(value, (list, tuple, set)):
        return any(
            _workspace_job_record_has_substantive_context(item, top_level=False)
            for item in value
        )

    if isinstance(value, str):
        text = _clean_text(value)
        if not text:
            return False
        if re.fullmatch(r"https?://\S+", text):
            return False
        return True

    return False


def _resolve_workspace_preview_job_record(
    job: Dict[str, Any],
    job_snapshot: Dict[str, Any],
) -> Dict[str, Any]:
    job_doc_id = _clean_text(
        (job_snapshot.get("job_doc_id", "") if isinstance(job_snapshot, dict) else "")
        or (job.get("job_doc_id", "") if isinstance(job, dict) else "")
    )

    if isinstance(job_snapshot, dict) and _workspace_job_record_has_substantive_context(job_snapshot):
        return dict(job_snapshot)

    if job_doc_id:
        corpus_record = _load_job_record_for_workspace_preview(job_doc_id)
        if _workspace_job_record_has_substantive_context(corpus_record):
            return corpus_record

    if isinstance(job, dict) and _workspace_job_record_has_substantive_context(job):
        return dict(job)

    raise ValueError(
        f"Workspace draft preview could not resolve a job record with substantive scoring context for job_doc_id: {job_doc_id or '<missing>'}"
    )

def _load_resume_evidence_for_workspace_preview(
    resume_name: str,
    *,
    owner_user_id: str = "",
):
    from src.resume.evidence_builder import build_resume_evidence

    safe_resume_name = _sanitize_resume_filename(resume_name)
    resume_pdf_path = planning_resume_preview_path(
        safe_resume_name,
        owner_user_id=owner_user_id,
    )
    resume_document = _new_scan_resume_document(
        resume_name=safe_resume_name,
        resume_file_path=str(resume_pdf_path),
        resume_text="",
    )

    if not _clean_text(resume_document.raw_text):
        raise ValueError(f"Could not load resume document text: {safe_resume_name}")

    return build_resume_evidence(resume_document)


def _workspace_preview_dimension_deltas(
    original_result: Any,
    projected_result: Any,
) -> Dict[str, float]:
    original_map = {
        str(dim.name): float(dim.weighted_score)
        for dim in list(getattr(original_result, "dimension_scores", []) or [])
    }
    projected_map = {
        str(dim.name): float(dim.weighted_score)
        for dim in list(getattr(projected_result, "dimension_scores", []) or [])
    }

    deltas: Dict[str, float] = {}
    for name in sorted(set(original_map) | set(projected_map)):
        delta = projected_map.get(name, 0.0) - original_map.get(name, 0.0)
        if abs(delta) > 1e-12:
            deltas[name] = round(delta, 6)

    return deltas


def _build_workspace_counterfactual_preview(
    original_result: Any,
    projected_result: Any,
    *,
    selected_candidate_ids: List[str],
    patch_specs: List[Dict[str, Any]],
    preview_note: str,
) -> Dict[str, Any]:
    dimension_deltas = _workspace_preview_dimension_deltas(
        original_result,
        projected_result,
    )

    original_missing = list(getattr(original_result.prefilter, "missing_requirements", []) or [])
    projected_missing = list(getattr(projected_result.prefilter, "missing_requirements", []) or [])
    original_matched = list(getattr(original_result.prefilter, "matched_terms", []) or [])
    projected_matched = list(getattr(projected_result.prefilter, "matched_terms", []) or [])

    overall_delta = float(projected_result.final_score) - float(original_result.final_score)

    return {
        "status": "scored",
        "note": preview_note,
        "original_final_score": round(float(original_result.final_score), 6),
        "projected_final_score": round(float(projected_result.final_score), 6),
        "projected_overall_delta": round(overall_delta, 6),
        "projected_dimension_deltas": dimension_deltas,
        "scorer_visible_evidence_changed": bool(
            dimension_deltas
            or original_missing != projected_missing
            or original_matched != projected_matched
            or abs(overall_delta) > 1e-12
        ),
        "selected_patch_count": len(selected_candidate_ids),
        "selected_candidate_ids": list(selected_candidate_ids),
        "requested_candidate_ids": list(selected_candidate_ids),
        "missing_candidate_ids": [],
        "ineligible_candidate_ids": [],
        "duplicate_source_bullet_ids": [],
        "selection_mode": "workspace_draft",
        "workspace_patch_count": len(patch_specs),
    }

def preview_tailoring_workspace_draft_payload(
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    *,
    tailoring_json_path: str = "",
    selected_resume: str = "",
    owner_user_id: str = "",
    selected_patch_candidate_ids: Any = None,
    manual_bullet_edits: Any = None,
    rewrite_review_decisions: Any = None,
) -> Dict[str, Any]:
    from src.matching.job_adapter import build_job_evidence
    from src.matching.scorer import score_resume_job_match
    from src.resume.evidence_builder import build_counterfactual_resume_evidence

    draft_response = load_tailoring_workspace_draft_payload(
        output_dir=output_dir,
        tailoring_json_path=tailoring_json_path,
        selected_resume=selected_resume,
    )
    draft = dict(draft_response.get("draft", {}) or {})

    artifact_path = _resolve_planning_artifact_path(
        tailoring_json_path,
        output_dir=output_dir,
    )

    if artifact_path.suffix.lower() != ".json":
        raise ValueError("Workspace draft preview requires a tailoring JSON artifact.")

    payload_data = _load_tailoring_json_artifact(artifact_path)

    if artifact_path.name.endswith("__tailoring.json"):
        payload_data = _rehydrate_legacy_tailoring_operator_payload(
            artifact_path,
            payload_data,
        )

    if payload_data.get("replacement_candidates") is not None:
        payload_data = _apply_saved_patch_selection_overlay(
            artifact_path,
            payload_data,
        )

    valid_candidate_ids = set(_tailoring_artifact_candidate_ids(payload_data))

    if selected_patch_candidate_ids is None:
        effective_selected_ids = _normalize_selected_patch_candidate_ids(
            draft.get("selected_patch_candidate_ids", [])
        )
    else:
        effective_selected_ids = _normalize_selected_patch_candidate_ids(
            selected_patch_candidate_ids
        )

    unknown_candidate_ids = [
        candidate_id
        for candidate_id in effective_selected_ids
        if candidate_id not in valid_candidate_ids
    ]
    if unknown_candidate_ids:
        raise ValueError(
            f"Unknown candidate IDs for this artifact: {', '.join(sorted(unknown_candidate_ids))}"
        )

    if manual_bullet_edits is None:
        effective_manual_edits = _normalize_workspace_manual_bullet_edits(
            draft.get("manual_bullet_edits", {})
        )
    else:
        effective_manual_edits = _normalize_workspace_manual_bullet_edits(
            manual_bullet_edits
        )

    review_decision_source = (
        draft.get("rewrite_review_decisions", {})
        if rewrite_review_decisions is None
        else rewrite_review_decisions
    )

    effective_review_decisions = _derive_workspace_rewrite_review_decisions(
        payload_data,
        selected_candidate_ids=effective_selected_ids,
        manual_bullet_edits=effective_manual_edits,
        rewrite_review_decisions=review_decision_source,
    )

    effective_review_telemetry = _build_workspace_rewrite_review_telemetry(
        payload_data,
        selected_candidate_ids=effective_selected_ids,
        manual_bullet_edits=effective_manual_edits,
        rewrite_review_decisions=effective_review_decisions,
    )

    selection = payload_data.get("selection", {}) or {}
    job = payload_data.get("job", {}) or {}
    job_snapshot = payload_data.get("job_snapshot", {}) or {}

    effective_selected_resume = (
        _sanitize_optional_resume_filename(selected_resume)
        or _sanitize_optional_resume_filename(draft.get("selected_resume"))
        or _sanitize_optional_resume_filename(selection.get("selected_resume"))
    )
    if not effective_selected_resume:
        raise ValueError("Workspace draft preview requires a selected resume.")

    manual_edit_count = len(effective_manual_edits)

    try:
        original_resume = _load_resume_evidence_for_workspace_preview(
            effective_selected_resume,
            owner_user_id=owner_user_id,
        )

        job_record = _resolve_workspace_preview_job_record(
            job,
            job_snapshot,
        )

        job_evidence = build_job_evidence(job_record)

        original_result = score_resume_job_match(original_resume, job_evidence)
        original_score = round(float(original_result.final_score), 6)
    except ValueError as exc:
        return {
            "ok": True,
            "preview_status": "preview_unavailable",
            "preview_note": str(exc),
            "tailoring_json_path": str(artifact_path),
            "draft_status": _clean_text(draft_response.get("draft_status")),
            "has_saved_draft": bool(draft_response.get("has_saved_draft", False)),
            "selected_patch_candidate_ids": effective_selected_ids,
            "manual_bullet_edits": effective_manual_edits,
            "manual_edit_count": manual_edit_count,
            "manual_edit_rescore_supported": False,
            "needs_full_draft_rescore": True,
            "original_score": None,
            "projected_score": None,
            "projected_delta": None,
            "selected_patch_set_counterfactual_preview": None,
            "unresolved_manual_edit_keys": [],
            "rewrite_review_decisions": effective_review_decisions,
            "rewrite_review_telemetry": effective_review_telemetry,
            "score_preview": _build_workspace_score_preview_contract(
                preview_status="preview_unavailable",
                preview_note=str(exc),
                original_score=None,
                projected_score=None,
                projected_delta=None,
                selected_candidate_ids=effective_selected_ids,
                manual_bullet_edits=effective_manual_edits,
                patch_specs=[],
                unresolved_manual_edit_keys=[],
            ),
            "draft_fragments": _build_workspace_draft_fragments_contract(
                preview_status="preview_unavailable",
                preview_note=str(exc),
                patch_specs=[],
                unresolved_manual_edit_keys=[],
            ),
        }

    patch_specs, unresolved_manual_keys = _build_tailoring_workspace_effective_patch_specs(
        payload_data,
        selected_candidate_ids=effective_selected_ids,
        manual_bullet_edits=effective_manual_edits,
    )

    if not patch_specs:
        if unresolved_manual_keys:
            preview_status = "workspace_draft_rescore_failed"
            preview_note = (
                "Workspace draft edits could not be mapped back to surfaced bullets for rescoring."
            )
        else:
            preview_status = "no_previewable_changes"
            preview_note = "No previewable workspace draft changes were found."

        return {
            "ok": True,
            "preview_status": preview_status,
            "preview_note": preview_note,
            "tailoring_json_path": str(artifact_path),
            "draft_status": _clean_text(draft_response.get("draft_status")),
            "has_saved_draft": bool(draft_response.get("has_saved_draft", False)),
            "selected_patch_candidate_ids": effective_selected_ids,
            "manual_bullet_edits": effective_manual_edits,
            "manual_edit_count": manual_edit_count,
            "manual_edit_rescore_supported": True,
            "needs_full_draft_rescore": False,
            "original_score": original_score,
            "projected_score": None,
            "projected_delta": None,
            "selected_patch_set_counterfactual_preview": None,
            "unresolved_manual_edit_keys": unresolved_manual_keys,
            "rewrite_review_decisions": effective_review_decisions,
            "rewrite_review_telemetry": effective_review_telemetry,
            "score_preview": _build_workspace_score_preview_contract(
                preview_status=preview_status,
                preview_note=preview_note,
                original_score=original_score,
                projected_score=None,
                projected_delta=None,
                selected_candidate_ids=effective_selected_ids,
                manual_bullet_edits=effective_manual_edits,
                patch_specs=[],
                unresolved_manual_edit_keys=unresolved_manual_keys,
            ),
            "draft_fragments": _build_workspace_draft_fragments_contract(
                preview_status=preview_status,
                preview_note=preview_note,
                patch_specs=[],
                unresolved_manual_edit_keys=unresolved_manual_keys,
            ),
        }

    counterfactual_resume = original_resume
    patch_status = "ok"

    for patch in patch_specs:
        counterfactual_resume, patch_status = build_counterfactual_resume_evidence(
            counterfactual_resume,
            source_bullet_id=str(patch.get("source_bullet_id", "") or ""),
            patch_text=str(patch.get("patch_text", "") or ""),
            source_raw_text=str(patch.get("source_raw_text", "") or ""),
        )
        if counterfactual_resume is None:
            break

    if counterfactual_resume is None:
        preview_note = f"Workspace draft rescoring failed: {patch_status}."
        if unresolved_manual_keys:
            preview_note += f" Ignored {len(unresolved_manual_keys)} unmapped manual edit key(s)."

        return {
            "ok": True,
            "preview_status": "workspace_draft_rescore_failed",
            "preview_note": preview_note,
            "tailoring_json_path": str(artifact_path),
            "draft_status": _clean_text(draft_response.get("draft_status")),
            "has_saved_draft": bool(draft_response.get("has_saved_draft", False)),
            "selected_patch_candidate_ids": effective_selected_ids,
            "manual_bullet_edits": effective_manual_edits,
            "manual_edit_count": manual_edit_count,
            "manual_edit_rescore_supported": True,
            "needs_full_draft_rescore": False,
            "original_score": original_score,
            "projected_score": None,
            "projected_delta": None,
            "selected_patch_set_counterfactual_preview": None,
            "patch_status": patch_status,
            "unresolved_manual_edit_keys": unresolved_manual_keys,
            "rewrite_review_decisions": effective_review_decisions,
            "rewrite_review_telemetry": effective_review_telemetry,
            "score_preview": _build_workspace_score_preview_contract(
                preview_status="workspace_draft_rescore_failed",
                preview_note=preview_note,
                original_score=original_score,
                projected_score=None,
                projected_delta=None,
                selected_candidate_ids=effective_selected_ids,
                manual_bullet_edits=effective_manual_edits,
                patch_specs=patch_specs,
                unresolved_manual_edit_keys=unresolved_manual_keys,
            ),
            "draft_fragments": _build_workspace_draft_fragments_contract(
                preview_status="workspace_draft_rescore_failed",
                preview_note=preview_note,
                patch_specs=patch_specs,
                unresolved_manual_edit_keys=unresolved_manual_keys,
            ),
        }

    projected_result = score_resume_job_match(counterfactual_resume, job_evidence)
    projected_score = round(float(projected_result.final_score), 6)
    projected_delta = round(projected_score - original_score, 6)

    note_parts: List[str] = []
    if effective_selected_ids:
        note_parts.append(f"{len(effective_selected_ids)} selected rewrite(s)")
    if manual_edit_count:
        note_parts.append(f"{manual_edit_count} manual edit(s)")

    if note_parts:
        preview_note = "Workspace draft rescored using " + " and ".join(note_parts) + "."
    else:
        preview_note = "Workspace draft rescored."

    if unresolved_manual_keys:
        preview_note += f" Ignored {len(unresolved_manual_keys)} unmapped manual edit key(s)."

    preview = _build_workspace_counterfactual_preview(
        original_result,
        projected_result,
        selected_candidate_ids=effective_selected_ids,
        patch_specs=patch_specs,
        preview_note=preview_note,
    )

    return {
        "ok": True,
        "preview_status": "workspace_draft_rescored",
        "preview_note": preview_note,
        "tailoring_json_path": str(artifact_path),
        "draft_status": _clean_text(draft_response.get("draft_status")),
        "has_saved_draft": bool(draft_response.get("has_saved_draft", False)),
        "selected_patch_candidate_ids": effective_selected_ids,
        "manual_bullet_edits": effective_manual_edits,
        "manual_edit_count": manual_edit_count,
        "manual_edit_rescore_supported": True,
        "needs_full_draft_rescore": False,
        "original_score": original_score,
        "projected_score": projected_score,
        "projected_delta": projected_delta,
        "selected_patch_set_counterfactual_preview": preview,
        "unresolved_manual_edit_keys": unresolved_manual_keys,
        "rewrite_review_decisions": effective_review_decisions,
        "rewrite_review_telemetry": effective_review_telemetry,
        "score_preview": _build_workspace_score_preview_contract(
            preview_status="workspace_draft_rescored",
            preview_note=preview_note,
            original_score=original_score,
            projected_score=projected_score,
            projected_delta=projected_delta,
            selected_candidate_ids=effective_selected_ids,
            manual_bullet_edits=effective_manual_edits,
            patch_specs=patch_specs,
            unresolved_manual_edit_keys=unresolved_manual_keys,
            dimension_deltas=preview.get("projected_dimension_deltas"),
        ),
        "draft_fragments": _build_workspace_draft_fragments_contract(
            preview_status="workspace_draft_rescored",
            preview_note=preview_note,
            patch_specs=patch_specs,
            unresolved_manual_edit_keys=unresolved_manual_keys,
        ),
    }


def _scan_phrase_signal_terms(
    *,
    guidance_text: str,
    supported_terms: List[str],
) -> List[str]:
    raw_terms: List[Any] = list(supported_terms or [])
    lead_match = re.search(
        r"\blead with\s+(.+?)\s+(?:in|within|then|and)\b",
        guidance_text,
        flags=re.IGNORECASE,
    )
    if lead_match:
        raw_terms.insert(0, lead_match.group(1))

    output: List[str] = []
    seen = set()
    for value in raw_terms:
        text = _scan_issue_display_label(value)
        text = re.sub(r"\bthis opening clause\b", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\bopening clause\b", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\s+", " ", text).strip(" ,.;:")
        if not text or len(text) > 48:
            continue
        key = text.lower()
        if key in seen:
            continue
        seen.add(key)
        output.append(text)
        if len(output) >= 3:
            break

    return output


def _scan_phrase_insert_after_lead_verb(current_text: str, term: str) -> str:
    text = _clean_text(current_text)
    clean_term = _clean_text(term)
    if not text or not clean_term:
        return text

    if _scan_issue_text_contains_term(text.lower(), clean_term):
        return text

    match = re.match(r"^([A-Z][A-Za-z-]+)\s+(.*)$", text)
    if not match:
        return f"{clean_term}: {text}"

    verb = match.group(1)
    rest = match.group(2).strip()
    lower_rest = rest.lower()
    if lower_rest.startswith(("a ", "an ", "the ")):
        return f"{verb} {clean_term}-focused {rest}"

    return f"{verb} {clean_term} {rest}"


def _scan_phrase_append_context(current_text: str, term: str) -> str:
    text = _clean_text(current_text).rstrip(".")
    clean_term = _clean_text(term)
    if not text or not clean_term:
        return text

    if _scan_issue_text_contains_term(text.lower(), clean_term):
        return text

    return f"{text}, surfacing {clean_term} relevance while preserving the original scope"


SCAN_PHRASE_OPTIONS_RESPONSE_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "properties": {
        "options": {
            "type": "array",
            "minItems": 1,
            "maxItems": 3,
            "items": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                    "text": {"type": "string"},
                    "reason": {"type": "string"},
                    "supported_terms": {
                        "type": "array",
                        "items": {"type": "string"},
                    },
                    "risk_flags": {
                        "type": "array",
                        "items": {"type": "string"},
                    },
                },
                "required": ["text", "reason", "supported_terms", "risk_flags"],
            },
        }
    },
    "required": ["options"],
}
SCAN_PHRASE_OPTIONS_SCHEMA_NAME = "scan_phrase_options_v1"
SCAN_PHRASE_PROMPT_VERSION = "v1"


def _scan_phrase_structured_output_contract() -> Dict[str, Any]:
    return {
        "name": SCAN_PHRASE_OPTIONS_SCHEMA_NAME,
        "strict": True,
        "schema": SCAN_PHRASE_OPTIONS_RESPONSE_SCHEMA,
    }

def _scan_phrase_default_provider() -> str:
    if os.getenv("OPENAI_API_KEY"):
        return "openai"
    return os.getenv(
        "TAILORING_REWRITE_PROVIDER",
        os.getenv(
            "PATCH_REFINEMENT_WRITER_PROVIDER",
            os.getenv("PATCH_REFINEMENT_PROVIDER", "groq"),
        ),
    )


def _scan_phrase_default_model(provider: str) -> str:
    if provider == "openai":
        return os.getenv("TAILORING_PHRASE_MODEL", "gpt-5-mini")
    if provider == "gemini":
        return "gemini-2.5-flash"
    return os.getenv(
        "TAILORING_PHRASE_MODEL",
        os.getenv(
            "TAILORING_REWRITE_MODEL",
            os.getenv(
                "PATCH_REFINEMENT_WRITER_MODEL",
                os.getenv("PATCH_REFINEMENT_MODEL", "llama-3.3-70b-versatile"),
            ),
        ),
    )


SCAN_PHRASE_PROVIDER = os.getenv(
    "SCAN_PHRASE_PROVIDER",
    os.getenv("TAILORING_SCAN_PHRASE_PROVIDER", _scan_phrase_default_provider()),
).strip().lower()
SCAN_PHRASE_MODEL = os.getenv(
    "SCAN_PHRASE_MODEL",
    os.getenv(
        "TAILORING_SCAN_PHRASE_MODEL",
        os.getenv("TAILORING_PHRASE_MODEL", _scan_phrase_default_model(SCAN_PHRASE_PROVIDER)),
    ),
).strip()
SCAN_PHRASE_FALLBACK_PROVIDER = os.getenv(
    "SCAN_PHRASE_FALLBACK_PROVIDER",
    "groq" if SCAN_PHRASE_PROVIDER != "groq" else "openai",
).strip().lower()
SCAN_PHRASE_FALLBACK_MODEL = os.getenv(
    "SCAN_PHRASE_FALLBACK_MODEL",
    _scan_phrase_default_model(SCAN_PHRASE_FALLBACK_PROVIDER),
).strip()
SCAN_PHRASE_LLM_FALLBACK_ENABLED = (
    os.getenv("SCAN_PHRASE_LLM_FALLBACK_ENABLED", "true").strip().lower() == "true"
)
SCAN_PHRASE_DETERMINISTIC_FALLBACK_ENABLED = (
    os.getenv("SCAN_PHRASE_DETERMINISTIC_FALLBACK_ENABLED", "false").strip().lower()
    == "true"
)


def _scan_phrase_extract_first_json_value(text: str) -> str:
    value = str(text or "").strip().replace("```json", "").replace("```", "").strip()
    if not value:
        return ""

    start_positions = [
        idx for idx in (value.find("{"), value.find("["))
        if idx >= 0
    ]
    if not start_positions:
        return value

    start = min(start_positions)
    opening = value[start]
    closing = "}" if opening == "{" else "]"
    depth = 0
    in_string = False
    escape = False

    for idx in range(start, len(value)):
        ch = value[idx]
        if in_string:
            if escape:
                escape = False
                continue
            if ch == "\\":
                escape = True
                continue
            if ch == '"':
                in_string = False
            continue

        if ch == '"':
            in_string = True
            continue
        if ch == opening:
            depth += 1
            continue
        if ch == closing:
            depth -= 1
            if depth == 0:
                return value[start : idx + 1].strip()

    return value[start:].strip()


def _scan_phrase_parse_options_payload(raw: Any) -> List[Any]:
    if isinstance(raw, dict):
        options = raw.get("options")
        return list(options or []) if isinstance(options, list) else []

    if isinstance(raw, list):
        return raw

    text = _clean_text(raw)
    if not text:
        return []

    candidates = [text, _scan_phrase_extract_first_json_value(text)]
    for candidate in candidates:
        if not candidate:
            continue
        try:
            parsed = json.loads(candidate)
        except Exception:
            continue
        return _scan_phrase_parse_options_payload(parsed)

    return []


def _scan_phrase_deterministic_options(
    *,
    current: str,
    terms: List[str],
) -> List[Dict[str, Any]]:
    primary_term = terms[0] if terms else ""
    raw_options = [
        _scan_phrase_insert_after_lead_verb(current, primary_term),
        _scan_phrase_append_context(current, primary_term),
        current,
    ]

    options: List[Dict[str, Any]] = []
    seen = set()
    for index, text in enumerate(raw_options, start=1):
        normalized = _normalize_tailoring_workspace_text_key(text)
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        options.append({
            "option_id": f"phrase_{index}",
            "text": text,
            "source": "deterministic_guidance_phrase",
            "reason": "Deterministic phrase draft from the guidance term and current bullet.",
            "supported_terms": terms,
            "risk_flags": [],
            "requires_review": True,
            "can_accept_directly": False,
        })

    return options[:3]


def _scan_phrase_llm_prompt(
    *,
    current: str,
    guidance: str,
    terms: List[str],
) -> str:
    term_line = ", ".join(terms) if terms else "none"
    return "\n".join([
        "Create up to three high-quality resume bullet rewrite drafts for a manual editor.",
        "",
        "Rewrite objective:",
        "- Apply the AI guidance as a concrete wording improvement, not a vague note.",
        "- Put the strongest supported job signal earlier in the bullet when truthful.",
        "- Preserve the original metric values, tools, employer context, and scope.",
        "- Do not invent tools, domains, ownership, impact, numbers, or responsibilities.",
        "- Keep each option as one resume bullet without a leading bullet symbol.",
        "- Return only JSON in this exact shape:",
        '{"options":[{"text":"...","reason":"...","supported_terms":["..."],"risk_flags":[]}]}',
        "",
        f"Supported terms: {term_line}",
        f"AI guidance: {guidance}",
        "",
        "Current bullet:",
        current,
    ])


def _scan_phrase_validate_llm_options(
    raw_options: Any,
    *,
    current: str,
    terms: List[str],
) -> List[Dict[str, Any]]:
    raw_option_rows = _scan_phrase_parse_options_payload(raw_options)
    if not isinstance(raw_option_rows, list):
        return []

    current_norm = _normalize_tailoring_workspace_text_key(current)
    options: List[Dict[str, Any]] = []
    seen = set()

    for index, row in enumerate(raw_option_rows, start=1):
        if not isinstance(row, dict):
            continue

        text = _clean_text(row.get("text"))
        normalized = _normalize_tailoring_workspace_text_key(text)
        if not text or not normalized or normalized in seen:
            continue

        if len(text) > max(280, len(current) + 120):
            continue

        seen.add(normalized)
        options.append({
            "option_id": f"phrase_llm_{index}",
            "text": text,
            "source": "llm_guidance_phrase",
            "reason": _clean_text(row.get("reason")) or "LLM phrase draft from guidance.",
            "supported_terms": _scan_issue_unique_display_labels(
                list(row.get("supported_terms", []) or []) or terms
            ),
            "risk_flags": list(row.get("risk_flags", []) or []),
            "requires_review": True,
            "can_accept_directly": False,
            "changed_from_current": normalized != current_norm,
        })

        if len(options) >= 3:
            break

    return options


def _generate_scan_phrase_options_with_llm(
    *,
    current: str,
    guidance: str,
    terms: List[str],
) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
    from src.ai.llm_client import run_chat_completion_with_metadata
    from src.tailoring import llm as tailoring_llm

    provider = SCAN_PHRASE_PROVIDER or tailoring_llm.PATCH_REFINEMENT_WRITER_PROVIDER
    model = SCAN_PHRASE_MODEL or (
        "gemini-2.5-flash"
        if provider == "gemini"
        else tailoring_llm.PATCH_REFINEMENT_WRITER_MODEL
    )
    fallback_provider = SCAN_PHRASE_FALLBACK_PROVIDER
    fallback_model = SCAN_PHRASE_FALLBACK_MODEL
    system_prompt = (
        "You generate conservative, truthful resume bullet rewrite options for manual editing. "
        "Return only JSON."
    )
    user_prompt = _scan_phrase_llm_prompt(
        current=current,
        guidance=guidance,
        terms=terms,
    )
    last_error = ""

    result = run_chat_completion_with_metadata(
        provider=provider,
        model=model,
        temperature=0,
        max_tokens=520,
        response_mime_type="application/json",
        response_schema=_scan_phrase_structured_output_contract()["schema"],
        return_parsed=True,
        thinking_budget=0,
        fallback_enabled=SCAN_PHRASE_LLM_FALLBACK_ENABLED,
        fallback_provider=fallback_provider,
        fallback_model=fallback_model,
        messages=[
            {
                "role": "system",
                "content": system_prompt,
            },
            {
                "role": "user",
                "content": user_prompt,
            },
        ],
    )

    content = result.get("content")
    options = _scan_phrase_validate_llm_options(
        content,
        current=current,
        terms=terms,
    )
    if not options:
        last_error = (
            "Structured phrase response had no valid options. Retrying with plain JSON parsing."
        )
        plain_result = run_chat_completion_with_metadata(
            provider=provider,
            model=model,
            temperature=0,
            max_tokens=520,
            response_mime_type=None,
            response_schema=None,
            return_parsed=False,
            thinking_budget=0,
            fallback_enabled=SCAN_PHRASE_LLM_FALLBACK_ENABLED,
            fallback_provider=fallback_provider,
            fallback_model=fallback_model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        )
        result = plain_result
        content = plain_result.get("content")
        options = _scan_phrase_validate_llm_options(
            content,
            current=current,
            terms=terms,
        )
    metadata = {
        "provider": _clean_text(result.get("provider")),
        "model": _clean_text(result.get("model")),
        "fallback_used": bool(result.get("fallback_used", False)),
        "requested_provider": provider,
        "requested_model": model,
        "prompt_version": SCAN_PHRASE_PROMPT_VERSION,
        "structured_output_requested": True,
        "structured_output_schema": _scan_phrase_structured_output_contract(),
        "plain_retry_used": bool(last_error),
    }
    if last_error and not options:
        metadata["last_error"] = last_error
    return options, metadata


def generate_tailoring_scan_phrase_payload(
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    *,
    tailoring_json_path: str = "",
    selected_resume: str = "",
    bullet_key: str = "",
    current_text: str = "",
    guidance_text: str = "",
    supported_terms: Any = None,
) -> Dict[str, Any]:
    artifact_path = _resolve_planning_artifact_path(
        tailoring_json_path,
        output_dir=output_dir,
    )
    if artifact_path.suffix.lower() != ".json":
        raise ValueError("Scan phrase generation requires a tailoring JSON artifact.")

    current = _clean_text(current_text)
    if not current:
        raise ValueError("Scan phrase generation requires current_text.")

    terms = _scan_phrase_signal_terms(
        guidance_text=_clean_text(guidance_text),
        supported_terms=list(supported_terms or []),
    )
    guidance = _clean_text(guidance_text)
    llm_options: List[Dict[str, Any]] = []
    llm_metadata: Dict[str, Any] = {}
    llm_error = ""

    if guidance or terms:
        try:
            llm_options, llm_metadata = _generate_scan_phrase_options_with_llm(
                current=current,
                guidance=guidance,
                terms=terms,
            )
        except Exception as exc:
            llm_error = str(exc)

    deterministic_options = (
        _scan_phrase_deterministic_options(
            current=current,
            terms=terms,
        )
        if SCAN_PHRASE_DETERMINISTIC_FALLBACK_ENABLED and not llm_options
        else []
    )

    options = llm_options or deterministic_options
    source = (
        "llm"
        if llm_options
        else "deterministic_fallback"
        if deterministic_options
        else "llm_unavailable"
    )

    if source == "llm":
        note = "LLM phrase drafts generated for manual edit and rescore."
    elif source == "deterministic_fallback":
        note = "Deterministic fallback phrase drafts generated because explicit fallback is enabled."
    else:
        note = (
            "No valid LLM phrase drafts were generated. Check the LLM provider/model settings "
            "or retry after the server has the latest code."
        )

    return {
        "ok": True,
        "version": "scan_phrase_options_v1",
        "tailoring_json_path": str(artifact_path),
        "selected_resume": _sanitize_optional_resume_filename(selected_resume),
        "bullet_key": _clean_text(bullet_key),
        "guidance_text": guidance,
        "supported_terms": terms,
        "options": options[:3],
        "source": source,
        "llm_metadata": llm_metadata,
        "llm_error": llm_error,
        "fallback_used": source == "deterministic_fallback",
        "note": note,
    }

def _normalize_workspace_export_format(value: Any) -> str:
    raw = _clean_text(value).lower()
    if raw == "pdf":
        return "pdf"
    if raw in {"word", "docx"}:
        return "word"
    raise ValueError("format must be either 'pdf' or 'word'.")


def _workspace_export_output_path(
    output_dir: Path,
    *,
    selected_resume: str,
    export_format: str,
) -> Tuple[Path, str, str]:
    export_dir = Path(output_dir) / "workspace_exports"
    export_dir.mkdir(parents=True, exist_ok=True)

    resume_stem = Path(_sanitize_resume_filename(selected_resume)).stem
    base_name = _slugify_text(resume_stem, max_len=80)

    if export_format == "word":
        filename = f"{base_name}__tailored_draft.docx"
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        filename = f"{base_name}__tailored_draft.pdf"
        media_type = "application/pdf"

    return export_dir / filename, filename, media_type


def _workspace_export_line_is_heading(
    text: str,
    *,
    font_size: float,
    is_bold: bool,
    left: float,
) -> bool:
    clean = _clean_text(text)
    if not clean:
        return False

    if len(clean) > 96:
        return False

    if clean.startswith(("●", "•", "▪", "◦", "·", "-")):
        return False

    if is_bold and font_size >= 10.5:
        return True

    return left <= 72 and font_size >= 12.5

def _workspace_export_span_style(span: Dict[str, Any]) -> Dict[str, Any]:
    raw_font_name = _clean_text(span.get("font"))
    raw_font_name_lower = raw_font_name.lower()
    flags = int(span.get("flags", 0) or 0)

    is_bold = "bold" in raw_font_name_lower or (flags & 16) != 0
    is_italic = "italic" in raw_font_name_lower or "oblique" in raw_font_name_lower

    return {
        "font_name": raw_font_name,
        "font_size": float(span.get("size", 0) or 11.0),
        "bold": bool(is_bold),
        "italic": bool(is_italic),
    }


def _workspace_export_style_key(style: Dict[str, Any]) -> Tuple[str, float, bool, bool]:
    return (
        _clean_text(style.get("font_name")),
        round(float(style.get("font_size", 11.0) or 11.0), 2),
        bool(style.get("bold", False)),
        bool(style.get("italic", False)),
    )


def _workspace_export_runs_from_spans(spans: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    runs: List[Dict[str, Any]] = []

    for span in list(spans or []):
        text = str(span.get("text", "") or "")
        if not text:
            continue

        style = _workspace_export_span_style(span)
        style_key = _workspace_export_style_key(style)

        if runs and runs[-1]["_style_key"] == style_key:
            runs[-1]["text"] += text
            continue

        runs.append(
            {
                "text": text,
                "font_name": style["font_name"],
                "font_size": style["font_size"],
                "bold": style["bold"],
                "italic": style["italic"],
                "_style_key": style_key,
            }
        )

    for run in runs:
        run.pop("_style_key", None)

    return runs


def _workspace_export_dominant_style(line_records: List[Dict[str, Any]]) -> Dict[str, Any]:
    style_weights: Dict[Tuple[str, float, bool, bool], int] = {}
    styles_by_key: Dict[Tuple[str, float, bool, bool], Dict[str, Any]] = {}

    for line in list(line_records or []):
        for run in list(line.get("runs", []) or []):
            text = str(run.get("text", "") or "")
            style = {
                "font_name": _clean_text(run.get("font_name")),
                "font_size": float(run.get("font_size", 11.0) or 11.0),
                "bold": bool(run.get("bold", False)),
                "italic": bool(run.get("italic", False)),
            }
            style_key = _workspace_export_style_key(style)
            weight = max(len(_clean_text(text)), len(text), 1)

            style_weights[style_key] = style_weights.get(style_key, 0) + weight
            styles_by_key[style_key] = style

    if not style_weights:
        return {
            "font_name": "",
            "font_size": 11.0,
            "bold": False,
            "italic": False,
        }

    dominant_key = max(
        style_weights.items(),
        key=lambda item: (item[1], item[0][1]),
    )[0]
    return dict(styles_by_key[dominant_key])


def _workspace_export_docx_font_name(raw_font_name: str) -> str:
    font = _clean_text(raw_font_name).lower()

    if "times" in font:
        return "Times New Roman"
    if "arial" in font or "helvetica" in font:
        return "Arial"
    if "calibri" in font:
        return "Calibri"
    if "courier" in font:
        return "Courier New"
    if "georgia" in font:
        return "Georgia"

    return "Calibri"


def _workspace_export_pdf_font_name(style: Dict[str, Any]) -> str:
    family = _workspace_export_docx_font_name(str(style.get("font_name", "")))
    is_bold = bool(style.get("bold", False))
    is_italic = bool(style.get("italic", False))

    if family == "Times New Roman":
        if is_bold and is_italic:
            return "Times-BoldItalic"
        if is_bold:
            return "Times-Bold"
        if is_italic:
            return "Times-Italic"
        return "Times-Roman"

    if family == "Courier New":
        if is_bold and is_italic:
            return "Courier-BoldOblique"
        if is_bold:
            return "Courier-Bold"
        if is_italic:
            return "Courier-Oblique"
        return "Courier"

    if is_bold and is_italic:
        return "Helvetica-BoldOblique"
    if is_bold:
        return "Helvetica-Bold"
    if is_italic:
        return "Helvetica-Oblique"
    return "Helvetica"


def _workspace_export_finalize_paragraph(
    line_records: List[Dict[str, Any]],
    *,
    previous_bottom: float | None,
    fallback_gap_before: float,
) -> Dict[str, Any]:
    if not line_records:
        return {}

    paragraph_text_parts: List[str] = []
    paragraph_runs: List[Dict[str, Any]] = []

    for line_index, line in enumerate(line_records):
        line_text = _clean_text(line.get("text"))
        if line_text:
            paragraph_text_parts.append(line_text)

        if line_index > 0 and paragraph_runs:
            paragraph_runs[-1]["text"] += " "

        for run in list(line.get("runs", []) or []):
            run_text = str(run.get("text", "") or "")
            if not run_text:
                continue

            normalized_run = {
                "text": run_text,
                "font_name": _clean_text(run.get("font_name")),
                "font_size": float(run.get("font_size", 11.0) or 11.0),
                "bold": bool(run.get("bold", False)),
                "italic": bool(run.get("italic", False)),
            }

            if (
                paragraph_runs
                and _workspace_export_style_key(paragraph_runs[-1])
                == _workspace_export_style_key(normalized_run)
            ):
                paragraph_runs[-1]["text"] += normalized_run["text"]
            else:
                paragraph_runs.append(normalized_run)

    text = " ".join(part for part in paragraph_text_parts if part).strip()
    style = _workspace_export_dominant_style(line_records)

    left = min(float(line.get("left", 0) or 0) for line in line_records)
    right = max(float(line.get("right", 0) or 0) for line in line_records)
    top = min(float(line.get("top", 0) or 0) for line in line_records)
    bottom = max(float(line.get("bottom", 0) or 0) for line in line_records)

    gap_before_raw = (
        top - float(previous_bottom)
        if previous_bottom is not None
        else float(fallback_gap_before)
    )
    gap_before = max(0.0, min(float(gap_before_raw), 28.0))

    is_heading = _workspace_export_line_is_heading(
        text,
        font_size=float(style.get("font_size", 11.0) or 11.0),
        is_bold=bool(style.get("bold", False)),
        left=float(left),
    )

    is_bullet = text.lstrip().startswith(("●", "•", "▪", "◦", "·", "-"))
    line_spacing = 1.16 if is_heading else (1.22 if len(line_records) > 1 else 1.12)

    return {
        "text": text,
        "runs": paragraph_runs,
        "style": style,
        "left": float(left),
        "right": float(right),
        "top": float(top),
        "bottom": float(bottom),
        "font_size": float(style.get("font_size", 11.0) or 11.0),
        "is_heading": bool(is_heading),
        "is_bullet": bool(is_bullet),
        "line_count": len(line_records),
        "gap_before": float(gap_before),
        "line_spacing": float(line_spacing),
        "left_indent_pt": max(0.0, min(float(left) - 36.0, 220.0)),
        "alignment": "left",
        "row_kind": "paragraph",
        "row_group_id": "",
        "row_side": "",
    }

def _workspace_export_same_row(
    left_paragraph: Dict[str, Any],
    right_paragraph: Dict[str, Any],
    *,
    y_tolerance: float = 6.0,
) -> bool:
    return abs(
        float(left_paragraph.get("top", 0.0) or 0.0)
        - float(right_paragraph.get("top", 0.0) or 0.0)
    ) <= y_tolerance

def _workspace_export_line_is_date_range(text: str) -> bool:
    clean = re.sub(r"\s+", " ", _clean_text(text))
    if not clean or len(clean) > 48:
        return False

    month = r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\.?"
    year = r"(?:19|20)\d{2}"
    point = rf"(?:{month}\s+{year}|{year}|Present|Current)"

    return bool(
        re.fullmatch(
            rf"{point}(?:\s*[-–—]\s*{point})?",
            clean,
            flags=re.IGNORECASE,
        )
    )


def _workspace_export_should_promote_to_paired_row(
    left_item: Dict[str, Any],
    right_item: Dict[str, Any],
    *,
    page_width: float,
) -> bool:
    left_left = float(left_item.get("left", 0.0) or 0.0)
    right_left = float(right_item.get("left", 0.0) or 0.0)
    right_edge = float(right_item.get("right", 0.0) or 0.0)

    has_clear_split = (right_left - left_left) >= 120.0
    right_is_right_aligned = right_edge >= (page_width - 54.0)
    if has_clear_split and right_is_right_aligned:
        return True

    right_text = _clean_text(right_item.get("text"))
    left_text = _clean_text(left_item.get("text"))
    left_style = dict(left_item.get("style", {}) or {})

    left_looks_like_role_row = (
        bool(left_item.get("is_heading"))
        or bool(left_style.get("bold", False))
        or float(left_item.get("font_size", 0.0) or 0.0) >= 10.5
    )

    fallback_split = (right_left - left_left) >= 64.0
    fallback_rightish = right_edge >= (page_width - 90.0)
    left_not_centered = not _workspace_export_is_centered_paragraph(
        left_item,
        page_width=page_width,
    )
    right_not_centered = not _workspace_export_is_centered_paragraph(
        right_item,
        page_width=page_width,
    )

    return (
        _workspace_export_line_is_date_range(right_text)
        and left_looks_like_role_row
        and left_not_centered
        and right_not_centered
        and fallback_split
        and fallback_rightish
        and len(left_text) <= 140
        and len(right_text) <= 48
    )

def _workspace_export_line_is_date_range(text: str) -> bool:
    clean = re.sub(r"\s+", " ", _clean_text(text))
    if not clean or len(clean) > 48:
        return False

    month = r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\.?"
    year = r"(?:19|20)\d{2}"
    point = rf"(?:{month}\s+{year}|{year}|Present|Current)"

    return bool(
        re.fullmatch(
            rf"{point}(?:\s*[-–—]\s*{point})?",
            clean,
            flags=re.IGNORECASE,
        )
    )


def _workspace_export_should_promote_to_paired_row(
    left_item: Dict[str, Any],
    right_item: Dict[str, Any],
    *,
    page_width: float,
) -> bool:
    left_left = float(left_item.get("left", 0.0) or 0.0)
    right_left = float(right_item.get("left", 0.0) or 0.0)
    right_edge = float(right_item.get("right", 0.0) or 0.0)

    has_clear_split = (right_left - left_left) >= 120.0
    right_is_right_aligned = right_edge >= (page_width - 54.0)
    if has_clear_split and right_is_right_aligned:
        return True

    right_text = _clean_text(right_item.get("text"))
    left_text = _clean_text(left_item.get("text"))
    left_style = dict(left_item.get("style", {}) or {})

    left_looks_like_role_row = (
        bool(left_item.get("is_heading"))
        or bool(left_style.get("bold", False))
        or float(left_item.get("font_size", 0.0) or 0.0) >= 10.5
    )

    fallback_split = (right_left - left_left) >= 64.0
    fallback_rightish = right_edge >= (page_width - 90.0)
    left_not_centered = not _workspace_export_is_centered_paragraph(
        left_item,
        page_width=page_width,
    )
    right_not_centered = not _workspace_export_is_centered_paragraph(
        right_item,
        page_width=page_width,
    )

    return (
        _workspace_export_line_is_date_range(right_text)
        and left_looks_like_role_row
        and left_not_centered
        and right_not_centered
        and fallback_split
        and fallback_rightish
        and len(left_text) <= 140
        and len(right_text) <= 48
    )

def _workspace_export_is_centered_paragraph(
    paragraph: Dict[str, Any],
    *,
    page_width: float,
) -> bool:
    if bool(paragraph.get("is_bullet")):
        return False

    if int(paragraph.get("line_count", 0) or 0) != 1:
        return False

    text = _clean_text(paragraph.get("text"))
    if not text or len(text) > 120:
        return False

    left = float(paragraph.get("left", 0.0) or 0.0)
    right = float(paragraph.get("right", 0.0) or 0.0)
    mid = (left + right) / 2.0
    page_mid = page_width / 2.0

    return abs(mid - page_mid) <= max(24.0, page_width * 0.04)


def _workspace_export_annotate_page_layout(
    *,
    page_width: float,
    page_blocks: List[Dict[str, Any]],
) -> None:
    flat_paragraphs: List[Dict[str, Any]] = []

    for block in list(page_blocks or []):
        for paragraph in list(block.get("paragraphs", []) or []):
            flat_paragraphs.append(paragraph)

    flat_paragraphs.sort(
        key=lambda paragraph: (
            round(float(paragraph.get("top", 0.0) or 0.0), 2),
            round(float(paragraph.get("left", 0.0) or 0.0), 2),
        )
    )

    row_index = 0
    idx = 0

    while idx < len(flat_paragraphs):
        current = flat_paragraphs[idx]
        row_group = [current]
        scan = idx + 1

        while scan < len(flat_paragraphs):
            candidate = flat_paragraphs[scan]
            if not _workspace_export_same_row(current, candidate):
                break
            row_group.append(candidate)
            scan += 1

        row_group.sort(key=lambda paragraph: float(paragraph.get("left", 0.0) or 0.0))

        if len(row_group) >= 2:
            left_item = row_group[0]
            right_item = row_group[-1]

            if _workspace_export_should_promote_to_paired_row(
                left_item,
                right_item,
                page_width=page_width,
            ):
                row_group_id = f"row_{row_index}"
                row_index += 1

                left_item["row_kind"] = "paired_row"
                left_item["row_group_id"] = row_group_id
                left_item["row_side"] = "left"
                left_item["alignment"] = "left"
                left_item["left_indent_pt"] = max(
                    0.0,
                    min(float(left_item.get("left", 0.0) or 0.0) - 36.0, 220.0),
                )

                right_item["row_kind"] = "paired_row"
                right_item["row_group_id"] = row_group_id
                right_item["row_side"] = "right"
                right_item["alignment"] = "right"
                right_item["left_indent_pt"] = 0.0

        for paragraph in row_group:
            if _clean_text(paragraph.get("row_kind")) == "paired_row":
                continue

            if _workspace_export_is_centered_paragraph(
                paragraph,
                page_width=page_width,
            ):
                paragraph["alignment"] = "center"
                paragraph["left_indent_pt"] = 0.0
            elif (
                not bool(paragraph.get("is_bullet"))
                and int(paragraph.get("line_count", 0) or 0) == 1
                and float(paragraph.get("right", 0.0) or 0.0) >= (page_width - 54.0)
            ):
                paragraph["alignment"] = "right"
                paragraph["left_indent_pt"] = 0.0
            else:
                paragraph["alignment"] = "left"

        idx = scan

def _workspace_export_line_starts_bullet(text: str) -> bool:
    clean = str(text or "").lstrip()
    return clean.startswith(("●", "•", "▪", "◦", "·", "-"))

def _workspace_export_line_starts_strong_label(text: str) -> bool:
    clean = _clean_text(text)
    if not clean or len(clean) > 220:
        return False

    return bool(
        re.match(
            r"^(?:[A-Z0-9][A-Za-z0-9+#/&().-]*)(?:\s+(?:&|[A-Z0-9][A-Za-z0-9+#/&().-]*)){1,6}:\s+\S",
            clean,
        )
    )

def _workspace_export_line_is_contact(text: str) -> bool:
    clean = _clean_text(text).lower()
    if not clean or len(clean) > 180:
        return False

    if any(
        token in clean
        for token in (
            "@",
            "linkedin",
            "github",
            "portfolio",
            "www.",
            "http://",
            "https://",
            ".com",
            "|",
        )
    ):
        return True

    return bool(re.search(r"\+?\d[\d(). \-]{6,}\d", clean))


def _workspace_export_line_is_centered(
    line: Dict[str, Any],
    *,
    page_width: float,
) -> bool:
    left = float(line.get("left", 0.0) or 0.0)
    right = float(line.get("right", 0.0) or 0.0)
    mid = (left + right) / 2.0
    page_mid = page_width / 2.0

    return abs(mid - page_mid) <= max(24.0, page_width * 0.04)


def _workspace_export_should_start_new_paragraph(
    previous_line: Dict[str, Any],
    current_line: Dict[str, Any],
    *,
    page_width: float,
) -> bool:
    previous_text = _clean_text(previous_line.get("text"))
    current_text = _clean_text(current_line.get("text"))

    if not previous_text or not current_text:
        return False

    previous_starts_bullet = _workspace_export_line_starts_bullet(previous_text)
    current_starts_bullet = _workspace_export_line_starts_bullet(current_text)
    previous_starts_strong_label = _workspace_export_line_starts_strong_label(previous_text)
    current_starts_strong_label = _workspace_export_line_starts_strong_label(current_text)

    if current_starts_bullet:
        return True

    if current_starts_strong_label:
        return True

    previous_font_size = float(previous_line.get("font_size", 11.0) or 11.0)
    current_font_size = float(current_line.get("font_size", 11.0) or 11.0)

    previous_is_heading = _workspace_export_line_is_heading(
        previous_text,
        font_size=previous_font_size,
        is_bold=bool(previous_line.get("is_bold", False)),
        left=float(previous_line.get("left", 0.0) or 0.0),
    )
    current_is_heading = _workspace_export_line_is_heading(
        current_text,
        font_size=current_font_size,
        is_bold=bool(current_line.get("is_bold", False)),
        left=float(current_line.get("left", 0.0) or 0.0),
    )

    if bool(previous_is_heading) != bool(current_is_heading):
        return True

    gap = float(current_line.get("top", 0.0) or 0.0) - float(previous_line.get("bottom", 0.0) or 0.0)
    if gap > max(8.0, previous_font_size * 0.9):
        return True

    previous_centered = _workspace_export_line_is_centered(
        previous_line,
        page_width=page_width,
    )
    current_centered = _workspace_export_line_is_centered(
        current_line,
        page_width=page_width,
    )

    if previous_centered and current_centered:
        if _workspace_export_line_is_contact(previous_text) != _workspace_export_line_is_contact(current_text):
            return True
        if abs(current_font_size - previous_font_size) >= 0.75:
            return True
        if bool(previous_line.get("is_bold", False)) != bool(current_line.get("is_bold", False)):
            return True

    previous_left = float(previous_line.get("left", 0.0) or 0.0)
    current_left = float(current_line.get("left", 0.0) or 0.0)
    indent_delta = current_left - previous_left

    if indent_delta <= -8.0:
        return True

    if previous_starts_bullet and indent_delta >= 10.0:
        return False

    if abs(indent_delta) > 32.0:
        return True

    return False

def _extract_resume_pdf_paragraph_pages_for_export(
    resume_pdf_path: Path,
) -> List[Dict[str, Any]]:
    import pymupdf as fitz

    doc = fitz.open(str(resume_pdf_path))
    try:
        exported_pages: List[Dict[str, Any]] = []

        for page_index, page in enumerate(doc, start=1):
            text_dict = page.get_text("dict", sort=True)
            page_blocks: List[Dict[str, Any]] = []
            previous_visible_bottom: float | None = None
            page_width = float(page.rect.width)

            text_blocks = [
                block
                for block in list(text_dict.get("blocks", []) or [])
                if int(block.get("type", -1)) == 0
            ]
            text_blocks.sort(
                key=lambda block: (
                    round(float((block.get("bbox") or [0, 0, 0, 0])[1]), 2),
                    round(float((block.get("bbox") or [0, 0, 0, 0])[0]), 2),
                )
            )

            for block_index, block in enumerate(text_blocks):
                block_bbox = tuple(block.get("bbox", (0, 0, 0, 0)))
                block_left = float(block_bbox[0] or 0)
                block_top = float(block_bbox[1] or 0)

                line_records: List[Dict[str, Any]] = []
                for line in list(block.get("lines", []) or []):
                    spans = list(line.get("spans", []) or [])
                    runs = _workspace_export_runs_from_spans(spans)
                    if not runs:
                        continue

                    line_text = "".join(str(run.get("text", "") or "") for run in runs).strip()
                    if not line_text:
                        continue

                    bbox = tuple(line.get("bbox", (0, 0, 0, 0)))
                    x0, y0, x1, y1 = bbox
                    dominant_style = _workspace_export_dominant_style(
                        [{"runs": runs}]
                    )

                    line_records.append(
                        {
                            "text": line_text,
                            "left": float(x0),
                            "top": float(y0),
                            "right": float(x1),
                            "bottom": float(y1),
                            "runs": runs,
                            "font_size": float(dominant_style.get("font_size", 11.0) or 11.0),
                            "is_bold": bool(dominant_style.get("bold", False)),
                        }
                    )

                if not line_records:
                    continue

                line_records.sort(
                    key=lambda row: (
                        round(float(row["top"]), 2),
                        round(float(row["left"]), 2),
                    )
                )

                paragraph_line_groups: List[List[Dict[str, Any]]] = []
                current_group: List[Dict[str, Any]] = []

                for line in line_records:
                    if not current_group:
                        current_group = [line]
                        continue

                    previous_line = current_group[-1]

                    should_start_new_paragraph = _workspace_export_should_start_new_paragraph(
                        previous_line,
                        line,
                        page_width=page_width,
                    )

                    if should_start_new_paragraph:
                        paragraph_line_groups.append(current_group)
                        current_group = [line]
                    else:
                        current_group.append(line)

                if current_group:
                    paragraph_line_groups.append(current_group)

                block_paragraphs: List[Dict[str, Any]] = []
                paragraph_previous_bottom = previous_visible_bottom

                for group_index, group in enumerate(paragraph_line_groups):
                    fallback_gap_before = (
                        float(group[0]["top"]) - float(block_top)
                        if group_index == 0 and previous_visible_bottom is None
                        else 0.0
                    )

                    paragraph = _workspace_export_finalize_paragraph(
                        group,
                        previous_bottom=paragraph_previous_bottom,
                        fallback_gap_before=fallback_gap_before,
                    )
                    if not paragraph or not _clean_text(paragraph.get("text")):
                        continue

                    block_paragraphs.append(paragraph)
                    paragraph_previous_bottom = float(paragraph["bottom"])

                if not block_paragraphs:
                    continue

                previous_visible_bottom = float(block_paragraphs[-1]["bottom"])

                page_blocks.append(
                    {
                        "block_index": block_index,
                        "left": block_left,
                        "top": block_top,
                        "paragraphs": block_paragraphs,
                    }
                )

            _workspace_export_annotate_page_layout(
                page_width=page_width,
                page_blocks=page_blocks,
            )

            exported_pages.append(
                {
                    "page_number": page_index,
                    "width": page_width,
                    "height": float(page.rect.height),
                    "blocks": page_blocks,
                }
            )

        return exported_pages
    finally:
        doc.close()

def _workspace_export_has_leading_bullet(text: str) -> bool:
    return str(text or "").lstrip().startswith(("●", "•", "▪", "◦", "·"))


def _workspace_export_preserve_bullet_prefix(
    original_text: str,
    patch_text: str,
) -> str:
    original = str(original_text or "")
    patched = str(patch_text or "").strip()

    if not patched:
        return patched

    if _workspace_export_has_leading_bullet(patched):
        return patched

    stripped = original.lstrip()
    leading_ws = original[: len(original) - len(stripped)]

    for marker in ("●", "•", "▪", "◦", "·"):
        if stripped.startswith(marker):
            return f"{leading_ws}{marker} {patched}"

    return patched


def _workspace_export_docx_has_merged_bullet_paragraphs(docx_path: Path) -> bool:
    from docx import Document

    document = Document(str(docx_path))

    for paragraph in list(document.paragraphs):
        text = _clean_text(paragraph.text)
        if not text:
            continue

        bullet_count = sum(text.count(marker) for marker in ("●", "•", "▪", "◦", "·"))
        if bullet_count >= 2:
            return True

    return False

def _workspace_export_match_score(source_text: str, candidate_text: str) -> int:
    source_norm = _normalize_tailoring_workspace_compare_text(source_text)
    candidate_norm = _normalize_tailoring_workspace_compare_text(candidate_text)

    if not source_norm or not candidate_norm:
        return 0

    if source_norm == candidate_norm:
        return 100000 + len(source_norm)

    if source_norm in candidate_norm or candidate_norm in source_norm:
        return 50000 + min(len(source_norm), len(candidate_norm))

    source_words = {word for word in source_norm.split(" ") if len(word) > 2}
    candidate_words = {word for word in candidate_norm.split(" ") if len(word) > 2}
    overlap = len(source_words & candidate_words)

    if not overlap:
        return 0

    return (overlap * 1000) - (abs(len(source_words) - len(candidate_words)) * 10)


def _apply_workspace_export_patch_specs(
    exported_pages: List[Dict[str, Any]],
    patch_specs: List[Dict[str, Any]],
) -> Dict[str, Any]:
    flat_targets: List[Tuple[int, int, int, Dict[str, Any]]] = []

    for page_index, page in enumerate(exported_pages):
        for block_index, block in enumerate(list(page.get("blocks", []) or [])):
            for paragraph_index, paragraph in enumerate(list(block.get("paragraphs", []) or [])):
                flat_targets.append((page_index, block_index, paragraph_index, paragraph))

    used_targets = set()
    applied_candidate_ids: List[str] = []
    unresolved_candidate_ids: List[str] = []

    for patch in list(patch_specs or []):
        source_raw_text = _clean_text(patch.get("source_raw_text"))
        patch_text = _clean_text(patch.get("patch_text"))
        candidate_id = _clean_text(patch.get("candidate_id"))

        if not source_raw_text or not patch_text:
            continue

        best_target: Tuple[int, int, int, Dict[str, Any]] | None = None
        best_score = 0

        for page_index, block_index, paragraph_index, paragraph in flat_targets:
            target_key = (page_index, block_index, paragraph_index)
            if target_key in used_targets:
                continue

            score = _workspace_export_match_score(
                source_raw_text,
                _clean_text(paragraph.get("text")),
            )
            if score > best_score:
                best_score = score
                best_target = (page_index, block_index, paragraph_index, paragraph)

        if best_target is None or best_score < 2000:
            if candidate_id:
                unresolved_candidate_ids.append(candidate_id)
            continue

        page_index, block_index, paragraph_index, paragraph = best_target
        updated = dict(paragraph)
        dominant_style = dict(updated.get("style", {}) or {})

        normalized_patch_text = _workspace_export_restore_original_lead_word(
            _clean_text(paragraph.get("text")),
            patch_text,
        )
        effective_patch_text = _workspace_export_preserve_bullet_prefix(
            _clean_text(paragraph.get("text")),
            normalized_patch_text,
        )

        updated["text"] = effective_patch_text
        updated["patched"] = True
        updated["patch_source"] = _clean_text(patch.get("patch_source"))
        updated["runs"] = _workspace_export_build_patched_runs(
            list(paragraph.get("runs", []) or []),
            effective_patch_text,
        )

        exported_pages[page_index]["blocks"][block_index]["paragraphs"][paragraph_index] = updated
        used_targets.add((page_index, block_index, paragraph_index))

        if candidate_id:
            applied_candidate_ids.append(candidate_id)

    return {
        "applied_candidate_ids": applied_candidate_ids,
        "unresolved_candidate_ids": unresolved_candidate_ids,
    }


def _workspace_export_apply_personal_details(
    exported_pages: List[Dict[str, Any]],
    personal_details: Any,
) -> bool:
    details = _normalize_workspace_personal_details(personal_details)
    if not any(details.values()) or not exported_pages:
        return False

    first_page = exported_pages[0]
    header_paragraphs: List[Dict[str, Any]] = []
    for block in list(first_page.get("blocks", []) or []):
        for paragraph in list(block.get("paragraphs", []) or []):
            if not _clean_text(paragraph.get("text")):
                continue
            if _clean_text(paragraph.get("alignment")).lower() != "center":
                continue
            header_paragraphs.append(paragraph)
            if len(header_paragraphs) >= 2:
                break
        if len(header_paragraphs) >= 2:
            break

    if not header_paragraphs:
        return False

    def _replace_paragraph(
        paragraph: Dict[str, Any],
        text: str,
        *,
        link_items: List[Dict[str, str]] | None = None,
    ) -> bool:
        clean_text = _workspace_export_clean_docx_text(_clean_text(text))
        if not clean_text:
            return False

        style = dict(paragraph.get("style", {}) or {})
        runs = list(paragraph.get("runs", []) or [])
        first_run = dict(runs[0]) if runs else {}

        paragraph["text"] = clean_text
        paragraph["runs"] = [
            {
                "text": clean_text,
                "font_name": _clean_text(first_run.get("font_name") or style.get("font_name")),
                "font_size": float(first_run.get("font_size", style.get("font_size", paragraph.get("font_size", 11.0))) or 11.0),
                "bold": bool(first_run.get("bold", style.get("bold", False))),
                "italic": bool(first_run.get("italic", style.get("italic", False))),
            }
        ]
        paragraph["patched"] = True
        paragraph["patch_source"] = "personal_details"
        if link_items is not None:
            paragraph["link_items"] = list(link_items or [])
        return True

    changed = False
    if _clean_text(details.get("name")):
        changed = _replace_paragraph(header_paragraphs[0], details["name"]) or changed

    contact_text = _workspace_personal_details_contact_text(details)
    if contact_text:
        if len(header_paragraphs) >= 2:
            contact_paragraph = header_paragraphs[1]
        else:
            contact_paragraph = dict(header_paragraphs[0])
            contact_paragraph["gap_before"] = 1.0
            contact_paragraph["font_size"] = 10.0
            first_page_blocks = list(first_page.get("blocks", []) or [])
            if not first_page_blocks:
                return changed
            first_page_blocks[0].setdefault("paragraphs", []).insert(1, contact_paragraph)

        changed = _replace_paragraph(
            contact_paragraph,
            contact_text,
            link_items=_workspace_personal_details_link_items(details),
        ) or changed

    return changed


def _workspace_export_is_section_heading_text(text: str) -> bool:
    clean = _clean_text(text)
    if not clean or len(clean) > 64:
        return False

    if clean.startswith(("●", "•", "▪", "◦", "·", "-")):
        return False

    letters_only = re.sub(r"[^A-Za-z]+", "", clean)
    if not letters_only:
        return False

    return clean == clean.upper()


def _workspace_export_apply_docx_section_rule(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph_xml = paragraph._p
    p_pr = paragraph_xml.get_or_add_pPr()

    existing_p_bdr = p_pr.find(qn("w:pBdr"))
    if existing_p_bdr is None:
        existing_p_bdr = OxmlElement("w:pBdr")
        p_pr.append(existing_p_bdr)

    for child in list(existing_p_bdr):
        if child.tag == qn("w:bottom"):
            existing_p_bdr.remove(child)

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "16")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "000000")
    existing_p_bdr.append(bottom)

def _workspace_export_pdf_contact_segments(
    contact_text: str,
    link_items: List[Dict[str, str]],
) -> List[Dict[str, str]]:
    segments: List[Dict[str, str]] = []
    normalized_links: List[Tuple[int, str, str]] = []

    for item in list(link_items or []):
        label = _clean_text(item.get("label"))
        uri = _clean_text(item.get("uri"))
        if not label or not uri:
            continue

        pos = contact_text.find(label)
        if pos < 0:
            continue

        normalized_links.append((pos, label, uri))

    normalized_links.sort(key=lambda item: item[0])

    if not normalized_links:
        return [{"kind": "text", "text": contact_text, "url": ""}]

    last_end = 0
    for pos, label, uri in normalized_links:
        if pos < last_end:
            continue

        prefix = contact_text[last_end:pos]
        if prefix:
            segments.append({"kind": "text", "text": prefix, "url": ""})

        segments.append({"kind": "link", "text": label, "url": uri})
        last_end = pos + len(label)

    suffix = contact_text[last_end:]
    if suffix:
        segments.append({"kind": "text", "text": suffix, "url": ""})

    return segments


def _workspace_export_draw_centered_pdf_contact_line(
    c,
    *,
    width: float,
    y: float,
    text: str,
    font_name: str,
    font_size: float,
    link_items: List[Dict[str, str]],
) -> float:
    from reportlab.pdfbase.pdfmetrics import stringWidth

    segments = _workspace_export_pdf_contact_segments(text, link_items)

    total_width = 0.0
    for segment in segments:
        total_width += stringWidth(
            str(segment.get("text", "") or ""),
            font_name,
            font_size,
        )

    x = max(36.0, (width - total_width) / 2.0)
    current_x = x

    for segment in segments:
        segment_text = str(segment.get("text", "") or "")
        if not segment_text:
            continue

        segment_width = stringWidth(segment_text, font_name, font_size)
        is_link = str(segment.get("kind", "")) == "link"
        url = _clean_text(segment.get("url"))

        if is_link and url:
            c.setFillColorRGB(5 / 255.0, 99 / 255.0, 193 / 255.0)
            c.setFont(font_name, font_size)
            c.drawString(current_x, y, segment_text)

            underline_y = y - 1.2
            c.setLineWidth(0.7)
            c.line(current_x, underline_y, current_x + segment_width, underline_y)

            c.linkURL(
                url,
                (
                    current_x,
                    y - (font_size * 0.35),
                    current_x + segment_width,
                    y + (font_size * 0.95),
                ),
                relative=0,
                thickness=0,
            )
            c.setFillColorRGB(0.0, 0.0, 0.0)
        else:
            c.setFont(font_name, font_size)
            c.drawString(current_x, y, segment_text)

        current_x += segment_width

    return total_width

def _build_workspace_export_pdf(
    exported_pages: List[Dict[str, Any]],
    output_path: Path,
    *,
    resume_pdf_path: Path | None = None,
) -> None:
    from reportlab.lib.utils import simpleSplit
    from reportlab.pdfbase.pdfmetrics import stringWidth
    from reportlab.pdfgen import canvas

    def _effective_font_size(
        paragraph: Dict[str, Any],
        *,
        page_index: int,
        page_emitted_paragraph_count: int,
        is_paired_row: bool = False,
    ) -> float:
        source_size = float(
            paragraph.get("font_size")
            or dict(paragraph.get("style", {}) or {}).get("font_size")
            or 11.0
        )

        alignment = _clean_text(paragraph.get("alignment")).lower()

        is_name_line = (
            page_index == 0
            and page_emitted_paragraph_count == 0
            and alignment == "center"
        )
        is_contact_line = (
            page_index == 0
            and page_emitted_paragraph_count == 1
            and alignment == "center"
        )

        return _workspace_export_docx_effective_font_size(
            source_size=source_size,
            is_name=is_name_line,
            is_contact=is_contact_line,
            is_heading=bool(paragraph.get("is_heading", False)),
            is_paired_row=is_paired_row,
        )

    def _left_content_x(paragraph: Dict[str, Any], width: float) -> float:
        return max(
            36.0,
            min(
                36.0 + float(paragraph.get("left_indent_pt") or 0.0),
                width - 144.0,
            ),
        )

    def _paragraph_text(paragraph: Dict[str, Any]) -> str:
        text = _workspace_export_clean_docx_text(
            _clean_text(paragraph.get("text"))
        )
        text = re.sub(r"\s+([,.;:])", r"\1", text)
        return text

    def _split_inline_label_prefix(text: str) -> Tuple[str, str]:
        clean = _clean_text(text)
        if ":" not in clean:
            return "", clean

        prefix, suffix = clean.split(":", 1)
        prefix = prefix.strip()
        suffix = suffix.lstrip()

        if not prefix:
            return "", clean

        if len(prefix) > 48:
            return "", clean

        return f"{prefix}:", suffix
    
    def _estimate_paragraph_block_height(
        paragraph: Dict[str, Any],
        *,
        width: float,
        page_index: int,
        page_emitted_paragraph_count: int,
    ) -> float:
        text = _paragraph_text(paragraph)
        if not text:
            return 0.0

        style = dict(paragraph.get("style", {}) or {})
        font_name = _workspace_export_pdf_font_name(style)
        font_size = _effective_font_size(
            paragraph,
            page_index=page_index,
            page_emitted_paragraph_count=page_emitted_paragraph_count,
            is_paired_row=False,
        )
        gap_before = max(0.0, min(float(paragraph.get("gap_before") or 0.0), 14.0))
        line_spacing = 1.0 if (
            page_index == 0
            and page_emitted_paragraph_count in {0, 1}
            and _clean_text(paragraph.get("alignment")).lower() == "center"
        ) else 1.08

        alignment = _clean_text(paragraph.get("alignment")).lower()
        is_header_line = (
            page_index == 0
            and page_emitted_paragraph_count in {0, 1}
            and alignment == "center"
        )
        tail_gap = 1.0 if is_header_line else 1.5

        is_label_value_line = (
            alignment == "left"
            and not bool(paragraph.get("is_heading"))
            and not bool(paragraph.get("is_bullet"))
            and ":" in text
        )

        if bool(paragraph.get("is_bullet")):
            body_text = _workspace_export_strip_leading_bullet_text(text).strip()
            content_x = _left_content_x(paragraph, width)
            body_x = content_x + 13.0
            available_width = max(120.0, width - body_x - 36.0)
            wrapped_lines = simpleSplit(body_text, font_name, font_size, available_width) or [body_text]
            line_height = font_size * 1.18
            bullet_tail_gap = 0.45
            return gap_before + (len(wrapped_lines) * line_height) + bullet_tail_gap

        if alignment in {"center", "right"}:
            available_width = max(120.0, width - 72.0)
        else:
            left_x = _left_content_x(paragraph, width)
            right_padding = 16.0 if is_label_value_line else 36.0
            available_width = max(120.0, width - left_x - right_padding)

        inline_label_prefix, inline_label_suffix = _split_inline_label_prefix(text)
        should_estimate_inline_label = (
            alignment == "left"
            and not bool(paragraph.get("is_heading"))
            and not bool(paragraph.get("is_bullet"))
            and bool(inline_label_prefix)
        )

        if should_estimate_inline_label:
            bold_style = dict(style)
            bold_style["bold"] = True
            prefix_font_name = _workspace_export_pdf_font_name(bold_style)
            prefix_text = inline_label_prefix + " "
            prefix_width = stringWidth(prefix_text, prefix_font_name, font_size)

            suffix_lines = _workspace_export_wrap_inline_label_value_text(
                inline_label_suffix,
                font_name=font_name,
                font_size=font_size,
                first_line_width=max(72.0, available_width - prefix_width),
                continuation_width=available_width,
            ) or [inline_label_suffix]

            line_height = font_size * 1.20
            return gap_before + (len(suffix_lines) * line_height) + tail_gap

        wrapped_lines = simpleSplit(text, font_name, font_size, available_width) or [text]
        line_height = font_size * (1.20 if is_label_value_line else line_spacing)
        extra_rule_space = 2.0 if _workspace_export_is_section_heading_text(text) else 0.0
        return gap_before + (len(wrapped_lines) * line_height) + extra_rule_space + tail_gap
    
    def _estimate_paired_row_height(
        left_item: Dict[str, Any],
        right_item: Dict[str, Any],
        *,
        page_index: int,
        page_emitted_paragraph_count: int,
    ) -> float:
        left_font_size = _effective_font_size(
            left_item,
            page_index=page_index,
            page_emitted_paragraph_count=page_emitted_paragraph_count,
            is_paired_row=True,
        )
        right_font_size = _effective_font_size(
            right_item,
            page_index=page_index,
            page_emitted_paragraph_count=page_emitted_paragraph_count,
            is_paired_row=True,
        )
        line_height = max(left_font_size, right_font_size) * 1.08
        gap_before = max(
            0.0,
            min(
                max(
                    float(left_item.get("gap_before") or 0.0),
                    float(right_item.get("gap_before") or 0.0),
                ),
                14.0,
            ),
        )
        paired_row_tail_gap = 1.0
        return gap_before + line_height + paired_row_tail_gap

    def _estimate_role_group_height(
        page_paragraphs: List[Dict[str, Any]],
        start_idx: int,
        *,
        width: float,
        page_index: int,
        page_emitted_paragraph_count: int,
    ) -> float:
        first = page_paragraphs[start_idx]
        row_group_id = _clean_text(first.get("row_group_id"))

        row_items = [first]
        scan = start_idx + 1
        while scan < len(page_paragraphs):
            candidate = page_paragraphs[scan]
            if _clean_text(candidate.get("row_group_id")) != row_group_id:
                break
            row_items.append(candidate)
            scan += 1

        left_item = next(
            (item for item in row_items if _clean_text(item.get("row_side")) == "left"),
            row_items[0],
        )
        right_item = next(
            (item for item in row_items if _clean_text(item.get("row_side")) == "right"),
            row_items[-1],
        )

        total = _estimate_paired_row_height(
            left_item,
            right_item,
            page_index=page_index,
            page_emitted_paragraph_count=page_emitted_paragraph_count,
        )

        while scan < len(page_paragraphs):
            candidate = page_paragraphs[scan]
            candidate_text = _paragraph_text(candidate)
            candidate_row_kind = _clean_text(candidate.get("row_kind"))

            if candidate_row_kind == "paired_row":
                break
            if _workspace_export_is_section_heading_text(candidate_text):
                break
            if not bool(candidate.get("is_bullet")):
                break

            total += _estimate_paragraph_block_height(
                candidate,
                width=width,
                page_index=page_index,
                page_emitted_paragraph_count=page_emitted_paragraph_count + (scan - start_idx),
            )
            scan += 1
            break

        return total

    c = canvas.Canvas(str(output_path))
    header_link_items = _workspace_export_extract_pdf_header_link_items(resume_pdf_path)

    first_page = exported_pages[0] if exported_pages else {"width": 612.0, "height": 792.0}
    width = max(612.0, float(first_page.get("width") or 0))
    height = max(792.0, float(first_page.get("height") or 0))
    c.setPageSize((width, height))

    y = height - 42.0
    page_emitted_paragraph_count = 0
    page_index = 0

    page_paragraphs: List[Dict[str, Any]] = []
    for source_page_index, page in enumerate(exported_pages):
        for block in list(page.get("blocks", []) or []):
            for paragraph in list(block.get("paragraphs", []) or []):
                paragraph_copy = dict(paragraph)
                paragraph_copy["_source_page_index"] = source_page_index
                page_paragraphs.append(paragraph_copy)

    page_paragraphs.sort(
        key=lambda paragraph: (
            int(paragraph.get("_source_page_index", 0) or 0),
            round(float(paragraph.get("top", 0.0) or 0.0), 2),
            round(float(paragraph.get("left", 0.0) or 0.0), 2),
            0 if _clean_text(paragraph.get("row_side")) == "left" else 1,
        )
    )

    idx = 0
    while idx < len(page_paragraphs):
            paragraph = page_paragraphs[idx]
            text = _paragraph_text(paragraph)
            if not text:
                idx += 1
                continue

            row_kind = _clean_text(paragraph.get("row_kind"))
            row_group_id = _clean_text(paragraph.get("row_group_id"))

            if row_kind == "paired_row" and row_group_id:
                role_group_height = _estimate_role_group_height(
                    page_paragraphs,
                    idx,
                    width=width,
                    page_index=page_index,
                    page_emitted_paragraph_count=page_emitted_paragraph_count,
                )
                fresh_page_capacity = height - 84.0

                if (
                    role_group_height <= fresh_page_capacity
                    and y - role_group_height < 42.0
                ):
                    c.showPage()
                    c.setPageSize((width, height))
                    y = height - 42.0

                row_items = [paragraph]
                scan = idx + 1

                while scan < len(page_paragraphs):
                    candidate = page_paragraphs[scan]
                    if _clean_text(candidate.get("row_group_id")) != row_group_id:
                        break
                    row_items.append(candidate)
                    scan += 1

                left_item = next(
                    (item for item in row_items if _clean_text(item.get("row_side")) == "left"),
                    row_items[0],
                )
                right_item = next(
                    (item for item in row_items if _clean_text(item.get("row_side")) == "right"),
                    row_items[-1],
                )

                left_style = dict(left_item.get("style", {}) or {})
                right_style = dict(right_item.get("style", {}) or {})

                left_font = _workspace_export_pdf_font_name(left_style)
                right_font = _workspace_export_pdf_font_name(right_style)

                left_font_size = _effective_font_size(
                    left_item,
                    page_index=page_index,
                    page_emitted_paragraph_count=page_emitted_paragraph_count,
                    is_paired_row=True,
                )
                right_font_size = _effective_font_size(
                    right_item,
                    page_index=page_index,
                    page_emitted_paragraph_count=page_emitted_paragraph_count,
                    is_paired_row=True,
                )

                line_height = max(left_font_size, right_font_size) * 1.08
                gap_before = max(
                    0.0,
                    min(
                        max(
                            float(left_item.get("gap_before") or 0.0),
                            float(right_item.get("gap_before") or 0.0),
                        ),
                        14.0,
                    ),
                )

                paired_row_tail_gap = 1.0
                if y - (gap_before + line_height + paired_row_tail_gap) < 42.0:
                    c.showPage()
                    c.setPageSize((width, height))
                    y = height - 42.0

                y -= gap_before

                left_text = _paragraph_text(left_item)
                right_text = _paragraph_text(right_item)

                left_x = _left_content_x(left_item, width)
                right_text_width = stringWidth(right_text, right_font, right_font_size)
                right_x = max(left_x + 120.0, width - 36.0 - right_text_width)

                c.setFont(left_font, left_font_size)
                c.drawString(left_x, y, left_text)

                c.setFont(right_font, right_font_size)
                c.drawString(right_x, y, right_text)

                paired_row_tail_gap = 1.0
                y -= line_height + paired_row_tail_gap
                page_emitted_paragraph_count += 1
                idx = scan
                continue

            style = dict(paragraph.get("style", {}) or {})
            font_name = _workspace_export_pdf_font_name(style)
            font_size = _effective_font_size(
                paragraph,
                page_index=page_index,
                page_emitted_paragraph_count=page_emitted_paragraph_count,
            )
            gap_before = max(0.0, min(float(paragraph.get("gap_before") or 0.0), 14.0))
            line_spacing = 1.0 if (
                page_index == 0
                and page_emitted_paragraph_count in {0, 1}
                and _clean_text(paragraph.get("alignment")).lower() == "center"
            ) else 1.08

            alignment = _clean_text(paragraph.get("alignment")).lower()
            is_header_line = (
                page_index == 0
                and page_emitted_paragraph_count in {0, 1}
                and alignment == "center"
            )
            tail_gap = 1.0 if is_header_line else 1.5

            is_label_value_line = (
                alignment == "left"
                and not bool(paragraph.get("is_heading"))
                and not bool(paragraph.get("is_bullet"))
                and ":" in text
            )

            if bool(paragraph.get("is_bullet")):
                body_text = _workspace_export_strip_leading_bullet_text(text).strip()
                content_x = _left_content_x(paragraph, width)
                body_x = content_x + 13.0
                bullet_center_x = content_x + 3.1
                available_width = max(120.0, width - body_x - 36.0)
                wrapped_lines = simpleSplit(body_text, font_name, font_size, available_width) or [body_text]
                line_height = font_size * 1.18
                bullet_tail_gap = 0.45
                block_height = gap_before + (len(wrapped_lines) * line_height) + bullet_tail_gap

                if y - block_height < 42.0:
                    c.showPage()
                    c.setPageSize((width, height))
                    y = height - 42.0

                y -= gap_before

                c.circle(
                    bullet_center_x,
                    y + (font_size * 0.27),
                    2.2,
                    stroke=0,
                    fill=1,
                )

                c.setFont(font_name, font_size)
                for line in wrapped_lines:
                    c.drawString(body_x, y, line)
                    y -= line_height

                y -= bullet_tail_gap
                page_emitted_paragraph_count += 1
                idx += 1
                continue

            is_centered_contact_line = (
                page_index == 0
                and page_emitted_paragraph_count == 1
                and alignment == "center"
                and _workspace_export_is_likely_contact_line(text)
            )

            if is_centered_contact_line:
                gap_before = min(gap_before, 0.5)
                line_height = font_size * line_spacing
                block_height = gap_before + line_height + 7.0

                if y - block_height < 42.0:
                    c.showPage()
                    c.setPageSize((width, height))
                    y = height - 42.0

                y -= gap_before

                _workspace_export_draw_centered_pdf_contact_line(
                    c,
                    width=width,
                    y=y,
                    text=text,
                    font_name=font_name,
                    font_size=font_size,
                    link_items=list(paragraph.get("link_items", []) or []) or header_link_items,
                )

                y -= line_height
                divider_y = y - 1.0
                c.setLineWidth(2.2)
                c.line(36.0, divider_y, width - 36.0, divider_y)
                y = divider_y - 5.0
                page_emitted_paragraph_count += 1
                idx += 1
                continue

            if alignment == "center":
                available_width = max(120.0, width - 72.0)
            elif alignment == "right":
                available_width = max(120.0, width - 72.0)
            else:
                left_x = _left_content_x(paragraph, width)
                right_padding = 16.0 if is_label_value_line else 36.0
                available_width = max(120.0, width - left_x - right_padding)

            inline_label_prefix, inline_label_suffix = _split_inline_label_prefix(text)
            should_draw_inline_label = (
                alignment == "left"
                and not bool(paragraph.get("is_heading"))
                and not bool(paragraph.get("is_bullet"))
                and bool(inline_label_prefix)
            )

            if should_draw_inline_label:
                bold_style = dict(style)
                bold_style["bold"] = True
                prefix_font_name = _workspace_export_pdf_font_name(bold_style)
                prefix_text = inline_label_prefix + " "
                prefix_width = stringWidth(prefix_text, prefix_font_name, font_size)

                suffix_lines = _workspace_export_wrap_inline_label_value_text(
                    inline_label_suffix,
                    font_name=font_name,
                    font_size=font_size,
                    first_line_width=max(72.0, available_width - prefix_width),
                    continuation_width=available_width,
                ) or [inline_label_suffix]

                line_height = font_size * 1.20
                block_height = gap_before + (len(suffix_lines) * line_height) + tail_gap

                if y - block_height < 42.0:
                    c.showPage()
                    c.setPageSize((width, height))
                    y = height - 42.0

                y -= gap_before

                x = _left_content_x(paragraph, width)
                c.setFont(prefix_font_name, font_size)
                c.drawString(x, y, inline_label_prefix)

                c.setFont(font_name, font_size)
                first_suffix = suffix_lines[0] if suffix_lines else ""
                if first_suffix:
                    c.drawString(x + prefix_width, y, first_suffix)

                y -= line_height

                for line in suffix_lines[1:]:
                    c.drawString(x, y, line)
                    y -= line_height

                y -= tail_gap
                page_emitted_paragraph_count += 1
                idx += 1
                continue

            wrapped_lines = simpleSplit(text, font_name, font_size, available_width) or [text]
            line_height = font_size * line_spacing
            block_height = gap_before + (len(wrapped_lines) * line_height) + tail_gap

            if y - block_height < 42.0:
                c.showPage()
                c.setPageSize((width, height))
                y = height - 42.0

            y -= gap_before

            for line in wrapped_lines:
                line_width = stringWidth(line, font_name, font_size)

                if alignment == "center":
                    x = max(36.0, (width - line_width) / 2.0)
                elif alignment == "right":
                    x = max(36.0, width - 36.0 - line_width)
                else:
                    x = _left_content_x(paragraph, width)

                c.setFont(font_name, font_size)
                c.drawString(x, y, line)
                y -= line_height

            if _workspace_export_is_section_heading_text(text):
                rule_y = y + 2.0
                c.setLineWidth(1.35)
                c.line(36.0, rule_y, width - 36.0, rule_y)
                y -= 2.0

            y -= tail_gap
            page_emitted_paragraph_count += 1
            idx += 1

    c.save()

def _workspace_export_clean_docx_text(value: Any) -> str:
    text = str(value or "")
    text = re.sub(r"[\u200b\u200c\u200d\ufeff]", "", text)
    text = text.replace("\xa0", " ")
    text = re.sub(r"\t+", "\t", text)
    text = re.sub(r" {2,}", " ", text)
    return text


def _workspace_export_docx_safe_font_name(raw_font_name: str = "") -> str:
    return "Arial"


def _workspace_export_docx_effective_font_size(
    *,
    source_size: float,
    is_name: bool = False,
    is_contact: bool = False,
    is_heading: bool = False,
    is_paired_row: bool = False,
) -> float:
    if is_name:
        return 14.0
    if is_contact:
        return 10.5
    if is_heading:
        return 10.5
    if is_paired_row:
        return 10.0
    return 10.0


def _workspace_export_configure_docx_section(section) -> None:
    from docx.shared import Inches

    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.42)
    section.right_margin = Inches(0.42)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)


def _workspace_export_configure_docx_document_defaults(document) -> None:
    from docx.shared import Pt

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11.0)

    for section in list(document.sections):
        _workspace_export_configure_docx_section(section)

def _build_workspace_export_docx(
    exported_pages: List[Dict[str, Any]],
    output_path: Path,
) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.shared import Pt

    def _append_runs(
        paragraph,
        runs: List[Dict[str, Any]],
        *,
        is_name: bool = False,
        is_contact: bool = False,
        is_heading: bool = False,
        is_paired_row: bool = False,
        trim_leading_whitespace: bool = False,
    ) -> None:
        first_visible_run = True

        for run_data in list(runs or []):
            run_text = _workspace_export_clean_docx_text(
                str(run_data.get("text", "") or "")
            )

            if trim_leading_whitespace and first_visible_run:
                run_text = run_text.lstrip()

            if not run_text.strip():
                continue

            run = paragraph.add_run(run_text)
            font = run.font
            font.name = _workspace_export_docx_safe_font_name(
                str(run_data.get("font_name", ""))
            )
            font.size = Pt(
                _workspace_export_docx_effective_font_size(
                    source_size=float(run_data.get("font_size", 11.0) or 11.0),
                    is_name=is_name,
                    is_contact=is_contact,
                    is_heading=is_heading,
                    is_paired_row=is_paired_row,
                )
            )
            font.bold = bool(run_data.get("bold", False)) or is_name
            font.italic = bool(run_data.get("italic", False)) and not is_contact
            first_visible_run = False

    document = Document()
    _workspace_export_configure_docx_document_defaults(document)
    use_seed_paragraph = True

    section = document.sections[0]
    usable_width_pt = float(
        section.page_width.pt - section.left_margin.pt - section.right_margin.pt
    )
    previous_emitted_row_kind = ""
    previous_paired_row_left_indent_pt = 0.0

    for page_index, page in enumerate(exported_pages):
        right_tab_stop_pt = max(360.0, usable_width_pt)
        page_emitted_paragraph_count = 0

        for block in list(page.get("blocks", []) or []):
            block_paragraphs = list(block.get("paragraphs", []) or [])
            idx = 0

            while idx < len(block_paragraphs):
                paragraph_data = block_paragraphs[idx]
                text = _clean_text(paragraph_data.get("text"))
                if not text:
                    idx += 1
                    continue

                row_kind = _clean_text(paragraph_data.get("row_kind"))
                row_group_id = _clean_text(paragraph_data.get("row_group_id"))

                if row_kind == "paired_row" and row_group_id:
                    row_items = [paragraph_data]
                    scan = idx + 1

                    while scan < len(block_paragraphs):
                        candidate = block_paragraphs[scan]
                        if _clean_text(candidate.get("row_group_id")) != row_group_id:
                            break
                        row_items.append(candidate)
                        scan += 1

                    left_candidates = [
                        item for item in row_items
                        if _clean_text(item.get("row_side")) == "left"
                    ]
                    right_candidates = [
                        item for item in row_items
                        if _clean_text(item.get("row_side")) == "right"
                    ]

                    if len(row_items) < 2 or not left_candidates or not right_candidates:
                        left_item = None
                        right_item = None
                    else:
                        left_item = left_candidates[0]
                        right_item = right_candidates[-1]

                        if left_item is not None and right_item is not None:
                            if use_seed_paragraph and document.paragraphs:
                                paragraph = document.paragraphs[0]
                                use_seed_paragraph = False
                            else:
                                paragraph = document.add_paragraph()

                            paragraph_format = paragraph.paragraph_format
                            paired_row_left_indent_pt = max(
                                0.0,
                                min(float(left_item.get("left_indent_pt") or 0.0), 220.0),
                            )
                            paragraph_format.left_indent = Pt(paired_row_left_indent_pt)
                            paragraph_format.right_indent = Pt(0.0)
                            paragraph_format.first_line_indent = Pt(0.0)
                            paragraph_format.space_before = Pt(
                                max(
                                    0.0,
                                    min(
                                        max(
                                            float(left_item.get("gap_before") or 0.0),
                                            float(right_item.get("gap_before") or 0.0),
                                        ),
                                        14.0,
                                    ),
                                )
                            )
                            paragraph_format.space_after = Pt(0.0)
                            paragraph_format.line_spacing = 1.08
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            paragraph_format.tab_stops.add_tab_stop(
                                Pt(right_tab_stop_pt),
                                WD_TAB_ALIGNMENT.RIGHT,
                                WD_TAB_LEADER.SPACES,
                            )

                            _append_runs(
                                paragraph,
                                list(left_item.get("runs", []) or []),
                                is_paired_row=True,
                                trim_leading_whitespace=True,
                            )
                            paragraph.add_run("\t")
                            _append_runs(
                                paragraph,
                                list(right_item.get("runs", []) or []),
                                is_paired_row=True,
                                trim_leading_whitespace=True,
                            )

                            previous_paired_row_left_indent_pt = paired_row_left_indent_pt
                            page_emitted_paragraph_count += 1
                            previous_emitted_row_kind = "paired_row"
                            idx = scan
                            continue

                if use_seed_paragraph and document.paragraphs:
                    paragraph = document.paragraphs[0]
                    use_seed_paragraph = False
                else:
                    paragraph = document.add_paragraph()

                paragraph_format = paragraph.paragraph_format
                alignment = _clean_text(paragraph_data.get("alignment")).lower()

                followup_subline_after_paired_row = (
                    previous_emitted_row_kind == "paired_row"
                    and alignment == "left"
                    and not bool(paragraph_data.get("is_bullet"))
                    and not bool(paragraph_data.get("is_heading"))
                    and int(paragraph_data.get("line_count", 0) or 0) == 1
                )
                bullet_after_paired_row = (
                    previous_emitted_row_kind == "paired_row"
                    and alignment == "left"
                    and bool(paragraph_data.get("is_bullet"))
                )

                if alignment == "center":
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph_format.left_indent = Pt(0.0)
                    paragraph_format.right_indent = Pt(0.0)
                elif alignment == "right":
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph_format.left_indent = Pt(0.0)
                    paragraph_format.right_indent = Pt(0.0)
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph_format.right_indent = Pt(0.0)
                    if followup_subline_after_paired_row or bullet_after_paired_row:
                        paragraph_format.left_indent = Pt(previous_paired_row_left_indent_pt)
                    else:
                        paragraph_format.left_indent = Pt(
                            max(0.0, min(float(paragraph_data.get("left_indent_pt") or 0.0), 220.0))
                        )

                paragraph_format.space_before = Pt(
                    0.0
                    if followup_subline_after_paired_row
                    else max(0.0, min(float(paragraph_data.get("gap_before") or 0.0), 14.0))
                )
                paragraph_format.space_after = Pt(1.0 if paragraph_data.get("is_heading") else 0.0)
                paragraph_format.line_spacing = 1.08

                if paragraph_data.get("is_bullet"):
                    paragraph_format.first_line_indent = Pt(-10.0)
                else:
                    paragraph_format.first_line_indent = Pt(0.0)

                runs = list(paragraph_data.get("runs", []) or [])
                if not runs:
                    style = dict(paragraph_data.get("style", {}) or {})
                    runs = [
                        {
                            "text": text,
                            "font_name": _clean_text(style.get("font_name")),
                            "font_size": float(style.get("font_size", paragraph_data.get("font_size", 11.0)) or 11.0),
                            "bold": bool(style.get("bold", paragraph_data.get("is_heading", False))),
                            "italic": bool(style.get("italic", False)),
                        }
                    ]

                is_name_line = (
                    page_index == 0
                    and page_emitted_paragraph_count == 0
                    and alignment == "center"
                )
                is_contact_line = (
                    page_index == 0
                    and page_emitted_paragraph_count == 1
                    and alignment == "center"
                )

                link_items = list(paragraph_data.get("link_items", []) or [])
                if is_contact_line and link_items:
                    style = dict(runs[0] if runs else paragraph_data.get("style", {}) or {})
                    _workspace_export_rebuild_docx_contact_paragraph(
                        paragraph,
                        text,
                        font_name=_clean_text(style.get("font_name")),
                        font_size_pt=_workspace_export_docx_effective_font_size(
                            source_size=float(style.get("font_size", paragraph_data.get("font_size", 10.0)) or 10.0),
                            is_contact=True,
                        ),
                        link_items=link_items,
                    )
                else:
                    _append_runs(
                        paragraph,
                        runs,
                        is_name=is_name_line,
                        is_contact=is_contact_line,
                        is_heading=bool(paragraph_data.get("is_heading", False)),
                        trim_leading_whitespace=True,
                    )

                if _workspace_export_is_section_heading_text(text):
                    paragraph.paragraph_format.space_after = Pt(2.0)
                    _workspace_export_apply_docx_section_rule(paragraph)

                page_emitted_paragraph_count += 1
                previous_emitted_row_kind = row_kind or "paragraph"
                idx += 1

    document.save(str(output_path))

def _workspace_export_is_likely_name_line(text: str) -> bool:
    clean = _clean_text(text)
    if not clean or len(clean) > 80:
        return False

    if any(
        token in clean.lower()
        for token in ("@", "linkedin", "github", "http://", "https://", ".com", "|")
    ):
        return False

    words = [part for part in clean.split() if part]
    if len(words) < 2 or len(words) > 6:
        return False

    return all(part[:1].isupper() for part in words if part[:1].isalpha())


def _workspace_export_is_likely_contact_line(text: str) -> bool:
    clean = _clean_text(text).lower()
    if not clean:
        return False

    if "|" in clean:
        return True

    if "@" in clean:
        return True

    if "linkedin" in clean or "github" in clean or ".com" in clean:
        return True

    return bool(re.search(r"\+?\d[\d(). \-]{6,}\d", clean))


def _workspace_export_split_merged_header_contact_text(text: str) -> Tuple[str, str]:
    clean = _workspace_export_clean_docx_text(text)
    clean = _clean_text(clean)
    if not clean:
        return "", ""

    pipe_index = clean.find("|")
    if pipe_index >= 0:
        prefix = clean[:pipe_index].rstrip()

        for location_pattern in (
            r"[A-Z][a-z]+ [A-Z][a-z]+ [A-Z][a-z]+,\s*[A-Z]{2}\s*$",
            r"[A-Z][a-z]+ [A-Z][a-z]+,\s*[A-Z]{2}\s*$",
            r"[A-Z][a-z]+,\s*[A-Z]{2}\s*$",
        ):
            location_match = re.search(location_pattern, prefix)
            if not location_match:
                continue

            left = _clean_text(prefix[:location_match.start()])
            right = _clean_text(clean[location_match.start():])

            if not left or not right:
                continue
            if not _workspace_export_is_likely_name_line(left):
                continue
            if not _workspace_export_is_likely_contact_line(right):
                continue

            return left, right

    fallback_match = re.search(
        r"(?:[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}.*|"
        r"\+?\d[\d(). \-]{6,}\d.*|"
        r"linkedin.*|github.*)",
        clean,
        flags=re.IGNORECASE,
    )
    if not fallback_match:
        return "", ""

    left = _clean_text(clean[:fallback_match.start()])
    right = _clean_text(clean[fallback_match.start():])

    if not left or not right:
        return "", ""
    if not _workspace_export_is_likely_name_line(left):
        return "", ""
    if not _workspace_export_is_likely_contact_line(right):
        return "", ""

    return left, right

def _workspace_export_is_bullet_paragraph_text(text: str) -> bool:
    return str(text or "").lstrip().startswith(("●", "•", "▪", "◦", "·"))


def _workspace_export_extract_pdf_header_link_items(
    resume_pdf_path: Path | None,
) -> List[Dict[str, str]]:
    import pymupdf as fitz

    if resume_pdf_path is None or not Path(resume_pdf_path).exists():
        return []

    doc = fitz.open(str(resume_pdf_path))
    try:
        if len(doc) == 0:
            return []

        page = doc[0]
        words = sorted(
            list(page.get_text("words") or []),
            key=lambda item: (round(float(item[1]), 2), round(float(item[0]), 2)),
        )

        out: List[Dict[str, str]] = []
        seen = set()

        for link in list(page.get_links() or []):
            uri = _clean_text(link.get("uri"))
            rect_raw = link.get("from")
            if not uri or rect_raw is None:
                continue

            rect = fitz.Rect(rect_raw)
            label_words: List[Tuple[float, str]] = []

            for word in words:
                word_rect = fitz.Rect(word[:4])
                if not word_rect.intersects(rect):
                    continue
                label_words.append((float(word[0]), str(word[4] or "")))

            label = " ".join(text for _, text in sorted(label_words)).strip()
            if not label:
                continue

            key = (label.lower(), uri)
            if key in seen:
                continue
            seen.add(key)

            out.append(
                {
                    "label": label,
                    "uri": uri,
                }
            )

        return out
    finally:
        doc.close()


def _workspace_export_append_plain_docx_run(
    paragraph,
    text: str,
    *,
    font_name: str,
    font_size_pt: float,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
) -> None:
    from docx.shared import Pt

    if not text:
        return

    run = paragraph.add_run(text)
    font = run.font
    font.name = _workspace_export_docx_font_name(font_name)
    font.size = Pt(font_size_pt)
    font.bold = bool(bold)
    font.italic = bool(italic)
    run.underline = bool(underline)


def _workspace_export_add_docx_hyperlink(
    paragraph,
    text: str,
    url: str,
    *,
    font_name: str,
    font_size_pt: float,
    bold: bool = False,
    italic: bool = False,
) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if not text or not url:
        return

    relationship_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
    r_id = paragraph.part.relate_to(url, relationship_type, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    resolved_font_name = _workspace_export_docx_font_name(font_name)

    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), resolved_font_name)
    r_fonts.set(qn("w:hAnsi"), resolved_font_name)
    r_fonts.set(qn("w:cs"), resolved_font_name)
    r_pr.append(r_fonts)

    half_points = str(int(round(float(font_size_pt) * 2.0)))

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), half_points)
    r_pr.append(sz)

    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), half_points)
    r_pr.append(sz_cs)

    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    if bold:
        b = OxmlElement("w:b")
        r_pr.append(b)

    if italic:
        i = OxmlElement("w:i")
        r_pr.append(i)

    run.append(r_pr)

    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _workspace_export_rebuild_docx_contact_paragraph(
    paragraph,
    contact_text: str,
    *,
    font_name: str,
    font_size_pt: float,
    link_items: List[Dict[str, str]],
) -> None:
    _workspace_export_clear_docx_paragraph_content(paragraph)

    normalized_links: List[Tuple[int, str, str]] = []
    for item in list(link_items or []):
        label = _clean_text(item.get("label"))
        uri = _clean_text(item.get("uri"))
        if not label or not uri:
            continue

        pos = contact_text.find(label)
        if pos < 0:
            continue

        normalized_links.append((pos, label, uri))

    normalized_links.sort(key=lambda item: item[0])

    if not normalized_links:
        _workspace_export_append_plain_docx_run(
            paragraph,
            contact_text,
            font_name=font_name,
            font_size_pt=font_size_pt,
            bold=False,
            italic=False,
            underline=False,
        )
        return

    last_end = 0
    for pos, label, uri in normalized_links:
        if pos < last_end:
            continue

        prefix = contact_text[last_end:pos]
        if prefix:
            _workspace_export_append_plain_docx_run(
                paragraph,
                prefix,
                font_name=font_name,
                font_size_pt=font_size_pt,
                bold=False,
                italic=False,
                underline=False,
            )

        _workspace_export_add_docx_hyperlink(
            paragraph,
            label,
            uri,
            font_name=font_name,
            font_size_pt=font_size_pt,
            bold=False,
            italic=False,
        )
        last_end = pos + len(label)

    suffix = contact_text[last_end:]
    if suffix:
        _workspace_export_append_plain_docx_run(
            paragraph,
            suffix,
            font_name=font_name,
            font_size_pt=font_size_pt,
            bold=False,
            italic=False,
            underline=False,
        )


def _workspace_export_normalize_docx_first_bullet_indent(docx_path: Path) -> bool:
    from docx import Document
    from docx.shared import Pt

    document = Document(str(docx_path))
    paragraphs = [
        paragraph
        for paragraph in list(document.paragraphs)
        if _clean_text(paragraph.text)
    ]

    changed = False

    for idx in range(len(paragraphs) - 2):
        lead = paragraphs[idx]
        first_bullet = paragraphs[idx + 1]
        second_bullet = paragraphs[idx + 2]

        if _workspace_export_is_bullet_paragraph_text(lead.text):
            continue
        if not _workspace_export_is_bullet_paragraph_text(first_bullet.text):
            continue
        if not _workspace_export_is_bullet_paragraph_text(second_bullet.text):
            continue

        first_left = (
            float(first_bullet.paragraph_format.left_indent.pt)
            if first_bullet.paragraph_format.left_indent is not None
            else 0.0
        )
        second_left = (
            float(second_bullet.paragraph_format.left_indent.pt)
            if second_bullet.paragraph_format.left_indent is not None
            else 0.0
        )

        if second_left - first_left < 4.0:
            continue

        first_bullet.paragraph_format.left_indent = Pt(second_left)

        if second_bullet.paragraph_format.first_line_indent is not None:
            first_bullet.paragraph_format.first_line_indent = Pt(
                float(second_bullet.paragraph_format.first_line_indent.pt)
            )

        if second_bullet.paragraph_format.right_indent is not None:
            first_bullet.paragraph_format.right_indent = Pt(
                float(second_bullet.paragraph_format.right_indent.pt)
            )

        changed = True

    if changed:
        document.save(str(docx_path))

    return changed

def _workspace_export_normalize_docx_bootstrap_header(
    docx_path: Path,
    *,
    resume_pdf_path: Path | None = None,
    link_items_override: List[Dict[str, str]] | None = None,
) -> bool:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document(str(docx_path))
    non_empty_paragraphs = [
        paragraph
        for paragraph in list(document.paragraphs)
        if _clean_text(paragraph.text)
    ]

    if not non_empty_paragraphs:
        return False

    style_source = non_empty_paragraphs[0]

    if (
        len(non_empty_paragraphs) >= 2
        and _workspace_export_is_likely_name_line(non_empty_paragraphs[0].text)
        and _workspace_export_is_likely_contact_line(non_empty_paragraphs[1].text)
    ):
        name_paragraph = non_empty_paragraphs[0]
        contact_paragraph = non_empty_paragraphs[1]
        name_text = _clean_text(name_paragraph.text)
        contact_text = _clean_text(contact_paragraph.text)
    else:
        merged_paragraph = non_empty_paragraphs[0]
        name_text, contact_text = _workspace_export_split_merged_header_contact_text(
            merged_paragraph.text
        )
        if not name_text or not contact_text:
            return False

        name_paragraph = merged_paragraph.insert_paragraph_before("")
        contact_paragraph = merged_paragraph

    style = _workspace_export_docx_first_run_style(style_source)
    resolved_font_name = str(style.get("font_name", "") or "")

    _workspace_export_clear_docx_paragraph_content(name_paragraph)
    _workspace_export_append_plain_docx_run(
        name_paragraph,
        _workspace_export_clean_docx_text(name_text),
        font_name=resolved_font_name,
        font_size_pt=14.0,
        bold=True,
        italic=False,
        underline=False,
    )
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_format = name_paragraph.paragraph_format
    name_format.left_indent = Pt(0.0)
    name_format.right_indent = Pt(0.0)
    name_format.first_line_indent = Pt(0.0)
    name_format.space_before = Pt(0.0)
    name_format.space_after = Pt(2.0)
    name_format.line_spacing = 1.0

    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_format = contact_paragraph.paragraph_format
    contact_format.left_indent = Pt(0.0)
    contact_format.right_indent = Pt(0.0)
    contact_format.first_line_indent = Pt(0.0)
    contact_format.space_before = Pt(0.0)
    contact_format.space_after = Pt(8.0)
    contact_format.line_spacing = 1.0

    link_items = (
        list(link_items_override or [])
        if link_items_override is not None
        else _workspace_export_extract_pdf_header_link_items(resume_pdf_path)
    )
    _workspace_export_rebuild_docx_contact_paragraph(
        contact_paragraph,
        _workspace_export_clean_docx_text(contact_text),
        font_name=resolved_font_name,
        font_size_pt=10.0,
        link_items=link_items,
    )

    document.save(str(docx_path))
    return True

def _workspace_export_docx_first_run_style(paragraph) -> Dict[str, Any]:
    for run in list(paragraph.runs):
        if not _clean_text(run.text):
            continue

        font = run.font
        font_size = 11.0
        if font.size is not None:
            try:
                font_size = float(font.size.pt)
            except Exception:
                font_size = 11.0

        return {
            "font_name": _clean_text(font.name),
            "font_size": font_size,
            "bold": bool(run.bold if run.bold is not None else font.bold),
            "italic": bool(run.italic if run.italic is not None else font.italic),
            "underline": bool(run.underline) if run.underline is not None else False,
        }

    return {
        "font_name": "",
        "font_size": 11.0,
        "bold": False,
        "italic": False,
        "underline": False,
    }


def _workspace_export_clear_docx_paragraph_content(paragraph) -> None:
    paragraph_xml = paragraph._p
    for child in list(paragraph_xml):
        if child.tag.endswith("}pPr"):
            continue
        paragraph_xml.remove(child)

def _workspace_export_is_bullet_only_run_text(text: str) -> bool:
    compact = re.sub(r"[\s\u200b\u200c\u200d\ufeff]+", "", str(text or ""))
    return compact in {"●", "•", "▪", "◦", "·"}


def _workspace_export_strip_leading_bullet_text(text: str) -> str:
    value = str(text or "")

    value = re.sub(r"^[\s\u200b\u200c\u200d\ufeff]+", "", value)
    value = re.sub(r"^[●•▪◦·][\s\u200b\u200c\u200d\ufeff]*", "", value, count=1)
    value = re.sub(r"^[\s\u200b\u200c\u200d\ufeff]+", "", value)

    return value

def _workspace_export_wrap_inline_label_value_text(
    text: str,
    *,
    font_name: str,
    font_size: float,
    first_line_width: float,
    continuation_width: float,
) -> List[str]:
    from reportlab.pdfbase.pdfmetrics import stringWidth

    clean = _clean_text(text)
    if not clean:
        return []

    words = clean.split()
    if not words:
        return [clean]

    def _wrap_words(source_words: List[str], max_width: float) -> Tuple[str, List[str]]:
        if not source_words:
            return "", []

        line_words: List[str] = []
        idx = 0

        while idx < len(source_words):
            candidate_words = [*line_words, source_words[idx]]
            candidate_text = " ".join(candidate_words)
            candidate_width = stringWidth(candidate_text, font_name, font_size)

            if line_words and candidate_width > max_width:
                break

            line_words = candidate_words
            idx += 1

            if candidate_width > max_width:
                break

        if not line_words:
            return source_words[0], source_words[1:]

        return " ".join(line_words), source_words[idx:]

    lines: List[str] = []

    first_width = max(72.0, float(first_line_width or 0.0))
    continuation_width = max(72.0, float(continuation_width or 0.0))

    first_line, remaining_words = _wrap_words(words, first_width)
    if first_line:
        lines.append(first_line)

    while remaining_words:
        next_line, remaining_words = _wrap_words(remaining_words, continuation_width)
        if not next_line:
            break
        lines.append(next_line)

    return lines or [clean]

def _workspace_export_restore_original_lead_word(
    original_text: str,
    patch_text: str,
) -> str:
    original_body = _workspace_export_strip_leading_bullet_text(original_text).strip()
    patch_body = _workspace_export_strip_leading_bullet_text(patch_text).strip()

    if not original_body or not patch_body:
        return patch_body or str(patch_text or "")

    if not re.match(r"^[a-z]", patch_body):
        return patch_body

    original_match = re.match(r"^([A-Z][a-z]+)\b", original_body)
    if not original_match:
        return patch_body

    original_lead = original_match.group(1)

    patch_match = re.match(r"^([A-Za-z]+)\b", patch_body)
    if patch_match and patch_match.group(1).lower() == original_lead.lower():
        return patch_body

    return f"{original_lead} {patch_body}"

def _workspace_export_run_style_payload(run_or_dict: Any) -> Dict[str, Any]:
    if isinstance(run_or_dict, dict):
        return {
            "font_name": _clean_text(run_or_dict.get("font_name")),
            "font_size": float(run_or_dict.get("font_size", 11.0) or 11.0),
            "bold": bool(run_or_dict.get("bold", False)),
            "italic": bool(run_or_dict.get("italic", False)),
            "underline": bool(run_or_dict.get("underline", False)),
        }

    font = run_or_dict.font
    font_size = 11.0
    if font.size is not None:
        try:
            font_size = float(font.size.pt)
        except Exception:
            font_size = 11.0

    return {
        "font_name": _clean_text(font.name),
        "font_size": font_size,
        "bold": bool(run_or_dict.bold if run_or_dict.bold is not None else font.bold),
        "italic": bool(run_or_dict.italic if run_or_dict.italic is not None else font.italic),
        "underline": bool(run_or_dict.underline) if run_or_dict.underline is not None else False,
    }


def _workspace_export_apply_docx_run_style(run, style: Dict[str, Any]) -> None:
    from docx.shared import Pt

    font = run.font
    font.name = _workspace_export_docx_font_name(str(style.get("font_name", "")))
    font.size = Pt(
        max(9.5, min(float(style.get("font_size", 11.0) or 11.0), 15.5))
    )
    font.bold = bool(style.get("bold", False))
    font.italic = bool(style.get("italic", False))
    run.underline = bool(style.get("underline", False))


def _workspace_export_build_patched_runs(
    original_runs: List[Dict[str, Any]],
    effective_patch_text: str,
) -> List[Dict[str, Any]]:
    normalized_runs = [
        {
            "text": str(run.get("text", "") or ""),
            "font_name": _clean_text(run.get("font_name")),
            "font_size": float(run.get("font_size", 11.0) or 11.0),
            "bold": bool(run.get("bold", False)),
            "italic": bool(run.get("italic", False)),
            "underline": bool(run.get("underline", False)),
        }
        for run in list(original_runs or [])
        if str(run.get("text", "") or "")
    ]

    if not normalized_runs:
        return [
            {
                "text": str(effective_patch_text or ""),
                "font_name": "",
                "font_size": 11.0,
                "bold": False,
                "italic": False,
                "underline": False,
            }
        ]

    first_run = normalized_runs[0]
    if not _workspace_export_is_bullet_only_run_text(first_run.get("text", "")):
        replacement = dict(first_run)
        replacement["text"] = str(effective_patch_text or "")
        return [replacement]

    body_style_source = next(
        (
            run
            for run in normalized_runs[1:]
            if _clean_text(run.get("text"))
        ),
        first_run,
    )

    stripped_patch_text = _workspace_export_strip_leading_bullet_text(
        effective_patch_text
    )

    bullet_run = dict(first_run)

    if not stripped_patch_text:
        return [bullet_run]

    body_run = dict(body_style_source)
    body_run["text"] = f" {stripped_patch_text}"
    return [bullet_run, body_run]

def _workspace_export_replace_docx_paragraph_text(paragraph, patch_text: str) -> None:
    normalized_patch_text = _workspace_export_restore_original_lead_word(
        paragraph.text,
        patch_text,
    )
    effective_patch_text = _workspace_export_preserve_bullet_prefix(
        paragraph.text,
        normalized_patch_text,
    )

    original_runs = [
        {
            "text": str(run.text or ""),
            **_workspace_export_run_style_payload(run),
        }
        for run in list(paragraph.runs)
        if str(run.text or "")
    ]

    patched_runs = _workspace_export_build_patched_runs(
        original_runs,
        effective_patch_text,
    )

    _workspace_export_clear_docx_paragraph_content(paragraph)

    for run_payload in patched_runs:
        run_text = _workspace_export_clean_docx_text(
            str(run_payload.get("text", "") or "")
        )
        if not run_text:
            continue

        run = paragraph.add_run(run_text)
        _workspace_export_apply_docx_run_style(run, run_payload)


def _apply_workspace_export_patch_specs_to_docx(
    docx_path: Path,
    patch_specs: List[Dict[str, Any]],
) -> Dict[str, Any]:
    from docx import Document

    document = Document(str(docx_path))

    table_count = len(list(document.tables))
    if table_count > 0:
        raise ValueError(
            f"pdf2docx bootstrap produced {table_count} table(s); falling back to native DOCX export for ATS safety."
        )

    paragraphs = [
        paragraph
        for paragraph in list(document.paragraphs)
        if _clean_text(paragraph.text)
    ]

    used_paragraph_indexes = set()
    applied_candidate_ids: List[str] = []
    unresolved_candidate_ids: List[str] = []

    for patch in list(patch_specs or []):
        source_raw_text = _clean_text(patch.get("source_raw_text"))
        patch_text = _clean_text(patch.get("patch_text"))
        candidate_id = _clean_text(patch.get("candidate_id"))

        if not source_raw_text or not patch_text:
            continue

        best_index = -1
        best_score = 0

        for idx, paragraph in enumerate(paragraphs):
            if idx in used_paragraph_indexes:
                continue

            score = _workspace_export_match_score(source_raw_text, paragraph.text)
            if score > best_score:
                best_score = score
                best_index = idx

        if best_index < 0 or best_score < 2000:
            if candidate_id:
                unresolved_candidate_ids.append(candidate_id)
            continue

        _workspace_export_replace_docx_paragraph_text(
            paragraphs[best_index],
            patch_text,
        )
        used_paragraph_indexes.add(best_index)

        if candidate_id:
            applied_candidate_ids.append(candidate_id)

    document.save(str(docx_path))

    return {
        "applied_candidate_ids": applied_candidate_ids,
        "unresolved_candidate_ids": unresolved_candidate_ids,
    }


def _build_workspace_export_docx_with_pdf2docx_bootstrap(
    *,
    resume_pdf_path: Path,
    patch_specs: List[Dict[str, Any]],
    output_path: Path,
) -> Dict[str, Any]:
    from pdf2docx import Converter

    temp_output_path: Path | None = None
    converter = None

    try:
        with tempfile.NamedTemporaryFile(
            prefix="workspace_export_bootstrap_",
            suffix=".docx",
            dir=str(output_path.parent),
            delete=False,
        ) as temp_file:
            temp_output_path = Path(temp_file.name)

        converter = Converter(str(resume_pdf_path))
        converter.convert(str(temp_output_path))

        patch_result = _apply_workspace_export_patch_specs_to_docx(
            temp_output_path,
            patch_specs,
        )

        _workspace_export_normalize_docx_bootstrap_header(
            temp_output_path,
            resume_pdf_path=resume_pdf_path,
        )

        if _workspace_export_docx_has_merged_bullet_paragraphs(temp_output_path):
            raise ValueError(
                "pdf2docx bootstrap merged multiple bullets into a single paragraph; falling back to native DOCX export."
            )

        temp_output_path.replace(output_path)

        return {
            "applied_candidate_ids": list(patch_result.get("applied_candidate_ids", []) or []),
            "unresolved_candidate_ids": list(patch_result.get("unresolved_candidate_ids", []) or []),
            "strategy": "pdf2docx_bootstrap",
        }
    finally:
        if converter is not None:
            try:
                converter.close()
            except Exception:
                pass

        if temp_output_path is not None and temp_output_path.exists():
            try:
                temp_output_path.unlink()
            except Exception:
                pass


def _workspace_export_header_link_items_from_pages(
    exported_pages: List[Dict[str, Any]],
) -> List[Dict[str, str]]:
    if not exported_pages:
        return []

    for block in list(exported_pages[0].get("blocks", []) or []):
        for paragraph in list(block.get("paragraphs", []) or []):
            link_items = list(paragraph.get("link_items", []) or [])
            if link_items:
                return link_items
    return []


def _workspace_export_has_personal_details_patch(
    exported_pages: List[Dict[str, Any]],
) -> bool:
    if not exported_pages:
        return False

    for block in list(exported_pages[0].get("blocks", []) or []):
        for paragraph in list(block.get("paragraphs", []) or []):
            if _clean_text(paragraph.get("patch_source")) == "personal_details":
                return True
    return False


def _workspace_export_find_soffice_binary() -> str:
    candidates = [
        _clean_text(os.environ.get("SOFFICE_BIN")),
        "soffice",
        "libreoffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]

    for candidate in candidates:
        if not candidate:
            continue

        candidate_path = Path(candidate)
        if candidate_path.exists():
            return str(candidate_path)

        resolved = shutil.which(candidate)
        if resolved:
            return resolved

    raise ValueError(
        "LibreOffice 'soffice' binary was not found. Install LibreOffice or set SOFFICE_BIN."
    )


def _build_workspace_export_finalized_docx(
    *,
    resume_pdf_path: Path,
    exported_pages: List[Dict[str, Any]],
    patch_specs: List[Dict[str, Any]],
    output_path: Path,
    base_patch_result: Dict[str, Any],
) -> Dict[str, Any]:
    patch_result = {
        "applied_candidate_ids": list(base_patch_result.get("applied_candidate_ids", []) or []),
        "unresolved_candidate_ids": list(base_patch_result.get("unresolved_candidate_ids", []) or []),
    }

    export_strategy = "native_docx"
    export_strategy_note = ""
    header_link_items = _workspace_export_header_link_items_from_pages(exported_pages)
    has_personal_details_patch = _workspace_export_has_personal_details_patch(exported_pages)

    try:
        if has_personal_details_patch:
            raise ValueError("Personal-details header edits require native DOCX export.")
        bootstrap_patch_result = _build_workspace_export_docx_with_pdf2docx_bootstrap(
            resume_pdf_path=resume_pdf_path,
            patch_specs=patch_specs,
            output_path=output_path,
        )
        patch_result = {
            "applied_candidate_ids": list(
                bootstrap_patch_result.get("applied_candidate_ids", []) or []
            ),
            "unresolved_candidate_ids": list(
                bootstrap_patch_result.get("unresolved_candidate_ids", []) or []
            ),
        }
        export_strategy = str(
            bootstrap_patch_result.get("strategy", "pdf2docx_bootstrap")
        )
    except Exception as exc:
        _build_workspace_export_docx(exported_pages, output_path)
        export_strategy = "native_docx_fallback"
        export_strategy_note = f"{exc.__class__.__name__}: {exc}"

    _workspace_export_normalize_docx_bootstrap_header(
        output_path,
        resume_pdf_path=resume_pdf_path,
        link_items_override=header_link_items if header_link_items else None,
    )
    _workspace_export_normalize_docx_first_bullet_indent(output_path)

    return {
        "patch_result": patch_result,
        "export_strategy": export_strategy,
        "export_strategy_note": export_strategy_note,
    }


def _convert_workspace_docx_to_pdf_with_soffice(
    *,
    docx_path: Path,
    output_path: Path,
) -> None:
    soffice_bin = _workspace_export_find_soffice_binary()

    with tempfile.TemporaryDirectory(
        prefix="workspace_export_pdf_convert_",
        dir=str(output_path.parent),
    ) as temp_dir_raw:
        temp_dir = Path(temp_dir_raw)
        profile_dir = temp_dir / "lo_profile"
        profile_dir.mkdir(parents=True, exist_ok=True)

        cmd = [
            soffice_bin,
            "--headless",
            f"-env:UserInstallation={profile_dir.resolve().as_uri()}",
            "--convert-to",
            "pdf:writer_pdf_Export",
            "--outdir",
            str(temp_dir),
            str(docx_path),
        ]

        completed = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            check=False,
        )

        if completed.returncode != 0:
            stderr = _clean_text(completed.stderr)
            stdout = _clean_text(completed.stdout)
            detail = stderr or stdout or "unknown soffice conversion failure"
            raise ValueError(f"LibreOffice PDF conversion failed: {detail}")

        converted_path = temp_dir / f"{docx_path.stem}.pdf"
        if not converted_path.exists():
            pdf_matches = list(temp_dir.glob("*.pdf"))
            if len(pdf_matches) == 1:
                converted_path = pdf_matches[0]
            else:
                raise ValueError("LibreOffice reported success but no converted PDF was found.")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        if output_path.exists():
            output_path.unlink()
        converted_path.replace(output_path)

def _build_tailoring_workspace_export_context(
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    *,
    tailoring_json_path: str = "",
    selected_resume: str = "",
    owner_user_id: str = "",
    selected_patch_candidate_ids: Any = None,
    manual_bullet_edits: Any = None,
    require_saved_draft: bool = False,
) -> Dict[str, Any]:
    draft_response = load_tailoring_workspace_draft_payload(
        output_dir=output_dir,
        tailoring_json_path=tailoring_json_path,
        selected_resume=selected_resume,
    )

    if require_saved_draft:
        if not bool(draft_response.get("has_saved_draft", False)):
            raise ValueError("Save a tailored draft before exporting.")
        if _clean_text(draft_response.get("draft_status")) != "saved":
            raise ValueError("Only a saved tailoring workspace draft can be exported.")

    draft = dict(draft_response.get("draft", {}) or {})
    artifact_path = _resolve_planning_artifact_path(
        tailoring_json_path,
        output_dir=output_dir,
    )

    payload_data = _load_tailoring_json_artifact(artifact_path)
    if artifact_path.name.endswith("__tailoring.json"):
        payload_data = _rehydrate_legacy_tailoring_operator_payload(
            artifact_path,
            payload_data,
        )

    if payload_data.get("replacement_candidates") is not None:
        payload_data = _apply_saved_patch_selection_overlay(
            artifact_path,
            payload_data,
        )

    effective_selected_resume = (
        _sanitize_optional_resume_filename(selected_resume)
        or _sanitize_optional_resume_filename(draft.get("selected_resume"))
    )
    if not effective_selected_resume:
        raise ValueError("Workspace draft is missing selected_resume.")

    effective_selected_candidate_ids = _normalize_selected_patch_candidate_ids(
        draft.get("selected_patch_candidate_ids", [])
        if selected_patch_candidate_ids is None
        else selected_patch_candidate_ids
    )
    if not effective_selected_candidate_ids:
        effective_selected_candidate_ids = _default_selected_candidate_ids_from_replacement_plan(
            payload_data
        )

    effective_manual_bullet_edits = _normalize_workspace_manual_bullet_edits(
        draft.get("manual_bullet_edits", {})
        if manual_bullet_edits is None
        else manual_bullet_edits
    )

    valid_candidate_ids = set(_tailoring_artifact_candidate_ids(payload_data))
    unknown_candidate_ids = [
        candidate_id
        for candidate_id in effective_selected_candidate_ids
        if candidate_id not in valid_candidate_ids
    ]
    if unknown_candidate_ids:
        raise ValueError(
            f"Unknown candidate IDs for this artifact: {', '.join(sorted(unknown_candidate_ids))}"
        )

    patch_specs, unresolved_manual_keys = _build_tailoring_workspace_effective_patch_specs(
        payload_data,
        selected_candidate_ids=effective_selected_candidate_ids,
        manual_bullet_edits=effective_manual_bullet_edits,
    )

    resume_pdf_path = planning_resume_preview_path(
        effective_selected_resume,
        owner_user_id=owner_user_id,
    )
    exported_pages = _extract_resume_pdf_paragraph_pages_for_export(resume_pdf_path)
    if not exported_pages:
        raise ValueError("Could not extract resume text for export.")

    _workspace_export_apply_personal_details(
        exported_pages,
        draft.get("personal_details", {}),
    )

    patch_result = _apply_workspace_export_patch_specs(
        exported_pages,
        patch_specs,
    )

    return {
        "draft_response": draft_response,
        "draft": draft,
        "artifact_path": artifact_path,
        "payload_data": payload_data,
        "selected_resume": effective_selected_resume,
        "selected_patch_candidate_ids": effective_selected_candidate_ids,
        "manual_bullet_edits": effective_manual_bullet_edits,
        "patch_specs": patch_specs,
        "unresolved_manual_edit_keys": list(unresolved_manual_keys or []),
        "resume_pdf_path": resume_pdf_path,
        "exported_pages": exported_pages,
        "patch_result": patch_result,
    }

def _workspace_export_preview_row_from_paragraph(
    paragraph: Dict[str, Any],
    *,
    text: str,
    gap_before: float | None = None,
) -> Dict[str, Any]:
    clean_text = _workspace_export_clean_docx_text(_clean_text(text))
    alignment = _clean_text(paragraph.get("alignment")).lower() or "left"

    return {
        "kind": "paragraph",
        "text": clean_text,
        "alignment": alignment,
        "gap_before": float(
            paragraph.get("gap_before") if gap_before is None else gap_before
        ),
        "left_indent_pt": float(paragraph.get("left_indent_pt") or 0.0),
        "is_heading": bool(paragraph.get("is_heading")),
        "is_section_heading": _workspace_export_is_section_heading_text(clean_text),
        "is_bullet": bool(paragraph.get("is_bullet")),
        "patched": bool(paragraph.get("patched")),
        "patch_source": _clean_text(paragraph.get("patch_source")),
        "link_items": list(paragraph.get("link_items", []) or []),
    }


def _workspace_export_split_skills_preview_rows(
    paragraph: Dict[str, Any],
) -> List[Dict[str, Any]]:
    def _build_split_rows(segment_texts: List[str]) -> List[Dict[str, Any]]:
        split_rows: List[Dict[str, Any]] = []

        for segment_index, segment_text in enumerate(segment_texts):
            clean_segment_text = _workspace_export_clean_docx_text(
                _clean_text(segment_text)
            )
            if not clean_segment_text or ":" not in clean_segment_text:
                return []

            split_rows.append(
                _workspace_export_preview_row_from_paragraph(
                    paragraph,
                    text=clean_segment_text,
                    gap_before=float(paragraph.get("gap_before") or 0.0)
                    if segment_index == 0
                    else 6.0,
                )
            )

        return split_rows

    raw_runs = list(paragraph.get("runs", []) or [])
    if raw_runs:
        merged_runs: List[Dict[str, Any]] = []
        for run in raw_runs:
            run_text = _workspace_export_clean_docx_text(
                str(run.get("text", "") or "")
            )
            if not run_text:
                continue

            is_bold = bool(run.get("bold", False))

            if merged_runs and bool(merged_runs[-1]["bold"]) == is_bold:
                merged_runs[-1]["text"] = f'{merged_runs[-1]["text"]}{run_text}'
            else:
                merged_runs.append({
                    "text": run_text,
                    "bold": is_bold,
                })

        segments: List[List[Dict[str, Any]]] = []
        current_segment: List[Dict[str, Any]] = []
        current_segment_has_value_text = False

        for run in merged_runs:
            run_text = _workspace_export_clean_docx_text(run["text"])
            is_bold = bool(run["bold"])

            starts_new_label = (
                is_bold
                and ":" in run_text
                and bool(current_segment)
                and current_segment_has_value_text
            )

            if starts_new_label:
                segments.append(current_segment)
                current_segment = [{
                    "text": run_text,
                    "bold": is_bold,
                }]
                current_segment_has_value_text = False
                continue

            current_segment.append({
                "text": run_text,
                "bold": is_bold,
            })

            if not is_bold and _clean_text(run_text):
                current_segment_has_value_text = True

        if current_segment:
            segments.append(current_segment)

        if len(segments) > 1:
            run_based_segment_texts = [
                _workspace_export_clean_docx_text(
                    "".join(str(item.get("text", "") or "") for item in segment)
                )
                for segment in segments
            ]
            run_based_rows = _build_split_rows(run_based_segment_texts)
            if run_based_rows:
                return run_based_rows

    clean_text = _workspace_export_clean_docx_text(
        _clean_text(paragraph.get("text"))
    )
    if not clean_text or ":" not in clean_text:
        return []

    label_pattern = re.compile(
        r"(?<!\S)([A-Z0-9][A-Za-z0-9+#/&().-]*(?:\s+(?:&|[A-Z0-9][A-Za-z0-9+#/&().-]*)){1,6}:)\s+"
    )
    matches = list(label_pattern.finditer(clean_text))

    if len(matches) <= 1:
        return []

    text_based_segment_texts: List[str] = []
    for idx, match in enumerate(matches):
        start = match.start()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(clean_text)
        segment_text = clean_text[start:end].strip()
        if segment_text:
            text_based_segment_texts.append(segment_text)

    return _build_split_rows(text_based_segment_texts)

def _workspace_export_split_header_like_preview_row(
    paragraph: Dict[str, Any],
) -> Dict[str, Any] | None:
    text = _workspace_export_clean_docx_text(
        _clean_text(paragraph.get("text"))
    )
    if not text:
        return None

    if bool(paragraph.get("is_bullet")):
        return None

    clean = re.sub(r"\s+", " ", text).strip()
    if len(clean) < 18:
        return None

    month = r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\.?"
    year = r"(?:19|20)\d{2}"
    point = rf"(?:{month}\s+{year}|{year}|Present|Current)"
    date_range_pattern = re.compile(
        rf"^(?P<left>.+?)\s+(?P<right>{point}(?:\s*[-–—]\s*{point})?)$",
        flags=re.IGNORECASE,
    )

    match = date_range_pattern.match(clean)
    if not match:
        return None

    left_text = _workspace_export_clean_docx_text(match.group("left"))
    right_text = _workspace_export_clean_docx_text(match.group("right"))

    if not left_text or not right_text:
        return None

    if len(right_text) > 48 or len(left_text) > 160:
        return None

    if ":" in left_text:
        return None

    if left_text.upper() == left_text and len(left_text) <= 40:
        return None

    paragraph_style = dict(paragraph.get("style", {}) or {})
    left_looks_like_header = (
        bool(paragraph.get("is_heading"))
        or bool(paragraph_style.get("bold", False))
        or float(paragraph.get("font_size", 0.0) or 0.0) >= 10.5
    )

    if not left_looks_like_header:
        return None

    return {
        "kind": "paired_row",
        "gap_before": float(paragraph.get("gap_before") or 0.0),
        "left_indent_pt": float(paragraph.get("left_indent_pt") or 0.0),
        "left_text": left_text,
        "right_text": right_text,
    }


def _workspace_export_section_supports_header_row_split(section_name: str) -> bool:
    normalized = _clean_text(section_name).upper()
    return normalized in {
        "PROFESSIONAL EXPERIENCE",
        "WORK EXPERIENCE",
        "EXPERIENCE",
        "EMPLOYMENT HISTORY",
        "EMPLOYMENT",
    }



def _workspace_export_line_is_date_range(text: str) -> bool:
    clean = re.sub(r"\s+", " ", _clean_text(text))
    if not clean or len(clean) > 48:
        return False

    month = r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\\.?"
    year = r"(?:19|20)\\d{2}"
    point = rf"(?:{month}\\s+{year}|{year}|Present|Current)"

    return bool(
        re.fullmatch(
            rf"{point}(?:\\s*[-–—]\\s*{point})?",
            clean,
            flags=re.IGNORECASE,
        )
    )


def _workspace_export_should_render_paired_row_inline(
    left_item: Dict[str, Any],
    right_item: Dict[str, Any],
    continuation_items: List[Dict[str, Any]],
    next_paragraph: Dict[str, Any] | None,
    *,
    current_section: str,
) -> bool:
    right_text = _clean_text(right_item.get("text"))
    if not _workspace_export_line_is_date_range(right_text):
        return False

    if _workspace_export_section_supports_header_row_split(current_section):
        return False

    left_text = _workspace_export_clean_docx_text(
        _clean_text(left_item.get("text"))
    )
    if not left_text:
        return False

    left_left = float(left_item.get("left", 0.0) or 0.0)
    left_font_size = float(left_item.get("font_size", 0.0) or 0.0)

    def _is_body_like_continuation(item: Dict[str, Any] | None) -> bool:
        if not item:
            return False

        if bool(item.get("is_bullet")):
            return False

        text = _workspace_export_clean_docx_text(
            _clean_text(item.get("text"))
        )
        if not text or _workspace_export_is_section_heading_text(text):
            return False

        style = dict(item.get("style", {}) or {})
        font_size = float(item.get("font_size", 0.0) or 0.0)

        if bool(item.get("is_heading")) or bool(style.get("bold", False)):
            return False

        if abs(float(item.get("left", 0.0) or 0.0) - left_left) > 96.0:
            return False

        if font_size > max(12.0, left_font_size + 0.5):
            return False

        return True

    grouped_continuation_exists = any(
        _is_body_like_continuation(item) for item in continuation_items
    )
    if grouped_continuation_exists:
        return True

    if next_paragraph is None:
        return False

    if not _is_body_like_continuation(next_paragraph):
        return False

    row_bottom = max(
        float(left_item.get("bottom", 0.0) or 0.0),
        float(right_item.get("bottom", 0.0) or 0.0),
    )
    next_top = float(next_paragraph.get("top", 0.0) or 0.0)
    gap_after_row = next_top - row_bottom

    return -2.0 <= gap_after_row <= 56.0

def _workspace_export_preview_pages_payload(
    exported_pages: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    pages_out: List[Dict[str, Any]] = []

    def _append_continuation_rows(
        rows: List[Dict[str, Any]],
        continuation_items: List[Dict[str, Any]],
    ) -> None:
        for extra in sorted(
            continuation_items,
            key=lambda item: (
                round(float(item.get("top", 0.0) or 0.0), 2),
                round(float(item.get("left", 0.0) or 0.0), 2),
            ),
        ):
            extra_text = _workspace_export_clean_docx_text(
                _clean_text(extra.get("text"))
            )
            if not extra_text:
                continue

            rows.append(
                _workspace_export_preview_row_from_paragraph(
                    extra,
                    text=extra_text,
                )
            )

    for page in list(exported_pages or []):
        page_width = float(page.get("width") or 612.0)
        page_height = float(page.get("height") or 792.0)

        page_paragraphs: List[Dict[str, Any]] = []
        for block in list(page.get("blocks", []) or []):
            for paragraph in list(block.get("paragraphs", []) or []):
                page_paragraphs.append(dict(paragraph))

        page_paragraphs.sort(
            key=lambda paragraph: (
                round(float(paragraph.get("top", 0.0) or 0.0), 2),
                round(float(paragraph.get("left", 0.0) or 0.0), 2),
                0 if _clean_text(paragraph.get("row_side")) == "left" else 1,
            )
        )

        rows: List[Dict[str, Any]] = []
        current_section = ""
        idx = 0

        while idx < len(page_paragraphs):
            paragraph = page_paragraphs[idx]
            text = _workspace_export_clean_docx_text(
                _clean_text(paragraph.get("text"))
            )
            if not text:
                idx += 1
                continue

            row_kind = _clean_text(paragraph.get("row_kind"))
            row_group_id = _clean_text(paragraph.get("row_group_id"))

            if row_kind == "paired_row" and row_group_id:
                row_items = [paragraph]
                scan = idx + 1
                while scan < len(page_paragraphs):
                    candidate = page_paragraphs[scan]
                    if _clean_text(candidate.get("row_group_id")) != row_group_id:
                        break
                    row_items.append(candidate)
                    scan += 1

                left_candidates = [
                    item for item in row_items
                    if _clean_text(item.get("row_side")) == "left"
                ]
                right_candidates = [
                    item for item in row_items
                    if _clean_text(item.get("row_side")) == "right"
                ]
                continuation_items = [
                    item for item in row_items
                    if _clean_text(item.get("row_side")) not in {"left", "right"}
                ]

                if len(row_items) < 2 or not left_candidates or not right_candidates:
                    rows.append(
                        _workspace_export_preview_row_from_paragraph(
                            paragraph,
                            text=text,
                        )
                    )
                    idx += 1
                    continue

                left_item = left_candidates[0]
                right_item = right_candidates[-1]

                left_text = _workspace_export_clean_docx_text(
                    _clean_text(left_item.get("text"))
                )
                right_text = _workspace_export_clean_docx_text(
                    _clean_text(right_item.get("text"))
                )

                keep_as_paired = (
                    _workspace_export_section_supports_header_row_split(current_section)
                    and _workspace_export_line_is_date_range(right_text)
                )

                if not keep_as_paired:
                    combined_text = _workspace_export_clean_docx_text(
                        f"{left_text} {right_text}"
                    )

                    rows.append(
                        _workspace_export_preview_row_from_paragraph(
                            left_item,
                            text=combined_text,
                            gap_before=float(
                                max(
                                    float(left_item.get("gap_before") or 0.0),
                                    float(right_item.get("gap_before") or 0.0),
                                )
                            ),
                        )
                    )

                    _append_continuation_rows(rows, continuation_items)
                    idx = scan
                    continue

                rows.append({
                    "kind": "paired_row",
                    "gap_before": float(
                        max(
                            float(left_item.get("gap_before") or 0.0),
                            float(right_item.get("gap_before") or 0.0),
                        )
                    ),
                    "left_indent_pt": float(left_item.get("left_indent_pt") or 0.0),
                    "left_text": left_text,
                    "right_text": right_text,
                })

                _append_continuation_rows(rows, continuation_items)
                idx = scan
                continue

            is_section_heading = _workspace_export_is_section_heading_text(text)

            if (
                current_section == "SKILLS"
                and not is_section_heading
                and not bool(paragraph.get("is_bullet"))
            ):
                split_skill_rows = _workspace_export_split_skills_preview_rows(paragraph)
                if split_skill_rows:
                    rows.extend(split_skill_rows)
                    idx += 1
                    continue

            split_header_row = None
            if _workspace_export_section_supports_header_row_split(current_section):
                split_header_row = _workspace_export_split_header_like_preview_row(
                    paragraph
                )

            if split_header_row:
                rows.append(split_header_row)
                if is_section_heading:
                    current_section = text.upper()
                idx += 1
                continue

            rows.append(
                _workspace_export_preview_row_from_paragraph(
                    paragraph,
                    text=text,
                )
            )

            if is_section_heading:
                current_section = text.upper()

            idx += 1

        pages_out.append({
            "page_number": int(page.get("page_number") or len(pages_out) + 1),
            "width": page_width,
            "height": page_height,
            "rows": rows,
        })

    return pages_out

def render_tailoring_workspace_draft_preview_payload(
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    *,
    tailoring_json_path: str = "",
    selected_resume: str = "",
    owner_user_id: str = "",
    selected_patch_candidate_ids: Any = None,
    manual_bullet_edits: Any = None,
) -> Dict[str, Any]:
    context = _build_tailoring_workspace_export_context(
        output_dir=output_dir,
        tailoring_json_path=tailoring_json_path,
        selected_resume=selected_resume,
        owner_user_id=owner_user_id,
        selected_patch_candidate_ids=selected_patch_candidate_ids,
        manual_bullet_edits=manual_bullet_edits,
        require_saved_draft=False,
    )

    pages_payload = _workspace_export_preview_pages_payload(
        context["exported_pages"]
    )

    return {
        "ok": True,
        "preview_status": "rendered",
        "draft_status": _clean_text(
            context["draft_response"].get("draft_status")
        ),
        "has_saved_draft": bool(
            context["draft_response"].get("has_saved_draft", False)
        ),
        "selected_resume": context["selected_resume"],
        "selected_patch_candidate_ids": list(
            context["selected_patch_candidate_ids"]
        ),
        "workspace_patch_count": len(context["patch_specs"]),
        "page_count": len(pages_payload),
        "pages": pages_payload,
        "unresolved_candidate_ids": list(
            context["patch_result"].get("unresolved_candidate_ids", []) or []
        ),
        "unresolved_manual_edit_keys": list(
            context["unresolved_manual_edit_keys"]
        ),
    }

def export_tailoring_workspace_draft_payload(
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    *,
    tailoring_json_path: str = "",
    selected_resume: str = "",
    owner_user_id: str = "",
    format: str = "pdf",
) -> Dict[str, Any]:
    export_format = _normalize_workspace_export_format(format)

    context = _build_tailoring_workspace_export_context(
        output_dir=output_dir,
        tailoring_json_path=tailoring_json_path,
        selected_resume=selected_resume,
        owner_user_id=owner_user_id,
        require_saved_draft=True,
    )

    draft_response = dict(context["draft_response"])
    draft = dict(context["draft"])
    patch_specs = list(context["patch_specs"])
    unresolved_manual_keys = list(context["unresolved_manual_edit_keys"])
    effective_selected_resume = _clean_text(context["selected_resume"])
    resume_pdf_path = Path(context["resume_pdf_path"])
    exported_pages = list(context["exported_pages"])
    patch_result = dict(context["patch_result"])

    output_path, filename, media_type = _workspace_export_output_path(
        output_dir,
        selected_resume=effective_selected_resume,
        export_format=export_format,
    )

    export_strategy = "native_pdf" if export_format == "pdf" else "native_docx"
    export_strategy_note = ""

    if export_format == "word":
        docx_build_result = _build_workspace_export_finalized_docx(
            resume_pdf_path=resume_pdf_path,
            exported_pages=exported_pages,
            patch_specs=patch_specs,
            output_path=output_path,
            base_patch_result=patch_result,
        )
        patch_result = dict(docx_build_result.get("patch_result", {}) or {})
        export_strategy = _clean_text(docx_build_result.get("export_strategy"))
        export_strategy_note = _clean_text(docx_build_result.get("export_strategy_note"))
    else:
        _build_workspace_export_pdf(
            exported_pages,
            output_path,
            resume_pdf_path=resume_pdf_path,
        )
        export_strategy = "native_pdf"
        export_strategy_note = ""

    unresolved_candidate_ids = list(patch_result.get("unresolved_candidate_ids", []) or [])
    unresolved_manual_edit_keys = list(unresolved_manual_keys or [])

    export_status = (
        "partial"
        if unresolved_candidate_ids or unresolved_manual_edit_keys
        else "complete"
    )

    warning_parts: List[str] = []
    if unresolved_candidate_ids:
        warning_parts.append(
            f"{len(unresolved_candidate_ids)} candidate mapping failure"
            f"{'' if len(unresolved_candidate_ids) == 1 else 's'}"
        )
    if unresolved_manual_edit_keys:
        warning_parts.append(
            f"{len(unresolved_manual_edit_keys)} manual edit key"
            f"{'' if len(unresolved_manual_edit_keys) == 1 else 's'} ignored"
        )

    warning_message = ""
    if warning_parts:
        warning_message = "Export completed with warnings: " + " and ".join(warning_parts) + "."

    return {
        "ok": True,
        "path": str(output_path),
        "filename": filename,
        "media_type": media_type,
        "format": export_format,
        "selected_resume": effective_selected_resume,
        "page_count": len(exported_pages),
        "workspace_patch_count": len(patch_specs),
        "export_status": export_status,
        "warning_message": warning_message,
        "export_strategy": export_strategy,
        "export_strategy_note": export_strategy_note,
        "applied_candidate_ids": patch_result.get("applied_candidate_ids", []),
        "unresolved_candidate_ids": unresolved_candidate_ids,
        "unresolved_manual_edit_keys": unresolved_manual_edit_keys,
    }

def record_application_action_payload(
    *,
    job_doc_id: str = "",
    job_url: str = "",
    job_company: str = "",
    job_title: str = "",
    application_status: str = "",
    source_view: str = "",
    note: str = "",
    owner_user_id: str = "",
) -> Dict[str, Any]:
    row = application_action_db_row(
        {
            "owner_user_id": _clean_text(owner_user_id),
            "action_timestamp": datetime.now(timezone.utc).isoformat(timespec="microseconds"),
            "job_doc_id": _clean_text(job_doc_id),
            "job_url": _clean_text(job_url),
            "job_company": _clean_text(job_company),
            "job_title": _clean_text(job_title),
            "application_status": _normalize_application_status(application_status),
            "source_view": _clean_text(source_view),
            "note": _clean_text(note),
        }
    )

    _validate_application_identity(row)
    postgres_write = _dual_write_application_action_postgres(row)

    return {
        "ok": True,
        "row": row,
        "overlay": _application_overlay_from_row(row),
        "postgres_write": postgres_write,
    }

def application_actions_payload(
    application_status: str = "",
    company_contains: str = "",
    title_contains: str = "",
    limit: int = 15,
    page: int = 1,
    owner_user_id: str = "",
) -> Dict[str, Any]:
    ja = _job_app()
    rows = _load_latest_application_actions(owner_user_id=owner_user_id)

    if application_status:
        status_target = _normalize_application_status(application_status)
        rows = [
            row for row in rows
            if _clean_text(row.get("application_status")) == status_target
        ]

    if company_contains:
        needle = ja._normalize_text(company_contains)
        rows = [
            row for row in rows
            if needle in ja._normalize_text(row.get("job_company", ""))
        ]

    if title_contains:
        needle = ja._normalize_text(title_contains)
        rows = [
            row for row in rows
            if needle in ja._normalize_text(row.get("job_title", ""))
        ]

    requested_limit = max(int(limit or 15), 1)
    current_page = max(int(page or 1), 1)
    page_size = 15

    rows = rows[:requested_limit]

    total_count = len(rows)
    total_pages = max((total_count + page_size - 1) // page_size, 1)
    current_page = min(current_page, total_pages)

    start_idx = (current_page - 1) * page_size
    end_idx = start_idx + page_size
    page_rows = rows[start_idx:end_idx]

    return {
        "filters": {
            "application_status": application_status,
            "company_contains": company_contains,
            "title_contains": title_contains,
            "limit": requested_limit,
            "page": current_page,
        },
        "rows": page_rows,
        "count": len(page_rows),
        "total_count": total_count,
        "page": current_page,
        "page_size": page_size,
        "total_pages": total_pages,
        "has_prev_page": current_page > 1,
        "has_next_page": current_page < total_pages,
    }

def applied_jobs_payload(
    company_contains: str = "",
    title_contains: str = "",
    limit: int = 15,
    page: int = 1,
    owner_user_id: str = "",
) -> Dict[str, Any]:
    return application_actions_payload(
        application_status="APPLIED",
        company_contains=company_contains,
        title_contains=title_contains,
        limit=limit,
        page=page,
        owner_user_id=owner_user_id,
    )

def jobs_search_lite_payload(
    request: str,
    top_k: int = 10,
) -> Dict[str, Any]:
    from src.rag.corpus_store import _load_job_corpus
    from src.rag.lexical_retriever import _lexical_search
    from src.rag.query_filters import _infer_metadata_filters

    inferred_filters = _infer_metadata_filters(request)
    lexical_results = _lexical_search(
        query=request,
        top_k=top_k,
        filters=inferred_filters or None,
    )

    compact_results = []
    for row in lexical_results:
        metadata = row.get("metadata", {}) or {}
        compact_results.append({
            "score": row.get("score"),
            "doc_id": metadata.get("doc_id", ""),
            "company": metadata.get("company", ""),
            "title": metadata.get("title", ""),
            "location": metadata.get("location", ""),
            "source": metadata.get("source", ""),
            "job_url": metadata.get("job_url", ""),
            "posted_at": metadata.get("posted_at", ""),
            "visa_sponsorship": metadata.get("visa_sponsorship", ""),
            "ai_fit_score": metadata.get("ai_fit_score"),
        })

    compact_results = _overlay_application_actions(compact_results)

    return {
        "ok": True,
        "request": request,
        "mode": "search_lite",
        "corpus_count": len(_load_job_corpus()),
        "inferred_filters": inferred_filters,
        "result_count": len(compact_results),
        "results": compact_results,
    }

def rag_search_payload(
    request: str,
    top_k: int = 5,
    fetch_k: int = 15,
    output_mode: str = "compact",
    include_diagnostics: bool = False,
) -> Dict[str, Any]:
    from src.rag.rag_executor import execute_rag_request

    payload = execute_rag_request(
        request=request,
        top_k=top_k,
        fetch_k=fetch_k,
        filters=None,
        output_mode=output_mode,
        include_diagnostics=include_diagnostics,
        intent_override="search_jobs",
    )

    if payload.get("ok") and isinstance(payload.get("response"), dict):
        response = dict(payload.get("response", {}))
        response["results"] = _overlay_application_actions(response.get("results", []) or [])
        payload = dict(payload)
        payload["response"] = response

    return payload

def rag_answer_payload(
    request: str,
    top_k: int = 5,
    fetch_k: int = 15,
    output_mode: str = "compact",
    include_diagnostics: bool = False,
) -> Dict[str, Any]:
    from src.rag.rag_executor import execute_rag_request

    payload = execute_rag_request(
        request=request,
        top_k=top_k,
        fetch_k=fetch_k,
        filters=None,
        output_mode=output_mode,
        include_diagnostics=include_diagnostics,
        intent_override="answer_job_query",
    )

    if payload.get("ok") and isinstance(payload.get("response"), dict):
        response = dict(payload.get("response", {}))
        response["sources"] = _overlay_application_actions(response.get("sources", []) or [])
        response["job_evidence"] = _overlay_application_actions(response.get("job_evidence", []) or [])
        payload = dict(payload)
        payload["response"] = response

    return payload
